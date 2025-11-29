"""
Format tools for effilocal MCP server.

This module extends word_document_server.tools.format_tools with:
1. Background highlighting with shading support
2. Run inspection utilities

Override Pattern:
- Import all upstream functions
- Add new formatting functions
"""

import os
from typing import Optional, Dict, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import all upstream functions (pass-through)
from word_document_server.tools.format_tools import (
    create_custom_style,
    format_text,
)

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension


# ============================================================================
# Helper Functions (shading and run inspection)
# ============================================================================

def _set_run_shading(run, fill_hex: Optional[str]):
    """
    Apply or remove background shading on a run via <w:shd>.
    
    NEW helper - applies background color via XML shading element.
    
    Args:
        run: The run to apply shading to
        fill_hex: e.g., '0000FF' or None to clear
    """
    rPr = run._element.get_or_add_rPr()
    shd = rPr.find(qn("w:shd"))
    if fill_hex:
        if shd is None:
            shd = OxmlElement("w:shd")
            rPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
    else:
        if shd is not None:
            rPr.remove(shd)


def _read_run_styles(run) -> Dict[str, Optional[str]]:
    """
    Read formatting properties from a run.
    
    NEW helper - extracts run formatting for inspection.
    
    Returns:
        Dictionary with run properties (bold, italic, font_name, font_size, color)
    """
    props = {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "font_name": run.font.name if run.font.name else None,
        "font_size": str(run.font.size) if run.font.size else None,
        "color": None,
    }
    
    # Extract color if present
    if run.font.color and run.font.color.rgb:
        props["color"] = str(run.font.color.rgb)
    
    return props


def _paragraph_runs_snapshot(paragraph) -> List[Dict[str, Optional[str]]]:
    """
    Capture a snapshot of all runs in a paragraph.
    
    NEW helper - used for run-level inspection and debugging.
    
    Returns:
        List of run property dictionaries
    """
    return [_read_run_styles(run) for run in paragraph.runs]


# ============================================================================
# NEW Functions (background highlighting)
# ============================================================================

async def set_background_highlight(
    filename: str,
    paragraph_index: int,
    start_pos: int,
    end_pos: int,
    color: str = "0000FF",
    use_shading: bool = True
) -> str:
    """
    Apply background to a span of text within a paragraph.
    
    NEW function - provides precise text highlighting with color control.
    
    - If use_shading=True: uses <w:shd> with a hex fill (e.g., '0000FF').
    - If use_shading=False: uses Word highlight palette via WD_COLOR_INDEX (e.g., 'YELLOW').
    
    Args:
        filename: Path to the Word document
        paragraph_index: 0-based paragraph index
        start_pos: Start character position in paragraph text
        end_pos: End character position (exclusive)
        color: Hex color code (e.g., '0000FF') or WD_COLOR_INDEX name
        use_shading: True for hex shading, False for highlight palette
    
    Returns:
        Human-readable status string
    """
    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        paragraph_index = int(paragraph_index)
        start_pos = int(start_pos)
        end_pos = int(end_pos)
    except (ValueError, TypeError):
        return "Invalid parameter: paragraph_index, start_pos, and end_pos must be integers"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)

        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."

        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text

        # Validate text positions
        if start_pos < 0 or end_pos > len(text) or start_pos > end_pos:
            return f"Invalid text positions. Paragraph has {len(text)} characters."

        # Rebuild runs as: [before][target][after]
        for run in paragraph.runs:
            run.clear()

        if start_pos > 0:
            paragraph.add_run(text[:start_pos])

        run_target = paragraph.add_run(text[start_pos:end_pos])

        if end_pos < len(text):
            paragraph.add_run(text[end_pos:])

        # Apply requested background mode
        if use_shading:
            # Normalize hex (accept '#RRGGBB' or 'RRGGBB')
            fill = color.strip().lstrip("#").upper()
            if len(fill) != 6 or any(c not in "0123456789ABCDEF" for c in fill):
                return f"Invalid hex color: {color}. Expected format: '0000FF' or '#0000FF'."
            _set_run_shading(run_target, fill)
        else:
            # Use Word's built-in highlight palette
            try:
                # Try to resolve color name to WD_COLOR_INDEX enum
                highlight_color = getattr(WD_COLOR_INDEX, color.upper(), None)
                if highlight_color is None:
                    return f"Invalid highlight color: {color}. Must be a valid WD_COLOR_INDEX name (e.g., 'YELLOW', 'BRIGHT_GREEN')."
                run_target.font.highlight_color = highlight_color
            except Exception as e:
                return f"Failed to apply highlight: {str(e)}"

        doc.save(filename)
        
        mode = "shading" if use_shading else "highlight"
        return f"Applied {mode} background (color={color}) to paragraph {paragraph_index}, positions {start_pos}:{end_pos} in {filename}"

    except Exception as e:
        return f"Failed to set background: {str(e)}"


def get_document_runs(filename: str, paragraph_index: int) -> Dict:
    """
    Return a JSON-serializable snapshot of all runs in a paragraph.
    
    NEW function - useful for debugging run-level formatting.
    
    Args:
        filename: Path to the Word document
        paragraph_index: 0-based paragraph index
    
    Returns:
        Dictionary with success flag and runs data
    """
    import json
    
    filename = ensure_docx_extension(filename)
    
    try:
        paragraph_index = int(paragraph_index)
    except (ValueError, TypeError):
        return {"success": False, "error": "paragraph_index must be an integer"}
    
    if not os.path.exists(filename):
        return {"success": False, "error": f"Document {filename} does not exist"}
    
    try:
        doc = Document(filename)
        
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return {
                "success": False,
                "error": f"Paragraph index {paragraph_index} out of range (0-{len(doc.paragraphs)-1})"
            }
        
        paragraph = doc.paragraphs[paragraph_index]
        runs_data = _paragraph_runs_snapshot(paragraph)
        
        return {
            "success": True,
            "paragraph_index": paragraph_index,
            "paragraph_text": paragraph.text,
            "run_count": len(runs_data),
            "runs": runs_data
        }
    
    except Exception as e:
        return {"success": False, "error": f"Failed to read runs: {str(e)}"}


# ============================================================================
# Re-export all functions for compatibility
# ============================================================================

__all__ = [
    # Pass-through functions (unchanged from upstream)
    'create_custom_style',
    'format_text',
    
    # New functions (effilocal-specific)
    'set_background_highlight',
    'get_document_runs',
]
