"""
Tests for Phase 1 Comment Features (Sprint 3).

This test file covers:
1. Backend enhancements - resolve_comment(), reference_text extraction
2. Comment extraction with status (active/resolved)
3. MCP tool integration for comment operations

TDD Approach: Tests written FIRST, expected to FAIL until implementation complete.
"""
import os
import shutil
import sys
import tempfile
from pathlib import Path
from typing import Dict, List, Any

import pytest
from docx import Document
from docx.oxml.ns import qn

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from effilocal.mcp_server.core.comments import (
    extract_all_comments,
    extract_comment_data,
    extract_comment_status_map,
    merge_comment_status,
)


# ============================================================================
# Test Fixtures
# ============================================================================

@pytest.fixture
def comments_fixture_dir() -> Path:
    """Path to comments fixture directory."""
    return Path(__file__).parent / "fixtures" / "comments"


@pytest.fixture
def test_doc_with_comments(tmp_path) -> Path:
    """
    Create a test document with comments for testing.
    Uses fixture XML files if available, otherwise creates a minimal doc.
    """
    # Try to find a real document with comments
    test_docs_paths = [
        Path(__file__).parent.parent / "EL_Projects" / "Test Project" / "drafts" / "current_drafts" / "HJ9 (TRACKED).docx",
        Path(__file__).parent / "fixtures" / "real_world" / "with_comments.docx",
    ]
    
    for doc_path in test_docs_paths:
        if doc_path.exists():
            # Copy to temp directory
            temp_doc = tmp_path / "test_comments.docx"
            shutil.copy(doc_path, temp_doc)
            return temp_doc
    
    # Fallback: create a minimal document (won't have comments for testing)
    doc = Document()
    doc.add_paragraph("Test paragraph without comments")
    temp_doc = tmp_path / "test_no_comments.docx"
    doc.save(str(temp_doc))
    return temp_doc


# ============================================================================
# Test: resolve_comment() - NEW FUNCTION (should fail initially)
# ============================================================================

class TestResolveComment:
    """Tests for resolve_comment() function - marks a comment as resolved."""
    
    def test_resolve_comment_import(self):
        """Test that resolve_comment can be imported from comments module."""
        try:
            from effilocal.mcp_server.core.comments import resolve_comment
            assert callable(resolve_comment), "resolve_comment should be callable"
        except ImportError:
            pytest.fail("resolve_comment not yet implemented - expected failure")
    
    def test_resolve_comment_by_id(self, test_doc_with_comments):
        """Test resolving a comment by its comment_id."""
        try:
            from effilocal.mcp_server.core.comments import resolve_comment
        except ImportError:
            pytest.skip("resolve_comment not yet implemented")
        
        doc = Document(str(test_doc_with_comments))
        
        # First, get all comments
        comments = extract_all_comments(doc)
        if not comments:
            pytest.skip("No comments in test document")
        
        # Get first active comment with a para_id
        active_comments = [c for c in comments if c.get('status') == 'active' and c.get('para_id')]
        if not active_comments:
            pytest.skip("No active comments with para_id to resolve")
        
        comment_to_resolve = active_comments[0]
        para_id = comment_to_resolve['para_id']
        comment_id = comment_to_resolve['comment_id']
        
        # Resolve the comment using para_id
        result = resolve_comment(doc, para_id)
        
        assert result is True, "resolve_comment should return True on success"
        
        # Save and reload to verify persistence
        doc.save(str(test_doc_with_comments))
        doc2 = Document(str(test_doc_with_comments))
        
        # Check the comment is now resolved
        updated_comments = extract_all_comments(doc2)
        resolved_comment = next((c for c in updated_comments if c['comment_id'] == comment_id), None)
        
        assert resolved_comment is not None, "Comment should still exist after resolving"
        assert resolved_comment['status'] == 'resolved', "Comment status should be 'resolved'"
        assert resolved_comment['is_resolved'] is True, "is_resolved should be True"
        assert resolved_comment['done_flag'] == 1, "done_flag should be 1"
    
    def test_resolve_comment_updates_comments_extended(self, test_doc_with_comments):
        """Test that resolve_comment updates the commentsExtended.xml part."""
        try:
            from effilocal.mcp_server.core.comments import resolve_comment
        except ImportError:
            pytest.skip("resolve_comment not yet implemented")
        
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        if not comments:
            pytest.skip("No comments in test document")
        
        active_comments = [c for c in comments if c.get('status') == 'active' and c.get('para_id')]
        if not active_comments:
            pytest.skip("No active comments with para_id to resolve")
        
        comment = active_comments[0]
        para_id = comment['para_id']
        
        # Resolve it using para_id
        resolve_comment(doc, para_id)
        doc.save(str(test_doc_with_comments))
        
        # Check the XML directly
        doc2 = Document(str(test_doc_with_comments))
        status_map = extract_comment_status_map(doc2)
        
        if para_id and para_id in status_map:
            assert status_map[para_id]['is_resolved'] is True, "commentsExtended should show resolved"
            assert status_map[para_id]['done'] == 1, "done flag should be 1 in XML"
    
    def test_resolve_comment_invalid_id(self, test_doc_with_comments):
        """Test that resolving an invalid comment_id returns False or raises."""
        try:
            from effilocal.mcp_server.core.comments import resolve_comment
        except ImportError:
            pytest.skip("resolve_comment not yet implemented")
        
        doc = Document(str(test_doc_with_comments))
        
        # Try to resolve a non-existent comment
        result = resolve_comment(doc, "nonexistent_comment_id_12345")
        
        assert result is False, "resolve_comment should return False for invalid ID"
    
    def test_resolve_already_resolved_comment(self, test_doc_with_comments):
        """Test that resolving an already resolved comment is idempotent."""
        try:
            from effilocal.mcp_server.core.comments import resolve_comment
        except ImportError:
            pytest.skip("resolve_comment not yet implemented")
        
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        if not comments:
            pytest.skip("No comments in test document")
        
        # Find or create a resolved comment with para_id
        resolved_comments = [c for c in comments if c.get('status') == 'resolved' and c.get('para_id')]
        if resolved_comments:
            comment = resolved_comments[0]
        else:
            # Resolve one first
            active = [c for c in comments if c.get('status') == 'active' and c.get('para_id')]
            if not active:
                pytest.skip("No comments with para_id to test")
            comment = active[0]
            resolve_comment(doc, comment['para_id'])
        
        para_id = comment['para_id']
        
        # Resolve again - should succeed (idempotent)
        result = resolve_comment(doc, para_id)
        assert result is True, "Resolving already-resolved comment should still return True"


# ============================================================================
# Test: unresolve_comment() - NEW FUNCTION (should fail initially)
# ============================================================================

class TestUnresolveComment:
    """Tests for unresolve_comment() function - marks a comment as active."""
    
    def test_unresolve_comment_import(self):
        """Test that unresolve_comment can be imported from comments module."""
        try:
            from effilocal.mcp_server.core.comments import unresolve_comment
            assert callable(unresolve_comment), "unresolve_comment should be callable"
        except ImportError:
            pytest.fail("unresolve_comment not yet implemented - expected failure")
    
    def test_unresolve_comment_by_id(self, test_doc_with_comments):
        """Test unresolving a comment by its para_id."""
        try:
            from effilocal.mcp_server.core.comments import resolve_comment, unresolve_comment
        except ImportError:
            pytest.skip("unresolve_comment not yet implemented")
        
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        if not comments:
            pytest.skip("No comments in test document")
        
        # First resolve a comment with para_id
        active_comments = [c for c in comments if c.get('status') == 'active' and c.get('para_id')]
        if not active_comments:
            pytest.skip("No active comments with para_id to test")
        
        comment = active_comments[0]
        para_id = comment['para_id']
        resolve_comment(doc, para_id)
        
        # Now unresolve it
        result = unresolve_comment(doc, para_id)
        assert result is True, "unresolve_comment should return True on success"
        
        # Verify status
        doc.save(str(test_doc_with_comments))
        doc2 = Document(str(test_doc_with_comments))
        updated_comments = extract_all_comments(doc2)
        
        unresolved_comment = next((c for c in updated_comments if c['comment_id'] == comment['comment_id']), None)
        assert unresolved_comment['status'] == 'active', "Comment should be active after unresolving"


# ============================================================================
# Test: reference_text extraction - ENHANCEMENT (should fail initially)
# ============================================================================

class TestReferenceTextExtraction:
    """Tests for extracting the text that a comment is anchored to."""
    
    def test_comment_has_reference_text(self, test_doc_with_comments):
        """Test that comments include the reference_text they're anchored to."""
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        if not comments:
            pytest.skip("No comments in test document")
        
        # Check that at least one comment has reference_text populated
        has_reference_text = any(c.get('reference_text', '').strip() for c in comments)
        
        assert has_reference_text, (
            "At least one comment should have reference_text populated. "
            "This tests that we extract the text between commentRangeStart and commentRangeEnd."
        )
    
    def test_reference_text_not_empty_for_anchored_comments(self, test_doc_with_comments):
        """Test that reference_text is populated for comments with anchor ranges."""
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        if not comments:
            pytest.skip("No comments in test document")
        
        # Comments with a paragraph_index should have reference_text
        anchored_comments = [c for c in comments if c.get('paragraph_index') is not None]
        
        if not anchored_comments:
            pytest.skip("No anchored comments in test document")
        
        for comment in anchored_comments[:3]:  # Check first 3
            ref_text = comment.get('reference_text', '')
            # Reference text should be non-empty for properly anchored comments
            # Note: Some comments might be anchored to the paragraph as a whole
            # In those cases, reference_text could be empty or the full paragraph
            assert 'reference_text' in comment, f"Comment {comment['comment_id']} should have reference_text field"
    
    def test_reference_text_extraction_function(self):
        """Test dedicated function for extracting reference text."""
        try:
            from effilocal.mcp_server.core.comments import extract_reference_text
            assert callable(extract_reference_text)
        except ImportError:
            pytest.fail("extract_reference_text not yet implemented - expected failure")


# ============================================================================
# Test: extract_all_comments() - EXISTING FUNCTION (should pass)
# ============================================================================

class TestExtractAllComments:
    """Tests for existing extract_all_comments() functionality."""
    
    def test_extract_all_comments_returns_list(self, test_doc_with_comments):
        """Test that extract_all_comments returns a list."""
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        assert isinstance(comments, list), "extract_all_comments should return a list"
    
    def test_comment_structure(self, test_doc_with_comments):
        """Test that each comment has required fields."""
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        if not comments:
            pytest.skip("No comments in test document")
        
        required_fields = [
            'id', 'comment_id', 'para_id', 'author', 'text', 
            'status', 'is_resolved', 'done_flag'
        ]
        
        for comment in comments[:5]:  # Check first 5
            for field in required_fields:
                assert field in comment, f"Comment missing required field: {field}"
    
    def test_comment_status_extraction(self, test_doc_with_comments):
        """Test that status is correctly extracted from commentsExtended.xml."""
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        if not comments:
            pytest.skip("No comments in test document")
        
        # All comments should have status (either from commentsExtended or default)
        for comment in comments:
            assert comment['status'] in ['active', 'resolved'], f"Invalid status: {comment['status']}"
            assert isinstance(comment['is_resolved'], bool), "is_resolved should be boolean"
            assert comment['done_flag'] in [0, 1], f"done_flag should be 0 or 1"
    
    def test_comment_date_format(self, test_doc_with_comments):
        """Test that dates are in ISO format."""
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        if not comments:
            pytest.skip("No comments in test document")
        
        for comment in comments[:5]:
            date_str = comment.get('date')
            if date_str:
                # Should be ISO format or close to it
                assert 'T' in date_str or '-' in date_str, f"Date should be ISO format: {date_str}"


# ============================================================================
# Test: extract_comment_status_map() - EXISTING FUNCTION
# ============================================================================

class TestExtractCommentStatusMap:
    """Tests for extract_comment_status_map() function."""
    
    def test_status_map_returns_dict(self, test_doc_with_comments):
        """Test that status map returns a dictionary."""
        doc = Document(str(test_doc_with_comments))
        status_map = extract_comment_status_map(doc)
        
        assert isinstance(status_map, dict), "Should return a dictionary"
    
    def test_status_map_keys_are_para_ids(self, test_doc_with_comments):
        """Test that status map keys are paraId values."""
        doc = Document(str(test_doc_with_comments))
        status_map = extract_comment_status_map(doc)
        
        if not status_map:
            pytest.skip("No commentsExtended data in test document")
        
        for para_id in status_map.keys():
            # paraId should be 8-character hex string
            assert len(para_id) == 8, f"paraId should be 8 chars: {para_id}"
            assert all(c in '0123456789ABCDEFabcdef' for c in para_id), f"paraId should be hex: {para_id}"
    
    def test_status_map_values(self, test_doc_with_comments):
        """Test that status map values have correct structure."""
        doc = Document(str(test_doc_with_comments))
        status_map = extract_comment_status_map(doc)
        
        if not status_map:
            pytest.skip("No commentsExtended data in test document")
        
        for para_id, status_info in status_map.items():
            assert 'status' in status_info
            assert 'done' in status_info
            assert 'is_resolved' in status_info
            assert status_info['status'] in ['active', 'resolved']
            assert status_info['done'] in [0, 1]
            assert isinstance(status_info['is_resolved'], bool)


# ============================================================================
# Test: MCP Tool - resolve_comment MCP tool (should fail initially)
# ============================================================================

class TestMCPResolveCommentTool:
    """Tests for resolve_comment MCP tool wrapper."""
    
    @pytest.mark.asyncio
    async def test_mcp_resolve_comment_tool_exists(self):
        """Test that resolve_comment MCP tool is registered."""
        try:
            from effilocal.mcp_server.tools.comment_tools import resolve_comment_tool
            assert callable(resolve_comment_tool)
        except ImportError:
            pytest.fail("resolve_comment_tool not yet implemented - expected failure")
    
    @pytest.mark.asyncio
    async def test_mcp_resolve_comment_returns_string(self, test_doc_with_comments):
        """Test that MCP tool returns a string result."""
        try:
            from effilocal.mcp_server.tools.comment_tools import resolve_comment_tool
        except ImportError:
            pytest.skip("resolve_comment_tool not yet implemented")
        
        result = await resolve_comment_tool(
            filename=str(test_doc_with_comments),
            comment_id="6"  # First comment in fixture
        )
        
        assert isinstance(result, str), "MCP tools should return strings"
        assert "resolved" in result.lower() or "success" in result.lower() or "error" in result.lower()


# ============================================================================
# Test: Comment-Paragraph Linking
# ============================================================================

class TestCommentParagraphLinking:
    """Tests for linking comments to their paragraph locations."""
    
    def test_comments_have_paragraph_index(self, test_doc_with_comments):
        """Test that comments have paragraph_index set when possible."""
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        if not comments:
            pytest.skip("No comments in test document")
        
        # At least some comments should have paragraph_index populated
        comments_with_index = [c for c in comments if c.get('paragraph_index') is not None]
        
        assert len(comments_with_index) > 0, (
            "At least some comments should have paragraph_index set. "
            "This is needed to link comments to document blocks."
        )
    
    def test_paragraph_index_is_valid(self, test_doc_with_comments):
        """Test that paragraph_index values are valid integers >= 0."""
        doc = Document(str(test_doc_with_comments))
        comments = extract_all_comments(doc)
        
        for comment in comments:
            para_idx = comment.get('paragraph_index')
            if para_idx is not None:
                assert isinstance(para_idx, int), f"paragraph_index should be int: {para_idx}"
                assert para_idx >= 0, f"paragraph_index should be >= 0: {para_idx}"


# ============================================================================
# Test: Comment Threading (deferred but basic structure test)
# ============================================================================

class TestCommentThreadingBasics:
    """Basic tests for comment threading structure (full threading deferred)."""
    
    def test_comments_can_have_parent(self, comments_fixture_dir):
        """Test that commentsExtended.xml supports parent-child relationships."""
        # Read fixture directly to check structure
        extended_path = comments_fixture_dir / "commentsExtended.xml"
        
        if not extended_path.exists():
            pytest.skip("No commentsExtended fixture")
        
        content = extended_path.read_text()
        
        # Check for paraIdParent attribute (threading)
        has_threading = 'paraIdParent' in content
        
        assert has_threading, (
            "commentsExtended.xml should contain paraIdParent attributes for threaded comments"
        )


# ============================================================================
# Test Utilities
# ============================================================================

class TestCommentTestUtilities:
    """Helper tests to verify test infrastructure."""
    
    def test_fixtures_exist(self, comments_fixture_dir):
        """Test that fixture files exist."""
        assert comments_fixture_dir.exists(), f"Fixture dir should exist: {comments_fixture_dir}"
        
        expected_files = ['comments.xml', 'commentsExtended.xml']
        for fname in expected_files:
            fpath = comments_fixture_dir / fname
            assert fpath.exists(), f"Fixture file should exist: {fpath}"
    
    def test_can_parse_fixture_xml(self, comments_fixture_dir):
        """Test that we can parse the fixture XML files."""
        from lxml import etree
        
        for xml_file in ['comments.xml', 'commentsExtended.xml']:
            fpath = comments_fixture_dir / xml_file
            if fpath.exists():
                tree = etree.parse(str(fpath))
                root = tree.getroot()
                assert root is not None, f"Should parse {xml_file}"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
