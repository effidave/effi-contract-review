
"""Table block builders for direct_docx."""

from __future__ import annotations

from typing import Any, Callable, List, Dict

from docx.table import Table

from effilocal.doc.amended_paragraph import AmendedParagraph
from effilocal.doc.blocks import TableCellBlock


def _get_amended_cell_content(cell) -> tuple[str, List[Dict[str, Any]]]:
    """
    Get amended text and runs for all paragraphs in a table cell.
    
    Combines text from all paragraphs with newlines. Runs are
    collected in order without position adjustment (text-based model).
    
    Returns:
        Tuple of (combined_text, combined_runs)
    """
    text_parts = []
    all_runs = []
    
    for para in cell.paragraphs:
        amended = AmendedParagraph(para)
        para_text = amended.amended_text
        para_runs = amended.amended_runs
        
        # Collect runs (no position adjustment needed with text-based model)
        all_runs.extend(para_runs)
        text_parts.append(para_text)
    
    combined_text = '\n'.join(text_parts)
    return combined_text, all_runs


def build_table_rows(
    table: Table,
    *,
    table_id: str,
    section_id: str | None,
    hash_provider: Callable[[str], str],
    fallback_heading_label: str | None,
    as_dataclass: bool = False,
) -> list[list[dict[str, Any] | TableCellBlock]]:
    """Return a list of table rows, each containing baseline cell blocks.
    
    Uses AmendedParagraph to correctly handle track changes in table cells:
    - amended_text includes insertions but excludes deletions
    - runs include formatting info with text content and deleted_text for deletions
    """

    rows: list[list[dict[str, Any] | TableCellBlock]] = []
    for row_index, row in enumerate(table.rows):
        row_blocks: list[dict[str, Any] | TableCellBlock] = []
        for col_index, cell in enumerate(row.cells):
            # Use amended content for track changes support
            text, runs = _get_amended_cell_content(cell)
            text = text.strip()
            
            if not text:
                continue
            
            # Ensure runs exist for non-empty text
            if not runs and text:
                runs = [{'text': text, 'formats': []}]
            
            style_name = cell.paragraphs[0].style.name if cell.paragraphs else ""
            heading_meta = (
                {
                    "text": fallback_heading_label,
                    "source": "none",
                    "fallback_label": fallback_heading_label,
                }
                if fallback_heading_label
                else None
            )
            block = {
                "id": None,  # ID assigned later by assign_block_ids()
                "type": "table_cell",
                "content_hash": hash_provider(text),
                "text": text,
                "style": style_name or "",
                "level": None,
                "section_id": section_id,  # None if not provided; caller handles
                "list": None,
                "table": {
                    "table_id": table_id,
                    "row": row_index,
                    "col": col_index,
                },
                "anchor": None,
                "metadata": None,
                "restart_group_id": None,
                "heading": heading_meta,
                "runs": runs,
            }
            block["indent"] = None
            if as_dataclass:
                row_blocks.append(
                    TableCellBlock(
                        id=block["id"],
                        type=block["type"],
                        content_hash=block["content_hash"],
                        text=block["text"],
                        style=block["style"],
                        level=block["level"],
                        section_id=block["section_id"],
                        table=block["table"],
                        indent=None,
                        heading=heading_meta,
                    )
                )
            else:
                row_blocks.append(block)

        if row_blocks:
            rows.append(row_blocks)

    return rows
