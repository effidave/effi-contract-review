"""UUID embedding via Word paragraph tags (w:tag elements).

This module provides functions to embed and extract UUIDs from Word documents
using the w:tag element within paragraph properties (w:pPr). UUIDs are stored as
tag values in the format "effi:block:<uuid>".

This approach uses standard WordprocessingML elements that:
- Survive save/close/reopen cycles in Word
- Don't interfere with python-docx's doc.paragraphs iteration
- Are simpler and less intrusive than SDT content controls

Supports embedding UUIDs in:
- Top-level body paragraphs (matched by para_idx)
- Paragraphs inside table cells (matched by table position: table_idx, row, col)
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import TYPE_CHECKING, Iterator, Union
from uuid import uuid4

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

if TYPE_CHECKING:
    from pathlib import Path

__all__ = [
    "embed_block_uuids",
    "extract_block_uuids",
    "extract_block_uuids_legacy",
    "remove_all_uuid_tags",
    "remove_all_uuid_controls",  # Backward compatibility alias
    "get_paragraph_uuid",
    "set_paragraph_uuid",
    "assign_block_ids",
    "EFFI_TAG_PREFIX",
    "BlockKey",
    "ParaKey",
    "TableCellKey",
]

# Prefix for effi block tags in paragraph properties
EFFI_TAG_PREFIX = "effi:block:"

# Word XML namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


@dataclass(frozen=True)
class ParaKey:
    """Key for a top-level paragraph block."""
    para_idx: int


@dataclass(frozen=True)
class TableCellKey:
    """Key for a table cell block."""
    table_idx: int
    row: int
    col: int


# Union type for block keys
BlockKey = Union[ParaKey, TableCellKey]


def _get_or_create_pPr(paragraph_element: etree._Element) -> etree._Element:
    """Get or create paragraph properties element.

    Args:
        paragraph_element: A w:p element

    Returns:
        The w:pPr element (created if it didn't exist)
    """
    pPr = paragraph_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        # Insert pPr as first child
        paragraph_element.insert(0, pPr)
    return pPr


def _get_tag_element(paragraph_element: etree._Element) -> etree._Element | None:
    """Get the w:tag element from a paragraph if it exists.

    Args:
        paragraph_element: A w:p element

    Returns:
        The w:tag element or None
    """
    pPr = paragraph_element.find(qn("w:pPr"))
    if pPr is None:
        return None
    return pPr.find(qn("w:tag"))


def _extract_uuid_from_tag(tag_element: etree._Element | None) -> str | None:
    """Extract UUID from a w:tag element if it has our prefix.

    Args:
        tag_element: A w:tag element or None

    Returns:
        The UUID string if the tag has our prefix, else None
    """
    if tag_element is None:
        return None
    
    val = tag_element.get(qn("w:val"), "")
    if val.startswith(EFFI_TAG_PREFIX):
        return val[len(EFFI_TAG_PREFIX):]
    
    return None


def _iter_paragraphs_with_index(document: Document) -> Iterator[tuple[int, etree._Element]]:
    """Iterate over top-level paragraph elements with their index in document order.
    
    This only yields top-level body paragraphs (not those in tables).
    For table cell paragraphs, use _iter_all_blocks().
    """
    body = document.element.body
    idx = 0
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield idx, child
            idx += 1


def _iter_all_blocks(document: Document) -> Iterator[tuple[BlockKey, etree._Element]]:
    """Iterate over ALL paragraph elements in the document, including those in tables.
    
    Yields tuples of (BlockKey, paragraph_element) where BlockKey is either:
    - ParaKey(para_idx) for top-level body paragraphs
    - TableCellKey(table_idx, row, col) for paragraphs inside table cells
    
    This traverses:
    - Top-level body paragraphs
    - Tables → rows → cells → paragraphs
    
    Note: Table cells may contain multiple paragraphs. Only the FIRST paragraph
    in each cell is yielded with the cell's key, as blocks.jsonl captures
    cell content as a single block.
    """
    body = document.element.body
    para_idx = 0
    table_idx = 0
    
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            # Top-level paragraph
            yield ParaKey(para_idx), child
            para_idx += 1
            
        elif child.tag == qn("w:tbl"):
            # Table - iterate rows and cells
            yield from _iter_table_cells(child, table_idx)
            table_idx += 1


def _iter_table_cells(
    table_element: etree._Element,
    table_idx: int,
) -> Iterator[tuple[TableCellKey, etree._Element]]:
    """Iterate over table cells, yielding the first paragraph in each cell.
    
    Args:
        table_element: A w:tbl element
        table_idx: The 0-based table index in the document
        
    Yields:
        Tuples of (TableCellKey, first_paragraph_element) for each cell
    """
    rows = table_element.findall(qn("w:tr"))
    for row_idx, row in enumerate(rows):
        cells = row.findall(qn("w:tc"))
        for col_idx, cell in enumerate(cells):
            # Find first paragraph in cell
            first_para = cell.find(qn("w:p"))
            if first_para is not None:
                yield TableCellKey(table_idx, row_idx, col_idx), first_para


def _block_to_key(block: dict) -> BlockKey | None:
    """Convert a block dict to its corresponding BlockKey.
    
    Args:
        block: A block dict from blocks.jsonl
        
    Returns:
        ParaKey for paragraphs, TableCellKey for table cells, None if invalid
    """
    # Check for table cell block
    table_info = block.get("table")
    if table_info is not None:
        # Table cell block: has table.table_id, table.row, table.col
        table_id = table_info.get("table_id", "")
        # Extract table index from table_id (e.g., "tbl_0" -> 0)
        if table_id.startswith("tbl_"):
            try:
                table_idx = int(table_id[4:])
            except ValueError:
                return None
        else:
            return None
        
        row = table_info.get("row")
        col = table_info.get("col")
        if row is not None and col is not None:
            return TableCellKey(table_idx, row, col)
        return None
    
    # Check for paragraph block
    para_idx = block.get("para_idx")
    if para_idx is not None:
        return ParaKey(para_idx)
    
    return None


def get_paragraph_uuid(paragraph_element: etree._Element) -> str | None:
    """Get the UUID from a paragraph's w:tag element if present.

    Args:
        paragraph_element: A w:p element

    Returns:
        The UUID string if the paragraph has an effi tag, else None
    """
    tag = _get_tag_element(paragraph_element)
    return _extract_uuid_from_tag(tag)


def set_paragraph_uuid(
    document: Document,
    paragraph_element: etree._Element,
    uuid_value: str,
) -> bool:
    """Set a UUID on a paragraph using the w:tag element.

    If the paragraph already has an effi tag, updates the value.
    Otherwise, creates a new w:tag element.

    Args:
        document: The Document object (for context, not currently used)
        paragraph_element: The w:p element to tag
        uuid_value: The UUID to embed

    Returns:
        True if successful
    """
    pPr = _get_or_create_pPr(paragraph_element)
    
    # Check for existing tag
    existing_tag = pPr.find(qn("w:tag"))
    
    if existing_tag is not None:
        # Check if it's our tag
        existing_val = existing_tag.get(qn("w:val"), "")
        if existing_val.startswith(EFFI_TAG_PREFIX):
            # Update existing effi tag
            existing_tag.set(qn("w:val"), f"{EFFI_TAG_PREFIX}{uuid_value}")
            return True
        else:
            # There's a non-effi tag - don't overwrite it
            # Could potentially use a different mechanism here
            return False
    
    # Create new tag
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), f"{EFFI_TAG_PREFIX}{uuid_value}")
    pPr.append(tag)
    
    return True


def embed_block_uuids(
    doc_or_path: Document | str | "Path",
    blocks: list[dict],
    *,
    overwrite: bool = False,
) -> dict[str, str]:
    """Embed UUIDs from blocks into the document as content controls.

    Each block's "id" field is used as the UUID. Blocks are matched to
    paragraphs by their location:
    - Top-level paragraphs: matched by "para_idx" field
    - Table cells: matched by "table" field (table_id, row, col)

    Args:
        doc_or_path: Document object or path to .docx file
        blocks: List of block dicts with "id" and either "para_idx" or "table" fields
        overwrite: If True, overwrite existing UUIDs; if False, skip already-wrapped paragraphs

    Returns:
        Mapping of block_id to location string for successfully embedded blocks
        (location is "para_N" for paragraphs or "tbl_T_r_R_c_C" for table cells)
    """
    from pathlib import Path

    if isinstance(doc_or_path, (str, Path)):
        document = Document(str(doc_or_path))
        save_path = Path(doc_or_path)
    else:
        document = doc_or_path
        save_path = None

    # Build index of BlockKey -> paragraph element
    key_to_element: dict[BlockKey, etree._Element] = {}
    for key, elem in _iter_all_blocks(document):
        key_to_element[key] = elem

    embedded: dict[str, str] = {}

    for block in blocks:
        block_id = block.get("id")
        if block_id is None:
            continue

        # Convert block to key
        block_key = _block_to_key(block)
        if block_key is None:
            continue

        para_elem = key_to_element.get(block_key)
        if para_elem is None:
            continue

        # Check if already has UUID
        existing = get_paragraph_uuid(para_elem)
        if existing is not None and not overwrite:
            embedded[block_id] = _key_to_location_string(block_key)
            continue

        # Embed UUID
        if set_paragraph_uuid(document, para_elem, block_id):
            embedded[block_id] = _key_to_location_string(block_key)

    # Save if we loaded from path
    if save_path is not None:
        document.save(str(save_path))

    return embedded


def _key_to_location_string(key: BlockKey) -> str:
    """Convert a BlockKey to a human-readable location string."""
    if isinstance(key, ParaKey):
        return f"para_{key.para_idx}"
    elif isinstance(key, TableCellKey):
        return f"tbl_{key.table_idx}_r_{key.row}_c_{key.col}"
    return "unknown"


def extract_block_uuids(doc_or_path: Document | str | "Path") -> dict[str, BlockKey]:
    """Extract UUID → BlockKey mapping from embedded content controls.

    Args:
        doc_or_path: Document object or path to .docx file

    Returns:
        Mapping of UUID string to BlockKey (ParaKey or TableCellKey)
    """
    from pathlib import Path

    if isinstance(doc_or_path, (str, Path)):
        document = Document(str(doc_or_path))
    else:
        document = doc_or_path

    uuid_to_key: dict[str, BlockKey] = {}

    for key, para_elem in _iter_all_blocks(document):
        uuid_val = get_paragraph_uuid(para_elem)
        if uuid_val is not None:
            uuid_to_key[uuid_val] = key

    return uuid_to_key


def extract_block_uuids_legacy(doc_or_path: Document | str | "Path") -> dict[str, int]:
    """Extract UUID → paragraph index mapping from embedded content controls.
    
    Legacy function for backward compatibility. Only returns top-level paragraphs.
    For full support including table cells, use extract_block_uuids().

    Args:
        doc_or_path: Document object or path to .docx file

    Returns:
        Mapping of UUID string to paragraph index (0-based)
    """
    from pathlib import Path

    if isinstance(doc_or_path, (str, Path)):
        document = Document(str(doc_or_path))
    else:
        document = doc_or_path

    uuid_to_idx: dict[str, int] = {}

    for idx, para_elem in _iter_paragraphs_with_index(document):
        uuid_val = get_paragraph_uuid(para_elem)
        if uuid_val is not None:
            uuid_to_idx[uuid_val] = idx

    return uuid_to_idx


def remove_all_uuid_tags(doc_or_path: Document | str | "Path") -> int:
    """Remove all effi UUID tags from the document.

    Only removes w:tag elements that have the effi:block: prefix.

    Args:
        doc_or_path: Document object or path to .docx file

    Returns:
        Number of tags removed
    """
    from pathlib import Path

    if isinstance(doc_or_path, (str, Path)):
        document = Document(str(doc_or_path))
        save_path = Path(doc_or_path)
    else:
        document = doc_or_path
        save_path = None

    removed = 0
    
    # Iterate all paragraphs (including in tables)
    for key, para_elem in _iter_all_blocks(document):
        tag = _get_tag_element(para_elem)
        if tag is not None:
            val = tag.get(qn("w:val"), "")
            if val.startswith(EFFI_TAG_PREFIX):
                # Remove the tag element
                pPr = para_elem.find(qn("w:pPr"))
                if pPr is not None:
                    pPr.remove(tag)
                    removed += 1

    if save_path is not None:
        document.save(str(save_path))

    return removed


# Keep old name as alias for backward compatibility
remove_all_uuid_controls = remove_all_uuid_tags


def assign_block_ids(
    blocks: list[dict],
    embedded_uuids: dict[str, BlockKey] | None = None,
    old_blocks: list[dict] | None = None,
    position_threshold: int = 5,
) -> dict[str, str]:
    """Assign IDs to blocks that have id=None.
    
    This is the single point of ID assignment for blocks. IDs are assigned
    using the following priority:
    
    1. Embedded UUID: If the block's position matches an embedded UUID from
       a content control in the document, use that UUID.
    2. Hash match: If old_blocks is provided, match by content_hash to preserve
       IDs for blocks with identical content.
    3. Position match: For unmatched blocks, match to nearby old blocks by
       position (within position_threshold paragraphs).
    4. Generate new: For truly new blocks, generate a fresh UUID.
    
    Args:
        blocks: List of block dicts. Blocks with id=None will be assigned IDs.
        embedded_uuids: Mapping of UUID -> BlockKey from extract_block_uuids().
            If provided, blocks matching these positions get the embedded UUID.
        old_blocks: Previous blocks.jsonl content for hash-based matching.
            Used as fallback for blocks without embedded UUIDs.
        position_threshold: Maximum distance (in paragraph indices) for position
            matching. Default is 5.
    
    Returns:
        Stats dict with counts: {"from_embedded": N, "from_hash": N, 
                                  "from_position": N, "generated": N}
    
    Mutates blocks in place to set the "id" field.
    """
    stats = {"from_embedded": 0, "from_hash": 0, "from_position": 0, "generated": 0}
    
    # Build reverse mapping: BlockKey -> UUID from embedded
    key_to_uuid: dict[BlockKey, str] = {}
    if embedded_uuids:
        for uuid_val, key in embedded_uuids.items():
            key_to_uuid[key] = uuid_val
    
    # Build hash -> ID mapping from old blocks for fallback matching
    hash_to_old_id: dict[str, str] = {}
    # Build position -> old block mapping for position-based fallback
    old_blocks_by_position: dict[int, dict] = {}
    if old_blocks:
        for old_block in old_blocks:
            old_id = old_block.get("id")
            old_hash = old_block.get("content_hash")
            if old_id and old_hash and old_hash not in hash_to_old_id:
                hash_to_old_id[old_hash] = old_id
            # Index by para_idx for position matching
            para_idx = old_block.get("para_idx")
            if para_idx is not None and old_id:
                old_blocks_by_position[para_idx] = old_block
    
    # Track which old IDs have been used (to avoid duplicates)
    used_ids: set[str] = set()
    
    # First pass: assign embedded UUIDs and hash matches
    unmatched_blocks: list[dict] = []
    for block in blocks:
        # Skip blocks that already have an ID
        if block.get("id") is not None:
            continue
        
        # Priority 1: Check embedded UUIDs
        block_key = _block_to_key(block)
        if block_key and block_key in key_to_uuid:
            block["id"] = key_to_uuid[block_key]
            used_ids.add(block["id"])
            stats["from_embedded"] += 1
            continue
        
        # Priority 2: Check hash match with old blocks
        content_hash = block.get("content_hash")
        if content_hash and content_hash in hash_to_old_id:
            old_id = hash_to_old_id[content_hash]
            if old_id not in used_ids:
                block["id"] = old_id
                used_ids.add(old_id)
                stats["from_hash"] += 1
                continue
        
        # Couldn't match yet - add to unmatched for position matching
        unmatched_blocks.append(block)
    
    # Second pass: position-based matching for remaining unmatched blocks
    for block in unmatched_blocks:
        if block.get("id") is not None:
            continue  # Already assigned in first pass
        
        para_idx = block.get("para_idx")
        if para_idx is not None and old_blocks_by_position:
            # Look for nearby old blocks within threshold
            best_match = None
            best_distance = position_threshold + 1
            
            for offset in range(position_threshold + 1):
                for check_idx in [para_idx + offset, para_idx - offset]:
                    if check_idx in old_blocks_by_position:
                        old_block = old_blocks_by_position[check_idx]
                        old_id = old_block.get("id")
                        # Only match if same type and ID not already used
                        if (old_id and 
                            old_id not in used_ids and
                            old_block.get("type") == block.get("type")):
                            distance = abs(check_idx - para_idx)
                            if distance < best_distance:
                                best_match = old_block
                                best_distance = distance
            
            if best_match:
                block["id"] = best_match["id"]
                used_ids.add(block["id"])
                stats["from_position"] += 1
                continue
        
        # Priority 4: Generate new UUID
        new_id = str(uuid4())
        block["id"] = new_id
        used_ids.add(new_id)
        stats["generated"] += 1
    
    return stats
