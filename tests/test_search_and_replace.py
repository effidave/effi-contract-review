"""
Tests for the search_and_replace functionality with snippet output.
"""
import os
import sys
import tempfile
import pytest
from pathlib import Path

# Add parent directory to path to import modules
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from word_document_server.utils.document_utils import find_and_replace_text
from word_document_server.tools.content_tools import search_and_replace


class TestFindAndReplaceText:
    """Test the find_and_replace_text utility function."""
    
    def test_simple_replacement(self):
        """Test simple text replacement in a paragraph."""
        doc = Document()
        doc.add_paragraph("This is a test document")
        
        count, snippets, split_matches = find_and_replace_text(doc, "test", "example")
        
        assert count == 1
        assert len(snippets) == 1
        assert "test" in snippets[0]["before"]
        assert "example" in snippets[0]["after"]
        assert snippets[0]["location"] == "paragraph"
    
    def test_multiple_replacements(self):
        """Test multiple replacements in different paragraphs."""
        doc = Document()
        doc.add_paragraph("The quick brown fox jumps")
        doc.add_paragraph("The lazy dog rests")
        
        count, snippets, split_matches = find_and_replace_text(doc, "The", "A")
        
        assert count == 2
        assert len(snippets) == 2
        assert snippets[0]["location"] == "paragraph"
        assert snippets[1]["location"] == "paragraph"
        assert "The" in snippets[0]["before"]
        assert "A" in snippets[0]["after"]
    
    def test_replacement_in_table(self):
        """Test replacement in table cells."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Find this text"
        table.cell(0, 1).text = "Other content"
        table.cell(1, 0).text = "More find this text"
        table.cell(1, 1).text = "More content"
        
        count, snippets, split_matches = find_and_replace_text(doc, "Find", "Replace")
        
        # The replacement might only happen once due to how runs are split
        assert count >= 1
        assert any(s["location"] == "table" for s in snippets)
    
    def test_no_matches(self):
        """Test when text is not found."""
        doc = Document()
        doc.add_paragraph("This is a test document")
        
        count, snippets, split_matches = find_and_replace_text(doc, "nonexistent", "replacement")
        
        assert count == 0
        assert len(snippets) == 0
    
    def test_context_extraction(self):
        """Test that context is properly extracted around replacement."""
        doc = Document()
        doc.add_paragraph("Before text Middle text After text")
        
        count, snippets, split_matches = find_and_replace_text(doc, "Middle", "REPLACED")
        
        assert count == 1
        assert "Middle" in snippets[0]["before"]
        assert "REPLACED" in snippets[0]["after"]
    
    def test_mixed_locations(self):
        """Test replacements in both paragraphs and tables."""
        doc = Document()
        doc.add_paragraph("Paragraph with marker")
        
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Table with marker"
        
        count, snippets, split_matches = find_and_replace_text(doc, "marker", "target")
        
        assert count == 2
        locations = [s["location"] for s in snippets]
        assert "paragraph" in locations
        assert "table" in locations
    
    def test_whole_word_only(self):
        """Test that whole_word_only prevents partial matches."""
        doc = Document()
        doc.add_paragraph("The WP and WPentifiable issues")
        
        # Without whole_word_only, both should be replaced
        count_all, _, _ = find_and_replace_text(doc, "WP", "Work Product", whole_word_only=False)
        assert count_all == 2  # "WP" and "WP" in "WPentifiable"
        
        # Create a fresh document
        doc2 = Document()
        doc2.add_paragraph("The WP and WPentifiable issues")
        
        # With whole_word_only, only the standalone "WP" should be replaced
        count_whole, snippets, _ = find_and_replace_text(doc2, "WP", "Work Product", whole_word_only=True)
        assert count_whole == 1  # Only the standalone "WP"
        assert "WPentifiable" in doc2.paragraphs[0].text  # Should remain unchanged
    
    def test_text_split_across_runs(self):
        """Test that text split across multiple runs is found and replaced correctly."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        doc = Document()
        para = doc.add_paragraph()
        
        # Manually create runs to simulate text split across boundaries
        # This simulates: "W" in run1, "P" in run2 (like in the real document issue)
        para.clear()
        
        # Add runs manually to create split text
        run1 = para.add_run("The ")
        run2 = para.add_run("W")
        run3 = para.add_run("P")
        run4 = para.add_run(" is important")
        
        # Verify the text is split
        assert para.text == "The WP is important"
        assert len(para.runs) == 4
        
        # Now test replacement - should find "WP" even though it's split across run2 and run3
        # The new behavior: split-run matches are REPORTED but NOT auto-replaced
        count, snippets, split_matches = find_and_replace_text(doc, "WP", "Work Product")
        
        # Count should be 0 because WP spans multiple runs (not in single run)
        assert count == 0, f"Expected 0 in-run replacements, got {count}"
        # But we should have 1 split match reported
        assert len(split_matches) == 1, f"Expected 1 split match, got {len(split_matches)}"
        
        # Verify the split match details
        match = split_matches[0]
        assert match['paragraph_index'] == 0
        assert match['full_match'] == "WP"
        assert len(match['runs']) == 2  # Spans 2 runs
        assert match['runs'][0]['run_index'] == 1  # run2 ("W")
        assert match['runs'][1]['run_index'] == 2  # run3 ("P")
        
        # Text should NOT be replaced (since we didn't auto-replace split matches)
        assert "WP" in para.text
        assert "Work Product" not in para.text
        assert para.text == "The WP is important"


class TestSearchAndReplaceTool:
    """Test the search_and_replace async tool."""
    
    def test_search_and_replace_basic(self):
        """Test basic search and replace with file I/O."""
        import asyncio
        
        async def run_test():
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            
            try:
                # Create test document
                doc = Document()
                doc.add_paragraph("Hello world, this is a test")
                doc.add_paragraph("Hello again, world!")
                doc.save(tmp_path)
                
                # Run search and replace
                result = await search_and_replace(tmp_path, "Hello", "Hi")
                
                assert "Replaced 2 occurrence(s)" in result
                assert "Replacement 1" in result
                assert "Replacement 2" in result
                assert "Before:" in result
                assert "After:" in result
                
                # Verify the document was actually modified
                modified_doc = Document(tmp_path)
                text = "\n".join([p.text for p in modified_doc.paragraphs])
                assert "Hi world" in text
                assert "Hello" not in text
            
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        
        asyncio.run(run_test())
    
    def test_search_and_replace_no_matches(self):
        """Test search and replace when text is not found."""
        import asyncio
        
        async def run_test():
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            
            try:
                # Create test document
                doc = Document()
                doc.add_paragraph("Some content here")
                doc.save(tmp_path)
                
                # Run search and replace
                result = await search_and_replace(tmp_path, "nonexistent", "replacement")
                
                assert "No occurrences of 'nonexistent' found" in result
            
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        
        asyncio.run(run_test())
    
    def test_search_and_replace_nonexistent_file(self):
        """Test search and replace on a file that doesn't exist."""
        import asyncio
        
        async def run_test():
            result = await search_and_replace("/nonexistent/path/file.docx", "test", "replacement")
            assert "does not exist" in result
        
        asyncio.run(run_test())
    
    def test_search_and_replace_with_table(self):
        """Test search and replace in a document with tables."""
        import asyncio
        
        async def run_test():
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            
            try:
                # Create test document with table
                doc = Document()
                doc.add_paragraph("Start of document")
                
                table = doc.add_table(rows=2, cols=2)
                table.cell(0, 0).text = "Target value in table"
                table.cell(0, 1).text = "Other cell"
                table.cell(1, 0).text = "Target value again"
                table.cell(1, 1).text = "More content"
                
                doc.add_paragraph("End of document")
                doc.save(tmp_path)
                
                # Run search and replace
                result = await search_and_replace(tmp_path, "Target", "Found")
                
                assert "Replaced" in result
                assert "(table)" in result
                
                # Verify modifications
                modified_doc = Document(tmp_path)
                all_text = "\n".join([p.text for p in modified_doc.paragraphs])
                for table in modified_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            all_text += "\n" + "\n".join([p.text for p in cell.paragraphs])
                
                assert "Found" in all_text
                assert "Target" not in all_text
            
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        
        asyncio.run(run_test())


if __name__ == "__main__":
    # Run tests with pytest
    pytest.main([__file__, "-v", "-s"])
