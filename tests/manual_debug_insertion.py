"""Manual test with debug output."""
from pathlib import Path
import sys
import shutil

sys.path.insert(0, str(Path(__file__).parent.parent))

from effilocal.doc import direct_docx
from docx import Document

# Copy fixture
fixture_path = Path(__file__).parent / "fixtures" / "real_world" / "schedules.docx"
temp_file = Path(__file__).parent / "debug_attachment_test.docx"
shutil.copy2(fixture_path, temp_file)

# Analyze with effilocal
blocks = list(direct_docx.iter_blocks(temp_file))

# Find Schedule 1
target_attachment_id = None
for block in blocks:
    attachment = block.get("attachment")
    if isinstance(attachment, dict) and attachment.get("ordinal") == "1":
        target_attachment_id = attachment["attachment_id"]
        break

# Find last block
last_block_para_idx = None
for i, block in enumerate(blocks):
    if block.get("attachment_id") == target_attachment_id:
        last_block_para_idx = block.get("para_idx")
        print(f"Block {i}, para_idx {block.get('para_idx')}: {block.get('text', '')[:60]}")

print(f"\nLast block para_idx: {last_block_para_idx}")

# Load with python-docx
doc = Document(temp_file)
print(f"Total paragraphs in docx: {len(doc.paragraphs)}")

# Get the paragraph
last_para = doc.paragraphs[last_block_para_idx]
print(f"Paragraph at index {last_block_para_idx}: {last_para.text[:60]}")

# Get XML position
target_p = last_para._p
parent = target_p.getparent()
target_position = list(parent).index(target_p)
print(f"XML position in parent: {target_position}")
print(f"Will insert at position: {target_position + 1}")

# Count how many paragraphs are before this position
all_elements = list(parent)
para_count_before = sum(1 for elem in all_elements[:target_position] if elem.tag.endswith('}p'))
print(f"Number of <w:p> elements before target: {para_count_before}")
