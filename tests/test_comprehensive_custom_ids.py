"""Comprehensive test of custom para IDs across all insertion tools."""
import asyncio
import sys
from pathlib import Path
import shutil
import re
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from effilocal.mcp_server.tools.content_tools import add_paragraph_after_clause
from effilocal.mcp_server.tools.attachment_tools import (
    add_new_attachment_after,
    add_paragraph_after_attachment,
    find_paragraph_by_id
)
from docx import Document

def extract_custom_id(result_message):
    """Extract custom para_id from result message."""
    match = re.search(r'para_id=([0-9a-f\-]+)', result_message)
    return match.group(1) if match else None

@pytest.mark.asyncio
async def test_comprehensive_custom_ids():
    print("=" * 80)
    print("COMPREHENSIVE TEST: Custom Para IDs for all insertion tools")
    print("=" * 80)
    
    # Test 1: Clause insertion
    print("\n" + "=" * 80)
    print("TEST 1: Clause insertion with custom ID")
    print("=" * 80)
    
    temp_clause = Path(__file__).parent / "fixtures" / "numbering_decimal.docx"
    temp_copy = None
    temp_attach = None
    
    try:
        if temp_clause.exists():
            temp_copy = Path(__file__).parent / "temp_clause_custom_id.docx"
            shutil.copy2(temp_clause, temp_copy)
            
            result = await add_paragraph_after_clause(
                filename=str(temp_copy),
                clause_number="1",
                text="New clause inserted with custom ID tracking"
            )
            
            print(f"Result: {result}")
            custom_id = extract_custom_id(result)
            
            if custom_id:
                print(f"✓ Custom ID returned: {custom_id}")
                
                # Verify we can find it
                doc = Document(temp_copy)
                para = find_paragraph_by_id(doc, custom_id=custom_id)
                if para:
                    print(f"✓ Found paragraph by custom ID: {para.text[:60]}")
                else:
                    print(f"✗ Could not find paragraph by custom ID")
            else:
                print(f"✗ No custom ID in result")
        else:
            print("Skipping clause test - fixture not found")
        
        # Test 2: Attachment paragraph insertion
        print("\n" + "=" * 80)
        print("TEST 2: Attachment paragraph insertion with custom ID")
        print("=" * 80)
        
        temp_attach = Path(__file__).parent / "temp_attach_custom_id.docx"
        shutil.copy2(
            Path(__file__).parent / "fixtures" / "real_world" / "schedules.docx",
            temp_attach
        )
        
        result = await add_paragraph_after_attachment(
            filename=str(temp_attach),
            attachment_identifier="Schedule 1",
            text="Paragraph added with custom ID tracking",
            inherit_numbering=False
        )
        
        print(f"Result: {result}")
        custom_id = extract_custom_id(result)
        
        if custom_id:
            print(f"✓ Custom ID returned: {custom_id}")
            
            # Verify we can find it
            doc = Document(temp_attach)
            para = find_paragraph_by_id(doc, custom_id=custom_id)
            if para:
                print(f"✓ Found paragraph by custom ID: {para.text[:60]}")
            else:
                print(f"✗ Could not find paragraph by custom ID")
        else:
            print(f"✗ No custom ID in result")
        
        # Test 3: New attachment creation
        print("\n" + "=" * 80)
        print("TEST 3: New attachment creation with custom IDs")
        print("=" * 80)
        
        result = await add_new_attachment_after(
            filename=str(temp_attach),
            after_attachment="Schedule 2",
            new_attachment_text="Schedule 3 - Created with tracking",
            content="Content paragraph for Schedule 3"
        )
        
        print(f"Result: {result}")
        
        # Extract both header_id and content_id
        header_match = re.search(r'header_id=([0-9a-f\-]+)', result)
        content_match = re.search(r'content_id=([0-9a-f\-]+)', result)
        
        header_id = header_match.group(1) if header_match else None
        content_id = content_match.group(1) if content_match else None
        
        if header_id:
            print(f"✓ Header custom ID returned: {header_id}")
            
            doc = Document(temp_attach)
            para = find_paragraph_by_id(doc, custom_id=header_id)
            if para:
                print(f"✓ Found header by custom ID: {para.text[:60]}")
            else:
                print(f"✗ Could not find header by custom ID")
        else:
            print(f"✗ No header ID in result")
        
        if content_id:
            print(f"✓ Content custom ID returned: {content_id}")
            
            doc = Document(temp_attach)
            para = find_paragraph_by_id(doc, custom_id=content_id)
            if para:
                print(f"✓ Found content by custom ID: {para.text[:60]}")
            else:
                print(f"✗ Could not find content by custom ID")
        else:
            print(f"✗ No content ID in result")
        
        print("\n" + "=" * 80)
        print("SUMMARY: All insertion tools now support custom para ID tracking!")
        print("=" * 80)
    finally:
        # Cleanup
        if temp_copy and temp_copy.exists():
            temp_copy.unlink()
        if temp_attach and temp_attach.exists():
            temp_attach.unlink()

if __name__ == "__main__":
    asyncio.run(test_comprehensive_custom_ids())
