from __future__ import annotations

import json
from pathlib import Path
from uuid import uuid4

from docx import Document

from effilocal.flows.analyze_doc import analyze


def _build_doc(tmp_path: Path) -> Path:
    doc = Document()

    heading = doc.add_paragraph("General")
    heading.style = "Heading 1"
    doc.add_paragraph("First general note.")
    doc.add_paragraph("Second general note.")

    heading2 = doc.add_paragraph("Other")
    heading2.style = "Heading 1"
    doc.add_paragraph("Different section note.")

    doc_path = tmp_path / "unnumbered.docx"
    doc.save(doc_path)
    return doc_path


def _load_blocks(out_dir: Path) -> list[dict[str, object]]:
    path = out_dir / "blocks.jsonl"
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def test_sequential_unnumbered_share_group(tmp_path: Path) -> None:
    doc_path = _build_doc(tmp_path)
    out_dir = tmp_path / "analysis"
    analyze(doc_path, doc_id=str(uuid4()), out_dir=out_dir)

    blocks = _load_blocks(out_dir)
    para1 = next(block for block in blocks if block["text"] == "First general note.")
    para2 = next(block for block in blocks if block["text"] == "Second general note.")
    para3 = next(block for block in blocks if block["text"] == "Different section note.")

    assert para1["clause_group_id"] == para2["clause_group_id"]
    assert para1.get("continuation_of") is None
    assert para2.get("continuation_of") is None
    assert para3["clause_group_id"] != para1["clause_group_id"]
