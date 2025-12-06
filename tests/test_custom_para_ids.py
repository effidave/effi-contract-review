"""Test paragraph tracking using native w14:paraId for newly created paragraphs."""
import asyncio
import sys
from pathlib import Path
import shutil
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from effilocal.mcp_server.tools.attachment_tools import (
    add_new_attachment_after,
    find_paragraph_by_id
)
from effilocal.doc.uuid_embedding import find_paragraph_by_para_id
from docx import Document
from docx.oxml.ns import qn

@pytest.mark.asyncio
async def test_custom_para_ids():
    """Test that newly created attachments have w14:paraId and can be found."""
    temp = Path(__file__).parent / "temp_custom_id_test.docx"
    shutil.copy2(Path(__file__).parent / "fixtures" / "real_world" / "schedules.docx", temp)
    
    try:
        print("=" * 80)
        print("TEST: Create Schedule 3 and track it with w14:paraId")
        print("=" * 80)
        
        # Add new Schedule 3
        result = await add_new_attachment_after(
            filename=str(temp),
            after_attachment="Schedule 2",
            new_attachment_text="Schedule 3 - Test Attachment",
            content="This is the content of Schedule 3"
        )
        
        print(f"\nResult: {result}")
        
        # Extract the para IDs from the result (now uses w14:paraId format)
        import re
        header_id_match = re.search(r'header_id=([0-9A-Fa-f]+)', result)
        content_id_match = re.search(r'content_id=([0-9A-Fa-f]+)', result)
        
        header_para_id = header_id_match.group(1) if header_id_match else None
        content_para_id = content_id_match.group(1) if content_id_match else None
        
        print(f"\nExtracted w14:paraIds:")
        print(f"  Header para_id: {header_para_id}")
        print(f"  Content para_id: {content_para_id}")
        
        # Reload the document and try to find the paragraphs
        print("\n" + "=" * 80)
        print("VERIFICATION: Finding paragraphs by w14:paraId")
        print("=" * 80)
        
        doc = Document(temp)
        
        # Find header by para_id
        if header_para_id:
            header_para = find_paragraph_by_para_id(doc, header_para_id)
            if header_para:
                print(f"\n✓ Found header paragraph by para_id:")
                print(f"  Text: {header_para.text[:80]}")
                print(f"  Style: {header_para.style.style_id if header_para.style else None}")
                
                # Check the w14:paraId in XML
                para_id = header_para._element.get(qn('w14:paraId'))
                print(f"  w14:paraId: {para_id}")
                
                assert header_para.text.startswith("Schedule 3"), f"Expected header text, got: {header_para.text}"
            else:
                print(f"\n✗ Could not find header paragraph by para_id")
                assert False, f"Header paragraph not found with para_id={header_para_id}"
        else:
            assert False, "No header_id returned in result"
        
        # Find content by para_id
        if content_para_id:
            content_para = find_paragraph_by_para_id(doc, content_para_id)
            if content_para:
                print(f"\n✓ Found content paragraph by para_id:")
                print(f"  Text: {content_para.text[:80]}")
                
                # Check the w14:paraId in XML
                para_id = content_para._element.get(qn('w14:paraId'))
                print(f"  w14:paraId: {para_id}")
                
                assert "Schedule 3" in content_para.text, f"Expected content text, got: {content_para.text}"
            else:
                print(f"\n✗ Could not find content paragraph by para_id")
                assert False, f"Content paragraph not found with para_id={content_para_id}"
        else:
            print("Note: No content_id returned (content might be merged with header)")
        
        # Test that we can perform operations on the found paragraph
        print("\n" + "=" * 80)
        print("TEST: Perform operations on found paragraph")
        print("=" * 80)
        
        if header_para_id:
            header_para = find_paragraph_by_para_id(doc, header_para_id)
            if header_para:
                # Try to add a new paragraph after it
                target_p = header_para._p
                parent = target_p.getparent()
                target_position = list(parent).index(target_p)
                
                from docx.oxml import OxmlElement
                new_p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = "New paragraph inserted after Schedule 3 header using custom ID"
                r.append(t)
                new_p.append(r)
                
                parent.insert(target_position + 1, new_p)
                doc.save(temp)
                
                print("✓ Successfully added a new paragraph after Schedule 3 header")
                
                # Reload and verify
                doc2 = Document(temp)
                for para in doc2.paragraphs:
                    if "New paragraph inserted after Schedule 3" in para.text:
                        print(f"✓ Verified: Found the newly inserted paragraph")
                        break
        
        print("\n" + "=" * 80)
    finally:
        # Cleanup
        if temp.exists():
            temp.unlink()

if __name__ == "__main__":
    asyncio.run(test_custom_para_ids())
