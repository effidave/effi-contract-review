#!/usr/bin/env python
"""Script to manage comments in Word documents.

Usage:
    manage_comments.py get_comments <docx_path>
    manage_comments.py resolve_comment <docx_path> <para_id>
    manage_comments.py unresolve_comment <docx_path> <para_id>

Outputs JSON to stdout.

Note: para_id is the w14:paraId from the comment's internal paragraph,
which is the stable identifier used for threading and status tracking.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from docx import Document

from effilocal.mcp_server.core.comments import (
    extract_all_comments,
    resolve_comment,
    unresolve_comment,
)


def get_comments(docx_path: Path) -> dict:
    """Extract all comments from a Word document."""
    if not docx_path.exists():
        return {"success": False, "error": f"File not found: {docx_path}"}
    
    try:
        doc = Document(str(docx_path))
        comments = extract_all_comments(doc)
        
        return {
            "success": True,
            "comments": comments,
            "total_comments": len(comments),
        }
    except Exception as e:
        return {"success": False, "error": f"Failed to extract comments: {str(e)}"}


def resolve_comment_cmd(docx_path: Path, para_id: str) -> dict:
    """Mark a comment as resolved using its para_id."""
    if not docx_path.exists():
        return {"success": False, "error": f"File not found: {docx_path}"}
    
    try:
        doc = Document(str(docx_path))
        success = resolve_comment(doc, para_id)
        
        if success:
            doc.save(str(docx_path))
            return {
                "success": True,
                "message": f"Comment resolved",
                "para_id": para_id,
            }
        else:
            return {
                "success": False,
                "error": f"Could not resolve comment. Comment may not exist or document may not have commentsExtended.xml.",
                "para_id": para_id,
            }
    except PermissionError:
        return {
            "success": False,
            "error": f"Permission denied - file '{docx_path.name}' may be open in Word",
            "para_id": para_id,
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to resolve comment: {str(e)}",
            "para_id": para_id,
        }


def unresolve_comment_cmd(docx_path: Path, para_id: str) -> dict:
    """Mark a comment as active (unresolve) using its para_id."""
    if not docx_path.exists():
        return {"success": False, "error": f"File not found: {docx_path}"}
    
    try:
        doc = Document(str(docx_path))
        success = unresolve_comment(doc, para_id)
        
        if success:
            doc.save(str(docx_path))
            return {
                "success": True,
                "message": f"Comment unresolved",
                "para_id": para_id,
            }
        else:
            return {
                "success": False,
                "error": f"Could not unresolve comment. Comment may not exist or document may not have commentsExtended.xml.",
                "para_id": para_id,
            }
    except PermissionError:
        return {
            "success": False,
            "error": f"Permission denied - file '{docx_path.name}' may be open in Word",
            "para_id": para_id,
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to unresolve comment: {str(e)}",
            "para_id": para_id,
        }


def main():
    parser = argparse.ArgumentParser(description="Manage comments in Word documents")
    parser.add_argument("command", choices=["get_comments", "resolve_comment", "unresolve_comment"],
                        help="Command to execute")
    parser.add_argument("docx_path", type=Path, help="Path to .docx file")
    parser.add_argument("para_id", nargs="?", help="Comment para_id (for resolve/unresolve)")

    args = parser.parse_args()

    if args.command == "get_comments":
        result = get_comments(args.docx_path)
    elif args.command == "resolve_comment":
        if not args.para_id:
            result = {"success": False, "error": "para_id is required for resolve_comment"}
        else:
            result = resolve_comment_cmd(args.docx_path, args.para_id)
    elif args.command == "unresolve_comment":
        if not args.para_id:
            result = {"success": False, "error": "para_id is required for unresolve_comment"}
        else:
            result = unresolve_comment_cmd(args.docx_path, args.para_id)
    else:
        result = {"success": False, "error": f"Unknown command: {args.command}"}

    print(json.dumps(result))


if __name__ == "__main__":
    main()
