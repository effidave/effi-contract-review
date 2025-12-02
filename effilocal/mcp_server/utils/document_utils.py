"""
Document utilities for effilocal MCP server.

This module extends word_document_server.utils.document_utils with:
1. Run-level text editing (edit_run_text)
2. Enhanced find_and_replace with whole_word_only support
3. Extended numbered list insertion
4. SDT-aware document iteration (for reading documents with content controls)

Override Pattern:
- Import all upstream functions
- Override functions with enhancements
- Add new utility functions
"""

import os
from typing import Dict, List, Optional, Any, Iterator
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

# Import upstream functions (pass-through where not overridden)
from word_document_server.utils.document_utils import (
    get_document_properties,
    extract_document_text,
    get_document_structure,
    get_document_xml,
    insert_header_near_text,
    insert_line_or_paragraph_near_text,
    replace_paragraph_block_below_header,
    replace_block_between_manual_anchors,
)

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension


# ============================================================================
# SDT-Aware Iteration Functions
# ============================================================================
# These functions handle documents that may contain SDT (Structured Document Tag)
# content controls. SDTs can wrap paragraphs/tables, hiding them from 
# doc.paragraphs/doc.tables. These iterators ensure we can read all content.

def iter_document_paragraphs(doc: DocxDocument) -> Iterator[Paragraph]:
    """Iterate over all paragraphs in a document, including those inside SDT content controls.
    
    SDT (Structured Document Tag) elements are content controls that may wrap
    paragraphs. This function unwraps them to yield the inner paragraphs.
    This is needed because doc.paragraphs doesn't see paragraphs wrapped in SDT elements.
    
    Note: We no longer CREATE SDT wrappers for UUID embedding (we use w:tag instead),
    but we still need to READ documents that may contain SDTs from other sources.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:sdt'):
            # SDT content control - look inside for paragraphs
            sdt_content = child.find(qn('w:sdtContent'))
            if sdt_content is not None:
                for inner in sdt_content.iterchildren():
                    if isinstance(inner, CT_P):
                        yield Paragraph(inner, doc)


def iter_document_tables(doc: DocxDocument) -> Iterator[Table]:
    """Iterate over all tables in a document, including those inside SDT content controls.
    
    Note: We no longer CREATE SDT wrappers for UUID embedding (we use w:tag instead),
    but we still need to READ documents that may contain SDTs from other sources.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_Tbl):
            yield Table(child, doc)
        elif child.tag == qn('w:sdt'):
            # SDT content control - look inside for tables
            sdt_content = child.find(qn('w:sdtContent'))
            if sdt_content is not None:
                for inner in sdt_content.iterchildren():
                    if isinstance(inner, CT_Tbl):
                        yield Table(inner, doc)


# ============================================================================
# Helper Functions
# ============================================================================

def get_paragraph_by_id(doc, para_id: str):
    """Find a paragraph by its w14:paraId (case-insensitive)."""
    target_id = para_id.upper()
    for p in doc.paragraphs:
        # Try standard w14:paraId
        pid = p._element.get('{http://schemas.microsoft.com/office/word/2010/wordml}paraId')
        if pid and pid.upper() == target_id:
            return p
    return None

def _find_and_replace_in_doc(doc, old_text: str, new_text: str, whole_word_only: bool = False):
    """
    Find and replace text in a Document object (used by tests).
    
    Returns:
        Tuple of (count, snippets, split_matches)
    """
    import re
    
    replacements = 0
    snippets = []
    split_matches = []
    
    # Process paragraphs
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text
        
        # Skip TOC paragraphs
        if paragraph.style and paragraph.style.name.startswith("TOC"):
            continue
        
        if old_text not in para_text:
            continue
        
        # Check for matches
        if whole_word_only:
            pattern = r'\b' + re.escape(old_text) + r'\b'
            matches = list(re.finditer(pattern, para_text))
        else:
            # Find all occurrences
            matches = []
            start = 0
            while True:
                pos = para_text.find(old_text, start)
                if pos == -1:
                    break
                matches.append(type('Match', (), {'start': lambda: pos, 'end': lambda: pos + len(old_text)})())
                start = pos + 1
        
        if not matches:
            continue
        
        # Check if text spans multiple runs
        if len(paragraph.runs) > 1:
            # Complex case - check if text actually spans runs
            # Build run map to detect split matches
            run_texts = [r.text for r in paragraph.runs]
            combined = ''.join(run_texts)
            
            if old_text in combined:
                # Text exists but might span runs - check each run individually
                found_in_single_run = any(old_text in r.text for r in paragraph.runs)
                
                if not found_in_single_run:
                    # Split across runs - find which runs contain the text
                    # Rebuild the text accumulating run by run to find where the match occurs
                    runs_info = []
                    for run_idx, run in enumerate(paragraph.runs):
                        # Check if this run contains any character from old_text by position matching
                        run_start = sum(len(r.text) for r in paragraph.runs[:run_idx])
                        run_end = run_start + len(run.text)
                        match_start = combined.find(old_text)
                        match_end = match_start + len(old_text)
                        
                        # If run overlaps with match
                        if run_start < match_end and run_end > match_start:
                            runs_info.append({'run_index': run_idx, 'text': run.text})
                    
                    split_matches.append({
                        'paragraph_index': para_idx,
                        'text': old_text,
                        'full_match': old_text,
                        'context': para_text,
                        'runs': runs_info
                    })
                    continue
        
        # Simple case: replace in runs
        for run in paragraph.runs:
            if old_text in run.text:
                before_text = run.text
                if whole_word_only:
                    run.text = re.sub(pattern, new_text, run.text)
                else:
                    run.text = run.text.replace(old_text, new_text)
                
                if run.text != before_text:
                    # Count actual replacements (how many times new_text appears more than before)
                    # OR count how many times old_text was in before but not in after
                    count = before_text.count(old_text) - run.text.count(old_text)
                    replacements += count
                    
                    # Add snippet
                    context_start = max(0, before_text.find(old_text) - 20)
                    context_end = min(len(before_text), before_text.find(old_text) + len(old_text) + 20)
                    
                    snippets.append({
                        'before': before_text[context_start:context_end],
                        'after': run.text[context_start:context_end],
                        'location': 'paragraph',
                        'paragraph_index': para_idx
                    })
    
    # Process tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    # Skip TOC paragraphs
                    if para.style and para.style.name.startswith("TOC"):
                        continue
                    
                    if old_text not in para.text:
                        continue
                    
                    for run in para.runs:
                        if old_text in run.text:
                            before_text = run.text
                            if whole_word_only:
                                pattern = r'\b' + re.escape(old_text) + r'\b'
                                run.text = re.sub(pattern, new_text, run.text)
                            else:
                                run.text = run.text.replace(old_text, new_text)
                            
                            if run.text != before_text:
                                # Count actual replacements
                                count = before_text.count(old_text) - run.text.count(old_text)
                                replacements += count
                                
                                context_start = max(0, before_text.find(old_text) - 20)
                                context_end = min(len(before_text), before_text.find(old_text) + len(old_text) + 20)
                                
                                snippets.append({
                                    'before': before_text[context_start:context_end],
                                    'after': run.text[context_start:context_end],
                                    'location': 'table',
                                    'table_index': table_idx,
                                    'row_index': row_idx,
                                    'cell_index': cell_idx
                                })
    
    return replacements, snippets, split_matches


# ============================================================================
# NEW / OVERRIDDEN Functions
# ============================================================================

def edit_run_text(
    doc_path: str, 
    paragraph_index: int, 
    run_index: int, 
    new_text: str, 
    start_offset: Optional[int] = None, 
    end_offset: Optional[int] = None
) -> str:
    """
    Edit text within a specific run of a paragraph.
    
    NEW function - provides precise text editing at run level.
    
    Args:
        doc_path: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        run_index: Index of the run within the paragraph (0-based)
        new_text: Text to insert/replace
        start_offset: Optional start position within the run (0-based)
        end_offset: Optional end position within the run (0-based, exclusive)
                   If not provided, replaces the entire run text
        
    Returns:
        Status message
    """
    doc_path = ensure_docx_extension(doc_path)
    
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    is_writeable, error_message = check_file_writeable(doc_path)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    try:
        doc = Document(doc_path)
        
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
        
        para = doc.paragraphs[paragraph_index]
        
        if run_index < 0 or run_index >= len(para.runs):
            return f"Invalid run index: {run_index}. Paragraph {paragraph_index} has {len(para.runs)} runs."
        
        run = para.runs[run_index]
        old_text = run.text
        
        # Default to replacing entire run
        if start_offset is None:
            start_offset = 0
        if end_offset is None:
            end_offset = len(old_text)
        
        # Validate offsets
        if start_offset < 0 or start_offset > len(old_text):
            return f"Invalid start_offset: {start_offset}. Run text length is {len(old_text)}."
        if end_offset < 0 or end_offset > len(old_text):
            return f"Invalid end_offset: {end_offset}. Run text length is {len(old_text)}."
        if start_offset > end_offset:
            return f"Invalid offsets: start_offset ({start_offset}) > end_offset ({end_offset})."
        
        # Perform the replacement
        run.text = old_text[:start_offset] + new_text + old_text[end_offset:]
        
        doc.save(doc_path)
        
        return f"Run {run_index} in paragraph {paragraph_index} updated: '{old_text[start_offset:end_offset]}' → '{new_text}'"
    
    except Exception as e:
        return f"Failed to edit run: {str(e)}"


def find_and_replace_text(
    doc_or_filename, 
    old_text: str, 
    new_text: str, 
    whole_word_only: bool = False
):
    """
    Find and replace text in a Word document.
    
    EXTENDED from upstream with whole_word_only parameter and dual signature support.
    
    Args:
        doc_or_filename: Either a Document object or path to document file
        old_text: Text to find
        new_text: Replacement text
        whole_word_only: If True, only match whole words (not substring matches)
    
    Returns:
        - If doc_or_filename is a Document object: tuple (count, snippets, split_matches)
        - If doc_or_filename is a string (filename): string status message
    """
    from docx.document import Document as DocumentClass
    
    # Determine if input is a Document object or filename
    if isinstance(doc_or_filename, DocumentClass):
        # Document object mode (for tests) - return tuple
        return _find_and_replace_in_doc(doc_or_filename, old_text, new_text, whole_word_only)
    
    # Filename mode (for MCP tools) - return string message
    filename = ensure_docx_extension(doc_or_filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(filename)
        
        # Use the helper function to do the replacement
        replacements, snippets, split_matches = _find_and_replace_in_doc(doc, old_text, new_text, whole_word_only)
        
        # Save the modified document
        doc.save(filename)
        
        # Build detailed result message
        if replacements == 0:
            return f"No occurrences of '{old_text}' found in {filename}"
        
        result_lines = [f"Replaced {replacements} occurrence(s) of '{old_text}' with '{new_text}' in {filename}"]
        result_lines.append("")  # Blank line
        
        # Add detailed snippets
        for i, snippet in enumerate(snippets, 1):
            result_lines.append(f"Replacement {i}:")
            result_lines.append(f"  Before: {snippet['before']}")
            result_lines.append(f"  After: {snippet['after']}")
            if snippet['location'] == 'table':
                result_lines.append(f"  Location: Table {snippet['table_index']}, Row {snippet['row_index']}, Cell {snippet['cell_index']} (table)")
            else:
                result_lines.append(f"  Location: Paragraph {snippet.get('paragraph_index', '?')}")
            result_lines.append("")  # Blank line between snippets
        
        if split_matches:
            result_lines.append(f"Note: {len(split_matches)} paragraph(s) contain '{old_text}' but it spans multiple runs (requires edit_run_text):")
            for match in split_matches:
                result_lines.append(f"  - Paragraph {match['paragraph_index']}: {match['context'][:100]}")
        
        return "\n".join(result_lines)
    
    except Exception as e:
        return f"Failed to find and replace: {str(e)}"


def insert_numbered_list_near_text(
    filename: str,
    target_text: Optional[str],
    list_items: Optional[List[str]],
    position: str = 'after',
    target_paragraph_index: Optional[int] = None,
    bullet_type: str = 'bullet'
) -> str:
    """
    Insert a numbered or bulleted list near target text or at a specific paragraph index.
    
    EXTENDED from upstream with enhanced parameter handling.
    
    Args:
        filename: Path to document
        target_text: Text to search for (or None if using paragraph_index)
        list_items: List of item texts
        position: 'before' or 'after' the target
        target_paragraph_index: Direct paragraph index (alternative to text search)
        bullet_type: 'bullet' or 'number'
    
    Returns:
        Status message
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    if not list_items:
        return "list_items parameter is required and must be non-empty"
    
    try:
        doc = Document(filename)
        
        # Find target paragraph
        target_para = None
        if target_paragraph_index is not None:
            if 0 <= target_paragraph_index < len(doc.paragraphs):
                target_para = doc.paragraphs[target_paragraph_index]
            else:
                return f"Invalid paragraph index: {target_paragraph_index}"
        elif target_text:
            for para in doc.paragraphs:
                if target_text in para.text:
                    target_para = para
                    break
        
        if target_para is None:
            return f"Target text '{target_text}' not found in document"
        
        # Find insertion point
        target_element = target_para._element
        parent = target_element.getparent()
        
        if position == 'before':
            insert_index = parent.index(target_element)
        else:  # after
            insert_index = parent.index(target_element) + 1
        
        # Insert list items
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        
        # Simple implementation: add paragraphs with list style
        for item_text in list_items:
            new_para = doc.add_paragraph(item_text, style='List Bullet' if bullet_type == 'bullet' else 'List Number')
            # Move to correct position
            para_element = new_para._element
            parent.remove(para_element)
            parent.insert(insert_index, para_element)
            insert_index += 1
        
        doc.save(filename)
        
        list_type = "bulleted" if bullet_type == 'bullet' else "numbered"
        return f"Inserted {len(list_items)}-item {list_type} list {position} target in {filename}"
    
    except Exception as e:
        return f"Failed to insert list: {str(e)}"


# ============================================================================
# Re-export all functions for compatibility
# ============================================================================

__all__ = [
    # Pass-through functions (unchanged from upstream)
    'get_document_properties',
    'extract_document_text',
    'get_document_structure',
    'get_document_xml',
    'insert_header_near_text',
    'insert_line_or_paragraph_near_text',
    'replace_paragraph_block_below_header',
    'replace_block_between_manual_anchors',
    
    # SDT-aware iteration (for reading documents with content controls)
    'iter_document_paragraphs',
    'iter_document_tables',
    
    # New / overridden functions
    'edit_run_text',
    'find_and_replace_text',
    'insert_numbered_list_near_text',
    'get_paragraph_by_id',
]
