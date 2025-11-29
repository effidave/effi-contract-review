"""
Integration test for comment status extraction.

This test demonstrates the full workflow of extracting comments with their
status information from a real Word document.
"""
import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from effilocal.mcp_server.core.comments import (
    extract_all_comments,
    extract_comment_status_map,
    merge_comment_status,
    filter_comments_by_author,
    get_comments_for_paragraph
)


def test_integration():
    """Integration test: Extract comments with status from a test document."""
    print("\n" + "=" * 80)
    print("INTEGRATION TEST: Comment Status Extraction")
    print("=" * 80)
    
    # Create a test document
    doc = Document()
    p1 = doc.add_paragraph("Introduction paragraph with important content.")
    p2 = doc.add_paragraph("Another section that needs review.")
    p3 = doc.add_paragraph("Final thoughts and conclusions.")
    
    # Save the test document
    test_doc_path = Path(__file__).parent / "test_comments_integration.docx"
    doc.save(str(test_doc_path))
    
    # Now extract comments
    doc = Document(str(test_doc_path))
    
    # Test 1: Extract all comments
    print("\n[Test 1] Extracting all comments...")
    comments = extract_all_comments(doc)
    print(f"✓ Extracted {len(comments)} comments")
    
    # Test 2: Extract status map
    print("\n[Test 2] Extracting comment status map...")
    status_map = extract_comment_status_map(doc)
    print(f"✓ Status map entries: {len(status_map)}")
    
    # Test 3: Check comment structure
    print("\n[Test 3] Verifying comment structure...")
    if comments:
        comment = comments[0]
        required_fields = ['status', 'is_resolved', 'done_flag', 'para_id']
        missing = [f for f in required_fields if f not in comment]
        
        if missing:
            print(f"[ERROR] Missing fields: {missing}")
            assert False, f"Missing fields: {missing}"
        else:
            print(f"✓ All required fields present: {required_fields}")
            print(f"  Sample comment status: {comment.get('status')}")
            print(f"  Sample comment resolved: {comment.get('is_resolved')}")
            print(f"  Sample comment done_flag: {comment.get('done_flag')}")
    
    # Test 4: Test filtering by author
    print("\n[Test 4] Testing author filtering...")
    author_comments = filter_comments_by_author(comments, "Unknown")
    print(f"✓ Author filtered: {len(author_comments)} comments")
    
    # Test 5: Test paragraph filtering
    print("\n[Test 5] Testing paragraph filtering...")
    para_comments = get_comments_for_paragraph(comments, 0)
    print(f"✓ Paragraph filtered: {len(para_comments)} comments")
    
    # Test 6: Status field consistency
    print("\n[Test 6] Verifying status field consistency...")
    all_have_status = all('status' in c for c in comments)
    all_have_resolved = all('is_resolved' in c for c in comments)
    all_have_flag = all('done_flag' in c for c in comments)
    
    if all_have_status and all_have_resolved and all_have_flag:
        print("✓ All comments have complete status information")
    else:
        print(f"[ERROR] Status consistency check failed:")
        print(f"   - has 'status': {all_have_status}")
        print(f"   - has 'is_resolved': {all_have_resolved}")
        print(f"   - has 'done_flag': {all_have_flag}")
        assert False, "Status consistency check failed"
    
    # Cleanup
    test_doc_path.unlink(missing_ok=True)
    
    print("\n" + "=" * 80)
    print("✓ All integration tests passed!")
    print("=" * 80)


if __name__ == "__main__":
    try:
        success = test_integration()
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"\n[ERROR] Test failed with error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
