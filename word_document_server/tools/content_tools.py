"""
Content tools for Word Document Server.

These tools add various types of content to Word documents,
including headings, paragraphs, tables, images, and page breaks.
"""
import os
from typing import List, Optional, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.document_utils import (
    find_and_replace_text,
    edit_run_text,
    insert_header_near_text,
    insert_numbered_list_near_text,
    insert_line_or_paragraph_near_text,
    replace_paragraph_block_below_header,
    replace_block_between_manual_anchors,
)
from word_document_server.core.styles import ensure_heading_style, ensure_table_style
from docx.oxml.ns import qn


def _parse_optional_bool(value: Optional[Any], field_name: str) -> Optional[bool]:
    """Normalize optional truthy fields supplied as bools/strings."""
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        normalized = value.strip().lower()
        if normalized in {"true", "1", "yes", "y", "on"}:
            return True
        if normalized in {"false", "0", "no", "n", "off"}:
            return False
    raise ValueError(f"Invalid parameter: {field_name} must be boolean-like")


def _normalize_color_string(color: Optional[str]) -> Optional[str]:
    """Return an uppercase hex RGB string or None."""
    if color is None:
        return None
    cleaned = str(color).strip()
    if cleaned.startswith("#"):
        cleaned = cleaned[1:]
    cleaned = cleaned.upper()
    if len(cleaned) != 6 or any(ch not in "0123456789ABCDEF" for ch in cleaned):
        raise ValueError("Invalid parameter: color must be a 6-digit hex value")
    return cleaned


def _apply_run_formatting(run, font_name: Optional[str], font_size_pt: Optional[float], bold: Optional[bool], italic: Optional[bool], color: Optional[str]) -> None:
    """Apply optional formatting attributes to a run."""
    if font_name:
        run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _apply_paragraph_bottom_border(paragraph) -> None:
    """Add a single bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = pBdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        pBdr.append(bottom)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')


async def add_heading(
    filename: str,
    text: str,
    level: int = 1,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    bold: Optional[Any] = None,
    italic: Optional[Any] = None,
    border_bottom: Optional[Any] = False,
    *,
    color: Optional[str] = None,
) -> str:
    """Add a heading to a Word document with optional formatting."""
    filename = ensure_docx_extension(filename)

    # Ensure level is converted to integer
    try:
        level = int(level)
    except (ValueError, TypeError):
        return "Invalid parameter: level must be an integer between 1 and 9"

    # Validate level range
    if level < 1 or level > 9:
        return f"Invalid heading level: {level}. Level must be between 1 and 9."

    text_value = "" if text is None else str(text)

    try:
        bold_value = _parse_optional_bool(bold, "bold")
    except ValueError as exc:
        return str(exc)

    try:
        italic_value = _parse_optional_bool(italic, "italic")
    except ValueError as exc:
        return str(exc)

    try:
        border_value = _parse_optional_bool(border_bottom, "border_bottom")
    except ValueError as exc:
        return str(exc)
    border_value = bool(border_value) if border_value is not None else False

    try:
        color_value = _normalize_color_string(color)
    except ValueError as exc:
        return str(exc)

    font_size_value = None
    if font_size is not None:
        try:
            font_size_value = float(font_size)
        except (TypeError, ValueError):
            return "Invalid parameter: font_size must be a positive number"
        if font_size_value <= 0:
            return "Invalid parameter: font_size must be greater than 0"

    font_name_value = font_name.strip() if isinstance(font_name, str) and font_name.strip() else None

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        ensure_heading_style(doc)

        fallback_used = False
        try:
            heading_para = doc.add_heading(text_value, level=level)
        except Exception:
            fallback_used = True
            heading_para = doc.add_paragraph()
            heading_para.style = doc.styles['Normal']

        heading_para.text = text_value
        if not heading_para.runs:
            heading_para.add_run(text_value)

        # Apply formatting to each run (usually a single run after setting text)
        effective_font_size = font_size_value
        if fallback_used and effective_font_size is None:
            effective_font_size = 16 if level == 1 else 14 if level == 2 else 12

        effective_bold = bold_value if bold_value is not None else (True if fallback_used else None)

        for run in heading_para.runs:
            _apply_run_formatting(run, font_name_value, effective_font_size, effective_bold, italic_value, color_value)

        if border_value:
            _apply_paragraph_bottom_border(heading_para)

        doc.save(filename)

        if fallback_used:
            return f"Heading '{text_value}' added to {filename} with direct formatting (style not available)"
        return f"Heading '{text_value}' (level {level}) added to {filename}"
    except Exception as exc:
        return f"Failed to add heading: {str(exc)}"


async def add_paragraph(
    filename: str,
    text: str,
    style: Optional[str] = None,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    bold: Optional[Any] = None,
    italic: Optional[Any] = None,
    color: Optional[str] = None,
) -> str:
    """Add a paragraph to a Word document with optional formatting."""
    filename = ensure_docx_extension(filename)

    text_value = "" if text is None else str(text)

    try:
        bold_value = _parse_optional_bool(bold, "bold")
    except ValueError as exc:
        return str(exc)

    try:
        italic_value = _parse_optional_bool(italic, "italic")
    except ValueError as exc:
        return str(exc)

    try:
        color_value = _normalize_color_string(color)
    except ValueError as exc:
        return str(exc)

    font_size_value = None
    if font_size is not None:
        try:
            font_size_value = float(font_size)
        except (TypeError, ValueError):
            return "Invalid parameter: font_size must be a positive number"
        if font_size_value <= 0:
            return "Invalid parameter: font_size must be greater than 0"

    font_name_value = font_name.strip() if isinstance(font_name, str) and font_name.strip() else None

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph()

        style_warning = None
        if style:
            try:
                paragraph.style = style
            except KeyError:
                style_warning = f"Style '{style}' not found; using default style."
                paragraph.style = doc.styles['Normal']

        run = paragraph.add_run(text_value)
        _apply_run_formatting(run, font_name_value, font_size_value, bold_value, italic_value, color_value)

        doc.save(filename)

        if style_warning:
            return f"Style '{style}' not found, paragraph added with default style to {filename}"
        return f"Paragraph added to {filename}"
    except Exception as exc:
        return f"Failed to add paragraph: {str(exc)}"


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    filename = ensure_docx_extension(filename)
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"


async def add_page_break(filename: str) -> str:
    """Add a page break to the document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        return f"Failed to add page break: {str(e)}"


async def add_table_of_contents(filename: str, title: str = "Table of Contents", max_level: int = 3) -> str:
    """Add a table of contents to a Word document based on heading styles.
    
    Args:
        filename: Path to the Word document
        title: Optional title for the table of contents
        max_level: Maximum heading level to include (1-9)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Ensure max_level is within valid range
        max_level = max(1, min(max_level, 9))
        
        doc = Document(filename)
        
        # Collect headings and their positions
        headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph style is a heading
            if paragraph.style and paragraph.style.name.startswith('Heading '):
                try:
                    # Extract heading level from style name
                    level = int(paragraph.style.name.split(' ')[1])
                    if level <= max_level:
                        headings.append({
                            'level': level,
                            'text': paragraph.text,
                            'position': i
                        })
                except (ValueError, IndexError):
                    # Skip if heading level can't be determined
                    pass
        
        if not headings:
            return f"No headings found in document {filename}. Table of contents not created."
        
        # Create a new document with the TOC
        toc_doc = Document()
        
        # Add title
        if title:
            toc_doc.add_heading(title, level=1)
        
        # Add TOC entries
        for heading in headings:
            # Indent based on level (using tab characters)
            indent = '    ' * (heading['level'] - 1)
            toc_doc.add_paragraph(f"{indent}{heading['text']}")
        
        # Add page break
        toc_doc.add_page_break()
        
        # Get content from original document
        for paragraph in doc.paragraphs:
            p = toc_doc.add_paragraph(paragraph.text)
            # Copy style if possible
            try:
                if paragraph.style:
                    p.style = paragraph.style.name
            except:
                pass
        
        # Copy tables
        for table in doc.tables:
            # Create a new table with the same dimensions
            new_table = toc_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            # Copy cell contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        new_table.cell(i, j).text = paragraph.text
        
        # Save the new document with TOC
        toc_doc.save(filename)
        
        return f"Table of contents with {len(headings)} entries added to {filename}"
    except Exception as e:
        return f"Failed to add table of contents: {str(e)}"


async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"


async def search_and_replace(filename: str, find_text: str, replace_text: str, whole_word_only: bool = False) -> str:
    """Search for text and replace all occurrences.
    
    Matches entirely within a single run are replaced automatically. Matches spanning multiple 
    runs are reported but not replaced; use edit_run_text to handle those cases precisely.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
        whole_word_only: If True, only replace whole words (not partial matches within words)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Perform find and replace
        count, snippets, split_matches = find_and_replace_text(doc, find_text, replace_text, whole_word_only=whole_word_only)
        
        # Save if we made any in-run replacements
        if count > 0:
            doc.save(filename)
        
        # Build response
        mode_info = " (whole word only)" if whole_word_only else ""
        response = ""
        
        if count > 0:
            response = f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'{mode_info}.\n"
            
            for i, snippet in enumerate(snippets, 1):
                response += f"\n--- Replacement {i} ({snippet['location']}) ---\n"
                response += f"Before: ...{snippet['before']}...\n"
                response += f"After:  ...{snippet['after']}...\n"
        
        # Report split matches
        if split_matches:
            if count == 0:
                response = f"No in-run occurrences of '{find_text}' found.\n"
            
            response += f"\n{len(split_matches)} match(es) span multiple runs and require manual editing:\n"
            
            for i, match in enumerate(split_matches, 1):
                para_idx = match['paragraph_index']
                run_info = ", ".join(
                    f"run {r['run_index']}[{r['offset_start']}:{r['offset_end']}]='{r['text']}'"
                    for r in match['runs']
                )
                response += f"\n  {i}. Paragraph {para_idx}: {run_info}\n"
            
            response += f"\nUse edit_run_text to modify these spans. Example for split match 1:\n"
            response += f"  edit_run_text(filename, {split_matches[0]['paragraph_index']}, {split_matches[0]['runs'][0]['run_index']}, new_text)\n"
        
        if count == 0 and not split_matches:
            response = f"No occurrences of '{find_text}' found."
        
        return response
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"


async def insert_header_near_text_tool(
    filename: str,
    target_text: Optional[str] = None,
    header_title: Optional[str] = None,
    position: str = 'after',
    header_style: str = 'Heading 1',
    target_paragraph_index: Optional[int] = None,
) -> str:
    """Insert a styled header relative to a target paragraph."""
    title_value = "" if header_title is None else str(header_title)
    normalized_position = (position or "after").lower()
    if normalized_position not in {"before", "after"}:
        return "position must be either 'before' or 'after'."

    style_value = header_style or 'Heading 1'
    return insert_header_near_text(
        filename,
        target_text,
        title_value,
        normalized_position,
        style_value,
        target_paragraph_index,
    )


async def insert_line_or_paragraph_near_text_tool(
    filename: str,
    target_text: Optional[str] = None,
    line_text: Optional[str] = None,
    position: str = 'after',
    line_style: Optional[str] = None,
    target_paragraph_index: Optional[int] = None,
) -> str:
    """Insert a paragraph relative to a target paragraph, optionally matching style."""
    text_value = "" if line_text is None else str(line_text)
    normalized_position = (position or "after").lower()
    if normalized_position not in {"before", "after"}:
        return "position must be either 'before' or 'after'."

    style_value = line_style.strip() if isinstance(line_style, str) and line_style.strip() else None

    return insert_line_or_paragraph_near_text(
        filename,
        target_text,
        text_value,
        normalized_position,
        style_value,
        target_paragraph_index,
    )


async def insert_numbered_list_near_text_tool(
    filename: str,
    target_text: Optional[str] = None,
    list_items: Optional[List[str]] = None,
    position: str = 'after',
    target_paragraph_index: Optional[int] = None,
    bullet_type: str = 'bullet',
) -> str:
    """Insert a bullet or numbered list relative to a target paragraph."""
    if list_items is None or not isinstance(list_items, list) or not list_items:
        return "list_items must be a non-empty list of strings."

    normalized_items = [str(item) for item in list_items]

    normalized_position = (position or "after").lower()
    if normalized_position not in {"before", "after"}:
        return "position must be either 'before' or 'after'."

    bullet_value = (bullet_type or 'bullet').lower()
    if bullet_value in {"bullet", "bullets"}:
        list_style = 'bullet'
    elif bullet_value in {"number", "numbered", "numbers"}:
        list_style = 'number'
    else:
        return "bullet_type must be 'bullet' or 'number'."

    return insert_numbered_list_near_text(
        filename,
        target_text,
        normalized_items,
        normalized_position,
        target_paragraph_index,
        list_style,
    )


async def insert_str_content_near_text(
    filename: str,
    target_text: Optional[str] = None,
    content_text: Optional[str] = None,
    position: str = 'after',
    content_style: Optional[str] = None,
    target_paragraph_index: Optional[int] = None,
) -> str:
    """
    Insert string content near a target paragraph. If the provided content_style appears to be a heading
    (defaults to Heading 1 when not specified), the content is inserted as a header. Otherwise a normal
    paragraph/line is inserted, with the optional paragraph style applied.
    """
    text_value = (content_text or "").strip()
    if not text_value:
        return "content_text must be provided and non-empty."

    position_normalized = (position or "after").lower()
    if position_normalized not in {"before", "after"}:
        return "position must be either 'before' or 'after'."

    style_value = (content_style or "").strip()
    is_heading = bool(style_value) and style_value.lower().startswith("heading")

    # If the caller explicitly requested a heading (style starts with 'heading'), use the header helper.
    if is_heading:
        header_style_value = style_value or "Heading 1"
        if header_style_value.lower().startswith("heading"):
            header_style_value = header_style_value.title()
        return await insert_header_near_text_tool(
            filename=filename,
            target_text=target_text,
            header_title=text_value,
            position=position_normalized,
            header_style=header_style_value,
            target_paragraph_index=target_paragraph_index,
        )

    # Non-heading content path inserts a paragraph/line with the provided style (if any).
    return await insert_line_or_paragraph_near_text_tool(
        filename=filename,
        target_text=target_text,
        line_text=text_value,
        position=position_normalized,
        line_style=style_value if style_value else None,
        target_paragraph_index=target_paragraph_index,
    )

#async def replace_paragraph_block_below_header_tool(filename: str, header_text: str, new_paragraphs: list, detect_block_end_fn=None) -> str:
#    """Reemplaza el bloque de párrafos debajo de un encabezado, evitando modificar TOC."""
#    return replace_paragraph_block_below_header(filename, header_text, new_paragraphs, detect_block_end_fn)

async def replace_block_between_manual_anchors_tool(filename: str, start_anchor_text: str, new_paragraphs: list, end_anchor_text: str = None, match_fn=None, new_paragraph_style: str = None) -> str:
    """Replace all content between start_anchor_text and end_anchor_text (or next logical header if not provided)."""
    return replace_block_between_manual_anchors(filename, start_anchor_text, new_paragraphs, end_anchor_text, match_fn, new_paragraph_style)


async def edit_run_text_tool(filename: str, paragraph_index: int, run_index: int, new_text: str, start_offset: Optional[int] = None, end_offset: Optional[int] = None) -> str:
    """
    Edit text within a specific run of a paragraph.
    
    Use this tool to handle matches that span multiple runs (as reported by search_and_replace).
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        run_index: Index of the run within the paragraph (0-based)
        new_text: Text to insert or replace
        start_offset: Optional start position within the run (0-based)
        end_offset: Optional end position within the run (0-based, exclusive).
                   If not provided, replaces the entire run text
    """
    return edit_run_text(filename, paragraph_index, run_index, new_text, start_offset, end_offset)


async def add_paragraph_after_clause(
    filename: str,
    clause_number: str,
    text: str,
    style: Optional[str] = None,
    inherit_numbering: bool = True,
) -> str:
    """
    Add a paragraph after a specific clause number, optionally inheriting its numbering.
    
    This tool uses the NumberingInspector to locate clauses by their rendered number
    (e.g., "1.2.3", "5", "7.1(a)") and inserts content after them.
    
    Args:
        filename: Path to the Word document
        clause_number: The rendered clause number to find (e.g., "1.2.3", "5.1")
        text: Text content to add
        style: Optional style name to apply to the new paragraph
        inherit_numbering: If True, inherit the numId and ilvl from the target clause
                          (creating a sibling clause at the same level)
    
    Returns:
        Success message or error description
    """
    from pathlib import Path
    
    try:
        from effilocal.doc.numbering_inspector import NumberingInspector
    except ImportError:
        return (
            "effilocal package not available. This tool requires effilocal for clause numbering analysis.\n"
            "The effilocal package should be in the same workspace."
        )
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        # Use NumberingInspector to analyze the document
        docx_path = Path(filename)
        inspector = NumberingInspector.from_docx(docx_path)
        rows, _ = inspector.analyze(debug=False)
        
        if not rows:
            return f"No paragraphs found in {filename}"
        
        # Find the clause with matching rendered number
        target_clause = None
        target_index = None
        
        for row in rows:
            rendered = row.get("rendered_number", "").strip()
            # Normalize for comparison (remove trailing dots)
            rendered_normalized = rendered.rstrip('.')
            clause_normalized = clause_number.strip().rstrip('.')
            
            if rendered_normalized == clause_normalized:
                target_clause = row
                # The idx in the row corresponds to the paragraph index in the document
                target_index = row.get("idx")
                break
        
        if target_clause is None:
            return f"Clause '{clause_number}' not found in {filename}"
        
        if target_index is None:
            return f"Could not determine paragraph index for clause '{clause_number}'"
        
        # Load the document
        doc = Document(filename)
        
        # Validate target_index
        if target_index < 0 or target_index >= len(doc.paragraphs):
            return f"Invalid paragraph index {target_index} for clause '{clause_number}'"
        
        # Get the target paragraph to extract numbering properties
        target_paragraph = doc.paragraphs[target_index]
        
        # Insert new paragraph after target
        # We need to insert into the XML structure directly
        target_p = target_paragraph._p
        parent = target_p.getparent()
        target_position = list(parent).index(target_p)
        
        # Create new paragraph element
        new_p = OxmlElement('w:p')
        
        # Add paragraph properties if needed
        pPr = OxmlElement('w:pPr')
        new_p.append(pPr)
        
        # Apply style if specified
        if style:
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style)
            pPr.append(pStyle)
        elif target_paragraph.style:
            # Inherit style from target if no style specified
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), target_paragraph.style.style_id)
            pPr.append(pStyle)
        
        # Inherit numbering properties if requested
        if inherit_numbering:
            num_id = target_clause.get("numId")
            ilvl = target_clause.get("ilvl")
            
            if num_id is not None and num_id != "":
                numPr = OxmlElement('w:numPr')
                
                # Add ilvl (level)
                ilvl_elem = OxmlElement('w:ilvl')
                ilvl_elem.set(qn('w:val'), str(ilvl if ilvl is not None else 0))
                numPr.append(ilvl_elem)
                
                # Add numId
                numId_elem = OxmlElement('w:numId')
                numId_elem.set(qn('w:val'), str(num_id))
                numPr.append(numId_elem)
                
                pPr.append(numPr)
        
        # Add text run
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        new_p.append(r)
        
        # Insert after target paragraph
        parent.insert(target_position + 1, new_p)
        
        doc.save(filename)
        
        numbering_info = ""
        if inherit_numbering:
            num_id = target_clause.get("numId")
            ilvl = target_clause.get("ilvl")
            if num_id:
                numbering_info = f" with inherited numbering (numId={num_id}, level={ilvl})"
        
        return f"Paragraph added after clause '{clause_number}'{numbering_info} in {filename}"
        
    except Exception as exc:
        return f"Failed to add paragraph after clause: {str(exc)}"


async def add_paragraphs_after_clause(
    filename: str,
    clause_number: str,
    paragraphs: List[str],
    style: Optional[str] = None,
    inherit_numbering: bool = True,
) -> str:
    """
    Add multiple paragraphs after a specific clause number.
    
    This tool adds multiple paragraphs sequentially after a clause, each inheriting
    the same numbering properties. This is useful for adding multiple items at the
    same level (e.g., adding 7.1(b), 7.1(c), 7.1(d) after 7.1(a)).
    
    Args:
        filename: Path to the Word document
        clause_number: The rendered clause number to find (e.g., "1.2.3", "7.1(a)")
        paragraphs: List of paragraph texts to add
        style: Optional style name to apply to all new paragraphs
        inherit_numbering: If True, all paragraphs inherit the numId and ilvl from the target clause
    
    Returns:
        Success message with count of paragraphs added
    """
    if not paragraphs or not isinstance(paragraphs, list):
        return "paragraphs parameter must be a non-empty list of strings"
    
    filename = ensure_docx_extension(filename)
    
    # Add first paragraph
    result = await add_paragraph_after_clause(
        filename, clause_number, paragraphs[0], style, inherit_numbering
    )
    
    if "Failed" in result or "not found" in result or "does not exist" in result:
        return result
    
    # If more paragraphs, add them sequentially
    # Each one is added after the previous, maintaining the same level
    if len(paragraphs) > 1:
        try:
            from pathlib import Path
            from effilocal.doc.numbering_inspector import NumberingInspector
            
            for i, text in enumerate(paragraphs[1:], 1):
                # Re-analyze document to find the last inserted paragraph
                docx_path = Path(filename)
                inspector = NumberingInspector.from_docx(docx_path)
                rows, _ = inspector.analyze(debug=False)
                
                # Find all paragraphs with the same numId and ilvl as target
                target_num_id = None
                target_ilvl = None
                
                for row in rows:
                    rendered = row.get("rendered_number", "").strip().rstrip('.')
                    clause_normalized = clause_number.strip().rstrip('.')
                    if rendered == clause_normalized:
                        target_num_id = row.get("numId")
                        target_ilvl = row.get("ilvl")
                        break
                
                if target_num_id is None:
                    break
                
                # Find the last paragraph with this numId and ilvl
                last_matching_idx = None
                for row in rows:
                    if (row.get("numId") == target_num_id and 
                        row.get("ilvl") == target_ilvl):
                        last_matching_idx = row.get("idx")
                
                if last_matching_idx is None:
                    continue
                
                # Insert after the last matching paragraph
                doc = Document(filename)
                if last_matching_idx >= len(doc.paragraphs):
                    break
                
                target_paragraph = doc.paragraphs[last_matching_idx]
                target_p = target_paragraph._p
                parent = target_p.getparent()
                target_position = list(parent).index(target_p)
                
                new_p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                new_p.append(pPr)
                
                if style:
                    pStyle = OxmlElement('w:pStyle')
                    pStyle.set(qn('w:val'), style)
                    pPr.append(pStyle)
                elif target_paragraph.style:
                    pStyle = OxmlElement('w:pStyle')
                    pStyle.set(qn('w:val'), target_paragraph.style.style_id)
                    pPr.append(pStyle)
                
                if inherit_numbering and target_num_id:
                    numPr = OxmlElement('w:numPr')
                    ilvl_elem = OxmlElement('w:ilvl')
                    ilvl_elem.set(qn('w:val'), str(target_ilvl if target_ilvl is not None else 0))
                    numPr.append(ilvl_elem)
                    numId_elem = OxmlElement('w:numId')
                    numId_elem.set(qn('w:val'), str(target_num_id))
                    numPr.append(numId_elem)
                    pPr.append(numPr)
                
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = text
                r.append(t)
                new_p.append(r)
                
                parent.insert(target_position + 1, new_p)
                doc.save(filename)
            
            return f"Added {len(paragraphs)} paragraph(s) after clause '{clause_number}' in {filename}"
            
        except Exception as exc:
            return f"Added 1 paragraph successfully, but failed to add remaining paragraphs: {str(exc)}"
    
    return result
