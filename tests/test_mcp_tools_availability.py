#!/usr/bin/env python3
"""Pytest suite verifying that every MCP tool in main.py responds via stdio."""

import json
import subprocess
import sys
import tempfile
import threading
from pathlib import Path

import pytest
from docx import Document

EXPECTED_TOOLS = {
    "create_document",
    "copy_document",
    "get_document_info",
    "get_document_text",
    "get_document_outline",
    "list_available_documents",
    "get_document_xml",
    "add_paragraph",
    "add_heading",
    "add_picture",
    "add_table",
    "add_page_break",
    "delete_paragraph",
    "search_and_replace",
    "insert_header_near_text",
    "insert_line_or_paragraph_near_text",
    "insert_numbered_list_near_text",
    "get_paragraph_text_from_document",
    "find_text_in_document",
    "edit_run_text",
    "replace_block_between_manual_anchors",
    "create_custom_style",
    "format_text",
    "format_table",
    "set_table_cell_shading",
    "apply_table_alternating_rows",
    "highlight_table_header",
    "merge_table_cells",
    "merge_table_cells_horizontal",
    "merge_table_cells_vertical",
    "set_table_cell_alignment",
    "set_table_alignment_all",
    "set_table_column_width",
    "set_table_column_widths",
    "set_table_width",
    "auto_fit_table_columns",
    "format_table_cell_text",
    "set_table_cell_padding",
    "protect_document",
    "unprotect_document",
    "add_footnote_to_document",
    "add_footnote_after_text",
    "add_footnote_before_text",
    "add_footnote_enhanced",
    "add_endnote_to_document",
    "customize_footnote_style",
    "delete_footnote_from_document",
    "add_footnote_robust",
    "validate_document_footnotes",
    "delete_footnote_robust",
    "get_all_comments",
    "get_comments_by_author",
    "get_comments_for_paragraph",
    "convert_to_pdf",
    "analyze_document_numbering",
    "get_numbering_summary",
    "extract_outline_structure",
    "add_paragraph_after_clause",
    "add_paragraphs_after_clause",
}


try:  # Used to relax expectations when Pillow is absent.
    import PIL.Image  # type: ignore # noqa: F401
except ImportError:
    PIL_AVAILABLE = False
else:
    PIL_AVAILABLE = True


class MCPToolTester:
    """Helper that manages the stdio MCP server and tracks tool calls."""

    def __init__(self):
        self.fixture_file: Path | None = None
        self.test_results: dict[str, dict] = {}
        self.called_tools: set[str] = set()
        self.mcp_process: subprocess.Popen | None = None

    def create_fixture_document(self) -> Path:
        # Use a temporary file with a unique name to avoid conflicts
        import tempfile
        fd, temp_path = tempfile.mkstemp(suffix=".docx", prefix="mcp_test_")
        import os
        os.close(fd)  # Close the file descriptor, we'll write with python-docx
        self.fixture_file = Path(temp_path)

        doc = Document()
        doc.add_heading("MCP Test Fixture Document", 0)
        doc.add_paragraph("This is a test document created for MCP tool testing.")
        doc.add_paragraph("Sample paragraph with some text to search and replace.")
        doc.add_paragraph("Another paragraph for testing various operations.")
        doc.add_heading("Section 1: Testing Content", level=1)
        doc.add_paragraph("Content under section 1.")
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(0, 2).text = "Header 3"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        table.cell(1, 2).text = "Data 3"
        doc.save(self.fixture_file)
        return self.fixture_file

    def cleanup_fixture(self) -> None:
        if self.fixture_file and self.fixture_file.exists():
            try:
                self.fixture_file.unlink()
            except PermissionError:
                # File is locked by another process, skip cleanup
                pass

    def _read_response_until_id(self, target_id: int, timeout_sec: float = 20.0) -> dict:
        import time

        start = time.time()
        while time.time() - start < timeout_sec:
            line = self.mcp_process.stdout.readline()
            if not line:
                time.sleep(0.05)
                continue
            line_stripped = line.strip()
            if not line_stripped:
                continue
            try:
                obj = json.loads(line_stripped)
            except Exception:
                continue
            if isinstance(obj, dict) and obj.get("id") == target_id:
                return obj
        return {"error": {"code": -32000, "message": f"Timeout waiting for response id {target_id}"}}

    def start_mcp_server(self) -> bool:
        try:
            self.mcp_process = subprocess.Popen(
                [sys.executable, "-m", "word_document_server.main"],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                bufsize=0,
            )

            def _stderr_drain(stream):
                for line in stream:
                    line = line.rstrip()
                    if line:
                        print(f"[MCP stderr] {line}")

            stderr_thread = threading.Thread(target=_stderr_drain, args=(self.mcp_process.stderr,), daemon=True)
            stderr_thread.start()
            return True
        except Exception as exc:
            print(f"Failed to start MCP server: {exc}")
            return False

    def stop_mcp_server(self) -> None:
        if self.mcp_process:
            self.mcp_process.terminate()
            self.mcp_process.wait()
            self.mcp_process = None

    def call_mcp_tool(self, tool_name: str, **kwargs) -> dict:
        if not self.mcp_process:
            response = {"success": False, "error": "MCP server not running"}
            self.test_results[tool_name] = response
            return response

        if not hasattr(self, "_initialized"):
            init_payload = {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "initialize",
                "params": {
                    "protocolVersion": "2024-11-05",
                    "capabilities": {},
                    "clientInfo": {"name": "pytest-client", "version": "1.0.0"},
                },
            }
            try:
                init_request = json.dumps(init_payload) + "\n"
                self.mcp_process.stdin.write(init_request)
                self.mcp_process.stdin.flush()
                init_result = self._read_response_until_id(1)
                if not init_result.get("result"):
                    response = {"success": False, "error": f"Initialization failed: {init_result}"}
                    self.test_results[tool_name] = response
                    return response
                tools_list_payload = {
                    "jsonrpc": "2.0",
                    "id": 2,
                    "method": "tools/list",
                    "params": {},
                }
                self.mcp_process.stdin.write(json.dumps(tools_list_payload) + "\n")
                self.mcp_process.stdin.flush()
                _ = self._read_response_until_id(2)
                self._initialized = True
            except Exception as exc:
                response = {"success": False, "error": f"Initialization failed: {exc}"}
                self.test_results[tool_name] = response
                return response

        if not hasattr(self, "_req_id"):
            self._req_id = 10
        self._req_id += 1
        req_id = self._req_id

        payload = {
            "jsonrpc": "2.0",
            "id": req_id,
            "method": "tools/call",
            "params": {"name": tool_name, "arguments": kwargs},
        }

        try:
            request_line = json.dumps(payload) + "\n"
            self.mcp_process.stdin.write(request_line)
            self.mcp_process.stdin.flush()
            result = self._read_response_until_id(req_id)
            if isinstance(result, dict) and "error" in result:
                response = {"success": False, "error": result["error"]}
            else:
                response = {"success": True, "result": result.get("result")}
        except json.JSONDecodeError as exc:
            response = {"success": False, "error": f"JSON decode error: {exc}"}
        except Exception as exc:
            response = {"success": False, "error": f"Communication error: {exc}"}

        self.test_results[tool_name] = response
        if response.get("success"):
            self.called_tools.add(tool_name)
        print(f"[MCP call] {tool_name} => {response}")
        return response


@pytest.fixture(scope="session")
def mcp_tester():
    tester = MCPToolTester()
    if not tester.start_mcp_server():
        pytest.fail("Failed to start MCP server via stdio")
    tester.create_fixture_document()
    yield tester
    tester.cleanup_fixture()
    tester.stop_mcp_server()


def test_document_tools(mcp_tester):
    tests = [
        ("create_document", {"filename": str(mcp_tester.fixture_file.with_suffix(".new.docx")), "title": "Test Doc", "author": "Test Author"}),
        ("copy_document", {"source_filename": str(mcp_tester.fixture_file), "destination_filename": str(mcp_tester.fixture_file.with_suffix(".copy.docx"))}),
        ("get_document_info", {"filename": str(mcp_tester.fixture_file)}),
        ("get_document_text", {"filename": str(mcp_tester.fixture_file)}),
        ("get_document_outline", {"filename": str(mcp_tester.fixture_file)}),
        ("list_available_documents", {"directory": str(mcp_tester.fixture_file.parent)}),
        ("get_document_xml", {"filename": str(mcp_tester.fixture_file)}),
        ("save_document_as_markdown", {"filename": str(mcp_tester.fixture_file)}),
    ]

    for tool_name, params in tests:
        result = mcp_tester.call_mcp_tool(tool_name, **params)
        assert result["success"], f"{tool_name} failed: {result.get('error')}"

    for suffix in [".new.docx", ".copy.docx", ".md"]:
        temp_path = mcp_tester.fixture_file.with_suffix(suffix)
        if temp_path.exists():
            temp_path.unlink()


def test_content_tools(mcp_tester):
    test_file = mcp_tester.fixture_file.with_suffix(".content_test.docx")
    copy_result = mcp_tester.call_mcp_tool(
        "copy_document",
        source_filename=str(mcp_tester.fixture_file),
        destination_filename=str(test_file),
    )
    assert copy_result["success"], f"copy_document failed: {copy_result.get('error')}"

    tests = [
        ("add_paragraph", {"filename": str(test_file), "text": "<<START ANCHOR>>", "bold": True}),
        ("add_heading", {"filename": str(test_file), "text": "Test Heading", "level": 2}),
        ("add_table", {"filename": str(test_file), "rows": 2, "cols": 2}),
        ("add_page_break", {"filename": str(test_file)}),
        ("insert_line_or_paragraph_near_text", {"filename": str(test_file), "target_text": "Test Heading", "line_text": "<<END ANCHOR>>", "position": "after"}),
        ("replace_block_between_manual_anchors", {"filename": str(test_file), "start_anchor_text": "<<START ANCHOR>>", "end_anchor_text": "<<END ANCHOR>>", "new_paragraphs": ["New content paragraph 1", "New content paragraph 2"]}),
        ("search_and_replace", {"filename": str(test_file), "find_text": "Sample", "replace_text": "modified"}),
        ("delete_paragraph", {"filename": str(test_file), "paragraph_index": 0}),
        ("insert_header_near_text", {"filename": str(test_file), "target_text": "MCP", "header_title": "New Header", "position": "after"}),
        ("insert_numbered_list_near_text", {"filename": str(test_file), "target_paragraph_index": 0, "list_items": ["Item 1", "Item 2", "Item 3"], "position": "after", "bullet_type": "bullet"}),
        ("get_paragraph_text_from_document", {"filename": str(test_file), "paragraph_index": 0}),
        ("find_text_in_document", {"filename": str(test_file), "text_to_find": "MCP"}),
        ("edit_run_text", {"filename": str(test_file), "paragraph_index": 0, "run_index": 0, "new_text": "Modified text"}),
    ]

    try:
        for tool_name, params in tests:
            result = mcp_tester.call_mcp_tool(tool_name, **params)
            assert result["success"], f"{tool_name} failed: {result.get('error')}"
    finally:
        if test_file.exists():
            test_file.unlink()


def test_format_tools(mcp_tester):
    test_file = mcp_tester.fixture_file.with_suffix(".format_test.docx")
    copy_result = mcp_tester.call_mcp_tool(
        "copy_document",
        source_filename=str(mcp_tester.fixture_file),
        destination_filename=str(test_file),
    )
    assert copy_result["success"], f"copy_document failed: {copy_result.get('error')}"

    tests = [
        ("create_custom_style", {"filename": str(test_file), "style_name": "TestStyle", "bold": True, "font_size": 14}),
        ("format_text", {"filename": str(test_file), "paragraph_index": 1, "start_pos": 0, "end_pos": 10, "bold": True}),
    ]

    try:
        for tool_name, params in tests:
            result = mcp_tester.call_mcp_tool(tool_name, **params)
            assert result["success"], f"{tool_name} failed: {result.get('error')}"
    finally:
        if test_file.exists():
            test_file.unlink()


def test_table_tools(mcp_tester):
    test_file = mcp_tester.fixture_file.with_suffix(".table_test.docx")
    copy_result = mcp_tester.call_mcp_tool(
        "copy_document",
        source_filename=str(mcp_tester.fixture_file),
        destination_filename=str(test_file),
    )
    assert copy_result["success"], f"copy_document failed: {copy_result.get('error')}"

    tests = [
        ("format_table", {"filename": str(test_file), "table_index": 0, "has_header_row": True}),
        ("set_table_cell_shading", {"filename": str(test_file), "table_index": 0, "row_index": 0, "col_index": 0, "fill_color": "FFFF00"}),
        ("apply_table_alternating_rows", {"filename": str(test_file), "table_index": 0}),
        ("highlight_table_header", {"filename": str(test_file), "table_index": 0}),
        ("merge_table_cells", {"filename": str(test_file), "table_index": 0, "start_row": 0, "start_col": 0, "end_row": 0, "end_col": 1}),
        ("merge_table_cells_horizontal", {"filename": str(test_file), "table_index": 0, "row_index": 1, "start_col": 0, "end_col": 1}),
        ("merge_table_cells_vertical", {"filename": str(test_file), "table_index": 0, "col_index": 2, "start_row": 0, "end_row": 1}),
        ("set_table_cell_alignment", {"filename": str(test_file), "table_index": 0, "row_index": 0, "col_index": 0, "horizontal": "center", "vertical": "center"}),
        ("set_table_alignment_all", {"filename": str(test_file), "table_index": 0, "horizontal": "center"}),
        ("set_table_column_width", {"filename": str(test_file), "table_index": 0, "col_index": 0, "width": 100, "width_type": "points"}),
        ("set_table_column_widths", {"filename": str(test_file), "table_index": 0, "widths": [80, 100, 120], "width_type": "points"}),
        ("set_table_width", {"filename": str(test_file), "table_index": 0, "width": 400, "width_type": "points"}),
        ("auto_fit_table_columns", {"filename": str(test_file), "table_index": 0}),
        ("format_table_cell_text", {"filename": str(test_file), "table_index": 0, "row_index": 0, "col_index": 0, "text_content": "Formatted Header", "bold": True, "color": "FF0000"}),
        ("set_table_cell_padding", {"filename": str(test_file), "table_index": 0, "row_index": 0, "col_index": 0, "top": 5, "bottom": 5, "left": 10, "right": 10}),
    ]

    try:
        for tool_name, params in tests:
            result = mcp_tester.call_mcp_tool(tool_name, **params)
            assert result["success"], f"{tool_name} failed: {result.get('error')}"
    finally:
        if test_file.exists():
            test_file.unlink()


def test_footnote_tools(mcp_tester):
    test_file = mcp_tester.fixture_file.with_suffix(".footnote_test.docx")
    copy_result = mcp_tester.call_mcp_tool(
        "copy_document",
        source_filename=str(mcp_tester.fixture_file),
        destination_filename=str(test_file),
    )
    assert copy_result["success"], f"copy_document failed: {copy_result.get('error')}"

    tests = [
        ("add_footnote_to_document", {"filename": str(test_file), "paragraph_index": 1, "footnote_text": "Test footnote"}),
        ("add_footnote_after_text", {"filename": str(test_file), "search_text": "Sample", "footnote_text": "Footnote after text"}),
        ("add_footnote_before_text", {"filename": str(test_file), "search_text": "paragraph", "footnote_text": "Footnote before text"}),
        ("add_footnote_enhanced", {"filename": str(test_file), "paragraph_index": 2, "footnote_text": "Enhanced footnote"}),
        ("add_endnote_to_document", {"filename": str(test_file), "paragraph_index": 3, "endnote_text": "Test endnote"}),
        ("customize_footnote_style", {"filename": str(test_file), "numbering_format": "a, b, c", "start_number": 1}),
        ("delete_footnote_from_document", {"filename": str(test_file), "search_text": "Sample"}),
        ("add_footnote_robust", {"filename": str(test_file), "search_text": "testing", "footnote_text": "Robust footnote test", "validate_location": True}),
        ("validate_document_footnotes", {"filename": str(test_file)}),
        ("delete_footnote_robust", {"filename": str(test_file), "footnote_id": 2, "clean_orphans": True}),
    ]

    try:
        for tool_name, params in tests:
            result = mcp_tester.call_mcp_tool(tool_name, **params)
            assert result["success"], f"{tool_name} failed: {result.get('error')}"
    finally:
        if test_file.exists():
            test_file.unlink()


def test_comment_tools(mcp_tester):
    tests = [
        ("get_all_comments", {"filename": str(mcp_tester.fixture_file)}),
        ("get_comments_by_author", {"filename": str(mcp_tester.fixture_file), "author": "TestAuthor"}),
        ("get_comments_for_paragraph", {"filename": str(mcp_tester.fixture_file), "paragraph_index": 0}),
    ]

    for tool_name, params in tests:
        result = mcp_tester.call_mcp_tool(tool_name, **params)
        assert result["success"], f"{tool_name} failed: {result.get('error')}"


def test_protection_tools(mcp_tester):
    test_file = mcp_tester.fixture_file.with_suffix(".protection_test.docx")
    copy_result = mcp_tester.call_mcp_tool(
        "copy_document",
        source_filename=str(mcp_tester.fixture_file),
        destination_filename=str(test_file),
    )
    assert copy_result["success"], f"copy_document failed: {copy_result.get('error')}"

    tests = [
        ("protect_document", {"filename": str(test_file), "password": "test123"}),
        ("unprotect_document", {"filename": str(test_file), "password": "test123"}),
    ]

    try:
        for tool_name, params in tests:
            result = mcp_tester.call_mcp_tool(tool_name, **params)
            assert result["success"], f"{tool_name} failed: {result.get('error')}"
    finally:
        if test_file.exists():
            test_file.unlink()


def test_utility_tools(mcp_tester):
    pdf_path = mcp_tester.fixture_file.with_suffix(".pdf")
    result = mcp_tester.call_mcp_tool(
        "convert_to_pdf",
        filename=str(mcp_tester.fixture_file),
        output_filename=str(pdf_path),
    )
    assert result["success"], f"convert_to_pdf failed: {result.get('error')}"
    if pdf_path.exists():
        pdf_path.unlink()


def test_image_tools(mcp_tester):
    Image = pytest.importorskip("PIL.Image", reason="Pillow is required for image insertion tests")

    test_image = mcp_tester.fixture_file.parent / "mcp_test_image.png"
    Image.new("RGB", (100, 100), color="red").save(test_image)

    test_file = mcp_tester.fixture_file.with_suffix(".image_test.docx")
    copy_result = mcp_tester.call_mcp_tool(
        "copy_document",
        source_filename=str(mcp_tester.fixture_file),
        destination_filename=str(test_file),
    )
    assert copy_result["success"], f"copy_document failed: {copy_result.get('error')}"

    try:
        result = mcp_tester.call_mcp_tool(
            "add_picture",
            filename=str(test_file),
            image_path=str(test_image),
            width=2.0,
        )
        assert result["success"], f"add_picture failed: {result.get('error')}"
    finally:
        if test_file.exists():
            test_file.unlink()
        if test_image.exists():
            test_image.unlink()


def test_numbering_tools(mcp_tester):
    # Use a document with known numbering
    fixtures_dir = Path(__file__).parent / "fixtures"
    numbered_doc = fixtures_dir / "numbering_decimal.docx"
    
    if not numbered_doc.exists():
        # Fallback to regular fixture
        numbered_doc = mcp_tester.fixture_file
    
    # Create a temporary copy to avoid modifying the fixture
    test_file = Path(tempfile.gettempdir()) / "test_numbering_tools.docx"
    import shutil
    shutil.copy2(numbered_doc, test_file)
    
    tests = [
        ("analyze_document_numbering", {"filename": str(test_file), "debug": False, "include_non_numbered": False}),
        ("get_numbering_summary", {"filename": str(test_file)}),
        ("extract_outline_structure", {"filename": str(test_file), "max_level": None}),
        ("add_paragraph_after_clause", {"filename": str(test_file), "clause_number": "1", "text": "Test paragraph", "inherit_numbering": True}),
        ("add_paragraphs_after_clause", {"filename": str(test_file), "clause_number": "1", "paragraphs": ["Test 1", "Test 2"], "inherit_numbering": True}),
    ]

    try:
        for tool_name, params in tests:
            result = mcp_tester.call_mcp_tool(tool_name, **params)
            assert result["success"], f"{tool_name} failed: {result.get('error')}"
    finally:
        if test_file.exists():
            test_file.unlink()


def test_all_expected_tools_were_called(mcp_tester):
    expected_tools = set(EXPECTED_TOOLS)
    if not PIL_AVAILABLE:
        expected_tools.discard("add_picture")

    missing = expected_tools - set(mcp_tester.test_results.keys())
    assert not missing, f"Missing tool coverage: {sorted(missing)}"

    failures = {
        name: result
        for name, result in mcp_tester.test_results.items()
        if name in expected_tools and not result.get("success")
    }
    assert not failures, f"Tools failed: {failures}"


def main():
    raise SystemExit(pytest.main([__file__]))


if __name__ == "__main__":
    main()
