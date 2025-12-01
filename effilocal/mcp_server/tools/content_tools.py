"""
Content tools for effilocal MCP server.

This module extends word_document_server.tools.content_tools with:
1. Color parameter support for add_heading
2. New contract-specific tools (add_paragraph_after_clause)
3. Run-level text editing (edit_run_text_tool)

Override Pattern:
- Import most upstream functions as pass-through
- Add new functions for contract-specific features
- Minimal wrappers for extended signatures
"""

import os
import uuid
from typing import List, Optional, Any
from pathlib import Path

from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import all upstream functions (pass-through)
from word_document_server.tools.content_tools import (
    add_table,
    add_picture, 
    add_page_break,
    add_table_of_contents,
    delete_paragraph,
)

# Import wrapper tools from upstream
from word_document_server.tools import content_tools as upstream_content_tools

# Import upstream add_heading and add_paragraph to wrap
from word_document_server.tools.content_tools import add_heading as upstream_add_heading
from word_document_server.tools.content_tools import add_paragraph as upstream_add_paragraph

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension

# Import NEW functions from our own utils (not in upstream)
from effilocal.mcp_server.utils.document_utils import edit_run_text, get_paragraph_by_id


# ============================================================================
# OVERRIDDEN Functions (minimal wrappers with extended signatures)
# ============================================================================

async def add_heading(
    filename: str,
    text: str,
    level: int = 1,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    border_bottom: bool = False,
    *,
    color: Optional[str] = None,
) -> str:
    """Add a heading to a Word document with optional formatting.
    
    EXTENDED from upstream with color parameter support.
    """
    # Call upstream for base functionality
    result = await upstream_add_heading(
        filename, text, level, font_name, 
        int(font_size) if font_size else None,
        bold, italic, border_bottom
    )
    
    # If color is specified and heading was added successfully, apply it
    if color and "added to" in result:
        try:
            doc = Document(filename)
            # Find the last paragraph (just added)
            last_para = doc.paragraphs[-1] if doc.paragraphs else None
            if last_para:
                # Normalize color
                color_hex = color.strip().lstrip('#').upper()
                if len(color_hex) == 6:
                    for run in last_para.runs:
                        run.font.color.rgb = RGBColor.from_string(color_hex)
                    doc.save(filename)
        except Exception:
            pass  # Silently fail color application, heading was still added
    
    return result


async def add_paragraph(
    filename: str,
    text: str,
    style: Optional[str] = None,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> str:
    """Add a paragraph to a Word document with optional formatting.
    
    Pass-through to upstream (upstream already supports all parameters).
    """
    return await upstream_add_paragraph(
        filename, text, style, font_name,
        int(font_size) if font_size else None,
        bold, italic, color
    )


async def search_and_replace(
    filename: str, 
    find_text: str, 
    replace_text: str, 
    whole_word_only: bool = False
) -> str:
    """Search and replace text in a Word document.
    
    EXTENDED - delegates to document_utils with whole_word_only support.
    """
    from effilocal.mcp_server.utils.document_utils import find_and_replace_text
    return find_and_replace_text(filename, find_text, replace_text, whole_word_only)


async def insert_str_content_near_text(
    filename: str,
    target_text: Optional[str] = None,
    content: str = "",
    position: str = 'after',
    target_paragraph_index: Optional[int] = None
) -> str:
    """Insert string content near target text or at a specific paragraph index.
    
    NEW function - provides flexible content insertion.
    """
    from word_document_server.utils.document_utils import insert_line_or_paragraph_near_text
    return insert_line_or_paragraph_near_text(
        filename, target_text, content, position, None, target_paragraph_index
    )


# ============================================================================
# NEW Functions (contract-specific tools)
# ============================================================================

async def edit_run_text_tool(
    filename: str, 
    paragraph_index: int, 
    run_index: int, 
    new_text: str, 
    start_offset: Optional[int] = None, 
    end_offset: Optional[int] = None
) -> str:
    """Edit text within a specific run of a paragraph.
    
    NEW function - handles matches that span multiple runs (as reported by search_and_replace).
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        run_index: Index of the run within the paragraph (0-based)
        new_text: Text to insert or replace
        start_offset: Optional start position within the run (0-based)
        end_offset: Optional end position within the run (0-based, exclusive).
                   If not provided, replaces the entire run text
    """
    return edit_run_text(filename, paragraph_index, run_index, new_text, start_offset, end_offset)


async def add_paragraph_after_clause(
    filename: str,
    clause_number: str,
    text: str,
    style: Optional[str] = None,
    inherit_numbering: bool = True,
) -> str:
    """Add a paragraph after a specific clause number, optionally inheriting its numbering.
    
    NEW function - contract-specific tool using effilocal.NumberingInspector.
    
    This tool uses the NumberingInspector to locate clauses by their rendered number
    (e.g., "1.2.3", "5", "7.1(a)") and inserts content after them.
    
    Args:
        filename: Path to the Word document
        clause_number: The rendered clause number to find (e.g., "1.2.3", "5.1")
        text: Text content to add
        style: Optional style name to apply to the new paragraph
        inherit_numbering: If True, inherit the numId and ilvl from the target clause
                          (creating a sibling clause at the same level)
    
    Returns:
        Success message or error description
    """
    try:
        from effilocal.doc.numbering_inspector import NumberingInspector
    except ImportError:
        return (
            "effilocal package not available. This tool requires effilocal for clause numbering analysis.\n"
            "The effilocal package should be in the same workspace."
        )
    
    # Import add_custom_para_id from attachment_tools (effilocal only - not in upstream)
    from effilocal.mcp_server.tools.attachment_tools import add_custom_para_id
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        # Use NumberingInspector to analyze the document
        docx_path = Path(filename)
        inspector = NumberingInspector.from_docx(docx_path)
        rows, _ = inspector.analyze(debug=False)
        
        if not rows:
            return f"No paragraphs found in {filename}"
        
        # Find the clause with matching rendered number
        target_clause = None
        target_index = None
        
        for row in rows:
            rendered = row.get("rendered_number", "").strip()
            # Normalize for comparison (remove trailing dots)
            rendered_normalized = rendered.rstrip('.')
            clause_normalized = clause_number.strip().rstrip('.')
            
            if rendered_normalized == clause_normalized:
                target_clause = row
                # The idx in the row corresponds to the paragraph index in the document
                target_index = row.get("idx")
                break
        
        if target_clause is None:
            return f"Clause '{clause_number}' not found in {filename}"
        
        if target_index is None:
            return f"Could not determine paragraph index for clause '{clause_number}'"
        
        # Find the last child clause (subclauses at deeper levels)
        target_ilvl = target_clause.get("ilvl")
        target_num_id = target_clause.get("numId")
        
        # Start from the target clause and look for any deeper-level clauses that follow
        last_child_index = target_index
        for i in range(target_index + 1, len(rows)):
            row = rows[i]
            row_ilvl = row.get("ilvl")
            row_num_id = row.get("numId")
            row_rendered = row.get("rendered_number", "").strip()
            
            # Check if this is a child clause (deeper level with same numId)
            if row_num_id == target_num_id and row_ilvl is not None and target_ilvl is not None:
                if row_ilvl > target_ilvl:
                    # This is a child clause, update the insertion point
                    last_child_index = row.get("idx", last_child_index)
                elif row_ilvl == target_ilvl:
                    # Same level clause - we've moved past the children
                    break
                else:
                    # Shallower level - we've moved past the section
                    break
            elif row_num_id != target_num_id and row_rendered:
                # Different numbering sequence, stop looking
                break
        
        # Now insert the paragraph after last_child_index
        doc = Document(filename)
        if last_child_index >= len(doc.paragraphs):
            return f"Invalid paragraph index {last_child_index} (document has {len(doc.paragraphs)} paragraphs)"
        
        target_paragraph = doc.paragraphs[last_child_index]
        target_p = target_paragraph._p
        parent = target_p.getparent()
        target_position = list(parent).index(target_p)
        
        # Create new paragraph element
        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        new_p.append(pPr)
        
        # Determine numbering source (paragraph-level numPr or style-based)
        numbering_source = target_clause.get("source", "paragraph")
        
        # Apply style
        if style:
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style)
            pPr.append(pStyle)
        elif inherit_numbering and numbering_source == "style":
            # Style-based numbering: inherit the style which contains numPr
            if target_paragraph.style:
                pStyle = OxmlElement('w:pStyle')
                pStyle.set(qn('w:val'), target_paragraph.style.style_id)
                pPr.append(pStyle)
        elif target_paragraph.style:
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), target_paragraph.style.style_id)
            pPr.append(pStyle)
        
        # Apply direct numPr only if source is 'paragraph' (not style-based)
        if inherit_numbering and numbering_source == "paragraph" and target_num_id:
            numPr = OxmlElement('w:numPr')
            ilvl_elem = OxmlElement('w:ilvl')
            ilvl_elem.set(qn('w:val'), str(target_ilvl if target_ilvl is not None else 0))
            numPr.append(ilvl_elem)
            numId_elem = OxmlElement('w:numId')
            numId_elem.set(qn('w:val'), str(target_num_id))
            numPr.append(numId_elem)
            pPr.append(numPr)
        
        # Add text run
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        new_p.append(r)
        
        # Insert into document
        parent.insert(target_position + 1, new_p)
        
        # Add custom para_id for future reference (function generates and returns UUID)
        custom_id = add_custom_para_id(new_p)
        
        doc.save(filename)
        
        numbering_info = ""
        if inherit_numbering:
            num_id = target_clause.get("numId")
            ilvl = target_clause.get("ilvl")
            numbering_source = target_clause.get("source", "paragraph")
            if num_id:
                numbering_info = f" with inherited numbering (source={numbering_source}, numId={num_id}, level={ilvl})"
            elif numbering_source == "style" and target_paragraph.style:
                numbering_info = f" with inherited numbering (source={numbering_source}, style={target_paragraph.style.style_id})"
        
        return f"Paragraph added after clause '{clause_number}'{numbering_info} in {filename} (para_id={custom_id})"
        
    except Exception as exc:
        return f"Failed to add paragraph after clause: {str(exc)}"


async def add_paragraphs_after_clause(
    filename: str,
    clause_number: str,
    paragraphs: List[str],
    style: Optional[str] = None,
    inherit_numbering: bool = True,
) -> str:
    """Add multiple paragraphs after a specific clause number.
    
    NEW function - contract-specific tool for adding multiple sibling clauses.
    
    This tool adds multiple paragraphs sequentially after a clause, each inheriting
    the same numbering properties. This is useful for adding multiple items at the
    same level (e.g., adding 7.1(b), 7.1(c), 7.1(d) after 7.1(a)).
    
    Args:
        filename: Path to the Word document
        clause_number: The rendered clause number to find (e.g., "1.2.3", "7.1(a)")
        paragraphs: List of paragraph texts to add
        style: Optional style name to apply to all new paragraphs
        inherit_numbering: If True, all paragraphs inherit the numId and ilvl from the target clause
    
    Returns:
        Success message with count of paragraphs added
    """
    if not paragraphs or not isinstance(paragraphs, list):
        return "paragraphs parameter must be a non-empty list of strings"
    
    filename = ensure_docx_extension(filename)
    
    # Add first paragraph
    result = await add_paragraph_after_clause(
        filename, clause_number, paragraphs[0], style, inherit_numbering
    )
    
    if "Failed" in result or "not found" in result or "does not exist" in result:
        return result
    
    # If more paragraphs, add them sequentially
    if len(paragraphs) > 1:
        try:
            from effilocal.doc.numbering_inspector import NumberingInspector
            
            for i, text in enumerate(paragraphs[1:], 1):
                # Re-analyze document to find the last inserted paragraph
                docx_path = Path(filename)
                inspector = NumberingInspector.from_docx(docx_path)
                rows, _ = inspector.analyze(debug=False)
                
                # Find target clause info
                target_num_id = None
                target_ilvl = None
                target_clause = None
                
                for row in rows:
                    rendered = row.get("rendered_number", "").strip().rstrip('.')
                    clause_normalized = clause_number.strip().rstrip('.')
                    if rendered == clause_normalized:
                        target_num_id = row.get("numId")
                        target_ilvl = row.get("ilvl")
                        target_clause = row
                        break
                
                if target_num_id is None:
                    break
                
                # Find the last paragraph with this numId and ilvl
                last_matching_idx = None
                for row in rows:
                    if (row.get("numId") == target_num_id and 
                        row.get("ilvl") == target_ilvl):
                        last_matching_idx = row.get("idx")
                
                if last_matching_idx is None:
                    continue
                
                # Insert after the last matching paragraph
                doc = Document(filename)
                if last_matching_idx >= len(doc.paragraphs):
                    break
                
                target_paragraph = doc.paragraphs[last_matching_idx]
                target_p = target_paragraph._p
                parent = target_p.getparent()
                target_position = list(parent).index(target_p)
                
                new_p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                new_p.append(pPr)
                
                # Determine numbering source
                numbering_source = target_clause.get("source", "paragraph") if target_clause else "paragraph"
                
                # Apply style
                if style:
                    pStyle = OxmlElement('w:pStyle')
                    pStyle.set(qn('w:val'), style)
                    pPr.append(pStyle)
                elif inherit_numbering and numbering_source == "style":
                    # Style-based numbering: inherit the style which contains numPr
                    if target_paragraph.style:
                        pStyle = OxmlElement('w:pStyle')
                        pStyle.set(qn('w:val'), target_paragraph.style.style_id)
                        pPr.append(pStyle)
                elif target_paragraph.style:
                    pStyle = OxmlElement('w:pStyle')
                    pStyle.set(qn('w:val'), target_paragraph.style.style_id)
                    pPr.append(pStyle)
                
                # Apply direct numPr only if source is 'paragraph' (not style-based)
                if inherit_numbering and numbering_source == "paragraph" and target_num_id:
                    numPr = OxmlElement('w:numPr')
                    ilvl_elem = OxmlElement('w:ilvl')
                    ilvl_elem.set(qn('w:val'), str(target_ilvl if target_ilvl is not None else 0))
                    numPr.append(ilvl_elem)
                    numId_elem = OxmlElement('w:numId')
                    numId_elem.set(qn('w:val'), str(target_num_id))
                    numPr.append(numId_elem)
                    pPr.append(numPr)
                
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = text
                r.append(t)
                new_p.append(r)
                
                parent.insert(target_position + 1, new_p)
                doc.save(filename)
            
            return f"Added {len(paragraphs)} paragraph(s) after clause '{clause_number}' in {filename}"
            
        except Exception as exc:
            return f"Failed to add additional paragraphs: {str(exc)}"
    
    return result


async def get_text_by_para_id(filename: str, para_id: str) -> str:
    """Get the text content of a paragraph identified by its paraId.
    
    Args:
        filename: Path to the Word document
        para_id: The 8-character hex paragraph ID (e.g., "3DD8236A")
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        doc = Document(filename)
        paragraph = get_paragraph_by_id(doc, para_id)
        
        if paragraph:
            return paragraph.text
        else:
            return f"Paragraph with ID {para_id} not found in {filename}"
            
    except Exception as e:
        return f"Failed to get text by para ID: {str(e)}"


async def replace_text_by_para_id(filename: str, para_id: str, new_text: str) -> str:
    """Replace the entire text content of a paragraph identified by its paraId.
    
    Args:
        filename: Path to the Word document
        para_id: The 8-character hex paragraph ID (e.g., "3DD8236A")
        new_text: The new text to set for the paragraph
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        paragraph = get_paragraph_by_id(doc, para_id)
        
        if not paragraph:
            return f"Paragraph with ID {para_id} not found in {filename}"
        
        old_text = paragraph.text
        paragraph.text = new_text
        
        doc.save(filename)
        return f"Replaced text in paragraph {para_id}.\nOld: {old_text[:50]}...\nNew: {new_text[:50]}..."
            
    except Exception as e:
        return f"Failed to replace text by para ID: {str(e)}"


# ============================================================================
# Re-export all functions for compatibility
# ============================================================================

__all__ = [
    # Overridden functions (extended from upstream)
    'add_heading',
    'add_paragraph',
    'search_and_replace',
    'insert_header_near_text_tool',
    'insert_line_or_paragraph_near_text_tool',
    'insert_numbered_list_near_text_tool',
    
    # Pass-through functions (unchanged from upstream)
    'add_table',
    'add_picture',
    'add_page_break',
    'add_table_of_contents',
    'delete_paragraph',
    
    # New functions (effilocal-specific)
    'edit_run_text_tool',
    'insert_str_content_near_text',
    'add_paragraph_after_clause',
    'add_paragraphs_after_clause',
    'get_text_by_para_id',
    'replace_text_by_para_id',
]

# Re-export wrapper tools from upstream for compatibility
# Import only what actually exists in current word_document_server
try:
    from word_document_server.tools.content_tools import (
        replace_block_between_manual_anchors_tool,
        insert_header_near_text_tool,
        insert_line_or_paragraph_near_text_tool,
        insert_numbered_list_near_text_tool,
    )
except ImportError:
    # Fallback: use the non-_tool versions
    from word_document_server.tools.content_tools import (
        insert_header_near_text_tool,
        insert_line_or_paragraph_near_text_tool,
        insert_numbered_list_near_text_tool,
    )
    replace_block_between_manual_anchors_tool = None
