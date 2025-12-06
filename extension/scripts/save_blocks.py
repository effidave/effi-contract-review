#!/usr/bin/env python3
"""
save_blocks.py - Save edited blocks back to a Word document

Sprint 2: WYSIWYG Editor save functionality

This script takes edited blocks from the webview editor and applies
the changes to the original .docx document using native w14:paraId
to identify which paragraphs to update.

Usage:
    python save_blocks.py <document_path> <blocks_json_path>

Input JSON format (text-based model):
    [
        {
            "id": "block-uuid",
            "para_id": "05C9333F",
            "text": "Updated paragraph text",
            "runs": [
                {"text": "Updated ", "formats": ["bold"]},
                {"text": "paragraph text", "formats": []}
            ]
        },
        ...
    ]

Note: Delete runs (with 'delete' in formats) are skipped during save
since they represent deleted content that shouldn't appear in the document.

Output JSON format:
    {
        "success": true,
        "block_count": 3,
        "message": "Saved 3 blocks to document"
    }
    or
    {
        "success": false,
        "error": "Error message"
    }
"""

import json
import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from effilocal.doc.uuid_embedding import (
    extract_block_uuids, 
    find_paragraph_by_para_id,
    generate_para_id,
    set_paragraph_para_id,
    collect_all_para_ids,
    ParaKey, 
    TableCellKey,
)


def get_paragraph_by_para_id(doc, para_id: str):
    """
    Get a paragraph element from the document by its w14:paraId.
    
    This uses Word's native paragraph IDs which are preserved across
    save/close/reopen cycles.
    
    Args:
        doc: Document object
        para_id: The 8-character hex paraId to find
        
    Returns:
        Paragraph object or None if not found
    """
    return find_paragraph_by_para_id(doc, para_id)


def get_paragraph_by_key(doc, key):
    """
    Get a paragraph element from the document by its BlockKey.
    
    This is a fallback for when UUID lookup fails.
    
    Args:
        doc: Document object
        key: ParaKey or TableCellKey
    
    Returns:
        Paragraph object or None if not found
    """
    if isinstance(key, ParaKey):
        if 0 <= key.para_idx < len(doc.paragraphs):
            return doc.paragraphs[key.para_idx]
    elif isinstance(key, TableCellKey):
        if 0 <= key.table_idx < len(doc.tables):
            table = doc.tables[key.table_idx]
            try:
                cell = table.cell(key.row, key.col)
                if cell.paragraphs:
                    return cell.paragraphs[0]
            except IndexError:
                pass
    return None


def apply_formatting(run, formats):
    """Apply formatting to a run based on format list."""
    run.bold = 'bold' in formats
    run.italic = 'italic' in formats
    run.underline = 'underline' in formats


def create_run_element(run_text, formats, author=None, date=None):
    """
    Create a w:r element with text and formatting.
    
    Returns the run element (not wrapped in ins/del).
    """
    r = OxmlElement('w:r')
    
    # Add run properties for formatting
    if formats:
        rPr = OxmlElement('w:rPr')
        if 'bold' in formats:
            b = OxmlElement('w:b')
            rPr.append(b)
        if 'italic' in formats:
            i = OxmlElement('w:i')
            rPr.append(i)
        if 'underline' in formats:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
        if len(rPr) > 0:
            r.append(rPr)
    
    t = OxmlElement('w:t')
    # Preserve spaces
    if run_text.startswith(' ') or run_text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = run_text
    r.append(t)
    
    return r


def create_del_run_element(deleted_text, formats, author=None, date=None):
    """
    Create a w:delText element inside w:r for deleted content.
    """
    r = OxmlElement('w:r')
    
    # Add run properties for formatting
    if formats:
        rPr = OxmlElement('w:rPr')
        if 'bold' in formats:
            b = OxmlElement('w:b')
            rPr.append(b)
        if 'italic' in formats:
            i = OxmlElement('w:i')
            rPr.append(i)
        if 'underline' in formats:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
        if len(rPr) > 0:
            r.append(rPr)
    
    dt = OxmlElement('w:delText')
    # Preserve spaces
    if deleted_text.startswith(' ') or deleted_text.endswith(' '):
        dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    dt.text = deleted_text
    r.append(dt)
    
    return r


def insert_new_paragraph(doc, block, after_para_id: str, existing_ids: set[str], use_para_id: str = None) -> tuple[str, bool]:
    """
    Insert a new paragraph into the document after the specified paragraph.
    
    Args:
        doc: Document object
        block: Block dict with text and runs
        after_para_id: The para_id of the paragraph to insert after
        existing_ids: Set of existing paraIds for collision checking
        use_para_id: Optional pre-generated para_id to use (from client-side generation)
        
    Returns:
        Tuple of (new_para_id, success)
    """
    # Find the paragraph to insert after
    after_para = find_paragraph_by_para_id(doc, after_para_id)
    if not after_para:
        return (None, False)
    
    # Create new paragraph element
    new_p = OxmlElement('w:p')
    
    # Use provided para_id or generate new one
    if use_para_id and use_para_id.upper() not in existing_ids:
        new_para_id = use_para_id.upper()
    else:
        new_para_id = generate_para_id(existing_ids)
    set_paragraph_para_id(new_p, new_para_id)
    existing_ids.add(new_para_id)
    
    # Copy style from the paragraph we're inserting after
    after_elem = after_para._element
    after_pPr = after_elem.find(qn('w:pPr'))
    if after_pPr is not None:
        # Clone the paragraph properties (excluding numPr to avoid numbering issues)
        new_pPr = OxmlElement('w:pPr')
        for child in after_pPr:
            if child.tag != qn('w:numPr'):  # Don't copy numbering
                # Create a copy of the element
                new_child = OxmlElement(child.tag.split('}')[-1] if '}' in child.tag else child.tag)
                for key, value in child.attrib.items():
                    new_child.set(key, value)
                new_pPr.append(new_child)
        if len(new_pPr) > 0:
            new_p.insert(0, new_pPr)
    
    # Add runs with text and formatting
    runs_data = block.get('runs', [])
    text = block.get('text', '')
    
    if not runs_data:
        runs_data = [{'text': text, 'formats': []}]
    
    # Default author/date for track changes in new paragraphs
    from datetime import datetime
    default_author = "User"
    default_date = datetime.now().isoformat()
    
    for run_data in runs_data:
        formats = run_data.get('formats', [])
        author = run_data.get('author', default_author)
        date = run_data.get('date', default_date)
        
        # Handle delete runs
        if 'delete' in formats:
            deleted_text = run_data.get('deleted_text', run_data.get('text', ''))
            if not deleted_text:
                continue
            
            # Create w:del wrapper with w:r containing w:delText
            del_elem = OxmlElement('w:del')
            del_elem.set(qn('w:author'), author)
            del_elem.set(qn('w:date'), date)
            
            r = create_del_run_element(deleted_text, [f for f in formats if f != 'delete'], author, date)
            del_elem.append(r)
            new_p.append(del_elem)
            continue
        
        run_text = run_data.get('text', '')
        if not run_text:
            continue
        
        # Handle insert runs
        if 'insert' in formats:
            ins_elem = OxmlElement('w:ins')
            ins_elem.set(qn('w:author'), author)
            ins_elem.set(qn('w:date'), date)
            
            r = create_run_element(run_text, [f for f in formats if f != 'insert'], author, date)
            ins_elem.append(r)
            new_p.append(ins_elem)
            continue
        
        # Normal run (no track changes)
        r = OxmlElement('w:r')
        
        # Add run properties for formatting
        if formats:
            rPr = OxmlElement('w:rPr')
            if 'bold' in formats:
                b = OxmlElement('w:b')
                rPr.append(b)
            if 'italic' in formats:
                i = OxmlElement('w:i')
                rPr.append(i)
            if 'underline' in formats:
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
            if len(rPr) > 0:
                r.append(rPr)
        
        t = OxmlElement('w:t')
        # Preserve spaces
        if run_text.startswith(' ') or run_text.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = run_text
        r.append(t)
        new_p.append(r)
    
    # Insert after the target paragraph
    parent = after_elem.getparent()
    target_position = list(parent).index(after_elem)
    parent.insert(target_position + 1, new_p)
    
    return (new_para_id, True)


def update_paragraph_content(paragraph, block):
    """
    Update a paragraph's content based on block data.
    
    Replaces all runs in the paragraph with new runs based on the
    block's text and runs formatting structure.
    
    Preserves track changes:
    - Insert runs are wrapped in <w:ins> elements
    - Delete runs are wrapped in <w:del> elements with <w:delText>
    
    Supports both text-based runs (new model) and position-based runs (legacy):
    - Text-based: run has 'text' or 'deleted_text' field
    - Position-based: run has 'start' and 'end' fields
    """
    from datetime import datetime
    
    text = block.get('text', '')
    runs_data = block.get('runs', [])
    
    # If no runs data, treat the whole text as unformatted
    if not runs_data:
        runs_data = [{'text': text, 'formats': []}]
    
    # Clear all existing content from paragraph
    p = paragraph._p
    # Remove all run elements (w:r), ins elements (w:ins), del elements (w:del)
    for child in list(p):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'ins', 'del', 'bookmarkStart', 'bookmarkEnd'):
            p.remove(child)
    
    # Track revision ID for ins/del elements
    revision_id = 0
    
    # Add new runs with formatting and track changes
    for run_data in runs_data:
        formats = run_data.get('formats', [])
        author = run_data.get('author', 'Unknown')
        date = run_data.get('date', datetime.now().isoformat())
        
        is_insert = 'insert' in formats
        is_delete = 'delete' in formats
        
        # Get run text
        if is_delete:
            # Delete runs use deleted_text field
            run_text = run_data.get('deleted_text', run_data.get('text', ''))
        elif 'text' in run_data:
            # Text-based model (new)
            run_text = run_data['text']
        elif 'start' in run_data and 'end' in run_data:
            # Position-based model (legacy/editor operations)
            start = run_data.get('start', 0)
            end = run_data.get('end', len(text))
            run_text = text[start:end]
        else:
            # No text content
            run_text = ''
        
        if not run_text:
            continue
        
        # Filter out insert/delete from formats for run properties
        run_formats = [f for f in formats if f not in ('insert', 'delete')]
        
        if is_delete:
            # Create <w:del> wrapper with <w:r><w:delText>
            del_elem = OxmlElement('w:del')
            del_elem.set(qn('w:id'), str(revision_id))
            del_elem.set(qn('w:author'), author)
            del_elem.set(qn('w:date'), date)
            revision_id += 1
            
            r = create_del_run_element(run_text, run_formats, author, date)
            del_elem.append(r)
            p.append(del_elem)
            
        elif is_insert:
            # Create <w:ins> wrapper with <w:r><w:t>
            ins_elem = OxmlElement('w:ins')
            ins_elem.set(qn('w:id'), str(revision_id))
            ins_elem.set(qn('w:author'), author)
            ins_elem.set(qn('w:date'), date)
            revision_id += 1
            
            r = create_run_element(run_text, run_formats, author, date)
            ins_elem.append(r)
            p.append(ins_elem)
            
        else:
            # Normal run - no track changes wrapper
            r = create_run_element(run_text, run_formats)
            p.append(r)


def update_blocks_jsonl(analysis_dir: Path, edited_blocks: list):
    """
    Update blocks.jsonl with edited block content and insert new blocks.
    
    Reads the existing blocks.jsonl, updates matching blocks with new text/runs,
    inserts new blocks at their correct positions, and writes back.
    
    Args:
        analysis_dir: Path to analysis directory containing blocks.jsonl
        edited_blocks: List of all block dicts with id, text, runs (includes new blocks)
    """
    import sys
    blocks_path = analysis_dir / 'blocks.jsonl'
    print(f"DEBUG update_blocks_jsonl: blocks_path = {blocks_path}", file=sys.stderr)
    print(f"DEBUG update_blocks_jsonl: edited_blocks count = {len(edited_blocks)}", file=sys.stderr)
    
    if not blocks_path.exists():
        print(f"DEBUG update_blocks_jsonl: blocks.jsonl does not exist, returning", file=sys.stderr)
        return
    
    # Create id -> edited block mapping
    edited_by_id = {b['id']: b for b in edited_blocks}
    
    # Read existing blocks
    existing_blocks = []
    with open(blocks_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if line:
                existing_blocks.append(json.loads(line))
    
    print(f"DEBUG update_blocks_jsonl: existing_blocks count = {len(existing_blocks)}", file=sys.stderr)
    
    # Track existing block IDs
    existing_ids = {b.get('id') for b in existing_blocks}
    
    # Identify new blocks (not in existing blocks.jsonl)
    new_blocks = [b for b in edited_blocks if b.get('id') not in existing_ids]
    print(f"DEBUG update_blocks_jsonl: new_blocks count = {len(new_blocks)}", file=sys.stderr)
    
    if new_blocks:
        for nb in new_blocks:
            runs = nb.get('runs', [])
            del_runs = [r for r in runs if r.get('deleted_text')]
            print(f"DEBUG update_blocks_jsonl: new block {nb.get('id', 'no-id')[:8]} has {len(del_runs)} deleted_text runs", file=sys.stderr)
    
    # Update matching blocks
    for block in existing_blocks:
        if block.get('id') in edited_by_id:
            edited = edited_by_id[block['id']]
            block['text'] = edited.get('text', block.get('text', ''))
            if 'runs' in edited:
                block['runs'] = edited['runs']
            # Update para_id if it was assigned during save
            if 'para_id' in edited:
                block['para_id'] = edited['para_id']
    
    # Insert new blocks at their correct positions
    # edited_blocks is in correct order from the UI, use it as reference
    if new_blocks:
        # Build ordered result based on edited_blocks order
        result_blocks = []
        existing_by_id = {b.get('id'): b for b in existing_blocks}
        
        for edited_block in edited_blocks:
            block_id = edited_block.get('id')
            if block_id in existing_by_id:
                # Use the updated existing block
                result_blocks.append(existing_by_id[block_id])
            else:
                # This is a new block - add it
                result_blocks.append(edited_block)
        
        existing_blocks = result_blocks
    
    # Write back
    with open(blocks_path, 'w', encoding='utf-8') as f:
        for block in existing_blocks:
            # Debug: log if block has runs with deleted_text
            runs = block.get('runs', [])
            has_del = any(r.get('deleted_text') for r in runs)
            if has_del:
                print(f"DEBUG update_blocks_jsonl: Block {block.get('id', 'no-id')[:8]} has deleted_text runs", file=sys.stderr)
            f.write(json.dumps(block, ensure_ascii=False) + '\n')


def save_blocks_to_document(doc_path: str, blocks_path: str):
    """
    Save edited blocks to a Word document.
    
    Args:
        doc_path: Path to the .docx document
        blocks_path: Path to JSON file containing edited blocks
    
    Returns:
        dict with success status and details
    """
    import sys
    
    try:
        # Load blocks
        with open(blocks_path, 'r', encoding='utf-8') as f:
            blocks = json.load(f)
        
        # Debug: log what we received
        print(f"DEBUG: Received {len(blocks)} blocks", file=sys.stderr)
        dirty_count = sum(1 for b in blocks if b.get('_isDirty', False))
        print(f"DEBUG: {dirty_count} blocks marked as dirty", file=sys.stderr)
        if blocks:
            print(f"DEBUG: First block id: {blocks[0].get('id', 'NO ID')}", file=sys.stderr)
            print(f"DEBUG: First block para_id: {blocks[0].get('para_id', 'NO PARA_ID')}", file=sys.stderr)
            print(f"DEBUG: First block _isDirty: {blocks[0].get('_isDirty', False)}", file=sys.stderr)
        
        if not blocks:
            return {"success": True, "block_count": 0, "message": "No blocks to save"}
        
        # Filter to only dirty blocks for document modification
        # But keep full blocks list for ordering/position reference
        dirty_blocks = [b for b in blocks if b.get('_isDirty', False)]
        
        if not dirty_blocks:
            return {"success": True, "block_count": 0, "message": "No dirty blocks to save"}
        
        # Load document
        doc = Document(doc_path)
        
        # Extract para_id -> paragraph mapping (uses native Word w14:paraId)
        para_id_map = extract_block_uuids(doc)
        
        # Collect all existing paraIds for collision checking when inserting new blocks
        existing_ids = collect_all_para_ids(doc)
        
        if not para_id_map:
            return {
                "success": False,
                "error": "No paragraph IDs found in document."
            }
        
        # Track updates and insertions
        updated_count = 0
        inserted_count = 0
        not_found = []
        
        # Separate DIRTY blocks into existing (para_id found in document) and new (para_idx=-1 or para_id not in doc)
        existing_blocks = []
        new_blocks = []
        
        for block in dirty_blocks:
            para_id = block.get('para_id', '')
            para_idx = block.get('para_idx', 0)
            
            # New blocks have para_idx=-1 (set by editor on split)
            # OR have a para_id that doesn't exist in the document yet
            if para_idx == -1:
                new_blocks.append(block)
            elif para_id and para_id.upper() not in existing_ids:
                # para_id was client-generated but not yet in document
                new_blocks.append(block)
            elif para_id:
                existing_blocks.append(block)
            else:
                # No para_id and non-negative para_idx - shouldn't happen but treat as update attempt
                existing_blocks.append(block)

        # First pass: Update existing paragraphs - find by para_id (w14:paraId)
        for block in existing_blocks:
            block_id = block.get('id', '')
            para_id = block.get('para_id', '')

            # Primary: look up by native Word para_id (8-char hex)
            paragraph = None
            if para_id:
                paragraph = get_paragraph_by_para_id(doc, para_id)

            if paragraph:
                update_paragraph_content(paragraph, block)
                updated_count += 1
            elif para_id and para_id in para_id_map:
                # Fall back to key-based lookup using position
                key = para_id_map[para_id]
                paragraph = get_paragraph_by_key(doc, key)
                if paragraph:
                    update_paragraph_content(paragraph, block)
                    updated_count += 1
                else:
                    not_found.append(f"{block_id} (para_id={para_id}, key={key})")
            else:
                not_found.append(f"{block_id} (para_id={para_id} not in document)")
        
        # Second pass: Insert new paragraphs (blocks without para_id)
        # These come from editor split operations
        for i, block in enumerate(new_blocks):
            block_id = block.get('id', '')
            
            # Find where to insert: look at previous block in the original blocks list
            block_index = next((j for j, b in enumerate(blocks) if b.get('id') == block_id), -1)
            
            if block_index <= 0:
                # Can't determine insertion point - insert at end
                not_found.append(f"{block_id} (new block, can't determine insertion point)")
                continue
            
            # Get the previous block's para_id
            prev_block = blocks[block_index - 1]
            prev_para_id = prev_block.get('para_id', '')
            
            if not prev_para_id:
                # Previous block also doesn't have para_id - skip for now
                not_found.append(f"{block_id} (new block, previous block also new)")
                continue
            
            # Insert after the previous paragraph, using client-generated para_id if available
            client_para_id = block.get('para_id', '')
            new_para_id, success = insert_new_paragraph(doc, block, prev_para_id, existing_ids, use_para_id=client_para_id)
            
            if success:
                inserted_count += 1
                # Update the block's para_id for future reference
                block['para_id'] = new_para_id
            else:
                not_found.append(f"{block_id} (failed to insert after {prev_para_id})")
        
        # Save document
        doc.save(doc_path)
        
        # Also update blocks.jsonl so the UI stays in sync
        # The temp blocks file is in the analysis directory
        temp_blocks_dir = Path(blocks_path).parent
        
        # Debug: Check if any blocks have deleted_text before updating blocks.jsonl
        for block in blocks:
            runs = block.get('runs', [])
            del_runs = [r for r in runs if r.get('deleted_text')]
            if del_runs:
                print(f"DEBUG save: Block {block.get('id', 'no-id')[:8]} has {len(del_runs)} deleted_text runs", file=sys.stderr)
                print(f"DEBUG save: First del run: {del_runs[0]}", file=sys.stderr)
        
        update_blocks_jsonl(temp_blocks_dir, blocks)
        
        total_count = updated_count + inserted_count
        message_parts = []
        if updated_count > 0:
            message_parts.append(f"updated {updated_count}")
        if inserted_count > 0:
            message_parts.append(f"inserted {inserted_count}")
        
        result = {
            "success": True,
            "block_count": total_count,
            "updated_count": updated_count,
            "inserted_count": inserted_count,
            "message": f"Saved {total_count} block(s) to document ({', '.join(message_parts)})" if message_parts else "No changes made"
        }
        
        if not_found:
            result["warnings"] = f"Could not find paragraphs for: {', '.join(not_found[:5])}"
            if len(not_found) > 5:
                result["warnings"] += f" and {len(not_found) - 5} more"
        
        return result
        
    except FileNotFoundError as e:
        return {"success": False, "error": f"File not found: {e}"}
    except json.JSONDecodeError as e:
        return {"success": False, "error": f"Invalid JSON: {e}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def main():
    if len(sys.argv) < 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: save_blocks.py <document_path> <blocks_json_path>"
        }))
        sys.exit(1)
    
    doc_path = sys.argv[1]
    blocks_path = sys.argv[2]
    
    result = save_blocks_to_document(doc_path, blocks_path)
    print(json.dumps(result))
    
    sys.exit(0 if result.get('success') else 1)


if __name__ == '__main__':
    main()
