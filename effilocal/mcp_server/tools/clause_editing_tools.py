"""MCP tools for editing clauses by ordinal number.

These tools allow LLM agents to modify document clauses by referring to their
clause numbers (ordinals) like "3.2.1" or "15(a)" instead of requiring paragraph IDs.
"""

from pathlib import Path
from typing import Optional

from effilocal.artifact_loader import ArtifactLoader
from effilocal.mcp_server.utils.document_utils import check_file_writeable


def replace_clause_text_by_ordinal(
    filename: str,
    clause_number: str,
    new_text: str,
    analysis_dir: Optional[str] = None
) -> str:
    """Replace the text of a clause identified by its ordinal number.
    
    Args:
        filename: Path to Word document
        clause_number: Clause ordinal (e.g., "3.2.1", "15.", "7(a)")
        new_text: New text to replace the clause with
        analysis_dir: Path to analysis directory (defaults to same directory as document)
    
    Returns:
        Success message with clause details
        
    Example:
        replace_clause_text_by_ordinal(
            filename="contract.docx",
            clause_number="5.2",
            new_text="The Supplier shall deliver within 30 days."
        )
    """
    from effilocal.mcp_server.utils.document_utils import find_and_replace_text
    
    check_file_writeable(filename)
    
    # Determine analysis directory
    if analysis_dir is None:
        doc_path = Path(filename).resolve()
        # Assume analysis is in ../analysis/ relative to document
        analysis_dir = doc_path.parent.parent / 'analysis'
    
    analysis_path = Path(analysis_dir)
    if not analysis_path.exists():
        return f"Error: Analysis directory not found: {analysis_dir}. Please analyze the document first."
    
    try:
        loader = ArtifactLoader(analysis_path)
    except Exception as e:
        return f"Error loading artifacts: {e}"
    
    # Find clause by ordinal
    clause = loader.find_clause_by_ordinal(clause_number)
    if not clause:
        available = [b.get('list', {}).get('ordinal', '') for b in loader.blocks if b.get('list')]
        available = [o for o in available if o][:20]  # First 20
        return f"Error: Clause {clause_number} not found. Available ordinals include: {', '.join(available)}"
    
    old_text = clause.get('text', '')
    if not old_text:
        return f"Error: Clause {clause_number} has no text content"
    
    # Replace text in document
    result = find_and_replace_text(
        filename=filename,
        find_text=old_text,
        replace_text=new_text,
        whole_word_only=False
    )
    
    # Check if replacement succeeded
    if result['count'] == 0:
        return f"Error: Text not found in document. Clause {clause_number} text may have changed since analysis."
    
    return (
        f"✓ Replaced clause {clause_number} text\n"
        f"Before: {old_text[:100]}{'...' if len(old_text) > 100 else ''}\n"
        f"After: {new_text[:100]}{'...' if len(new_text) > 100 else ''}\n"
        f"Replacements: {result['count']}"
    )


def insert_paragraph_after_clause(
    filename: str,
    clause_number: str,
    text: str,
    style: str = "Normal",
    inherit_numbering: bool = False,
    analysis_dir: Optional[str] = None
) -> str:
    """Insert a new paragraph after a clause identified by ordinal.
    
    Args:
        filename: Path to Word document
        clause_number: Clause ordinal to insert after (e.g., "3.2", "15.", "7(a)")
        text: Text for the new paragraph
        style: Word style name (default: "Normal")
        inherit_numbering: If True, create a sibling clause at same level
        analysis_dir: Path to analysis directory (defaults to same directory as document)
    
    Returns:
        Success message with new clause number if numbered
        
    Example:
        insert_paragraph_after_clause(
            filename="contract.docx",
            clause_number="8.2",
            text="The Recipient shall not disclose Confidential Information.",
            inherit_numbering=True
        )
        # Creates clause 8.3
    """
    from effilocal.mcp_server.tools.content_tools import add_paragraph_after_clause
    
    check_file_writeable(filename)
    
    # Determine analysis directory
    if analysis_dir is None:
        doc_path = Path(filename).resolve()
        analysis_dir = doc_path.parent.parent / 'analysis'
    
    analysis_path = Path(analysis_dir)
    if not analysis_path.exists():
        return f"Error: Analysis directory not found: {analysis_dir}"
    
    try:
        loader = ArtifactLoader(analysis_path)
    except Exception as e:
        return f"Error loading artifacts: {e}"
    
    # Find clause by ordinal
    clause = loader.find_clause_by_ordinal(clause_number)
    if not clause:
        return f"Error: Clause {clause_number} not found"
    
    # Use existing MCP tool
    result = add_paragraph_after_clause(
        filename=filename,
        clause_number=clause_number,
        text=text,
        style=style,
        inherit_numbering=inherit_numbering
    )
    
    return result


def delete_clause_by_ordinal(
    filename: str,
    clause_number: str,
    analysis_dir: Optional[str] = None
) -> str:
    """Delete a clause identified by its ordinal number.
    
    Deletes the main clause and any continuation paragraphs that belong to it.
    
    Args:
        filename: Path to Word document
        clause_number: Clause ordinal (e.g., "3.2.1", "15.", "7(a)")
        analysis_dir: Path to analysis directory (defaults to same directory as document)
    
    Returns:
        Success message with deleted text preview
        
    Example:
        delete_clause_by_ordinal(
            filename="contract.docx",
            clause_number="12.5"
        )
    """
    from docx import Document
    from effilocal.mcp_server.utils.file_utils import ensure_docx_extension
    
    filename = ensure_docx_extension(filename)
    check_file_writeable(filename)
    
    # Determine analysis directory
    if analysis_dir is None:
        doc_path = Path(filename).resolve()
        analysis_dir = doc_path.parent.parent / 'analysis'
    
    analysis_path = Path(analysis_dir)
    if not analysis_path.exists():
        return f"Error: Analysis directory not found: {analysis_dir}"
    
    try:
        loader = ArtifactLoader(analysis_path)
    except Exception as e:
        return f"Error loading artifacts: {e}"
    
    # Find clause by ordinal
    clause = loader.find_clause_by_ordinal(clause_number)
    if not clause:
        return f"Error: Clause {clause_number} not found"
    
    # Get clause group (main + continuations)
    clause_group = loader.get_clause_group(clause['id'])
    if not clause_group:
        return f"Error: Could not find clause group for {clause_number}"
    
    # Get paragraph IDs to delete
    para_ids = [b.get('para_id') for b in clause_group if b.get('para_id')]
    if not para_ids:
        return f"Error: No paragraph IDs found for clause {clause_number}"
    
    # Load document
    doc = Document(filename)
    
    # Find and delete paragraphs
    deleted_count = 0
    deleted_texts = []
    
    for paragraph in doc.paragraphs[:]:  # [:] creates a copy to iterate safely
        para_id = paragraph._element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId')
        if para_id in para_ids:
            deleted_texts.append(paragraph.text[:60])
            # Delete paragraph by removing from parent
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
            deleted_count += 1
    
    if deleted_count == 0:
        return f"Error: Could not find paragraphs to delete for clause {clause_number}"
    
    # Save document
    doc.save(filename)
    
    preview = '\n'.join(f"  - {t}..." for t in deleted_texts[:3])
    if len(deleted_texts) > 3:
        preview += f"\n  ... and {len(deleted_texts) - 3} more"
    
    return (
        f"✓ Deleted clause {clause_number} ({deleted_count} paragraph(s))\n"
        f"Deleted text:\n{preview}"
    )


def get_clause_text_by_ordinal(
    filename: str,
    clause_number: str,
    include_continuations: bool = True,
    analysis_dir: Optional[str] = None
) -> str:
    """Get the text of a clause by its ordinal number.
    
    Args:
        filename: Path to Word document (used to locate analysis)
        clause_number: Clause ordinal (e.g., "3.2.1", "15.", "7(a)")
        include_continuations: Include continuation paragraphs (default: True)
        analysis_dir: Path to analysis directory (defaults to same directory as document)
    
    Returns:
        Clause text or error message
        
    Example:
        text = get_clause_text_by_ordinal(
            filename="contract.docx",
            clause_number="8.2"
        )
    """
    # Determine analysis directory
    if analysis_dir is None:
        doc_path = Path(filename).resolve()
        analysis_dir = doc_path.parent.parent / 'analysis'
    
    analysis_path = Path(analysis_dir)
    if not analysis_path.exists():
        return f"Error: Analysis directory not found: {analysis_dir}"
    
    try:
        loader = ArtifactLoader(analysis_path)
    except Exception as e:
        return f"Error loading artifacts: {e}"
    
    # Find clause by ordinal
    clause = loader.find_clause_by_ordinal(clause_number)
    if not clause:
        return f"Error: Clause {clause_number} not found"
    
    if include_continuations:
        # Get full clause group
        clause_group = loader.get_clause_group(clause['id'])
        texts = [b.get('text', '') for b in clause_group]
        full_text = '\n'.join(texts)
        return (
            f"Clause {clause_number}:\n"
            f"{full_text}\n\n"
            f"({len(clause_group)} paragraph(s))"
        )
    else:
        return f"Clause {clause_number}:\n{clause.get('text', '')}"


def list_all_clause_numbers(
    filename: str,
    analysis_dir: Optional[str] = None
) -> str:
    """List all clause ordinals in the document.
    
    Useful for discovering what clause numbers exist before editing.
    
    Args:
        filename: Path to Word document (used to locate analysis)
        analysis_dir: Path to analysis directory (defaults to same directory as document)
    
    Returns:
        Formatted list of all clause ordinals
        
    Example:
        ordinals = list_all_clause_numbers(filename="contract.docx")
    """
    # Determine analysis directory
    if analysis_dir is None:
        doc_path = Path(filename).resolve()
        analysis_dir = doc_path.parent.parent / 'analysis'
    
    analysis_path = Path(analysis_dir)
    if not analysis_path.exists():
        return f"Error: Analysis directory not found: {analysis_dir}"
    
    try:
        loader = ArtifactLoader(analysis_path)
    except Exception as e:
        return f"Error loading artifacts: {e}"
    
    # Collect all ordinals
    ordinals = []
    for block in loader.blocks:
        list_meta = block.get('list')
        if list_meta and list_meta.get('ordinal'):
            ordinal = list_meta['ordinal']
            text_preview = block.get('text', '')[:50]
            ordinals.append(f"{ordinal:8} {text_preview}...")
    
    if not ordinals:
        return "No numbered clauses found in document"
    
    return (
        f"Found {len(ordinals)} numbered clauses:\n\n" +
        '\n'.join(ordinals[:100])  # Limit to first 100
    )
