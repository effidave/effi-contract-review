"""Debug what's in the temp file after insertion."""
from pathlib import Path
import sys
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from effilocal.doc import direct_docx

temp_file = Path(__file__).parent / "temp_test_attachment.docx"

if temp_file.exists():
    # Check with python-docx what paragraphs exist
    doc = Document(temp_file)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # Look for our test text and Schedule 1
    for i, para in enumerate(doc.paragraphs):
        text = para.text[:80]
        if "Test paragraph" in text or "Schedule 1" in text:
            print(f"Para {i}: {text}")
    
    print("\n" + "="*80)
    print("Now checking with effilocal:")
    
    blocks = list(direct_docx.iter_blocks(temp_file))
    
    # Find blocks around Schedule 1
    for i, b in enumerate(blocks):
        text = b.get('text', '')[:80]
        if "Test paragraph" in text or "Schedule 1" in text:
            attachment = b.get('attachment')
            attachment_id = b.get('attachment_id')
            para_idx = b.get('para_idx')
            print(f"Block {i}, Para {para_idx}: {text}")
            if attachment:
                print(f"   ^ Attachment header")
            elif attachment_id:
                print(f"   ^ Part of attachment: {attachment_id[:20]}...")
else:
    print(f"Temp file not found: {temp_file}")
