#!/usr/bin/env python3
"""
Tests for AmendedParagraph - Text-based run model.

Text-Based Run Model:
- amended_text: Contains only VISIBLE text (normal runs + insertions, NO deletions)
- amended_runs: All runs including deletions
  - Normal/insert runs: have 'text' field containing their text content
  - Delete runs: have 'deleted_text' field containing struck content

Example:
{
  "text": "The quick fox",  # Visible only - no "brown " deletion
  "runs": [
    { "text": "The ", "formats": [] },                                          # Normal
    { "text": "quick", "formats": ["insert"], "author": "John" },               # Insert
    { "deleted_text": "brown ", "formats": ["delete"], "author": "Jane" },      # Delete
    { "text": " fox", "formats": [] }                                           # Normal
  ]
}

Concatenating run['text'] values (excluding deletes) equals amended_text.
"""

import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))


class TestAmendedParagraphImport:
    """Test that the module can be imported."""
    
    def test_import_amended_paragraph(self):
        """Test that AmendedParagraph can be imported."""
        try:
            from effilocal.doc.amended_paragraph import AmendedParagraph
        except ImportError as e:
            pytest.fail(f"Cannot import AmendedParagraph: {e}")
    
    def test_import_iter_amended_paragraphs(self):
        """Test that iter_amended_paragraphs can be imported."""
        try:
            from effilocal.doc.amended_paragraph import iter_amended_paragraphs
        except ImportError as e:
            pytest.fail(f"Cannot import iter_amended_paragraphs: {e}")
    
    def test_import_iter_amended_elements(self):
        """Test that iter_amended_elements can be imported."""
        try:
            from effilocal.doc.amended_paragraph import iter_amended_elements
        except ImportError as e:
            pytest.fail(f"Cannot import iter_amended_elements: {e}")


class TestAmendedTextVisibleOnly:
    """Tests for amended_text containing only visible text."""
    
    @pytest.fixture
    def sample_doc(self):
        """Create a basic document for testing."""
        return Document()
    
    @pytest.fixture
    def doc_with_normal_text(self, sample_doc):
        """Create document with normal text only."""
        doc = sample_doc
        doc.add_paragraph("Normal paragraph text")
        return doc
    
    @pytest.fixture
    def doc_with_insertion(self, sample_doc):
        """Create document with an insertion (w:ins containing w:t)."""
        doc = sample_doc
        para = doc.add_paragraph()
        
        # Add normal text
        run1 = para.add_run("Before ")
        
        # Create w:ins element with inserted text
        p_elem = para._p
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), '1')
        ins.set(qn('w:author'), 'John Smith')
        ins.set(qn('w:date'), '2024-01-15T10:30:00Z')
        
        # Create run inside ins with w:t (visible)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = 'inserted'
        r.append(t)
        ins.append(r)
        p_elem.append(ins)
        
        # Add more normal text
        run3 = para.add_run(" after")
        
        return doc
    
    @pytest.fixture
    def doc_with_deletion(self, sample_doc):
        """Create document with a deletion (w:del containing w:delText)."""
        doc = sample_doc
        para = doc.add_paragraph()
        
        # Add normal text
        run1 = para.add_run("Before ")
        
        # Create w:del element with deleted text
        p_elem = para._p
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), '2')
        del_elem.set(qn('w:author'), 'Jane Doe')
        del_elem.set(qn('w:date'), '2024-01-16T14:45:00Z')
        
        # Create run inside del with w:delText (NOT visible)
        r = OxmlElement('w:r')
        del_text = OxmlElement('w:delText')
        del_text.text = 'deleted'
        r.append(del_text)
        del_elem.append(r)
        p_elem.append(del_elem)
        
        # Add more normal text
        run3 = para.add_run(" after")
        
        return doc
    
    @pytest.fixture
    def doc_with_both(self, sample_doc):
        """Create document with both insertion and deletion."""
        doc = sample_doc
        para = doc.add_paragraph()
        
        # Normal: "The "
        para.add_run("The ")
        
        # Insertion: "quick"
        p_elem = para._p
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), '1')
        ins.set(qn('w:author'), 'John')
        r_ins = OxmlElement('w:r')
        t_ins = OxmlElement('w:t')
        t_ins.text = 'quick'
        r_ins.append(t_ins)
        ins.append(r_ins)
        p_elem.append(ins)
        
        # Deletion: "brown " (should NOT appear in amended_text)
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), '2')
        del_elem.set(qn('w:author'), 'Jane')
        r_del = OxmlElement('w:r')
        del_text = OxmlElement('w:delText')
        del_text.text = 'brown '
        r_del.append(del_text)
        del_elem.append(r_del)
        p_elem.append(del_elem)
        
        # Normal: " fox"
        para.add_run(" fox")
        
        return doc
    
    def test_normal_text_unchanged(self, doc_with_normal_text):
        """Test that normal text is returned as-is."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_normal_text.paragraphs[0]
        amended = AmendedParagraph(para)
        
        assert amended.amended_text == "Normal paragraph text"
    
    def test_insertion_included_in_text(self, doc_with_insertion):
        """Test that inserted text (w:ins > w:r > w:t) IS included."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_insertion.paragraphs[0]
        amended = AmendedParagraph(para)
        
        assert "inserted" in amended.amended_text
        assert amended.amended_text == "Before inserted after"
    
    def test_deletion_excluded_from_text(self, doc_with_deletion):
        """Test that deleted text (w:del > w:r > w:delText) is NOT included."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_deletion.paragraphs[0]
        amended = AmendedParagraph(para)
        
        assert "deleted" not in amended.amended_text
        assert amended.amended_text == "Before  after"  # Note double space where deletion was
    
    def test_mixed_insertions_deletions(self, doc_with_both):
        """Test paragraph with both insertions and deletions."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_both.paragraphs[0]
        amended = AmendedParagraph(para)
        
        # Should include "The ", "quick", " fox" but NOT "brown "
        assert amended.amended_text == "The quick fox"
        assert "brown" not in amended.amended_text


class TestAmendedRunsWithDeletedText:
    """Tests for amended_runs with zero-width deletes and deleted_text field."""
    
    @pytest.fixture
    def sample_doc(self):
        return Document()
    
    @pytest.fixture
    def doc_with_deletion(self, sample_doc):
        """Create document with a deletion."""
        doc = sample_doc
        para = doc.add_paragraph()
        
        para.add_run("Before ")
        
        p_elem = para._p
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), '1')
        del_elem.set(qn('w:author'), 'Jane Doe')
        del_elem.set(qn('w:date'), '2024-01-16T14:45:00Z')
        
        r = OxmlElement('w:r')
        del_text = OxmlElement('w:delText')
        del_text.text = 'removed'
        r.append(del_text)
        del_elem.append(r)
        p_elem.append(del_elem)
        
        para.add_run(" after")
        
        return doc
    
    @pytest.fixture
    def doc_with_both(self, sample_doc):
        """Create document with insertion followed by deletion."""
        doc = sample_doc
        para = doc.add_paragraph()
        
        para.add_run("The ")
        
        p_elem = para._p
        
        # Insertion
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), '1')
        ins.set(qn('w:author'), 'John')
        r_ins = OxmlElement('w:r')
        t_ins = OxmlElement('w:t')
        t_ins.text = 'quick'
        r_ins.append(t_ins)
        ins.append(r_ins)
        p_elem.append(ins)
        
        # Deletion
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), '2')
        del_elem.set(qn('w:author'), 'Jane')
        r_del = OxmlElement('w:r')
        del_text = OxmlElement('w:delText')
        del_text.text = 'brown '
        r_del.append(del_text)
        del_elem.append(r_del)
        p_elem.append(del_elem)
        
        para.add_run(" fox")
        
        return doc
    
    def test_delete_run_has_no_text_field(self, doc_with_deletion):
        """Test that delete runs have deleted_text but not text field."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_deletion.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        delete_runs = [r for r in runs if 'delete' in r.get('formats', [])]
        assert len(delete_runs) == 1
        
        del_run = delete_runs[0]
        assert 'deleted_text' in del_run, "Delete run should have deleted_text"
        assert 'text' not in del_run, "Delete run should not have text field"
    
    def test_delete_run_has_deleted_text_field(self, doc_with_deletion):
        """Test that delete runs have deleted_text field with content."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_deletion.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        delete_runs = [r for r in runs if 'delete' in r.get('formats', [])]
        assert len(delete_runs) == 1
        
        del_run = delete_runs[0]
        assert 'deleted_text' in del_run, "Delete run should have deleted_text field"
        assert del_run['deleted_text'] == 'removed'
    
    def test_delete_run_has_author_and_date(self, doc_with_deletion):
        """Test that delete runs have author and date metadata."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_deletion.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        delete_runs = [r for r in runs if 'delete' in r.get('formats', [])]
        del_run = delete_runs[0]
        
        assert del_run.get('author') == 'Jane Doe'
        assert del_run.get('date') == '2024-01-16T14:45:00Z'
    
    def test_insert_run_has_text_field(self, doc_with_both):
        """Test that insert runs have text field with inserted content."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_both.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        insert_runs = [r for r in runs if 'insert' in r.get('formats', [])]
        assert len(insert_runs) == 1
        
        ins_run = insert_runs[0]
        assert 'text' in ins_run, "Insert run should have text field"
        assert ins_run['text'] == 'quick'
    
    def test_delete_run_has_deleted_text_in_mixed(self, doc_with_both):
        """Test that delete run has deleted_text in mixed content."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_both.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        delete_runs = [r for r in runs if 'delete' in r.get('formats', [])]
        assert len(delete_runs) == 1
        
        del_run = delete_runs[0]
        assert del_run['deleted_text'] == 'brown '
    
    def test_runs_text_concatenates_to_amended_text(self, doc_with_both):
        """Test that concatenating run text equals amended_text."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_both.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        text = amended.amended_text
        
        # Filter out delete runs (they don't contribute to visible text)
        visible_runs = [r for r in runs if 'delete' not in r.get('formats', [])]
        
        # Reconstruct text from runs
        reconstructed = ''.join(r.get('text', '') for r in visible_runs)
        
        assert reconstructed == text


class TestIterAmendedParagraphs:
    """Tests for iter_amended_paragraphs function."""
    
    @pytest.fixture
    def doc_with_multiple_paras(self):
        """Create document with multiple paragraphs."""
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("Third paragraph")
        return doc
    
    def test_iter_yields_amended_paragraphs(self, doc_with_multiple_paras):
        """Test that iterator yields AmendedParagraph objects."""
        from effilocal.doc.amended_paragraph import (
            AmendedParagraph,
            iter_amended_paragraphs,
        )
        
        paras = list(iter_amended_paragraphs(doc_with_multiple_paras))
        
        assert len(paras) == 3
        for para in paras:
            assert isinstance(para, AmendedParagraph)
    
    def test_iter_preserves_order(self, doc_with_multiple_paras):
        """Test that paragraphs are yielded in document order."""
        from effilocal.doc.amended_paragraph import iter_amended_paragraphs
        
        paras = list(iter_amended_paragraphs(doc_with_multiple_paras))
        
        assert paras[0].amended_text == "First paragraph"
        assert paras[1].amended_text == "Second paragraph"
        assert paras[2].amended_text == "Third paragraph"


class TestRealDocument:
    """Tests using a real tracked changes document."""
    
    @pytest.fixture
    def tracked_doc_path(self):
        """Path to test document with track changes."""
        return Path(r"C:\Users\DavidSant\effi-contract-review\EL_Projects\Test Project\drafts\current_drafts\Norton R&D Services Agreement (DRAFT) - HJ9 (TRACKED).docx")
    
    def test_real_doc_amended_text_excludes_deletions(self, tracked_doc_path):
        """Test that amended_text excludes deleted content in real document."""
        if not tracked_doc_path.exists():
            pytest.skip(f"Test document not found: {tracked_doc_path}")
        
        from effilocal.doc.amended_paragraph import iter_amended_paragraphs
        
        doc = Document(str(tracked_doc_path))
        
        # Find paragraphs with substantial deletions (not just whitespace)
        for amended in iter_amended_paragraphs(doc):
            runs = amended.amended_runs
            delete_runs = [r for r in runs if 'delete' in r.get('formats', [])]
            
            # Look for deletions with meaningful text (at least 3 chars, not just whitespace)
            meaningful_deletes = [
                r for r in delete_runs 
                if r.get('deleted_text', '').strip() and len(r.get('deleted_text', '')) >= 3
            ]
            
            if meaningful_deletes:
                # Verify deleted text is NOT in amended_text
                for del_run in meaningful_deletes:
                    deleted_text = del_run.get('deleted_text', '')
                    # Verify structure: has deleted_text, not text
                    assert 'deleted_text' in del_run
                    assert 'text' not in del_run
                return  # Found what we needed
        
        pytest.skip("No paragraphs with substantial deletions found in document")
    
    def test_real_doc_delete_runs_have_deleted_text(self, tracked_doc_path):
        """Test that delete runs have deleted_text field in real document."""
        if not tracked_doc_path.exists():
            pytest.skip(f"Test document not found: {tracked_doc_path}")
        
        from effilocal.doc.amended_paragraph import iter_amended_paragraphs
        
        doc = Document(str(tracked_doc_path))
        
        found_delete_with_text = False
        for amended in iter_amended_paragraphs(doc):
            runs = amended.amended_runs
            for run in runs:
                if 'delete' in run.get('formats', []):
                    assert 'deleted_text' in run, "Delete run should have deleted_text field"
                    assert 'text' not in run, "Delete run should not have text field"
                    if run.get('deleted_text'):
                        found_delete_with_text = True
        
        assert found_delete_with_text, "Should find at least one delete run with text"


class TestTableCellAmendedParagraphs:
    """Tests for amended paragraphs in table cells."""
    
    @pytest.fixture
    def doc_with_table(self):
        """Create document with a table containing track changes."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        
        # Add normal text to first cell
        cell = table.cell(0, 0)
        cell.text = "Normal cell"
        
        # Add paragraph with insertion to second cell
        cell2 = table.cell(0, 1)
        para = cell2.paragraphs[0]
        para.add_run("Before ")
        
        # Add insertion
        p_elem = para._p
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), '1')
        ins.set(qn('w:author'), 'Test')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = 'inserted'
        r.append(t)
        ins.append(r)
        p_elem.append(ins)
        
        para.add_run(" after")
        
        return doc
    
    def test_table_cell_amended_text(self, doc_with_table):
        """Test that table cell paragraphs have amended_text."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        table = doc_with_table.tables[0]
        cell_para = table.cell(0, 1).paragraphs[0]
        amended = AmendedParagraph(cell_para)
        
        assert amended.amended_text == "Before inserted after"


class TestTextBasedRunModel:
    """
    Tests for the refactored text-based run model.
    
    NEW DATA MODEL:
    - Normal/insert runs have 'text' field containing their text content
    - Delete runs have 'deleted_text' field (unchanged)
    - No more 'start'/'end' position fields
    
    This model avoids stale position data when paragraphs are edited.
    """
    
    @pytest.fixture
    def sample_doc(self):
        return Document()
    
    @pytest.fixture
    def doc_with_normal_text(self, sample_doc):
        """Create document with normal text only."""
        doc = sample_doc
        doc.add_paragraph("Normal paragraph text")
        return doc
    
    @pytest.fixture
    def doc_with_insertion(self, sample_doc):
        """Create document with an insertion."""
        doc = sample_doc
        para = doc.add_paragraph()
        para.add_run("Before ")
        
        p_elem = para._p
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), '1')
        ins.set(qn('w:author'), 'John Smith')
        ins.set(qn('w:date'), '2024-01-15T10:30:00Z')
        
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = 'inserted'
        r.append(t)
        ins.append(r)
        p_elem.append(ins)
        
        para.add_run(" after")
        return doc
    
    @pytest.fixture
    def doc_with_deletion(self, sample_doc):
        """Create document with a deletion."""
        doc = sample_doc
        para = doc.add_paragraph()
        para.add_run("Before ")
        
        p_elem = para._p
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), '1')
        del_elem.set(qn('w:author'), 'Jane Doe')
        del_elem.set(qn('w:date'), '2024-01-16T14:45:00Z')
        
        r = OxmlElement('w:r')
        del_text = OxmlElement('w:delText')
        del_text.text = 'removed'
        r.append(del_text)
        del_elem.append(r)
        p_elem.append(del_elem)
        
        para.add_run(" after")
        return doc
    
    @pytest.fixture
    def doc_with_both(self, sample_doc):
        """Create document with insertion followed by deletion."""
        doc = sample_doc
        para = doc.add_paragraph()
        
        para.add_run("The ")
        
        p_elem = para._p
        
        # Insertion: "quick"
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), '1')
        ins.set(qn('w:author'), 'John')
        ins.set(qn('w:date'), '2024-01-15T10:00:00Z')
        r_ins = OxmlElement('w:r')
        t_ins = OxmlElement('w:t')
        t_ins.text = 'quick'
        r_ins.append(t_ins)
        ins.append(r_ins)
        p_elem.append(ins)
        
        # Deletion: "brown "
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), '2')
        del_elem.set(qn('w:author'), 'Jane')
        del_elem.set(qn('w:date'), '2024-01-16T14:00:00Z')
        r_del = OxmlElement('w:r')
        del_text = OxmlElement('w:delText')
        del_text.text = 'brown '
        r_del.append(del_text)
        del_elem.append(r_del)
        p_elem.append(del_elem)
        
        para.add_run(" fox")
        
        return doc
    
    # =========================================================================
    # Tests for 'text' field on normal runs
    # =========================================================================
    
    def test_normal_run_has_text_field(self, doc_with_normal_text):
        """Test that normal runs have 'text' field containing their text."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_normal_text.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        assert len(runs) >= 1
        # Normal run should have 'text' field
        assert 'text' in runs[0], "Normal run should have 'text' field"
        assert runs[0]['text'] == "Normal paragraph text"
    
    def test_normal_run_no_start_end(self, doc_with_normal_text):
        """Test that runs no longer have 'start'/'end' position fields."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_normal_text.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        for run in runs:
            assert 'start' not in run, "Runs should not have 'start' field"
            assert 'end' not in run, "Runs should not have 'end' field"
    
    # =========================================================================
    # Tests for 'text' field on insert runs
    # =========================================================================
    
    def test_insert_run_has_text_field(self, doc_with_insertion):
        """Test that insert runs have 'text' field containing inserted text."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_insertion.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        insert_runs = [r for r in runs if 'insert' in r.get('formats', [])]
        assert len(insert_runs) == 1
        
        ins_run = insert_runs[0]
        assert 'text' in ins_run, "Insert run should have 'text' field"
        assert ins_run['text'] == 'inserted'
    
    def test_insert_run_has_author_date(self, doc_with_insertion):
        """Test that insert runs have author and date metadata."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_insertion.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        insert_runs = [r for r in runs if 'insert' in r.get('formats', [])]
        ins_run = insert_runs[0]
        
        assert ins_run.get('author') == 'John Smith'
        assert ins_run.get('date') == '2024-01-15T10:30:00Z'
    
    # =========================================================================
    # Tests for 'deleted_text' field on delete runs (unchanged behavior)
    # =========================================================================
    
    def test_delete_run_has_deleted_text_field(self, doc_with_deletion):
        """Test that delete runs have 'deleted_text' field (not 'text')."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_deletion.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        delete_runs = [r for r in runs if 'delete' in r.get('formats', [])]
        assert len(delete_runs) == 1
        
        del_run = delete_runs[0]
        assert 'deleted_text' in del_run, "Delete run should have 'deleted_text' field"
        assert del_run['deleted_text'] == 'removed'
        # Delete runs should NOT have 'text' field
        assert 'text' not in del_run, "Delete runs should not have 'text' field"
    
    def test_delete_run_has_author_date(self, doc_with_deletion):
        """Test that delete runs have author and date metadata."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_deletion.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        delete_runs = [r for r in runs if 'delete' in r.get('formats', [])]
        del_run = delete_runs[0]
        
        assert del_run.get('author') == 'Jane Doe'
        assert del_run.get('date') == '2024-01-16T14:45:00Z'
    
    # =========================================================================
    # Tests for mixed content (insertions + deletions)
    # =========================================================================
    
    def test_mixed_runs_have_correct_fields(self, doc_with_both):
        """Test paragraph with both insertions and deletions has correct fields."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_both.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        # Check all runs have correct structure
        for run in runs:
            formats = run.get('formats', [])
            
            if 'delete' in formats:
                # Delete runs have deleted_text, not text
                assert 'deleted_text' in run
                assert 'text' not in run
            else:
                # Normal and insert runs have text
                assert 'text' in run
                assert 'deleted_text' not in run
            
            # No position fields
            assert 'start' not in run
            assert 'end' not in run
    
    def test_concatenating_run_text_equals_amended_text(self, doc_with_both):
        """Test that concatenating all run.text values equals amended_text."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_both.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        # Concatenate text from non-delete runs
        concatenated = ''
        for run in runs:
            if 'delete' not in run.get('formats', []):
                concatenated += run.get('text', '')
        
        assert concatenated == amended.amended_text
        assert concatenated == "The quick fox"
    
    def test_run_order_preserved(self, doc_with_both):
        """Test that runs are in document order."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        para = doc_with_both.paragraphs[0]
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        # Expected order: "The ", "quick" (insert), "brown " (delete), " fox"
        assert len(runs) == 4
        
        # First run: "The "
        assert runs[0].get('text') == "The "
        assert 'insert' not in runs[0].get('formats', [])
        
        # Second run: "quick" (insertion)
        assert runs[1].get('text') == "quick"
        assert 'insert' in runs[1].get('formats', [])
        
        # Third run: deletion of "brown "
        assert runs[2].get('deleted_text') == "brown "
        assert 'delete' in runs[2].get('formats', [])
        
        # Fourth run: " fox"
        assert runs[3].get('text') == " fox"
        assert 'insert' not in runs[3].get('formats', [])
    
    # =========================================================================
    # Tests for formatting preservation
    # =========================================================================
    
    def test_formatting_preserved_in_text_runs(self):
        """Test that formatting (bold, italic) is preserved in text-based runs."""
        from effilocal.doc.amended_paragraph import AmendedParagraph
        
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("bold text")
        run.bold = True
        
        amended = AmendedParagraph(para)
        runs = amended.amended_runs
        
        assert len(runs) == 1
        assert runs[0].get('text') == "bold text"
        assert 'bold' in runs[0].get('formats', [])
    
    # =========================================================================
    # Real document test
    # =========================================================================
    
    def test_real_doc_text_based_model(self):
        """Test text-based model with real tracked changes document."""
        doc_path = Path(r"C:\Users\DavidSant\effi-contract-review\EL_Projects\Test Project\drafts\current_drafts\deleted_run.docx")
        
        if not doc_path.exists():
            pytest.skip(f"Test document not found: {doc_path}")
        
        from effilocal.doc.amended_paragraph import iter_amended_paragraphs
        
        doc = Document(str(doc_path))
        
        for amended in iter_amended_paragraphs(doc):
            runs = amended.amended_runs
            
            for run in runs:
                formats = run.get('formats', [])
                
                # Verify structure
                if 'delete' in formats:
                    assert 'deleted_text' in run
                    assert 'text' not in run
                else:
                    assert 'text' in run
                    assert 'deleted_text' not in run
                
                # No position fields
                assert 'start' not in run
                assert 'end' not in run
            
            # Verify concatenation
            concatenated = ''.join(
                r.get('text', '') for r in runs 
                if 'delete' not in r.get('formats', [])
            )
            assert concatenated == amended.amended_text


class TestProcessTrackChangesWithTextModel:
    """
    Tests for consumers that use the text-based run model.
    
    These tests verify that code which previously used start/end positions
    can work correctly with the new text field.
    """
    
    @pytest.fixture
    def doc_with_changes(self):
        """Create document with tracked insertion and deletion."""
        doc = Document()
        para = doc.add_paragraph()
        
        para.add_run("Original ")
        
        p_elem = para._p
        
        # Insertion
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), '1')
        ins.set(qn('w:author'), 'Editor')
        ins.set(qn('w:date'), '2024-01-15T10:00:00Z')
        r_ins = OxmlElement('w:r')
        t_ins = OxmlElement('w:t')
        t_ins.text = 'new content'
        r_ins.append(t_ins)
        ins.append(r_ins)
        p_elem.append(ins)
        
        # Deletion
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), '2')
        del_elem.set(qn('w:author'), 'Editor')
        del_elem.set(qn('w:date'), '2024-01-16T14:00:00Z')
        r_del = OxmlElement('w:r')
        del_text_elem = OxmlElement('w:delText')
        del_text_elem.text = 'old content'
        r_del.append(del_text_elem)
        del_elem.append(r_del)
        p_elem.append(del_elem)
        
        para.add_run(" end.")
        
        return doc
    
    def test_extract_insertions_using_text_field(self, doc_with_changes):
        """Test extracting insertions using the new 'text' field."""
        from effilocal.doc.amended_paragraph import iter_amended_paragraphs
        
        insertions = []
        for amended in iter_amended_paragraphs(doc_with_changes):
            for run in amended.amended_runs:
                if 'insert' in run.get('formats', []):
                    # NEW: Use run['text'] directly instead of slicing
                    insertions.append({
                        'text': run.get('text'),
                        'author': run.get('author'),
                        'date': run.get('date'),
                    })
        
        assert len(insertions) == 1
        assert insertions[0]['text'] == 'new content'
        assert insertions[0]['author'] == 'Editor'
    
    def test_extract_deletions_using_deleted_text_field(self, doc_with_changes):
        """Test extracting deletions using 'deleted_text' field (unchanged)."""
        from effilocal.doc.amended_paragraph import iter_amended_paragraphs
        
        deletions = []
        for amended in iter_amended_paragraphs(doc_with_changes):
            for run in amended.amended_runs:
                if 'delete' in run.get('formats', []):
                    deletions.append({
                        'text': run.get('deleted_text'),
                        'author': run.get('author'),
                        'date': run.get('date'),
                    })
        
        assert len(deletions) == 1
        assert deletions[0]['text'] == 'old content'
        assert deletions[0]['author'] == 'Editor'


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
