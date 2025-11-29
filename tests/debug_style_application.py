"""Debug style application."""
import sys
from pathlib import Path
import shutil

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Create a test
temp = Path(__file__).parent / "debug_style_test.docx"
shutil.copy2(Path(__file__).parent / "fixtures" / "real_world" / "schedules.docx", temp)

doc = Document(temp)

# Find Schedule 2
schedule_2_para = None
for para in doc.paragraphs:
    if "Schedule 2 - MSA" in para.text:
        schedule_2_para = para
        break

if schedule_2_para:
    print(f"Schedule 2 style: {schedule_2_para.style.style_id}")
    
    # Create a new paragraph with the same style
    last_para = doc.paragraphs[-1]
    target_p = last_para._p
    parent = target_p.getparent()
    target_position = list(parent).index(target_p)
    
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    new_p.append(pPr)
    
    # Set style
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Schedule')
    pPr.append(pStyle)
    
    # Add text
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "TEST Schedule with Schedule style"
    r.append(t)
    new_p.append(r)
    
    parent.insert(target_position + 1, new_p)
    
    doc.save(temp)
    print(f"Saved test document to {temp}")
    
    # Reload and check
    doc2 = Document(temp)
    for para in doc2.paragraphs:
        if "TEST Schedule" in para.text:
            print(f"New paragraph style after save/reload: {para.style.style_id}")
