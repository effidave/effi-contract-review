"""
Example demonstrating comment status extraction from Word documents.

This example shows how to extract comments with their status (active/resolved)
using the extended comment tools.
"""
import json
import sys
from pathlib import Path

# Add the parent directory to the path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from word_document_server.core.comments import (
    extract_all_comments,
    extract_comment_status_map,
    merge_comment_status
)


def example_extract_comments_with_status():
    """
    Example: Extract all comments from a document with their status.
    """
    print("=" * 70)
    print("Example: Extracting Comments with Status Information")
    print("=" * 70)
    
    # In a real scenario, this would be a user-provided document path
    # For testing, we can create a minimal example
    
    try:
        # Create a sample document
        doc = Document()
        doc.add_paragraph("This is a sample document for demonstrating comment status.")
        
        # Save it
        sample_doc_path = Path(__file__).parent / "sample_comments_doc.docx"
        doc.save(str(sample_doc_path))
        
        # Load and extract comments
        doc = Document(str(sample_doc_path))
        
        # Extract all comments with status information
        comments = extract_all_comments(doc)
        
        print(f"\nExtracted {len(comments)} comments from document")
        
        if comments:
            print("\nComment Structure (includes status fields):")
            print("-" * 70)
            
            # Show the first comment as an example
            comment = comments[0]
            
            print(f"  ID:          {comment.get('id')}")
            print(f"  Comment ID:  {comment.get('comment_id')}")
            print(f"  Para ID:     {comment.get('para_id')}")
            print(f"  Author:      {comment.get('author')}")
            print(f"  Initials:    {comment.get('initials')}")
            print(f"  Date:        {comment.get('date')}")
            print(f"  Text:        {comment.get('text')}")
            print(f"  Status:      {comment.get('status')} (NEW FIELD)")
            print(f"  Resolved:    {comment.get('is_resolved')} (NEW FIELD)")
            print(f"  Done Flag:   {comment.get('done_flag')} (NEW FIELD)")
            
            print("\n  Status Information Explanation:")
            print("    - status:     'active' (unresolved) or 'resolved'")
            print("    - is_resolved: boolean True/False")
            print("    - done_flag:   integer 0=active, 1=resolved (from w15:done)")
        
        # Clean up
        sample_doc_path.unlink(missing_ok=True)
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()


def example_status_map():
    """
    Example: Show how the status map is built.
    """
    print("\n" + "=" * 70)
    print("Example: Comment Status Map Structure")
    print("=" * 70)
    
    print("\nThe status_map is built from commentsExtended.xml:")
    print("-" * 70)
    
    status_map_example = {
        "128F7C96": {
            "status": "active",
            "done": 0,
            "is_resolved": False
        },
        "115FC64C": {
            "status": "resolved",
            "done": 1,
            "is_resolved": True
        }
    }
    
    print(json.dumps(status_map_example, indent=2))
    
    print("\nMapping Process:")
    print("  1. Extract w15:commentEx elements from commentsExtended.xml")
    print("  2. For each element, extract:")
    print("     - w15:paraId (links to w14:paraId in comments.xml)")
    print("     - w15:done (0=active, 1=resolved)")
    print("  3. Merge into comment dictionaries using para_id as the key")


def example_filtering_by_status():
    """
    Example: How to filter comments by status.
    """
    print("\n" + "=" * 70)
    print("Example: Filtering Comments by Status")
    print("=" * 70)
    
    print("\nAfter extracting comments with status, you can filter them:")
    print("-" * 70)
    
    sample_comments = [
        {
            "id": "comment_1",
            "author": "David Sant",
            "text": "This needs revision",
            "status": "active",
            "is_resolved": False
        },
        {
            "id": "comment_2",
            "author": "David Sant",
            "text": "Good point about the formatting",
            "status": "resolved",
            "is_resolved": True
        },
        {
            "id": "comment_3",
            "author": "David Sant",
            "text": "Please clarify this section",
            "status": "active",
            "is_resolved": False
        }
    ]
    
    # Filter active comments
    active_comments = [c for c in sample_comments if c['status'] == 'active']
    resolved_comments = [c for c in sample_comments if c['status'] == 'resolved']
    
    print(f"\nActive (unresolved) comments: {len(active_comments)}")
    for comment in active_comments:
        print(f"  - {comment['id']}: \"{comment['text']}\"")
    
    print(f"\nResolved comments: {len(resolved_comments)}")
    for comment in resolved_comments:
        print(f"  - {comment['id']}: \"{comment['text']}\"")


if __name__ == "__main__":
    print("\n" + "=" * 70)
    print("COMMENT STATUS EXTRACTION - IMPLEMENTATION EXAMPLES")
    print("=" * 70)
    
    example_extract_comments_with_status()
    example_status_map()
    example_filtering_by_status()
    
    print("\n" + "=" * 70)
    print("SUMMARY: Using Comment Status in Tools")
    print("=" * 70)
    
    print("""
The extended comment tools now support three new fields in comment data:

1. status (str):     'active' or 'resolved'
2. is_resolved (bool): True if comment is resolved, False otherwise  
3. done_flag (int):  0 for active, 1 for resolved (matches w15:done)

These fields are automatically populated from commentsExtended.xml when:
- Extracting all comments via extract_all_comments()
- Getting comments by author via filter_comments_by_author()
- Getting comments for a paragraph via get_comments_for_paragraph()

The linking mechanism uses w14:paraId (in comments.xml) matched with
w15:paraId (in commentsExtended.xml) to associate each comment with its status.

For documents without commentsExtended.xml (older Word versions), all
comments default to status='active', is_resolved=False, done_flag=0.
    """)
    
    print("\n✓ Example demonstration complete!\n")
