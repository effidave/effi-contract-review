from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import patch
from uuid import uuid4

from docx import Document

from effilocal.doc import direct_docx
from effilocal.doc.trackers.numbering_tracker import NumberingTracker
from effilocal.flows.analyze_doc import analyze


def _build_attachment_doc(tmp_path: Path) -> Path:
    doc = Document()

    schedule = doc.add_paragraph("Schedule 1 - Data Processing Agreement")
    schedule.style = "Heading 1"
    doc.add_paragraph("Schedule introduction paragraph.")
    schedule_item = doc.add_paragraph("Schedule obligation")
    schedule_item.style = "List Number"

    annex = doc.add_paragraph("Annex A - Security Measures")
    annex.style = "Normal"
    annex_item = doc.add_paragraph("Annex obligation")
    annex_item.style = "List Number"
    doc.add_paragraph("Annex body content.")

    heading = doc.add_paragraph("1 Definitions")
    heading.style = "Heading 1"
    doc.add_paragraph("Definitions clause text.")

    doc_path = tmp_path / "attachment_example.docx"
    doc.save(doc_path)
    return doc_path


def test_attachment_blocks_include_metadata(tmp_path: Path) -> None:
    doc_path = _build_attachment_doc(tmp_path)

    original_reset = NumberingTracker.reset_for_attachment
    reset_count = {"value": 0}

    def _wrapped_reset(self) -> None:
        reset_count["value"] += 1
        return original_reset(self)

    with patch.object(NumberingTracker, "reset_for_attachment", _wrapped_reset):
        blocks = list(direct_docx.iter_blocks(doc_path))

    attachment_blocks = [block for block in blocks if isinstance(block.get("attachment"), dict)]

    assert len(attachment_blocks) == 2, "Expected one schedule and one annex anchor."
    schedule_block = attachment_blocks[0]
    annex_block = attachment_blocks[1]

    schedule_meta = schedule_block["attachment"]
    annex_meta = annex_block["attachment"]

    assert schedule_meta["type"] == "schedule"
    assert schedule_meta["ordinal"] == "1"
    assert schedule_meta["title"] == "Data Processing Agreement"
    assert schedule_meta["parent_attachment_id"] is None

    assert annex_meta["type"] == "annex"
    assert annex_meta["ordinal"] == "A"
    assert annex_meta["parent_attachment_id"] == schedule_meta["attachment_id"]

    intro = next(block for block in blocks if block["text"] == "Schedule introduction paragraph.")
    assert intro["attachment_id"] == schedule_meta["attachment_id"]

    annex_body = next(block for block in blocks if block["text"] == "Annex body content.")
    assert annex_body["attachment_id"] == annex_meta["attachment_id"]

    main_heading = next(block for block in blocks if block["text"] == "1 Definitions")
    assert main_heading["attachment_id"] is None

    schedule_item_block = next(block for block in blocks if block["text"] == "Schedule obligation")
    annex_item_block = next(block for block in blocks if block["text"] == "Annex obligation")

    assert schedule_item_block["attachment_id"] == schedule_meta["attachment_id"]
    assert annex_item_block["attachment_id"] == annex_meta["attachment_id"]
    assert reset_count["value"] >= 2


def test_manifest_records_attachments(tmp_path: Path) -> None:
    doc_path = _build_attachment_doc(tmp_path)
    out_dir = tmp_path / "analysis"
    doc_id = str(uuid4())

    analyze(doc_path, doc_id=doc_id, out_dir=out_dir)

    manifest_path = out_dir / "manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    attachments = manifest.get("attachments")

    assert isinstance(attachments, list)
    assert len(attachments) == 2

    schedule_entry, annex_entry = attachments
    assert schedule_entry["type"] == "schedule"
    assert annex_entry["type"] == "annex"
    assert annex_entry["parent_attachment_id"] == schedule_entry["attachment_id"]

