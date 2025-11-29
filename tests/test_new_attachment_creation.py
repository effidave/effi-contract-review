"""Test adding new attachments with style inheritance."""
import asyncio
import sys
from pathlib import Path
import shutil
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from effilocal.mcp_server.tools.attachment_tools import add_new_attachment_after
from effilocal.doc import direct_docx
from docx import Document
from docx.oxml.ns import qn

@pytest.mark.asyncio
async def test_new_attachment():
    temp = Path(__file__).parent / "temp_new_attachment.docx"
    shutil.copy2(Path(__file__).parent / "fixtures" / "real_world" / "schedules.docx", temp)
    
    try:
        print("=" * 80)
        print("TEST: Add new Schedule 3 after Schedule 2")
        print("=" * 80)
        
        # Get original Schedule 2 style
        doc = Document(temp)
        blocks = list(direct_docx.iter_blocks(temp))
        
        schedule_2_block = None
        for b in blocks:
            att = b.get("attachment")
            if att and att.get("ordinal") == "2":
                schedule_2_block = b
                break
        
        if schedule_2_block:
            para_id = schedule_2_block.get("para_id")
            schedule_2_para = None
            for para in doc.paragraphs:
                if para._element.get(qn('w14:paraId')) == para_id:
                    schedule_2_para = para
                    break
            
            if schedule_2_para:
                print(f"\nSchedule 2 original style:")
                print(f"  Style: {schedule_2_para.style.name if schedule_2_para.style else 'None'}")
                print(f"  Style ID: {schedule_2_para.style.style_id if schedule_2_para.style else 'None'}")
                if schedule_2_para.runs:
                    run = schedule_2_para.runs[0]
                    print(f"  Font: {run.font.name}")
                    print(f"  Size: {run.font.size}")
                    print(f"  Bold: {run.font.bold}")
        
        # Add new Schedule 3
        result = await add_new_attachment_after(
            filename=str(temp),
            after_attachment="Schedule 2",
            new_attachment_text="Schedule 3 - Additional Terms and Conditions",
            content="[INSERT SCHEDULE 3 CONTENT HERE]"
        )
        
        print(f"\nResult: {result}")
        
        # Verify the new schedule was added
        print("\n" + "=" * 80)
        print("VERIFICATION: Checking document structure")
        print("=" * 80)
        
        blocks_after = list(direct_docx.iter_blocks(temp))
        
        # Find all schedules
        schedules = []
        for i, b in enumerate(blocks_after):
            att = b.get("attachment")
            if att and att.get("type") == "schedule":
                schedules.append({
                    'idx': i,
                    'ordinal': att.get("ordinal"),
                    'title': att.get("title"),
                    'text': b.get("text", "")[:80]
                })
        
        print(f"\nFound {len(schedules)} schedules:")
        for s in schedules:
            print(f"  Schedule {s['ordinal']}: {s['text']}")
        
        # Check if Schedule 3 exists
        schedule_3 = [s for s in schedules if s['ordinal'] == '3']
        if schedule_3:
            print(f"\n✓ Schedule 3 successfully created!")
            print(f"  Text: {schedule_3[0]['text']}")
        else:
            print(f"\n✗ Schedule 3 NOT found")
        
        # Check the style of Schedule 3
        print("\n" + "=" * 80)
        print("STYLE VERIFICATION: Comparing Schedule 2 vs Schedule 3")
        print("=" * 80)
        
        doc_after = Document(temp)
        
        # Find Schedule 3 paragraph (the actual attachment header, not TOC)
        # Look for paragraphs with "Schedule 3" and check which one has the Schedule style
        schedule_3_candidates = []
        for i, para in enumerate(doc_after.paragraphs):
            if "Schedule 3" in para.text and "Additional Terms" in para.text:
                schedule_3_candidates.append({
                    'index': i,
                    'para': para,
                    'style_id': para.style.style_id if para.style else None,
                    'text': para.text[:60]
                })
        
        print(f"Found {len(schedule_3_candidates)} paragraphs with 'Schedule 3':")
        for cand in schedule_3_candidates:
            print(f"  Para {cand['index']}: style={cand['style_id']}, text={cand['text']}")
        
        # Use the one with "Schedule" style (the actual header)
        schedule_3_para = None
        for cand in schedule_3_candidates:
            if cand['style_id'] == 'Schedule':
                schedule_3_para = cand['para']
                print(f"\n→ Using para {cand['index']} as the Schedule 3 header")
                break
        
        if schedule_3_para:
            print(f"\nSchedule 3 style:")
            print(f"  Style: {schedule_3_para.style.name if schedule_3_para.style else 'None'}")
            print(f"  Style ID: {schedule_3_para.style.style_id if schedule_3_para.style else 'None'}")
            if schedule_3_para.runs:
                run = schedule_3_para.runs[0]
                print(f"  Font: {run.font.name}")
                print(f"  Size: {run.font.size}")
                print(f"  Bold: {run.font.bold}")
            
            # Compare with Schedule 2
            schedule_2_para_after = None
            for para in doc_after.paragraphs:
                if "Schedule 2 - MSA" in para.text:
                    schedule_2_para_after = para
                    break
            
            if schedule_2_para_after:
                match_style = (schedule_2_para_after.style == schedule_3_para.style)
                print(f"\n{'✓' if match_style else '✗'} Styles match: {match_style}")
        
        print("\n" + "=" * 80)
    finally:
        # Cleanup
        if temp.exists():
            temp.unlink()

if __name__ == "__main__":
    asyncio.run(test_new_attachment())
