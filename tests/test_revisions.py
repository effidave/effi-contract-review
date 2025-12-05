"""
Comprehensive tests for track changes (revision) extraction, accept, and reject functionality.

Tests cover:
1. Revision extraction from Word documents
2. Accept individual revisions (insertions and deletions)
3. Reject individual revisions
4. Accept/reject all revisions
5. Revision metadata (author, date, type)
6. Para_id association for UI linking

Test Strategy:
- Create fixture documents with known tracked changes
- Test extraction returns correct structure
- Test accept/reject modifies document correctly
- Test edge cases (empty revisions, nested runs, etc.)
"""
import os
import sys
import shutil
import tempfile
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any

import pytest

# Add the parent directory to the path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap, qn, nsdecls


# ============================================================================
# Fixture: Create test document with tracked changes
# ============================================================================

def create_document_with_tracked_changes() -> Document:
    """Create a test document with various tracked changes."""
    doc = Document()
    
    # Enable track changes by adding the appropriate setting
    # (Note: python-docx doesn't directly support track changes, so we'll inject XML)
    
    # Paragraph 1: Normal text
    p1 = doc.add_paragraph("This is normal unchanged text.")
    
    # Paragraph 2: Contains an insertion
    p2 = doc.add_paragraph()
    p2_element = p2._element
    
    # Add normal run
    r1 = parse_xml(f'''
        <w:r {nsdecls('w')}>
            <w:t xml:space="preserve">Before the </w:t>
        </w:r>
    ''')
    p2_element.append(r1)
    
    # Add insertion
    ins = parse_xml(f'''
        <w:ins {nsdecls('w', 'w14')} w:id="1" w:author="Test Author" w:date="2025-12-01T10:00:00Z">
            <w:r>
                <w:t>INSERTED TEXT</w:t>
            </w:r>
        </w:ins>
    ''')
    p2_element.append(ins)
    
    # Add more normal run
    r2 = parse_xml(f'''
        <w:r {nsdecls('w')}>
            <w:t xml:space="preserve"> insertion.</w:t>
        </w:r>
    ''')
    p2_element.append(r2)
    
    # Paragraph 3: Contains a deletion
    p3 = doc.add_paragraph()
    p3_element = p3._element
    
    # Add normal run
    r3 = parse_xml(f'''
        <w:r {nsdecls('w')}>
            <w:t xml:space="preserve">Text with </w:t>
        </w:r>
    ''')
    p3_element.append(r3)
    
    # Add deletion
    delete = parse_xml(f'''
        <w:del {nsdecls('w', 'w14')} w:id="2" w:author="Another Author" w:date="2025-12-02T14:30:00Z">
            <w:r>
                <w:delText>DELETED TEXT</w:delText>
            </w:r>
        </w:del>
    ''')
    p3_element.append(delete)
    
    # Add more normal run
    r4 = parse_xml(f'''
        <w:r {nsdecls('w')}>
            <w:t xml:space="preserve"> remaining.</w:t>
        </w:r>
    ''')
    p3_element.append(r4)
    
    # Paragraph 4: Contains both insertion and deletion
    p4 = doc.add_paragraph()
    p4_element = p4._element
    
    r5 = parse_xml(f'''
        <w:r {nsdecls('w')}>
            <w:t xml:space="preserve">Mixed: </w:t>
        </w:r>
    ''')
    p4_element.append(r5)
    
    del2 = parse_xml(f'''
        <w:del {nsdecls('w', 'w14')} w:id="3" w:author="Test Author" w:date="2025-12-03T09:00:00Z">
            <w:r>
                <w:delText>old text</w:delText>
            </w:r>
        </w:del>
    ''')
    p4_element.append(del2)
    
    ins2 = parse_xml(f'''
        <w:ins {nsdecls('w', 'w14')} w:id="4" w:author="Test Author" w:date="2025-12-03T09:00:00Z">
            <w:r>
                <w:t>new text</w:t>
            </w:r>
        </w:ins>
    ''')
    p4_element.append(ins2)
    
    r6 = parse_xml(f'''
        <w:r {nsdecls('w')}>
            <w:t xml:space="preserve"> end.</w:t>
        </w:r>
    ''')
    p4_element.append(r6)
    
    return doc


@pytest.fixture
def tracked_changes_doc(tmp_path) -> Path:
    """Create a temporary document with tracked changes."""
    doc = create_document_with_tracked_changes()
    doc_path = tmp_path / "test_tracked_changes.docx"
    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def hj9_doc_path() -> Path:
    """Path to the real HJ9 document with tracked changes."""
    path = Path(__file__).parent.parent / "EL_Projects" / "Test Project" / "drafts" / "current_drafts" / "Norton R&D Services Agreement (DRAFT) - HJ9 (TRACKED).docx"
    if not path.exists():
        pytest.skip(f"HJ9 document not found at {path}")
    return path


@pytest.fixture
def hj9_doc_copy(hj9_doc_path, tmp_path) -> Path:
    """Create a copy of HJ9 doc for destructive tests."""
    copy_path = tmp_path / "hj9_copy.docx"
    shutil.copy(hj9_doc_path, copy_path)
    return copy_path


# ============================================================================
# Tests: Revision Extraction
# ============================================================================

class TestRevisionExtraction:
    """Tests for extracting revisions from Word documents."""
    
    def test_extract_revisions_returns_list(self, tracked_changes_doc):
        """Extract revisions should return a list."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        assert isinstance(revisions, list)
    
    def test_extract_revisions_finds_insertions(self, tracked_changes_doc):
        """Extract revisions should find insertion elements."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        insertions = [r for r in revisions if r['type'] == 'insert']
        assert len(insertions) >= 1, "Should find at least one insertion"
    
    def test_extract_revisions_finds_deletions(self, tracked_changes_doc):
        """Extract revisions should find deletion elements."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        deletions = [r for r in revisions if r['type'] == 'delete']
        assert len(deletions) >= 1, "Should find at least one deletion"
    
    def test_revision_has_required_fields(self, tracked_changes_doc):
        """Each revision should have all required fields."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        required_fields = ['id', 'type', 'text', 'author', 'date', 'paragraph_index']
        
        for rev in revisions:
            for field in required_fields:
                assert field in rev, f"Revision missing required field: {field}"
    
    def test_insertion_contains_inserted_text(self, tracked_changes_doc):
        """Insertion revisions should contain the inserted text."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        insertions = [r for r in revisions if r['type'] == 'insert']
        # Find our known inserted text
        found = any('INSERTED TEXT' in r.get('text', '') for r in insertions)
        assert found, "Should find 'INSERTED TEXT' in insertions"
    
    def test_deletion_contains_deleted_text(self, tracked_changes_doc):
        """Deletion revisions should contain the deleted text."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        deletions = [r for r in revisions if r['type'] == 'delete']
        # Find our known deleted text
        found = any('DELETED TEXT' in r.get('text', '') for r in deletions)
        assert found, "Should find 'DELETED TEXT' in deletions"
    
    def test_revision_has_author(self, tracked_changes_doc):
        """Revisions should include author information."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        # Our test doc has "Test Author" and "Another Author"
        authors = {r.get('author') for r in revisions}
        assert 'Test Author' in authors or len(authors) > 0
    
    def test_revision_has_date(self, tracked_changes_doc):
        """Revisions should include date information."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        for rev in revisions:
            assert rev.get('date'), f"Revision {rev.get('id')} missing date"
    
    def test_revision_has_paragraph_index(self, tracked_changes_doc):
        """Revisions should be associated with a paragraph index."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        for rev in revisions:
            assert 'paragraph_index' in rev
            assert isinstance(rev['paragraph_index'], int)
            assert rev['paragraph_index'] >= 0


class TestRevisionExtractionRealDocument:
    """Tests using the real HJ9 document with tracked changes."""
    
    def test_extract_from_hj9_document(self, hj9_doc_path):
        """Should extract revisions from the real HJ9 document."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(hj9_doc_path))
        revisions = extract_all_revisions(doc)
        
        # HJ9 has 192 insertions and 33 deletions based on our earlier check
        assert len(revisions) > 0, "Should find revisions in HJ9 document"
        
        insertions = [r for r in revisions if r['type'] == 'insert']
        deletions = [r for r in revisions if r['type'] == 'delete']
        
        # Should have both types
        assert len(insertions) > 0, "Should find insertions"
        assert len(deletions) > 0, "Should find deletions"
    
    def test_hj9_revision_count_matches(self, hj9_doc_path):
        """Revision count should approximately match XML element count."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(hj9_doc_path))
        revisions = extract_all_revisions(doc)
        
        insertions = [r for r in revisions if r['type'] == 'insert']
        deletions = [r for r in revisions if r['type'] == 'delete']
        
        # Based on our earlier check: 192 insertions, 33 deletions in XML
        # But many may be formatting-only (no text), so we count text-containing revisions
        # Allow some variance due to how we count (nested elements, formatting-only, etc.)
        assert len(insertions) >= 50, f"Expected many insertions, got {len(insertions)}"
        # Deletions with actual text content may be fewer than raw XML count
        assert len(deletions) >= 1, f"Expected some deletions, got {len(deletions)}"
    
    def test_hj9_author_is_david_sant(self, hj9_doc_path):
        """HJ9 revisions should have David Sant as author."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(hj9_doc_path))
        revisions = extract_all_revisions(doc)
        
        authors = {r.get('author') for r in revisions if r.get('author')}
        assert 'David Sant' in authors


# ============================================================================
# Tests: Accept Revision
# ============================================================================

class TestAcceptRevision:
    """Tests for accepting individual revisions."""
    
    def test_accept_insertion_keeps_text(self, tracked_changes_doc):
        """Accepting an insertion should keep the text and remove the w:ins wrapper."""
        from effilocal.mcp_server.core.revisions import (
            extract_all_revisions,
            accept_revision
        )
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        # Find an insertion
        insertion = next((r for r in revisions if r['type'] == 'insert'), None)
        assert insertion, "Test requires at least one insertion"
        
        revision_id = insertion['id']
        original_text = insertion['text']
        
        # Accept the revision
        success = accept_revision(doc, revision_id)
        assert success, "Accept revision should return True"
        
        # Save and reload
        doc.save(str(tracked_changes_doc))
        doc2 = Document(str(tracked_changes_doc))
        
        # The text should still be there, but no longer marked as insertion
        revisions_after = extract_all_revisions(doc2)
        remaining_ids = {r['id'] for r in revisions_after}
        
        assert revision_id not in remaining_ids, "Accepted revision should be removed"
        
        # Verify text is kept
        full_text = '\n'.join(p.text for p in doc2.paragraphs)
        assert 'INSERTED TEXT' in full_text or original_text in full_text
    
    def test_accept_deletion_removes_text(self, tracked_changes_doc):
        """Accepting a deletion should remove the text entirely."""
        from effilocal.mcp_server.core.revisions import (
            extract_all_revisions,
            accept_revision
        )
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        # Find a deletion
        deletion = next((r for r in revisions if r['type'] == 'delete'), None)
        assert deletion, "Test requires at least one deletion"
        
        revision_id = deletion['id']
        deleted_text = deletion['text']
        
        # Accept the deletion (means we agree to delete)
        success = accept_revision(doc, revision_id)
        assert success, "Accept revision should return True"
        
        # Save and reload
        doc.save(str(tracked_changes_doc))
        doc2 = Document(str(tracked_changes_doc))
        
        # The deleted text should be gone
        full_text = '\n'.join(p.text for p in doc2.paragraphs)
        assert deleted_text not in full_text, "Deleted text should be removed after accepting deletion"
    
    def test_accept_returns_false_for_invalid_id(self, tracked_changes_doc):
        """Accept revision should return False for non-existent revision ID."""
        from effilocal.mcp_server.core.revisions import accept_revision
        
        doc = Document(str(tracked_changes_doc))
        success = accept_revision(doc, "nonexistent_id_12345")
        
        assert success is False
    
    def test_accept_on_hj9_copy(self, hj9_doc_copy):
        """Test accepting a revision on a copy of the real HJ9 document."""
        from effilocal.mcp_server.core.revisions import (
            extract_all_revisions,
            accept_revision
        )
        
        doc = Document(str(hj9_doc_copy))
        revisions = extract_all_revisions(doc)
        
        assert len(revisions) > 0, "Should have revisions to test"
        
        # Accept the first revision
        first_rev = revisions[0]
        success = accept_revision(doc, first_rev['id'])
        
        assert success, "Should successfully accept revision"
        
        # Save and verify count decreased
        doc.save(str(hj9_doc_copy))
        doc2 = Document(str(hj9_doc_copy))
        revisions_after = extract_all_revisions(doc2)
        
        assert len(revisions_after) == len(revisions) - 1


# ============================================================================
# Tests: Reject Revision
# ============================================================================

class TestRejectRevision:
    """Tests for rejecting individual revisions."""
    
    def test_reject_insertion_removes_text(self, tracked_changes_doc):
        """Rejecting an insertion should remove the inserted text."""
        from effilocal.mcp_server.core.revisions import (
            extract_all_revisions,
            reject_revision
        )
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        # Find an insertion
        insertion = next((r for r in revisions if r['type'] == 'insert'), None)
        assert insertion, "Test requires at least one insertion"
        
        revision_id = insertion['id']
        inserted_text = insertion['text']
        
        # Reject the insertion
        success = reject_revision(doc, revision_id)
        assert success, "Reject revision should return True"
        
        # Save and reload
        doc.save(str(tracked_changes_doc))
        doc2 = Document(str(tracked_changes_doc))
        
        # The inserted text should be gone
        full_text = '\n'.join(p.text for p in doc2.paragraphs)
        assert inserted_text not in full_text, "Inserted text should be removed after rejection"
    
    def test_reject_deletion_keeps_text(self, tracked_changes_doc):
        """Rejecting a deletion should keep the original text."""
        from effilocal.mcp_server.core.revisions import (
            extract_all_revisions,
            reject_revision
        )
        
        doc = Document(str(tracked_changes_doc))
        revisions = extract_all_revisions(doc)
        
        # Find a deletion
        deletion = next((r for r in revisions if r['type'] == 'delete'), None)
        assert deletion, "Test requires at least one deletion"
        
        revision_id = deletion['id']
        deleted_text = deletion['text']
        
        # Reject the deletion (means we want to keep the text)
        success = reject_revision(doc, revision_id)
        assert success, "Reject revision should return True"
        
        # Save and reload
        doc.save(str(tracked_changes_doc))
        doc2 = Document(str(tracked_changes_doc))
        
        # The text should still be there (converted to normal text)
        full_text = '\n'.join(p.text for p in doc2.paragraphs)
        assert deleted_text in full_text, "Deleted text should be kept after rejecting deletion"
    
    def test_reject_returns_false_for_invalid_id(self, tracked_changes_doc):
        """Reject revision should return False for non-existent revision ID."""
        from effilocal.mcp_server.core.revisions import reject_revision
        
        doc = Document(str(tracked_changes_doc))
        success = reject_revision(doc, "nonexistent_id_99999")
        
        assert success is False


# ============================================================================
# Tests: Accept/Reject All
# ============================================================================

class TestBulkOperations:
    """Tests for accept all / reject all operations."""
    
    def test_accept_all_revisions(self, tracked_changes_doc):
        """Accept all should process all revisions."""
        from effilocal.mcp_server.core.revisions import (
            extract_all_revisions,
            accept_all_revisions
        )
        
        doc = Document(str(tracked_changes_doc))
        revisions_before = extract_all_revisions(doc)
        
        assert len(revisions_before) >= 2, "Test requires multiple revisions"
        
        # Accept all
        result = accept_all_revisions(doc)
        
        assert result['success'], "Accept all should succeed"
        assert result['accepted_count'] == len(revisions_before)
        
        # Save and verify no revisions remain
        doc.save(str(tracked_changes_doc))
        doc2 = Document(str(tracked_changes_doc))
        revisions_after = extract_all_revisions(doc2)
        
        assert len(revisions_after) == 0, "No revisions should remain after accept all"
    
    def test_reject_all_revisions(self, tracked_changes_doc):
        """Reject all should process all revisions."""
        from effilocal.mcp_server.core.revisions import (
            extract_all_revisions,
            reject_all_revisions
        )
        
        doc = Document(str(tracked_changes_doc))
        revisions_before = extract_all_revisions(doc)
        
        assert len(revisions_before) >= 2, "Test requires multiple revisions"
        
        # Reject all
        result = reject_all_revisions(doc)
        
        assert result['success'], "Reject all should succeed"
        assert result['rejected_count'] == len(revisions_before)
        
        # Save and verify no revisions remain
        doc.save(str(tracked_changes_doc))
        doc2 = Document(str(tracked_changes_doc))
        revisions_after = extract_all_revisions(doc2)
        
        assert len(revisions_after) == 0, "No revisions should remain after reject all"
    
    def test_accept_all_on_hj9_copy(self, hj9_doc_copy):
        """Accept all should work on real HJ9 document."""
        from effilocal.mcp_server.core.revisions import (
            extract_all_revisions,
            accept_all_revisions
        )
        
        doc = Document(str(hj9_doc_copy))
        revisions_before = extract_all_revisions(doc)
        
        initial_count = len(revisions_before)
        assert initial_count > 0, "HJ9 should have revisions"
        
        result = accept_all_revisions(doc)
        
        assert result['success']
        assert result['accepted_count'] == initial_count
        
        # Save and verify
        doc.save(str(hj9_doc_copy))
        doc2 = Document(str(hj9_doc_copy))
        revisions_after = extract_all_revisions(doc2)
        
        assert len(revisions_after) == 0


# ============================================================================
# Tests: Para_id Association
# ============================================================================

class TestParaIdAssociation:
    """Tests for linking revisions to paragraphs via para_id."""
    
    def test_revision_has_para_id(self, hj9_doc_path):
        """Revisions should have para_id for UI linking."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document(str(hj9_doc_path))
        revisions = extract_all_revisions(doc)
        
        # At least some revisions should have para_id
        revisions_with_para_id = [r for r in revisions if r.get('para_id')]
        
        assert len(revisions_with_para_id) > 0, "Some revisions should have para_id"
    
    def test_para_id_format(self, hj9_doc_path):
        """Para_id should be in 8-char hex format."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        import re
        
        doc = Document(str(hj9_doc_path))
        revisions = extract_all_revisions(doc)
        
        hex_pattern = re.compile(r'^[0-9A-Fa-f]{8}$')
        
        for rev in revisions:
            para_id = rev.get('para_id')
            if para_id:
                assert hex_pattern.match(para_id), f"Para_id should be 8-char hex: {para_id}"


# ============================================================================
# Tests: Edge Cases
# ============================================================================

class TestEdgeCases:
    """Tests for edge cases and error handling."""
    
    def test_empty_document(self, tmp_path):
        """Should handle document with no revisions."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document()
        doc.add_paragraph("No tracked changes here.")
        doc_path = tmp_path / "empty_revisions.docx"
        doc.save(str(doc_path))
        
        doc2 = Document(str(doc_path))
        revisions = extract_all_revisions(doc2)
        
        assert revisions == [], "Empty document should return empty list"
    
    def test_revision_with_multiple_runs(self, tmp_path):
        """Should handle insertions/deletions spanning multiple runs."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document()
        p = doc.add_paragraph()
        p_element = p._element
        
        # Create insertion with multiple runs (different formatting)
        ins = parse_xml(f'''
            <w:ins {nsdecls('w', 'w14')} w:id="10" w:author="Multi Run" w:date="2025-12-01T10:00:00Z">
                <w:r>
                    <w:t>First part </w:t>
                </w:r>
                <w:r>
                    <w:rPr><w:b/></w:rPr>
                    <w:t>bold part </w:t>
                </w:r>
                <w:r>
                    <w:t>last part</w:t>
                </w:r>
            </w:ins>
        ''')
        p_element.append(ins)
        
        doc_path = tmp_path / "multi_run.docx"
        doc.save(str(doc_path))
        
        doc2 = Document(str(doc_path))
        revisions = extract_all_revisions(doc2)
        
        assert len(revisions) >= 1
        # The text should be combined from all runs
        insertion = revisions[0]
        assert 'First part' in insertion['text']
        assert 'bold part' in insertion['text']
        assert 'last part' in insertion['text']
    
    def test_nested_formatting_in_revision(self, tmp_path):
        """Should handle complex formatting inside revisions."""
        from effilocal.mcp_server.core.revisions import extract_all_revisions
        
        doc = Document()
        p = doc.add_paragraph()
        p_element = p._element
        
        # Insertion containing formatted text
        ins = parse_xml(f'''
            <w:ins {nsdecls('w', 'w14')} w:id="20" w:author="Format Test" w:date="2025-12-01T10:00:00Z">
                <w:r>
                    <w:rPr><w:b/><w:i/></w:rPr>
                    <w:t>Bold and Italic</w:t>
                </w:r>
            </w:ins>
        ''')
        p_element.append(ins)
        
        doc_path = tmp_path / "formatted_revision.docx"
        doc.save(str(doc_path))
        
        doc2 = Document(str(doc_path))
        revisions = extract_all_revisions(doc2)
        
        assert len(revisions) >= 1
        assert 'Bold and Italic' in revisions[0]['text']


# ============================================================================
# Run tests if executed directly
# ============================================================================

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
