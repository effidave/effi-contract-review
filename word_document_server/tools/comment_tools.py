"""
Comment extraction tools for Word Document Server.

These tools provide high-level interfaces for extracting and analyzing
comments from Word documents through the MCP protocol.
"""
import os
import json
from typing import Dict, List, Optional, Any
from docx import Document
from docx.text.run import Run
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.comments import CommentsPart


from word_document_server.utils.file_utils import ensure_docx_extension
from word_document_server.core.comments import (
    extract_all_comments,
    filter_comments_by_author,
    get_comments_for_paragraph,
    extract_comment_status_map,
    merge_comment_status
)


def _get_or_add_comments_part(doc_part):
    """Return the comments part; create it if it doesn't exist.
    This follows the WordprocessingML spec: create /word/comments.xml and relate it.
    """
    try:
        return doc_part.part_related_by(RT.COMMENTS)
    except KeyError:
        # Let python-docx provide a correctly wired default part
        package = doc_part.package
        assert package is not None
        comments_part = CommentsPart.default(package)
        doc_part.relate_to(comments_part, RT.COMMENTS)
        return comments_part


def _wrap_run_with_comment(r: Run, cid: int):
    """Insert comment range start/end markers around the given Run and add a commentReference run."""
    r_el = r._r

    # commentRangeStart before the run
    start = OxmlElement('w:commentRangeStart')
    start.set(qn('w:id'), str(cid))
    r_el.addprevious(start)

    # commentRangeEnd after the run
    end = OxmlElement('w:commentRangeEnd')
    end.set(qn('w:id'), str(cid))
    r_el.addnext(end)

    # add a reference run right after the end
    ref_run = OxmlElement('w:r')
    comment_ref = OxmlElement('w:commentReference')
    comment_ref.set(qn('w:id'), str(cid))
    ref_run.append(comment_ref)
    end.addnext(ref_run)


# ----------------------- NEW PUBLIC TOOLS -----------------------

async def add_comment_after_text(
    filename: str,
    search_text: str,
    comment_text: str,
    author: Optional[str] = None,
    initials: Optional[str] = None,
):
    """Add a Word comment to the first occurrence of `search_text`.

    Returns a JSON payload with success flag and metadata.
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"}, indent=2)

    if not search_text:
        return json.dumps({"success": False, "error": "search_text cannot be empty"}, indent=2)

    try:
        doc = Document(filename)

        target_run = None
        for p in doc.paragraphs:
            for run in p.runs:
                if search_text in run.text:
                    target_run = run
                    break
            if target_run:
                break

        if target_run is None:
            return json.dumps({
                "success": False,
                "error": f"Text '{search_text}' not found in document."
            }, indent=2)

        comments_part = _get_or_add_comments_part(doc.part)
        comment = comments_part.comments.add_comment(
            text=comment_text or "",
            author=author or "",
            initials=initials if initials is not None else "",
        )
        cid = comment.comment_id
        _wrap_run_with_comment(target_run, cid)

        doc.save(filename)
        return json.dumps({
            "success": True,
            "action": "add_comment_after_text",
            "filename": filename,
            "search_text": search_text,
            "comment_id": cid,
            "author": author,
            "initials": initials,
        }, indent=2)

    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to add comment: {str(e)}"}, indent=2)


async def add_comment_for_paragraph(
    filename: str,
    paragraph_index: int,
    comment_text: str,
    author: Optional[str] = None,
    initials: Optional[str] = None,
):
    """Add a Word comment anchored to the entire paragraph at `paragraph_index` (0-based)."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"}, indent=2)

    try:
        doc = Document(filename)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return json.dumps({
                "success": False,
                "error": f"Paragraph index {paragraph_index} is out of range (0..{len(doc.paragraphs)-1})."
            }, indent=2)

        p = doc.paragraphs[paragraph_index]
        # If paragraph has no runs, create an empty run so we have an anchor
        if not p.runs:
            r = p.add_run("")
        else:
            r = p.runs[0]

        comments_part = _get_or_add_comments_part(doc.part)
        comment = comments_part.comments.add_comment(
            text=comment_text or "",
            author=author or "",
            initials=initials if initials is not None else "",
        )
        cid = comment.comment_id
        _wrap_run_with_comment(r, cid)

        doc.save(filename)
        return json.dumps({
            "success": True,
            "action": "add_comment_for_paragraph",
            "filename": filename,
            "paragraph_index": paragraph_index,
            "comment_id": cid,
            "author": author,
            "initials": initials,
        }, indent=2)

    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to add comment: {str(e)}"}, indent=2)


async def get_all_comments(filename: str) -> str:
    """
    Extract all comments from a Word document.
    
    Args:
        filename: Path to the Word document
        
    Returns:
        JSON string containing all comments with metadata
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Extract all comments
        comments = extract_all_comments(doc)
        
        # Return results
        return json.dumps({
            'success': True,
            'comments': comments,
            'total_comments': len(comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)


async def get_comments_by_author(filename: str, author: str) -> str:
    """
    Extract comments from a specific author in a Word document.
    
    Args:
        filename: Path to the Word document
        author: Name of the comment author to filter by
        
    Returns:
        JSON string containing filtered comments
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    if not author or not author.strip():
        return json.dumps({
            'success': False,
            'error': 'Author name cannot be empty'
        }, indent=2)
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Extract all comments
        all_comments = extract_all_comments(doc)
        
        # Filter by author
        author_comments = filter_comments_by_author(all_comments, author)
        
        # Return results
        return json.dumps({
            'success': True,
            'author': author,
            'comments': author_comments,
            'total_comments': len(author_comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)


async def get_comments_for_paragraph(filename: str, paragraph_index: int) -> str:
    """
    Extract comments for a specific paragraph in a Word document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        
    Returns:
        JSON string containing comments for the specified paragraph
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    if paragraph_index < 0:
        return json.dumps({
            'success': False,
            'error': 'Paragraph index must be non-negative'
        }, indent=2)
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Check if paragraph index is valid
        if paragraph_index >= len(doc.paragraphs):
            return json.dumps({
                'success': False,
                'error': f'Paragraph index {paragraph_index} is out of range. Document has {len(doc.paragraphs)} paragraphs.'
            }, indent=2)
        
        # Extract all comments
        all_comments = extract_all_comments(doc)
        
        # Filter for the specific paragraph
        from word_document_server.core.comments import get_comments_for_paragraph as core_get_comments_for_paragraph
        para_comments = core_get_comments_for_paragraph(all_comments, paragraph_index)
        
        # Get the paragraph text for context
        paragraph_text = doc.paragraphs[paragraph_index].text
        
        # Return results
        return json.dumps({
            'success': True,
            'paragraph_index': paragraph_index,
            'paragraph_text': paragraph_text,
            'comments': para_comments,
            'total_comments': len(para_comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)
