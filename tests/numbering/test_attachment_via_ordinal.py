"""Tests for attachment detection via ordinal patterns (e.g., 'Schedule %1')."""

import pytest
from pathlib import Path
from docx import Document

from effilocal.doc.pipeline import AnalysisPipeline
from effilocal.doc.numbering import load_numbering
from effilocal.doc.numbering_inspector import NumberingInspector


@pytest.fixture
def norton_docx_path() -> Path:
    """Path to the Norton R&D document with Schedule numbering."""
    return Path(
        "EL_Projects/Norton - R&D/drafts/current_drafts/"
        "Norton R&D Services Agreement (DRAFT) - HJ9 (TRACKED).docx"
    )


@pytest.fixture
def schedule_fixture_path(tmp_path: Path) -> Path:
    """Create a minimal docx with Schedule-style numbering for testing."""
    from docx.oxml.ns import qn
    from docx.oxml import register_element_cls
    from docx.oxml.numbering import CT_Num, CT_NumLvl
    
    doc = Document()
    
    # We can't easily create custom numbering in python-docx,
    # so we'll use a real document that has Schedule numbering
    # For unit testing, we rely on the Norton document fixture
    return None


def test_schedule_blocks_get_attachment_id(norton_docx_path: Path) -> None:
    """Verify that blocks with 'Schedule %1' pattern get attachment_id assigned.
    
    This is a regression test for the bug where numbering events were stored
    with block['id'] (which is None during iteration) instead of para_id.
    """
    if not norton_docx_path.exists():
        pytest.skip("Norton document not available")
    
    doc = Document(str(norton_docx_path))
    numbering_defs = load_numbering(norton_docx_path)
    numbering_inspector = NumberingInspector.from_docx(norton_docx_path)
    
    pipeline = AnalysisPipeline(
        numbering_defs=numbering_defs,
        numbering_inspector=numbering_inspector,
        fallback_heading_label="Body (no heading)",
    )
    
    # Process all paragraphs
    blocks = []
    for para in doc.paragraphs:
        block = pipeline.process_paragraph(para)
        if block:
            blocks.append(block)
    
    # Find Schedule blocks (those with pattern containing "Schedule")
    schedule_blocks = [
        b for b in blocks 
        if b.get('list') and 'Schedule' in str(b['list'].get('pattern', ''))
    ]
    
    # There should be Schedule blocks in the Norton document
    assert len(schedule_blocks) >= 1, "Expected at least one Schedule block"
    
    # Each Schedule block should have an attachment_id
    for block in schedule_blocks:
        assert block.get('attachment_id') is not None, (
            f"Schedule block with ordinal '{block['list'].get('ordinal')}' "
            f"at para_idx={block.get('para_idx')} should have attachment_id"
        )
        
        # The attachment should also be set with proper metadata
        attachment = block.get('attachment')
        assert attachment is not None, "Schedule block should have attachment metadata"
        assert attachment.get('type') == 'schedule', "Attachment type should be 'schedule'"
        assert attachment.get('ordinal') is not None, "Attachment should have ordinal"


def test_attachment_event_uses_para_id_not_block_id() -> None:
    """Verify that numbering events use para_id for matching, not block['id'].
    
    This tests the core fix: since block['id'] is None during iteration,
    we must use para_id for event storage and lookup.
    """
    from effilocal.doc.trackers.attachment_tracker import AttachmentTracker
    from effilocal.doc.numbering_context import NumberingEvent
    
    tracker = AttachmentTracker()
    
    # Simulate a block with para_id but no block id (as during iteration)
    block = {
        'id': None,  # Not yet assigned during iteration
        'para_id': '12345ABC',  # Available from Word's w14:paraId
        'type': 'list_item',
        'text': 'Services',
        'list': {
            'num_id': 27,
            'abstract_num_id': 50,
            'level': 0,
            'counters': [1],
            'ordinal': 'Schedule 1',
            'format': 'decimal',
            'pattern': 'Schedule %1',
            'is_legal': False,
            'restart_boundary': True,
        }
    }
    
    # Create and send numbering event (should use para_id)
    event = NumberingEvent(
        block_id=block.get('para_id') or str(block.get('id')),  # Mimics pipeline behavior
        payload=block['list'],
        restart_group_id='some-group-id',
    )
    tracker.on_numbering_event(event)
    
    # Verify event was stored with para_id, not 'None'
    assert '12345ABC' in tracker._pending_numbering_events
    assert 'None' not in tracker._pending_numbering_events
    
    # Apply should find the event using para_id
    section_id, changed = tracker.apply(block, 'initial-section')
    
    assert changed is True, "Attachment should be detected"
    assert block.get('attachment_id') is not None, "attachment_id should be set"
    assert block.get('attachment', {}).get('type') == 'schedule'


def test_annex_pattern_detected() -> None:
    """Verify that Annex patterns are also detected via ordinal."""
    from effilocal.doc.trackers.attachment_tracker import AttachmentTracker
    from effilocal.doc.numbering_context import NumberingEvent
    
    tracker = AttachmentTracker()
    
    block = {
        'id': None,
        'para_id': 'ANNEX001',
        'type': 'list_item',
        'text': 'Technical Specifications',
        'list': {
            'level': 0,
            'ordinal': 'Annex A',
            'pattern': 'Annex %1',
        }
    }
    
    event = NumberingEvent(
        block_id='ANNEX001',
        payload=block['list'],
        restart_group_id='annex-group',
    )
    tracker.on_numbering_event(event)
    
    section_id, changed = tracker.apply(block, 'section')
    
    assert changed is True
    assert block.get('attachment', {}).get('type') == 'annex'


def test_exhibit_pattern_detected() -> None:
    """Verify that Exhibit patterns are also detected via ordinal."""
    from effilocal.doc.trackers.attachment_tracker import AttachmentTracker
    from effilocal.doc.numbering_context import NumberingEvent
    
    tracker = AttachmentTracker()
    
    block = {
        'id': None,
        'para_id': 'EXHIBIT01',
        'type': 'list_item',
        'text': 'Form of Agreement',
        'list': {
            'level': 0,
            'ordinal': 'Exhibit 1',
            'pattern': 'Exhibit %1',
        }
    }
    
    event = NumberingEvent(
        block_id='EXHIBIT01',
        payload=block['list'],
        restart_group_id='exhibit-group',
    )
    tracker.on_numbering_event(event)
    
    section_id, changed = tracker.apply(block, 'section')
    
    assert changed is True
    assert block.get('attachment', {}).get('type') == 'exhibit'


def test_multiple_schedules_each_get_unique_attachment_id() -> None:
    """Verify that each Schedule block gets a unique attachment_id.
    
    Regression test: Previously all schedules might get the same ID
    or no ID at all due to group_id processing logic.
    """
    from effilocal.doc.trackers.attachment_tracker import AttachmentTracker
    from effilocal.doc.numbering_context import NumberingEvent
    
    tracker = AttachmentTracker()
    attachment_ids = []
    
    # Simulate 6 Schedule blocks
    for i in range(1, 7):
        block = {
            'id': None,
            'para_id': f'SCHED{i:03d}',
            'type': 'list_item',
            'text': f'Schedule {i} Content',
            'list': {
                'level': 0,
                'ordinal': f'Schedule {i}',
                'pattern': 'Schedule %1',
                'restart_boundary': i == 1,  # Only first has restart
            }
        }
        
        event = NumberingEvent(
            block_id=f'SCHED{i:03d}',
            payload=block['list'],
            restart_group_id='schedule-group',
        )
        tracker.on_numbering_event(event)
        
        section_id, changed = tracker.apply(block, 'section')
        
        assert changed is True, f"Schedule {i} should be detected as attachment"
        assert block.get('attachment_id') is not None, f"Schedule {i} should have attachment_id"
        attachment_ids.append(block['attachment_id'])
    
    # All attachment IDs should be unique
    assert len(set(attachment_ids)) == 6, "Each Schedule should have a unique attachment_id"


def test_event_not_found_when_using_wrong_key() -> None:
    """Verify that events stored with para_id cannot be found with block['id'].
    
    This documents the bug we fixed: if events were stored with 'None' as key,
    they wouldn't be found when looking up by para_id.
    """
    from effilocal.doc.trackers.attachment_tracker import AttachmentTracker
    from effilocal.doc.numbering_context import NumberingEvent
    
    tracker = AttachmentTracker()
    
    # Store event with the WRONG key (simulating the old buggy behavior)
    bad_event = NumberingEvent(
        block_id='None',  # Wrong! This is what the old code did
        payload={'level': 0, 'pattern': 'Schedule %1', 'ordinal': 'Schedule 1'},
        restart_group_id='group',
    )
    tracker.on_numbering_event(bad_event)
    
    # Block has para_id but lookup will fail because event is keyed by 'None'
    block = {
        'id': None,
        'para_id': 'REAL_PARA_ID',
        'type': 'list_item',
        'text': 'Content',
        'list': {'level': 0, 'pattern': 'Schedule %1'},
    }
    
    # The event won't be found (simulates the old bug)
    assert tracker._pending_numbering_events.get('REAL_PARA_ID') is None
    assert tracker._pending_numbering_events.get('None') is not None
