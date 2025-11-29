"""Tools for inserting paragraphs after attachments (Schedules, Annexes, Exhibits, etc.)."""
import os
import uuid
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from word_document_server.utils.file_utils import (
    ensure_docx_extension,
    check_file_writeable,
)


def add_custom_para_id(paragraph_element):
    """
    Add a custom paragraph ID using w:tag element for tracking paragraphs we create.
    
    Word doesn't assign w14:paraId until the document is opened and saved in Word.
    We use w:tag (a custom XML tag) to store a UUID that we can use to find paragraphs
    we've created before Word assigns them a paraId.
    
    Args:
        paragraph_element: The w:p element to add the ID to
        
    Returns:
        The UUID string that was added
    """
    # Generate a UUID
    para_uuid = str(uuid.uuid4())
    
    # Get or create paragraph properties
    pPr = paragraph_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        # Insert pPr as first child
        paragraph_element.insert(0, pPr)
    
    # Add w:tag element with our UUID
    tag = OxmlElement('w:tag')
    tag.set(qn('w:val'), f'effi-para-id:{para_uuid}')
    pPr.append(tag)
    
    return para_uuid


def find_paragraph_by_id(doc, para_id=None, custom_id=None):
    """
    Find a paragraph by either w14:paraId or our custom effi-para-id.
    
    Args:
        doc: The Document object
        para_id: The w14:paraId to search for (optional)
        custom_id: The custom effi-para-id UUID to search for (optional)
        
    Returns:
        The paragraph object if found, None otherwise
    """
    for para in doc.paragraphs:
        # Check w14:paraId first (if Word has assigned one)
        if para_id:
            doc_para_id = para._element.get(qn('w14:paraId'))
            if doc_para_id == para_id:
                return para
        
        # Check our custom tag
        if custom_id:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                tag = pPr.find(qn('w:tag'))
                if tag is not None:
                    tag_val = tag.get(qn('w:val'))
                    if tag_val == f'effi-para-id:{custom_id}':
                        return para
    
    return None


async def add_paragraph_after_attachment(
    filename: str,
    attachment_identifier: str,
    text: str,
    style: Optional[str] = None,
    inherit_numbering: bool = True,
) -> str:
    """
    Add a paragraph after a specific attachment (Schedule, Annex, Exhibit, etc.).
    
    This tool locates attachments by their identifier (e.g., "Schedule 1", "Annex A", 
    "Exhibit B") and inserts content after the entire attachment section.
    
    Args:
        filename: Path to the Word document
        attachment_identifier: Attachment to find (e.g., "Schedule 1", "Annex A", "Exhibit B")
        text: Text content to add
        style: Optional style name to apply to the new paragraph
        inherit_numbering: If True, inherit numbering from the attachment header
    
    Returns:
        Success message or error description
        
    Examples:
        - "Schedule 1" -> finds Schedule 1 and inserts after its content
        - "Annex A" -> finds Annex A and inserts after its content
        - "Exhibit B" -> finds Exhibit B and inserts after its content
    """
    try:
        from effilocal.doc import direct_docx
    except ImportError:
        return (
            "effilocal package not available. This tool requires effilocal for attachment detection.\n"
            "The effilocal package should be in the same workspace."
        )
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    try:
        # Use effilocal to analyze the document and find attachments
        docx_path = Path(filename)
        blocks = list(direct_docx.iter_blocks(docx_path))
        
        # Find attachment blocks
        attachment_blocks = []
        for block in blocks:
            attachment_meta = block.get("attachment")
            if isinstance(attachment_meta, dict):
                block_type = attachment_meta.get("type", "")
                ordinal = attachment_meta.get("ordinal", "")
                title = attachment_meta.get("title") or ""
                
                # Construct the attachment identifier
                full_identifier = f"{block_type.title()} {ordinal}".strip()
                
                # Check if this matches what the user is looking for
                if (attachment_identifier.lower() == full_identifier.lower() or
                    attachment_identifier.lower() in full_identifier.lower()):
                    attachment_blocks.append({
                        "block": block,
                        "meta": attachment_meta,
                        "identifier": full_identifier
                    })
        
        if not attachment_blocks:
            return f"Attachment '{attachment_identifier}' not found in {filename}"
        
        # Use the first matching attachment
        target_attachment = attachment_blocks[0]
        target_attachment_id = target_attachment["meta"]["attachment_id"]
        
        # Find the last block belonging to this attachment
        last_block_para_id = None
        
        for i, block in enumerate(blocks):
            if block.get("attachment_id") == target_attachment_id:
                last_block_para_id = block.get("para_id")
        
        if last_block_para_id is None:
            return f"Could not find content for attachment '{attachment_identifier}'"
        
        # Load the document
        doc = Document(filename)
        from docx.oxml.ns import qn
        
        # Find the paragraph with matching para_id (w14:paraId)
        # This is more reliable than matching by text or para_idx
        last_paragraph = None
        
        for para in doc.paragraphs:
            para_id = para._element.get(qn('w14:paraId'))
            if para_id == last_block_para_id:
                last_paragraph = para
                break
        
        if last_paragraph is None:
            return f"Could not find paragraph with ID {last_block_para_id} for attachment '{attachment_identifier}'"
        
        # Insert new paragraph after the attachment
        target_p = last_paragraph._p
        parent = target_p.getparent()
        target_position = list(parent).index(target_p)
        
        # Create new paragraph element
        new_p = OxmlElement('w:p')
        
        # Add paragraph properties
        pPr = OxmlElement('w:pPr')
        new_p.append(pPr)
        
        # Apply style if specified
        if style:
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style)
            pPr.append(pStyle)
        elif last_paragraph.style:
            # Inherit style from last paragraph of attachment
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), last_paragraph.style.style_id)
            pPr.append(pStyle)
        
        # Optionally inherit numbering
        if inherit_numbering:
            # Check if the last paragraph has numbering
            last_block = blocks[last_block_idx]
            list_meta = last_block.get("list")
            if isinstance(list_meta, dict):
                num_id = list_meta.get("num_id")
                level = list_meta.get("level", 0)
                
                if num_id is not None:
                    numPr = OxmlElement('w:numPr')
                    
                    # Add ilvl (level)
                    ilvl_elem = OxmlElement('w:ilvl')
                    ilvl_elem.set(qn('w:val'), str(level))
                    numPr.append(ilvl_elem)
                    
                    # Add numId
                    numId_elem = OxmlElement('w:numId')
                    numId_elem.set(qn('w:val'), str(num_id))
                    numPr.append(numId_elem)
                    
                    pPr.append(numPr)
        
        # Add text run
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        new_p.append(r)
        
        # Add custom para ID for tracking
        custom_id = add_custom_para_id(new_p)
        
        # Insert after last paragraph of attachment
        parent.insert(target_position + 1, new_p)
        
        doc.save(filename)
        
        return f"Paragraph added after {target_attachment['identifier']} in {filename} (para_id={custom_id})"
        
    except Exception as exc:
        return f"Failed to add paragraph after attachment: {str(exc)}"


async def add_paragraphs_after_attachment(
    filename: str,
    attachment_identifier: str,
    paragraphs: List[str],
    style: Optional[str] = None,
    inherit_numbering: bool = True,
) -> str:
    """
    Add multiple paragraphs after a specific attachment.
    
    This tool adds multiple paragraphs sequentially after an attachment section.
    
    Args:
        filename: Path to the Word document
        attachment_identifier: Attachment to find (e.g., "Schedule 1", "Annex A")
        paragraphs: List of paragraph texts to add
        style: Optional style name to apply to all new paragraphs
        inherit_numbering: If True, inherit numbering from the attachment
    
    Returns:
        Success message with count of paragraphs added
    """
    if not paragraphs or not isinstance(paragraphs, list):
        return "paragraphs parameter must be a non-empty list of strings"
    
    filename = ensure_docx_extension(filename)
    
    # Add first paragraph
    result = await add_paragraph_after_attachment(
        filename, attachment_identifier, paragraphs[0], style, inherit_numbering
    )
    
    if "Failed" in result or "not found" in result or "does not exist" in result:
        return result
    
    # For additional paragraphs, add them sequentially
    # Note: We could enhance this to re-analyze and find the new insertion point,
    # similar to add_paragraphs_after_clause, but for now keep it simple
    if len(paragraphs) > 1:
        try:
            doc = Document(filename)
            # Find the last paragraph we just added
            last_added_para = None
            for para in doc.paragraphs:
                if paragraphs[0] in para.text:
                    last_added_para = para
            
            if last_added_para:
                target_p = last_added_para._p
                parent = target_p.getparent()
                base_position = list(parent).index(target_p)
                
                for i, text in enumerate(paragraphs[1:], 1):
                    new_p = OxmlElement('w:p')
                    pPr = OxmlElement('w:pPr')
                    new_p.append(pPr)
                    
                    if style:
                        pStyle = OxmlElement('w:pStyle')
                        pStyle.set(qn('w:val'), style)
                        pPr.append(pStyle)
                    elif last_added_para.style:
                        pStyle = OxmlElement('w:pStyle')
                        pStyle.set(qn('w:val'), last_added_para.style.style_id)
                        pPr.append(pStyle)
                    
                    # Add text
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = text
                    r.append(t)
                    new_p.append(r)
                    
                    # Insert after previous
                    parent.insert(base_position + i, new_p)
                
                doc.save(filename)
            
            return f"Added {len(paragraphs)} paragraph(s) after {attachment_identifier} in {filename}"
            
        except Exception as exc:
            return f"Added 1 paragraph successfully, but failed to add remaining paragraphs: {str(exc)}"
    
    return result


async def add_new_attachment_after(
    filename: str,
    after_attachment: str,
    new_attachment_text: str,
    content: Optional[str] = None,
) -> str:
    """
    Add a new attachment (Schedule, Annex, Exhibit) after an existing attachment.
    
    This creates a new attachment header with the same style and formatting as the 
    reference attachment, automatically incrementing the ordinal (e.g., after "Schedule 4",
    creates "Schedule 5").
    
    Args:
        filename: Path to the Word document
        after_attachment: Existing attachment to insert after (e.g., "Schedule 4")
        new_attachment_text: Text for the new attachment header (e.g., "Schedule 5 - New Title")
        content: Optional content paragraph to add below the attachment header
    
    Returns:
        Success message or error description
        
    Examples:
        - add_new_attachment_after("doc.docx", "Schedule 4", "Schedule 5 - New Services")
        - add_new_attachment_after("doc.docx", "Annex A", "Annex B - Additional Terms")
    """
    try:
        from effilocal.doc import direct_docx
    except ImportError:
        return (
            "effilocal package not available. This tool requires effilocal for attachment detection.\n"
            "The effilocal package should be in the same workspace."
        )
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    try:
        # Use effilocal to analyze the document and find attachments
        docx_path = Path(filename)
        blocks = list(direct_docx.iter_blocks(docx_path))
        
        # Find the reference attachment
        reference_attachment = None
        reference_block = None
        
        for block in blocks:
            attachment_meta = block.get("attachment")
            if isinstance(attachment_meta, dict):
                block_type = attachment_meta.get("type", "")
                ordinal = attachment_meta.get("ordinal", "")
                
                full_identifier = f"{block_type.title()} {ordinal}".strip()
                
                if (after_attachment.lower() == full_identifier.lower() or
                    after_attachment.lower() in full_identifier.lower()):
                    reference_attachment = attachment_meta
                    reference_block = block
                    break
        
        if not reference_attachment:
            return f"Attachment '{after_attachment}' not found in {filename}"
        
        # Find the last block of the reference attachment
        reference_att_id = reference_attachment["attachment_id"]
        last_block_para_id = None
        
        for block in blocks:
            if block.get("attachment_id") == reference_att_id:
                last_block_para_id = block.get("para_id")
        
        if last_block_para_id is None:
            return f"Could not find content for attachment '{after_attachment}'"
        
        # Load the document
        doc = Document(filename)
        from docx.oxml.ns import qn
        
        # Find the reference attachment header paragraph (the one with the "attachment" metadata)
        reference_header_para = None
        reference_header_para_id = reference_block.get("para_id")
        
        for para in doc.paragraphs:
            para_id = para._element.get(qn('w14:paraId'))
            if para_id == reference_header_para_id:
                reference_header_para = para
                break
        
        if not reference_header_para:
            return f"Could not find header paragraph for '{after_attachment}'"
        
        # Find the last paragraph of the reference attachment
        last_paragraph = None
        for para in doc.paragraphs:
            para_id = para._element.get(qn('w14:paraId'))
            if para_id == last_block_para_id:
                last_paragraph = para
                break
        
        if not last_paragraph:
            return f"Could not find last paragraph for '{after_attachment}'"
        
        # Get the style from the reference header
        header_style = reference_header_para.style
        header_style_id = header_style.style_id if header_style else None
        
        # Insert new attachment header
        target_p = last_paragraph._p
        parent = target_p.getparent()
        target_position = list(parent).index(target_p)
        
        # Create new attachment header paragraph
        new_header_p = OxmlElement('w:p')
        
        # Add paragraph properties
        pPr = OxmlElement('w:pPr')
        new_header_p.append(pPr)
        
        # Apply the same style as reference attachment header
        if header_style_id:
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), header_style_id)
            pPr.append(pStyle)
        
        # Add text for new attachment
        r = OxmlElement('w:r')
        
        # Copy run properties from reference if available
        if reference_header_para.runs:
            ref_run = reference_header_para.runs[0]
            if ref_run._element.rPr is not None:
                rPr = OxmlElement('w:rPr')
                for child in ref_run._element.rPr:
                    rPr.append(child)
                r.append(rPr)
        
        t = OxmlElement('w:t')
        t.text = new_attachment_text
        r.append(t)
        new_header_p.append(r)
        
        # Add custom para ID for tracking
        header_custom_id = add_custom_para_id(new_header_p)
        
        # Insert the new header
        parent.insert(target_position + 1, new_header_p)
        
        # Optionally add content paragraph
        content_custom_id = None
        if content:
            new_content_p = OxmlElement('w:p')
            
            # Add paragraph properties (use Normal style or inherit from last content paragraph)
            content_pPr = OxmlElement('w:pPr')
            new_content_p.append(content_pPr)
            
            # Try to match the style of the reference attachment's content
            if last_paragraph.style and last_paragraph != reference_header_para:
                content_style = OxmlElement('w:pStyle')
                content_style.set(qn('w:val'), last_paragraph.style.style_id)
                content_pPr.append(content_style)
            
            # Add content text
            content_r = OxmlElement('w:r')
            content_t = OxmlElement('w:t')
            content_t.text = content
            content_r.append(content_t)
            new_content_p.append(content_r)
            
            # Add custom para ID for content tracking
            content_custom_id = add_custom_para_id(new_content_p)
            
            # Insert content after header
            parent.insert(target_position + 2, new_content_p)
        
        doc.save(filename)
        
        content_msg = f" with content" if content else ""
        ids_msg = f" (header_id={header_custom_id}"
        if content_custom_id:
            ids_msg += f", content_id={content_custom_id}"
        ids_msg += ")"
        
        return f"New attachment '{new_attachment_text}' added after {after_attachment}{content_msg} in {filename}{ids_msg}"
        
    except Exception as exc:
        return f"Failed to add new attachment: {str(exc)}"
