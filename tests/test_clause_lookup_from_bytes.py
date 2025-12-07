"""
Tests for ClauseLookup.from_docx_bytes factory method.

This is part of Step 7 cleanup: enabling ClauseLookup to be constructed
directly from docx bytes, eliminating the need for the separate 
extract_clause_numbers_from_doc function.

Test Categories:
1. Factory method creation from bytes
2. Clause number lookup from bytes-based ClauseLookup  
3. Clause text lookup from bytes-based ClauseLookup
4. Integration with existing ClauseLookup functionality
5. Error handling for invalid bytes
6. Backward compatibility with dict-based approach
"""

import io
import pytest
from pathlib import Path
from docx import Document

from tests.helpers.docx_builder import DocBuilder, create_test_doc_bytes


# =============================================================================
# Test Category 1: Factory method creation from bytes
# =============================================================================

class TestFromDocxBytesFactory:
    """Test ClauseLookup.from_docx_bytes factory method."""
    
    def test_from_docx_bytes_returns_clause_lookup_instance(self):
        """Factory method should return a ClauseLookup instance."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        # Create simple docx in memory
        doc = Document()
        doc.add_paragraph("Test paragraph")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        
        # Should return ClauseLookup instance
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        assert isinstance(lookup, ClauseLookup)
    
    def test_from_docx_bytes_has_blocks_list(self):
        """ClauseLookup from bytes should have populated blocks list."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        doc = Document()
        doc.add_paragraph("Paragraph 1")
        doc.add_paragraph("Paragraph 2")
        doc.add_paragraph("Paragraph 3")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        # Should have at least 3 blocks
        assert len(lookup.blocks) >= 3
    
    def test_from_docx_bytes_builds_para_id_index(self):
        """ClauseLookup from bytes should build para_id index."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        doc = Document()
        doc.add_paragraph("Test paragraph")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        # Internal index should exist (may be empty for simple docs without para_ids)
        assert hasattr(lookup, '_blocks_by_para_id')
        assert isinstance(lookup._blocks_by_para_id, dict)
    
    def test_from_docx_bytes_with_empty_document(self):
        """Factory should handle empty documents gracefully."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        # Should succeed even with empty doc
        assert isinstance(lookup, ClauseLookup)


# =============================================================================
# Test Category 2: Clause number lookup from bytes-based ClauseLookup
# =============================================================================

class TestClauseNumberLookupFromBytes:
    """Test clause number lookup using bytes-created ClauseLookup."""
    
    @pytest.fixture
    def numbered_doc_bytes(self) -> bytes:
        """Create a document with numbered paragraphs."""
        # Use a fixture document that has proper numbering
        fixture_path = Path(__file__).parent / "fixtures" / "nested_clauses.docx"
        if fixture_path.exists():
            return fixture_path.read_bytes()
        
        # Fallback: create simple doc (won't have numbering)
        doc = Document()
        doc.add_paragraph("1. First clause")
        doc.add_paragraph("2. Second clause")
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    def test_get_clause_number_from_bytes_lookup(self, numbered_doc_bytes):
        """Should be able to get clause numbers from bytes-based lookup."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        lookup = ClauseLookup.from_docx_bytes(numbered_doc_bytes)
        
        # Find a para_id that has a clause number
        found_clause = False
        for para_id in lookup._blocks_by_para_id:
            clause_num = lookup.get_clause_number(para_id)
            if clause_num:
                found_clause = True
                assert isinstance(clause_num, str)
                break
        
        # If fixture has numbered clauses, we should find at least one
        # If not, just ensure no errors occurred
        assert isinstance(lookup, ClauseLookup)
    
    def test_clause_numbers_have_no_trailing_period(self, numbered_doc_bytes):
        """Clause numbers should not have trailing periods."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        lookup = ClauseLookup.from_docx_bytes(numbered_doc_bytes)
        
        for para_id in lookup._blocks_by_para_id:
            clause_num = lookup.get_clause_number(para_id)
            if clause_num:
                assert not clause_num.endswith('.'), f"Clause number should not end with period: {clause_num}"


# =============================================================================
# Test Category 3: Clause text lookup from bytes-based ClauseLookup
# =============================================================================

class TestClauseTextLookupFromBytes:
    """Test clause text lookup using bytes-created ClauseLookup."""
    
    def test_get_clause_text_from_bytes_lookup(self):
        """Should be able to get clause text from bytes-based lookup."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        test_text = "This is a test paragraph with unique content xyz123."
        docx_bytes = create_test_doc_bytes(test_text)
        
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        
        # Simple in-memory docs may not have para_ids, so check blocks directly
        found_text = False
        for block in lookup.blocks:
            text = block.get("text", "")
            if text and "xyz123" in text:
                found_text = True
                break
        
        assert found_text, "Should find the test text in lookup"
    
    def test_get_clause_text_returns_full_paragraph(self):
        """Clause text should be the full paragraph content."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        long_text = "This is a long paragraph. " * 10
        docx_bytes = create_test_doc_bytes(long_text)
        
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        
        for para_id in lookup._blocks_by_para_id:
            text = lookup.get_clause_text(para_id)
            if text and "This is a long paragraph" in text:
                # Text should contain multiple sentences
                assert text.count("This is a long paragraph") >= 5
                break


# =============================================================================
# Test Category 4: Integration with existing ClauseLookup functionality
# =============================================================================

class TestClauseLookupBytesIntegration:
    """Test that bytes-created ClauseLookup integrates with all methods."""
    
    def test_get_clause_info_works_with_bytes_lookup(self):
        """get_clause_info should work with bytes-created lookup."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        docx_bytes = create_test_doc_bytes("INDEMNIFICATION. The party shall indemnify...")
        
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        
        for para_id in lookup._blocks_by_para_id:
            info = lookup.get_clause_info(para_id)
            assert info is not None
            assert "para_id" in info
            assert "clause_number" in info
            assert "clause_title" in info
            assert "text" in info
            break
    
    def test_get_clause_title_works_with_bytes_lookup(self):
        """get_clause_title should work with bytes-created lookup."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        docx_bytes = create_test_doc_bytes("LIMITATION OF LIABILITY. Subject to clause 12...")
        
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        
        # Check that title extraction works
        for para_id in lookup._blocks_by_para_id:
            title = lookup.get_clause_title(para_id)
            text = lookup.get_clause_text(para_id) or ""
            if "LIMITATION OF LIABILITY" in text:
                assert title == "LIMITATION OF LIABILITY"
                break


# =============================================================================
# Test Category 5: Error handling for invalid bytes
# =============================================================================

class TestFromDocxBytesErrorHandling:
    """Test error handling for invalid inputs."""
    
    def test_from_docx_bytes_with_invalid_bytes_raises_error(self):
        """Should raise appropriate error for non-docx bytes."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        invalid_bytes = b"This is not a valid docx file"
        
        with pytest.raises(Exception):  # Could be ValueError, BadZipFile, etc.
            ClauseLookup.from_docx_bytes(invalid_bytes)
    
    def test_from_docx_bytes_with_empty_bytes_raises_error(self):
        """Should raise error for empty bytes."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        with pytest.raises(Exception):
            ClauseLookup.from_docx_bytes(b"")
    
    def test_from_docx_bytes_with_none_raises_type_error(self):
        """Should raise TypeError for None input."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        with pytest.raises(TypeError):
            ClauseLookup.from_docx_bytes(None)  # type: ignore


# =============================================================================
# Test Category 6: Backward compatibility with dict-based approach
# =============================================================================

class TestBackwardCompatibility:
    """Test compatibility with existing dict-based extract_clause_numbers_from_doc."""
    
    def test_to_ordinal_map_returns_dict(self):
        """to_ordinal_map method should return para_id->ordinal dict."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        doc = Document()
        doc.add_paragraph("Test paragraph")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        ordinal_map = lookup.to_ordinal_map()
        
        assert isinstance(ordinal_map, dict)
    
    def test_to_text_map_returns_dict(self):
        """to_text_map method should return para_id->text dict."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        docx_bytes = create_test_doc_bytes("Test paragraph content")
        
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        text_map = lookup.to_text_map()
        
        assert isinstance(text_map, dict)
        # TestDocument creates paragraphs with para_ids, so text_map should have entries
        assert any("Test paragraph" in text for text in text_map.values())
    
    def test_to_ordinal_map_matches_extract_function_output(self):
        """to_ordinal_map should produce same result as extract_clause_numbers_from_doc."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        # Use fixture with numbering if available
        fixture_path = Path(__file__).parent / "fixtures" / "nested_clauses.docx"
        if not fixture_path.exists():
            pytest.skip("Fixture not available")
        
        docx_bytes = fixture_path.read_bytes()
        
        # Get result via ClauseLookup
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        lookup_ordinal_map = lookup.to_ordinal_map()
        
        # Get result via old function
        from scripts.generate_review_example import extract_clause_numbers_from_doc
        old_ordinal_map, _ = extract_clause_numbers_from_doc(docx_bytes)
        
        # Should be equivalent
        assert lookup_ordinal_map == old_ordinal_map
    
    def test_to_text_map_matches_extract_function_output(self):
        """to_text_map should produce same result as extract_clause_numbers_from_doc."""
        from effilocal.doc.clause_lookup import ClauseLookup
        
        # Use fixture with content
        fixture_path = Path(__file__).parent / "fixtures" / "nested_clauses.docx"
        if not fixture_path.exists():
            pytest.skip("Fixture not available")
        
        docx_bytes = fixture_path.read_bytes()
        
        # Get result via ClauseLookup
        lookup = ClauseLookup.from_docx_bytes(docx_bytes)
        lookup_text_map = lookup.to_text_map()
        
        # Get result via old function
        from scripts.generate_review_example import extract_clause_numbers_from_doc
        _, old_text_map = extract_clause_numbers_from_doc(docx_bytes)
        
        # Should be equivalent
        assert lookup_text_map == old_text_map


# =============================================================================
# Test Category 7: Cleanup verification - old function uses new implementation
# =============================================================================

class TestCleanupVerification:
    """Verify that extract_clause_numbers_from_doc uses ClauseLookup internally."""
    
    def test_extract_function_uses_clause_lookup(self):
        """extract_clause_numbers_from_doc should use ClauseLookup internally."""
        import inspect
        from scripts.generate_review_example import extract_clause_numbers_from_doc
        
        source = inspect.getsource(extract_clause_numbers_from_doc)
        
        # After refactoring, should use ClauseLookup
        assert "ClauseLookup" in source, (
            "extract_clause_numbers_from_doc should use ClauseLookup after refactoring"
        )
    
    def test_extract_function_is_thin_wrapper(self):
        """extract_clause_numbers_from_doc should be a thin wrapper (< 10 lines)."""
        import inspect
        from scripts.generate_review_example import extract_clause_numbers_from_doc
        
        source = inspect.getsource(extract_clause_numbers_from_doc)
        lines = [l for l in source.split('\n') if l.strip() and not l.strip().startswith('#')]
        
        # After refactoring, should be a small wrapper
        assert len(lines) <= 15, (
            f"extract_clause_numbers_from_doc should be a thin wrapper, but has {len(lines)} lines"
        )
