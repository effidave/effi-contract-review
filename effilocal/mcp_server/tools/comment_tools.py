"""
Comment tools for effilocal MCP server.

This module extends word_document_server.tools.comment_tools with:
1. Add comment functionality (add_comment_after_text, add_comment_for_paragraph)
2. Comment status support (active/resolved from commentsExtended.xml)

Override Pattern:
- Import all upstream functions
- Add new comment creation functions
- Use enhanced core.comments with status support
"""

import os
import json
from typing import Dict, List, Optional, Any
from docx import Document
from docx.text.run import Run
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.comments import CommentsPart

from word_document_server.utils.file_utils import ensure_docx_extension
from effilocal.mcp_server.utils.document_utils import iter_document_paragraphs

# Import from effilocal core.comments (which has status support)
from effilocal.mcp_server.core.comments import (
    extract_all_comments,
    filter_comments_by_author,
    get_comments_for_paragraph,
    extract_comment_status_map,
    merge_comment_status,
    resolve_comment,
    unresolve_comment,
)

# Note: We don't import tool wrappers from upstream since we define our own
# The tool wrappers delegate to core functions


# ============================================================================
# Tool Wrapper Functions (delegate to core)
# ============================================================================

async def get_all_comments(filename: str) -> str:
    """Extract all comments from a Word document including status information."""
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({'success': False, 'error': f'Document {filename} does not exist'}, indent=2)
    
    try:
        doc = Document(filename)
        comments = extract_all_comments(doc)
        return json.dumps({'success': True, 'comments': comments, 'total_comments': len(comments)}, indent=2)
    except Exception as e:
        return json.dumps({'success': False, 'error': f'Failed to extract comments: {str(e)}'}, indent=2)


async def get_comments_by_author(filename: str, author: str) -> str:
    """Extract comments from a specific author."""
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({'success': False, 'error': f'Document {filename} does not exist'}, indent=2)
    
    if not author or not author.strip():
        return json.dumps({'success': False, 'error': 'Author name cannot be empty'}, indent=2)
    
    try:
        doc = Document(filename)
        all_comments = extract_all_comments(doc)
        author_comments = filter_comments_by_author(all_comments, author)
        return json.dumps({'success': True, 'author': author, 'comments': author_comments, 'total_comments': len(author_comments)}, indent=2)
    except Exception as e:
        return json.dumps({'success': False, 'error': f'Failed to extract comments: {str(e)}'}, indent=2)


async def get_paragraph_comments(filename: str, paragraph_index: int) -> str:
    """Extract comments for a specific paragraph."""
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({'success': False, 'error': f'Document {filename} does not exist'}, indent=2)
    
    if paragraph_index < 0:
        return json.dumps({'success': False, 'error': 'Paragraph index must be non-negative'}, indent=2)
    
    try:
        doc = Document(filename)
        if paragraph_index >= len(doc.paragraphs):
            return json.dumps({'success': False, 'error': f'Paragraph index {paragraph_index} is out of range. Document has {len(doc.paragraphs)} paragraphs.'}, indent=2)
        
        all_comments = extract_all_comments(doc)
        para_comments = get_comments_for_paragraph(all_comments, paragraph_index)
        paragraph_text = doc.paragraphs[paragraph_index].text
        
        return json.dumps({'success': True, 'paragraph_index': paragraph_index, 'paragraph_text': paragraph_text, 'comments': para_comments, 'total_comments': len(para_comments)}, indent=2)
    except Exception as e:
        return json.dumps({'success': False, 'error': f'Failed to extract comments: {str(e)}'}, indent=2)


# ============================================================================
# Helper Functions (for comment creation)
# ============================================================================

def _get_or_add_comments_part(doc_part):
    """Return the comments part; create it if it doesn't exist.
    
    NEW helper - follows WordprocessingML spec to create /word/comments.xml.
    """
    try:
        return doc_part.part_related_by(RT.COMMENTS)
    except KeyError:
        # Let python-docx provide a correctly wired default part
        package = doc_part.package
        assert package is not None
        comments_part = CommentsPart.default(package)
        doc_part.relate_to(comments_part, RT.COMMENTS)
        return comments_part


def _wrap_run_with_comment(r: Run, cid: int):
    """Insert comment range start/end markers around the given Run and add a commentReference run.
    
    NEW helper - implements Word comment anchoring at run level.
    """
    r_el = r._r

    # commentRangeStart before the run
    start = OxmlElement('w:commentRangeStart')
    start.set(qn('w:id'), str(cid))
    r_el.addprevious(start)

    # commentRangeEnd after the run
    end = OxmlElement('w:commentRangeEnd')
    end.set(qn('w:id'), str(cid))
    r_el.addnext(end)

    # add a reference run right after the end
    ref_run = OxmlElement('w:r')
    comment_ref = OxmlElement('w:commentReference')
    comment_ref.set(qn('w:id'), str(cid))
    ref_run.append(comment_ref)
    end.addnext(ref_run)


# ============================================================================
# NEW Functions (comment creation)
# ============================================================================

async def add_comment_after_text(
    filename: str,
    search_text: str,
    comment_text: str,
    author: Optional[str] = None,
    initials: Optional[str] = None,
) -> str:
    """Add a Word comment to the first occurrence of `search_text`.
    
    NEW function - creates comment anchored to specific text.
    
    Args:
        filename: Path to the Word document
        search_text: Text to find and anchor comment to
        comment_text: Comment content
        author: Comment author name
        initials: Comment author initials
    
    Returns:
        JSON string with success flag and metadata
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"}, indent=2)

    if not search_text:
        return json.dumps({"success": False, "error": "search_text cannot be empty"}, indent=2)

    try:
        doc = Document(filename)

        target_run = None
        for p in doc.paragraphs:
            for run in p.runs:
                if search_text in run.text:
                    target_run = run
                    break
            if target_run:
                break

        if target_run is None:
            return json.dumps({
                "success": False,
                "error": f"Text '{search_text}' not found in document."
            }, indent=2)

        comments_part = _get_or_add_comments_part(doc.part)
        comment = comments_part.comments.add_comment(
            text=comment_text or "",
            author=author or "",
            initials=initials if initials is not None else "",
        )
        cid = comment.comment_id
        _wrap_run_with_comment(target_run, cid)

        doc.save(filename)
        return json.dumps({
            "success": True,
            "action": "add_comment_after_text",
            "filename": filename,
            "search_text": search_text,
            "comment_id": cid,
            "author": author,
            "initials": initials,
        }, indent=2)

    except PermissionError:
        return json.dumps({
            "success": False, 
            "error": f"Permission denied: The file '{os.path.basename(filename)}' is likely open in Word. Please close it and try again."
        }, indent=2)
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to add comment: {str(e)}"}, indent=2)


async def add_comment_for_paragraph(
    filename: str,
    paragraph_index: int,
    comment_text: str,
    author: Optional[str] = None,
    initials: Optional[str] = None,
) -> str:
    """Add a Word comment anchored to the entire paragraph at `paragraph_index` (0-based).
    
    NEW function - creates comment anchored to specific paragraph.
    
    Args:
        filename: Path to the Word document
        paragraph_index: 0-based paragraph index
        comment_text: Comment content
        author: Comment author name
        initials: Comment author initials
    
    Returns:
        JSON string with success flag and metadata
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"}, indent=2)

    try:
        doc = Document(filename)
        # Use SDT-aware iterator to get all paragraphs including those in content controls
        all_paragraphs = list(iter_document_paragraphs(doc))
        if paragraph_index < 0 or paragraph_index >= len(all_paragraphs):
            return json.dumps({
                "success": False,
                "error": f"Paragraph index {paragraph_index} is out of range (0..{len(all_paragraphs)-1})."
            }, indent=2)

        p = all_paragraphs[paragraph_index]
        # If paragraph has no runs, create an empty run so we have an anchor
        if not p.runs:
            r = p.add_run("")
        else:
            r = p.runs[0]

        comments_part = _get_or_add_comments_part(doc.part)
        comment = comments_part.comments.add_comment(
            text=comment_text or "",
            author=author or "",
            initials=initials if initials is not None else "",
        )
        cid = comment.comment_id
        _wrap_run_with_comment(r, cid)

        doc.save(filename)
        return json.dumps({
            "success": True,
            "action": "add_comment_for_paragraph",
            "filename": filename,
            "paragraph_index": paragraph_index,
            "comment_id": cid,
            "author": author,
            "initials": initials,
        }, indent=2)

    except PermissionError:
        return json.dumps({
            "success": False, 
            "error": f"Permission denied: The file '{os.path.basename(filename)}' is likely open in Word. Please close it and try again."
        }, indent=2)
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to add comment: {str(e)}"}, indent=2)


async def update_comment(
    filename: str,
    comment_id: str,
    new_text: str
) -> str:
    """Update the text of an existing comment.
    
    NEW function - updates comment content.
    
    Args:
        filename: Path to the Word document
        comment_id: The ID of the comment to update (hex string or int string)
        new_text: New text content for the comment
    
    Returns:
        JSON string with success flag
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"}, indent=2)

    try:
        doc = Document(filename)
        comments_part = doc.part.part_related_by(RT.COMMENTS)
        
        # Find comment by ID
        # comment_id might be passed as "1" or "0" or hex
        # The XML expects the ID string exactly as it appears in w:id
        
        # Try exact match first
        comment_els = comments_part.element.xpath(f'.//w:comment[@w:id="{comment_id}"]')
        
        if not comment_els:
            return json.dumps({"success": False, "error": f"Comment with ID {comment_id} not found"}, indent=2)
            
        comment_el = comment_els[0]
        
        # Find w:t elements
        t_els = comment_el.xpath('.//w:t')
        if t_els:
            # Update first text element
            t_els[0].text = new_text
            # Clear others if any (to avoid mixed content)
            for t in t_els[1:]:
                t.text = ''
        else:
            # If no text element exists (empty comment?), we need to add one
            # Structure: w:comment -> w:p -> w:r -> w:t
            # This is a bit complex to construct from scratch if missing, 
            # but usually comments have at least a paragraph.
            p_els = comment_el.xpath('.//w:p')
            if p_els:
                p = p_els[0]
                # We can use python-docx wrappers if we wrap the element, 
                # but here we are doing low-level XML.
                # Let's just try to find a run or create one.
                r_els = p.xpath('.//w:r')
                if r_els:
                    r = r_els[0]
                else:
                    r = OxmlElement('w:r')
                    p.append(r)
                
                t = OxmlElement('w:t')
                t.text = new_text
                r.append(t)
            else:
                # Very empty comment
                return json.dumps({"success": False, "error": "Comment structure too simple to update safely"}, indent=2)
            
        doc.save(filename)
        return json.dumps({
            "success": True, 
            "action": "update_comment",
            "comment_id": comment_id
        }, indent=2)

    except PermissionError:
        return json.dumps({
            "success": False, 
            "error": f"Permission denied: The file '{os.path.basename(filename)}' is likely open in Word. Please close it and try again."
        }, indent=2)
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to update comment: {str(e)}"}, indent=2)


# ============================================================================
# NEW Functions (resolve/unresolve - Sprint 3 Phase 1)
# ============================================================================

async def resolve_comment_tool(
    filename: str,
    comment_id: str
) -> str:
    """Mark a comment as resolved in a Word document.
    
    This updates the commentsExtended.xml part to set the comment's status
    to resolved (done="1"), which is how Word tracks resolved comments.
    
    Args:
        filename: Path to the Word document
        comment_id: The ID of the comment to resolve (the w:id attribute)
    
    Returns:
        String result indicating success or failure
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Error: Document {filename} does not exist"

    try:
        doc = Document(filename)
        success = resolve_comment(doc, comment_id)
        
        if success:
            doc.save(filename)
            return f"Successfully resolved comment {comment_id} in {os.path.basename(filename)}"
        else:
            return f"Error: Could not resolve comment {comment_id}. Comment may not exist or document may not have commentsExtended.xml."

    except PermissionError:
        return f"Error: Permission denied - The file '{os.path.basename(filename)}' is likely open in Word. Please close it and try again."
    except Exception as e:
        return f"Error: Failed to resolve comment: {str(e)}"


async def unresolve_comment_tool(
    filename: str,
    comment_id: str
) -> str:
    """Mark a comment as active (unresolve it) in a Word document.
    
    This updates the commentsExtended.xml part to set the comment's status
    to active (done="0"), reversing a previous resolve operation.
    
    Args:
        filename: Path to the Word document
        comment_id: The ID of the comment to unresolve (the w:id attribute)
    
    Returns:
        String result indicating success or failure
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Error: Document {filename} does not exist"

    try:
        doc = Document(filename)
        success = unresolve_comment(doc, comment_id)
        
        if success:
            doc.save(filename)
            return f"Successfully unresolved comment {comment_id} in {os.path.basename(filename)}"
        else:
            return f"Error: Could not unresolve comment {comment_id}. Comment may not exist or document may not have commentsExtended.xml."

    except PermissionError:
        return f"Error: Permission denied - The file '{os.path.basename(filename)}' is likely open in Word. Please close it and try again."
    except Exception as e:
        return f"Error: Failed to unresolve comment: {str(e)}"


# ============================================================================
# Re-export all functions for compatibility
# ============================================================================

__all__ = [
    # Pass-through functions (unchanged from upstream)
    'get_all_comments',
    'get_comments_by_author',
    'get_paragraph_comments',
    
    # New functions (comment creation)
    'add_comment_after_text',
    'add_comment_for_paragraph',
    'update_comment',
    
    # New functions (resolve/unresolve - Sprint 3 Phase 1)
    'resolve_comment_tool',
    'unresolve_comment_tool',
]
