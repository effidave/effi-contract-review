from __future__ import annotations

import json
from pathlib import Path
from uuid import uuid4

from docx import Document

from effilocal.flows.analyze_doc import analyze


def _build_doc(tmp_path: Path) -> Path:
    doc = Document()

    heading = doc.add_paragraph("1 Definitions")
    heading.style = "Heading 1"

    doc.add_paragraph('"Affiliate" means any entity that controls, is controlled by, or is under common control with the Company.')
    doc.add_paragraph("For the purposes of this definition, control refers to ownership of more than fifty percent of the voting interests.")
    doc.add_paragraph('"Business Day" – any day other than a Saturday, Sunday, or public holiday in London.')

    table_heading = doc.add_paragraph("Definitions Table")
    table_heading.style = "Heading 2"

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Customer Data"
    table.cell(0, 1).text = "All information provided by or on behalf of the Customer to the Supplier."
    table.cell(1, 0).text = "Force Majeure Event"
    table.cell(1, 1).text = "Any event beyond a party's reasonable control which prevents performance."

    next_heading = doc.add_paragraph("2 Scope")
    next_heading.style = "Heading 1"
    doc.add_paragraph("This clause is outside the definitions section and should not be tagged.")

    doc_path = tmp_path / "definitions.docx"
    doc.save(doc_path)
    return doc_path


def test_definition_paragraphs_and_tables_tagged(tmp_path: Path) -> None:
    doc_path = _build_doc(tmp_path)
    out_dir = tmp_path / "analysis"
    analyze(doc_path, doc_id=str(uuid4()), out_dir=out_dir)

    blocks_path = out_dir / "blocks.jsonl"
    blocks = [
        json.loads(line)
        for line in blocks_path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]

    def _find(text: str) -> dict[str, object]:
        for block in blocks:
            if block["text"] == text:
                return block
        raise AssertionError(f"Block with text {text!r} not found")

    affiliate_block = _find(
        '"Affiliate" means any entity that controls, is controlled by, or is under common control with the Company.'
    )
    assert affiliate_block["definition"]["role"] == "term"
    assert affiliate_block["definition"]["term"] == "Affiliate"
    assert affiliate_block["definition"]["inline_definition"].startswith("any entity")

    affiliate_body = _find(
        "For the purposes of this definition, control refers to ownership of more than fifty percent of the voting interests."
    )
    assert affiliate_body["definition"]["role"] == "body"
    assert (
        affiliate_block["definition"]["definition_id"]
        == affiliate_body["definition"]["definition_id"]
    )

    business_day_block = _find(
        '"Business Day" – any day other than a Saturday, Sunday, or public holiday in London.'
    )
    assert business_day_block["definition"]["role"] == "term"
    assert business_day_block["definition"]["term"] == "Business Day"
    assert business_day_block["definition"]["inline_definition"].startswith("any day")

    table_term = _find("Customer Data")
    table_body = _find("All information provided by or on behalf of the Customer to the Supplier.")
    assert table_term["definition"]["role"] == "term"
    assert table_term["definition"]["term"] == "Customer Data"
    assert table_body["definition"]["role"] == "body"
    assert (
        table_term["definition"]["definition_id"]
        == table_body["definition"]["definition_id"]
    )

    outside_definition = _find("This clause is outside the definitions section and should not be tagged.")
    assert outside_definition.get("definition") is None
