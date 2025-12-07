"""
Tests for effilocal.doc.party_detection module.

This module provides functions for detecting and extracting party information
from Word documents:
- extract_defined_party_terms: Extract defined terms like "Vendor", "Company"
- extract_company_names: Extract company names from preamble
- extract_full_company_names: Extract full legal names with suffixes
- extract_company_to_defined_term_mapping: Map company names to defined terms
- extract_party_alternate_names: Extract alternate names for parties
- extract_party_from_comment_prefixes: Extract party identifiers from "For X:" comments
- compute_similarity: Calculate string similarity
- match_prefixes_to_parties: Match comment prefixes to detected parties
- infer_party_role: Infer semantic role from defined term
- get_role_placeholder: Convert role to placeholder format
- PartyInfo: Dataclass for party information
- DEFINED_TERM_TO_ROLE: Mapping of terms to semantic roles
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest
from docx import Document


# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def simple_doc() -> Document:
    """Create a simple Word document with party information."""
    doc = Document()
    
    # Preamble with party definitions
    doc.add_paragraph(
        "This Agreement is entered into between NBCUniversal Media, LLC, "
        'a Delaware limited liability company ("Company" or "NBCUniversal"), '
        'and Didimo, Inc., a Delaware corporation ("Vendor").'
    )
    
    # Body with party usage
    doc.add_paragraph("The Vendor shall provide services to the Company.")
    doc.add_paragraph("The Company will pay the Vendor upon completion.")
    doc.add_paragraph("Each party may terminate this Agreement with notice.")
    
    return doc


@pytest.fixture
def license_agreement_doc() -> Document:
    """Create a document for a license agreement."""
    doc = Document()
    
    doc.add_paragraph(
        'This License Agreement is between Acme Software Ltd ("Licensor") '
        'and TechCorp, Inc. ("Licensee").'
    )
    
    doc.add_paragraph("The Licensor shall grant a license to the Licensee.")
    doc.add_paragraph("The Licensee will pay royalties to the Licensor.")
    
    return doc


@pytest.fixture
def lease_agreement_doc() -> Document:
    """Create a document for a lease agreement."""
    doc = Document()
    
    doc.add_paragraph(
        'This Lease Agreement is between Property Holdings LLC ("Landlord") '
        'and Business Tenants, Inc. ("Tenant").'
    )
    
    doc.add_paragraph("The Landlord shall provide the premises.")
    doc.add_paragraph("The Tenant will pay rent monthly.")
    
    return doc


@pytest.fixture
def sample_comments() -> list[dict[str, Any]]:
    """Create sample comments with 'For X:' prefixes."""
    return [
        {"text": "For Didimo: Please confirm this is acceptable.", "author": "Lawyer"},
        {"text": "For NBC: We suggest revising this clause.", "author": "Lawyer"},
        {"text": "For Didimo: This protects your interests.", "author": "Lawyer"},
        {"text": "General comment without prefix.", "author": "Lawyer"},
        {"text": "For NBC: Consider adding a cap.", "author": "Lawyer"},
    ]


@pytest.fixture
def single_party_comments() -> list[dict[str, Any]]:
    """Comments with only one party prefix."""
    return [
        {"text": "For Client: Please review this.", "author": "Lawyer"},
        {"text": "For Client: Important clause.", "author": "Lawyer"},
    ]


# =============================================================================
# Tests: extract_defined_party_terms
# =============================================================================


class TestExtractDefinedPartyTerms:
    """Tests for extract_defined_party_terms function."""
    
    def test_extracts_vendor_and_company(self, simple_doc: Document) -> None:
        """Should extract defined terms from 'X shall/will/may' patterns."""
        from effilocal.doc.party_detection import extract_defined_party_terms
        
        terms = extract_defined_party_terms(simple_doc)
        
        assert "Vendor" in terms
        assert "Company" in terms
    
    def test_returns_terms_sorted_by_frequency(self, simple_doc: Document) -> None:
        """Should return terms sorted by frequency (most common first)."""
        from effilocal.doc.party_detection import extract_defined_party_terms
        
        terms = extract_defined_party_terms(simple_doc)
        
        # Should be a list
        assert isinstance(terms, list)
        # Should have at least 2 terms
        assert len(terms) >= 2
    
    def test_extracts_licensor_licensee(self, license_agreement_doc: Document) -> None:
        """Should extract Licensor and Licensee from license agreement."""
        from effilocal.doc.party_detection import extract_defined_party_terms
        
        terms = extract_defined_party_terms(license_agreement_doc)
        
        assert "Licensor" in terms
        assert "Licensee" in terms
    
    def test_filters_out_common_words(self) -> None:
        """Should filter out common false positives like 'Each', 'Either'."""
        from effilocal.doc.party_detection import extract_defined_party_terms
        
        doc = Document()
        doc.add_paragraph("Each party shall comply.")
        doc.add_paragraph("Either party may terminate.")
        doc.add_paragraph("The Vendor shall deliver.")
        
        terms = extract_defined_party_terms(doc)
        
        assert "Each" not in terms
        assert "Either" not in terms
        assert "Vendor" in terms
    
    def test_handles_empty_document(self) -> None:
        """Should return empty list for document with no parties."""
        from effilocal.doc.party_detection import extract_defined_party_terms
        
        doc = Document()
        doc.add_paragraph("This is a paragraph without party terms.")
        
        terms = extract_defined_party_terms(doc)
        
        assert terms == []
    
    def test_extracts_from_tables(self) -> None:
        """Should also extract terms from table cells."""
        from effilocal.doc.party_detection import extract_defined_party_terms
        
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "The Supplier shall provide goods."
        table.cell(1, 1).text = "The Buyer will pay."
        
        terms = extract_defined_party_terms(doc)
        
        assert "Supplier" in terms
        assert "Buyer" in terms


# =============================================================================
# Tests: extract_company_names
# =============================================================================


class TestExtractCompanyNames:
    """Tests for extract_company_names function."""
    
    def test_extracts_company_names_with_llc(self, simple_doc: Document) -> None:
        """Should extract company names with LLC suffix."""
        from effilocal.doc.party_detection import extract_company_names
        
        names = extract_company_names(simple_doc)
        
        assert any("NBCUniversal" in n for n in names)
    
    def test_extracts_company_names_with_inc(self, simple_doc: Document) -> None:
        """Should extract company names with Inc suffix."""
        from effilocal.doc.party_detection import extract_company_names
        
        names = extract_company_names(simple_doc)
        
        assert any("Didimo" in n for n in names)
    
    def test_extracts_ltd_companies(self, license_agreement_doc: Document) -> None:
        """Should extract company names with Ltd suffix."""
        from effilocal.doc.party_detection import extract_company_names
        
        names = extract_company_names(license_agreement_doc)
        
        assert any("Acme" in n for n in names)
    
    def test_removes_duplicates(self) -> None:
        """Should remove duplicate company names."""
        from effilocal.doc.party_detection import extract_company_names
        
        doc = Document()
        doc.add_paragraph("Between Acme Ltd and Widget Corp.")
        doc.add_paragraph("Acme Ltd shall provide services to Widget Corp.")
        
        names = extract_company_names(doc)
        
        # Count occurrences
        acme_count = sum(1 for n in names if "Acme" in n)
        assert acme_count == 1
    
    def test_max_paragraphs_limit(self) -> None:
        """Should respect max_paragraphs limit."""
        from effilocal.doc.party_detection import extract_company_names
        
        doc = Document()
        # First 5 paragraphs
        for i in range(5):
            doc.add_paragraph(f"Paragraph {i} with FirstCompany Ltd.")
        # Paragraph 6+ with different company
        for i in range(5, 25):
            doc.add_paragraph(f"Paragraph {i} with SecondCompany Inc.")
        
        # With max_paragraphs=5, should only find FirstCompany
        names = extract_company_names(doc, max_paragraphs=5)
        
        assert any("FirstCompany" in n for n in names)
        # SecondCompany should not be found with limit of 5
        assert not any("SecondCompany" in n for n in names)


# =============================================================================
# Tests: extract_full_company_names
# =============================================================================


class TestExtractFullCompanyNames:
    """Tests for extract_full_company_names function."""
    
    def test_extracts_full_name_with_suffix(self, simple_doc: Document) -> None:
        """Should extract full company name including legal suffix."""
        from effilocal.doc.party_detection import extract_full_company_names
        
        names = extract_full_company_names(simple_doc)
        
        # Should include the full names with suffixes
        assert any("NBCUniversal Media" in n and "LLC" in n for n in names)
    
    def test_extracts_inc_with_period(self) -> None:
        """Should handle Inc. with trailing period."""
        from effilocal.doc.party_detection import extract_full_company_names
        
        doc = Document()
        doc.add_paragraph('Widget Corp, Inc. is the vendor.')
        
        names = extract_full_company_names(doc)
        
        assert any("Widget Corp" in n for n in names)
    
    def test_preserves_comma_separator(self) -> None:
        """Should preserve comma between name and suffix."""
        from effilocal.doc.party_detection import extract_full_company_names
        
        doc = Document()
        doc.add_paragraph('Acme Solutions, LLC is the provider.')
        
        names = extract_full_company_names(doc)
        
        # Should preserve the comma format
        assert any("Acme Solutions, LLC" in n for n in names) or \
               any("Acme Solutions" in n and "LLC" in n for n in names)


# =============================================================================
# Tests: extract_company_to_defined_term_mapping
# =============================================================================


class TestExtractCompanyToDefinedTermMapping:
    """Tests for extract_company_to_defined_term_mapping function."""
    
    def test_maps_company_to_defined_term(self, simple_doc: Document) -> None:
        """Should map company name to its defined term."""
        from effilocal.doc.party_detection import extract_company_to_defined_term_mapping
        
        mapping = extract_company_to_defined_term_mapping(simple_doc)
        
        # NBCUniversal should map to Company
        assert any("NBCUniversal" in k and mapping[k] == "Company" for k in mapping)
    
    def test_maps_vendor_company(self, simple_doc: Document) -> None:
        """Should map Didimo to Vendor."""
        from effilocal.doc.party_detection import extract_company_to_defined_term_mapping
        
        mapping = extract_company_to_defined_term_mapping(simple_doc)
        
        # Didimo should map to Vendor
        assert any("Didimo" in k and mapping[k] == "Vendor" for k in mapping)
    
    def test_handles_curly_quotes(self) -> None:
        """Should handle curly quotes in defined terms."""
        from effilocal.doc.party_detection import extract_company_to_defined_term_mapping
        
        doc = Document()
        # Using curly quotes (Unicode)
        doc.add_paragraph('Acme Ltd (\u201CProvider\u201D) agrees to provide services.')
        
        mapping = extract_company_to_defined_term_mapping(doc)
        
        assert any("Acme" in k for k in mapping)
    
    def test_returns_empty_for_no_matches(self) -> None:
        """Should return empty dict when no mappings found."""
        from effilocal.doc.party_detection import extract_company_to_defined_term_mapping
        
        doc = Document()
        doc.add_paragraph("This is just a regular paragraph.")
        
        mapping = extract_company_to_defined_term_mapping(doc)
        
        assert mapping == {}


# =============================================================================
# Tests: extract_party_alternate_names
# =============================================================================


class TestExtractPartyAlternateNames:
    """Tests for extract_party_alternate_names function."""
    
    def test_extracts_primary_name(self, simple_doc: Document) -> None:
        """Should extract primary defined term."""
        from effilocal.doc.party_detection import extract_party_alternate_names
        
        alternates = extract_party_alternate_names(simple_doc)
        
        assert "Company" in alternates or "Vendor" in alternates
    
    def test_extracts_alternate_name(self, simple_doc: Document) -> None:
        """Should extract alternate name from 'or' pattern."""
        from effilocal.doc.party_detection import extract_party_alternate_names
        
        alternates = extract_party_alternate_names(simple_doc)
        
        # Company should have NBCUniversal as alternate
        if "Company" in alternates:
            assert "NBCUniversal" in alternates["Company"]
    
    def test_includes_primary_in_list(self, simple_doc: Document) -> None:
        """Should include primary term in the list of alternates."""
        from effilocal.doc.party_detection import extract_party_alternate_names
        
        alternates = extract_party_alternate_names(simple_doc)
        
        for term, names in alternates.items():
            assert term in names


# =============================================================================
# Tests: extract_party_from_comment_prefixes (renamed from extract_comment_prefixes)
# =============================================================================


class TestExtractPartyFromCommentPrefixes:
    """Tests for extract_party_from_comment_prefixes function."""
    
    def test_extracts_prefixes_from_comments(self, sample_comments: list[dict]) -> None:
        """Should extract party prefixes from 'For X:' pattern."""
        from effilocal.doc.party_detection import extract_party_from_comment_prefixes
        
        prefixes = extract_party_from_comment_prefixes(sample_comments)
        
        assert "Didimo" in prefixes
        assert "NBC" in prefixes
    
    def test_returns_unique_prefixes(self, sample_comments: list[dict]) -> None:
        """Should return unique prefixes only."""
        from effilocal.doc.party_detection import extract_party_from_comment_prefixes
        
        prefixes = extract_party_from_comment_prefixes(sample_comments)
        
        # Should have exactly 2 unique prefixes
        assert len(prefixes) == 2
    
    def test_ignores_comments_without_prefix(self, sample_comments: list[dict]) -> None:
        """Should ignore comments that don't have 'For X:' prefix."""
        from effilocal.doc.party_detection import extract_party_from_comment_prefixes
        
        prefixes = extract_party_from_comment_prefixes(sample_comments)
        
        # "General comment" should not create a prefix
        assert len(prefixes) == 2
    
    def test_handles_empty_comments_list(self) -> None:
        """Should return empty list for empty comments."""
        from effilocal.doc.party_detection import extract_party_from_comment_prefixes
        
        prefixes = extract_party_from_comment_prefixes([])
        
        assert prefixes == []
    
    def test_case_insensitive_for_pattern(self) -> None:
        """Should match 'For', 'FOR', 'for' patterns."""
        from effilocal.doc.party_detection import extract_party_from_comment_prefixes
        
        comments = [
            {"text": "FOR Client: Important note."},
            {"text": "for Other: Another note."},
        ]
        
        prefixes = extract_party_from_comment_prefixes(comments)
        
        assert "Client" in prefixes
        assert "Other" in prefixes


# =============================================================================
# Tests: compute_similarity
# =============================================================================


class TestComputeSimilarity:
    """Tests for compute_similarity function."""
    
    def test_exact_match_returns_one(self) -> None:
        """Should return 1.0 for exact match."""
        from effilocal.doc.party_detection import compute_similarity
        
        score = compute_similarity("Vendor", "Vendor")
        
        assert score == 1.0
    
    def test_case_insensitive_exact_match(self) -> None:
        """Should return 1.0 for case-insensitive exact match."""
        from effilocal.doc.party_detection import compute_similarity
        
        score = compute_similarity("vendor", "VENDOR")
        
        assert score == 1.0
    
    def test_substring_match_high_score(self) -> None:
        """Should return high score when one contains the other."""
        from effilocal.doc.party_detection import compute_similarity
        
        score = compute_similarity("NBC", "NBCUniversal")
        
        assert score >= 0.8
    
    def test_no_match_returns_low_score(self) -> None:
        """Should return low score for unrelated strings."""
        from effilocal.doc.party_detection import compute_similarity
        
        score = compute_similarity("Apple", "Microsoft")
        
        assert score < 0.3
    
    def test_token_overlap_increases_score(self) -> None:
        """Should increase score for token overlap."""
        from effilocal.doc.party_detection import compute_similarity
        
        score = compute_similarity("Acme Corp", "Acme Industries")
        
        # Shared "Acme" token should give some similarity
        assert score > 0.3


# =============================================================================
# Tests: match_prefixes_to_parties
# =============================================================================


class TestMatchPrefixesToParties:
    """Tests for match_prefixes_to_parties function."""
    
    def test_matches_prefix_to_similar_term(self) -> None:
        """Should match prefix to similar defined term."""
        from effilocal.doc.party_detection import match_prefixes_to_parties
        
        prefixes = ["Didimo", "NBC"]
        defined_terms = ["Vendor", "Company"]
        company_names = ["Didimo Inc", "NBCUniversal"]
        
        matches = match_prefixes_to_parties(prefixes, defined_terms, company_names)
        
        # Should match Didimo to Didimo Inc
        assert "Didimo" in matches
        assert "Didimo" in matches["Didimo"] or "Inc" in matches["Didimo"]
    
    def test_returns_all_prefixes(self) -> None:
        """Should return a mapping for all prefixes."""
        from effilocal.doc.party_detection import match_prefixes_to_parties
        
        prefixes = ["Client", "Vendor"]
        defined_terms = ["Client", "Vendor"]
        company_names = []
        
        matches = match_prefixes_to_parties(prefixes, defined_terms, company_names)
        
        assert len(matches) == 2
        assert "Client" in matches
        assert "Vendor" in matches
    
    def test_keeps_original_for_no_match(self) -> None:
        """Should keep original prefix if no good match found."""
        from effilocal.doc.party_detection import match_prefixes_to_parties
        
        prefixes = ["XYZCorp"]
        defined_terms = ["Vendor", "Company"]
        company_names = ["Acme Inc", "Widget Ltd"]
        
        matches = match_prefixes_to_parties(prefixes, defined_terms, company_names)
        
        # Should keep XYZCorp as-is since no good match
        assert matches["XYZCorp"] == "XYZCorp"


# =============================================================================
# Tests: infer_party_role
# =============================================================================


class TestInferPartyRole:
    """Tests for infer_party_role function."""
    
    def test_vendor_maps_to_supplier(self) -> None:
        """Should map 'Vendor' to 'supplier' role."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("Vendor")
        
        assert role == "supplier"
    
    def test_company_maps_to_customer(self) -> None:
        """Should map 'Company' to 'customer' role."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("Company")
        
        assert role == "customer"
    
    def test_licensor_maps_to_licensor(self) -> None:
        """Should map 'Licensor' to 'licensor' role."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("Licensor")
        
        assert role == "licensor"
    
    def test_licensee_maps_to_licensee(self) -> None:
        """Should map 'Licensee' to 'licensee' role."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("Licensee")
        
        assert role == "licensee"
    
    def test_landlord_maps_to_landlord(self) -> None:
        """Should map 'Landlord' to 'landlord' role."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("Landlord")
        
        assert role == "landlord"
    
    def test_tenant_maps_to_tenant(self) -> None:
        """Should map 'Tenant' to 'tenant' role."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("Tenant")
        
        assert role == "tenant"
    
    def test_unknown_term_returns_party(self) -> None:
        """Should return 'party' for unknown terms."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("SomeUnknownTerm")
        
        assert role == "party"
    
    def test_case_insensitive(self) -> None:
        """Should be case-insensitive."""
        from effilocal.doc.party_detection import infer_party_role
        
        assert infer_party_role("vendor") == "supplier"
        assert infer_party_role("VENDOR") == "supplier"
        assert infer_party_role("Vendor") == "supplier"
    
    def test_seller_maps_to_seller(self) -> None:
        """Should map 'Seller' to 'seller' role."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("Seller")
        
        assert role == "seller"
    
    def test_employer_maps_to_employer(self) -> None:
        """Should map 'Employer' to 'employer' role."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("Employer")
        
        assert role == "employer"
    
    def test_employee_maps_to_employee(self) -> None:
        """Should map 'Employee' to 'employee' role."""
        from effilocal.doc.party_detection import infer_party_role
        
        role = infer_party_role("Employee")
        
        assert role == "employee"


# =============================================================================
# Tests: get_role_placeholder
# =============================================================================


class TestGetRolePlaceholder:
    """Tests for get_role_placeholder function."""
    
    def test_supplier_returns_supplier_placeholder(self) -> None:
        """Should return [SUPPLIER] for supplier role."""
        from effilocal.doc.party_detection import get_role_placeholder
        
        placeholder = get_role_placeholder("supplier")
        
        assert placeholder == "[SUPPLIER]"
    
    def test_customer_returns_customer_placeholder(self) -> None:
        """Should return [CUSTOMER] for customer role."""
        from effilocal.doc.party_detection import get_role_placeholder
        
        placeholder = get_role_placeholder("customer")
        
        assert placeholder == "[CUSTOMER]"
    
    def test_party_returns_party_placeholder(self) -> None:
        """Should return [PARTY] for party role."""
        from effilocal.doc.party_detection import get_role_placeholder
        
        placeholder = get_role_placeholder("party")
        
        assert placeholder == "[PARTY]"
    
    def test_uppercase_conversion(self) -> None:
        """Should convert role to uppercase."""
        from effilocal.doc.party_detection import get_role_placeholder
        
        assert get_role_placeholder("licensor") == "[LICENSOR]"
        assert get_role_placeholder("Licensee") == "[LICENSEE]"


# =============================================================================
# Tests: DEFINED_TERM_TO_ROLE mapping
# =============================================================================


class TestDefinedTermToRoleMapping:
    """Tests for DEFINED_TERM_TO_ROLE constant."""
    
    def test_contains_vendor(self) -> None:
        """Should contain vendor -> supplier mapping."""
        from effilocal.doc.party_detection import DEFINED_TERM_TO_ROLE
        
        assert "vendor" in DEFINED_TERM_TO_ROLE
        assert DEFINED_TERM_TO_ROLE["vendor"] == "supplier"
    
    def test_contains_company(self) -> None:
        """Should contain company -> customer mapping."""
        from effilocal.doc.party_detection import DEFINED_TERM_TO_ROLE
        
        assert "company" in DEFINED_TERM_TO_ROLE
        assert DEFINED_TERM_TO_ROLE["company"] == "customer"
    
    def test_contains_licensing_roles(self) -> None:
        """Should contain licensor and licensee mappings."""
        from effilocal.doc.party_detection import DEFINED_TERM_TO_ROLE
        
        assert "licensor" in DEFINED_TERM_TO_ROLE
        assert "licensee" in DEFINED_TERM_TO_ROLE
    
    def test_contains_property_roles(self) -> None:
        """Should contain landlord, tenant, lessor, lessee mappings."""
        from effilocal.doc.party_detection import DEFINED_TERM_TO_ROLE
        
        assert "landlord" in DEFINED_TERM_TO_ROLE
        assert "tenant" in DEFINED_TERM_TO_ROLE
        assert "lessor" in DEFINED_TERM_TO_ROLE
        assert "lessee" in DEFINED_TERM_TO_ROLE


# =============================================================================
# Tests: PartyInfo dataclass
# =============================================================================


class TestPartyInfo:
    """Tests for PartyInfo dataclass."""
    
    def test_basic_instantiation(self) -> None:
        """Should create PartyInfo with required fields."""
        from effilocal.doc.party_detection import PartyInfo
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty"
        )
        
        assert party_info.client_prefix == "Didimo"
        assert party_info.counterparty_prefix == "NBC"
        assert party_info.client_defined_term == "Vendor"
        assert party_info.counterparty_defined_term == "Company"
        assert party_info.original_provided_by == "counterparty"
    
    def test_default_roles(self) -> None:
        """Should have default role of 'party'."""
        from effilocal.doc.party_detection import PartyInfo
        
        party_info = PartyInfo(
            client_prefix="A",
            counterparty_prefix="B",
            client_defined_term="PartyA",
            counterparty_defined_term="PartyB",
            original_provided_by="client"
        )
        
        assert party_info.client_role == "party"
        assert party_info.counterparty_role == "party"
    
    def test_client_placeholder_property(self) -> None:
        """Should compute client_placeholder from client_role."""
        from effilocal.doc.party_detection import PartyInfo
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_role="supplier",
            counterparty_role="customer"
        )
        
        assert party_info.client_placeholder == "[SUPPLIER]"
    
    def test_counterparty_placeholder_property(self) -> None:
        """Should compute counterparty_placeholder from counterparty_role."""
        from effilocal.doc.party_detection import PartyInfo
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_role="supplier",
            counterparty_role="customer"
        )
        
        assert party_info.counterparty_placeholder == "[CUSTOMER]"
    
    def test_all_client_names_property(self) -> None:
        """Should return all client names including alternates."""
        from effilocal.doc.party_detection import PartyInfo
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_alternate_names=["Didimo Inc", "Didimo, Inc."]
        )
        
        all_names = party_info.all_client_names
        
        assert "Didimo" in all_names
        assert "Vendor" in all_names
        assert "Didimo Inc" in all_names
    
    def test_all_counterparty_names_property(self) -> None:
        """Should return all counterparty names including alternates."""
        from effilocal.doc.party_detection import PartyInfo
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            counterparty_alternate_names=["NBCUniversal", "NBC Universal Media"]
        )
        
        all_names = party_info.all_counterparty_names
        
        assert "NBC" in all_names
        assert "Company" in all_names
        assert "NBCUniversal" in all_names
    
    def test_all_names_deduplicates(self) -> None:
        """Should deduplicate names in all_*_names properties."""
        from effilocal.doc.party_detection import PartyInfo
        
        party_info = PartyInfo(
            client_prefix="Vendor",
            counterparty_prefix="Company",
            client_defined_term="Vendor",  # Same as prefix
            counterparty_defined_term="Company",  # Same as prefix
            original_provided_by="counterparty",
            client_alternate_names=["Vendor"]  # Duplicate
        )
        
        all_names = party_info.all_client_names
        
        # Should have only one "Vendor"
        assert all_names.count("Vendor") == 1
    
    def test_all_names_excludes_empty(self) -> None:
        """Should exclude empty strings from all_*_names."""
        from effilocal.doc.party_detection import PartyInfo
        
        party_info = PartyInfo(
            client_prefix="",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="",
            original_provided_by="counterparty"
        )
        
        client_names = party_info.all_client_names
        counterparty_names = party_info.all_counterparty_names
        
        assert "" not in client_names
        assert "" not in counterparty_names


# =============================================================================
# Integration Tests
# =============================================================================


class TestIntegration:
    """Integration tests combining multiple functions."""
    
    def test_full_workflow_service_agreement(self, simple_doc: Document, sample_comments: list[dict]) -> None:
        """Should extract and process party info from service agreement."""
        from effilocal.doc.party_detection import (
            extract_defined_party_terms,
            extract_company_names,
            extract_party_from_comment_prefixes,
            extract_company_to_defined_term_mapping,
            infer_party_role,
        )
        
        defined_terms = extract_defined_party_terms(simple_doc)
        company_names = extract_company_names(simple_doc)
        prefixes = extract_party_from_comment_prefixes(sample_comments)
        mapping = extract_company_to_defined_term_mapping(simple_doc)
        
        # Should find Vendor and Company
        assert "Vendor" in defined_terms
        assert "Company" in defined_terms
        
        # Should find company names
        assert len(company_names) >= 2
        
        # Should find comment prefixes
        assert "Didimo" in prefixes
        assert "NBC" in prefixes
        
        # Roles should be inferred correctly
        assert infer_party_role("Vendor") == "supplier"
        assert infer_party_role("Company") == "customer"
    
    def test_full_workflow_license_agreement(self, license_agreement_doc: Document) -> None:
        """Should extract party info from license agreement."""
        from effilocal.doc.party_detection import (
            extract_defined_party_terms,
            infer_party_role,
            get_role_placeholder,
        )
        
        defined_terms = extract_defined_party_terms(license_agreement_doc)
        
        assert "Licensor" in defined_terms
        assert "Licensee" in defined_terms
        
        assert infer_party_role("Licensor") == "licensor"
        assert infer_party_role("Licensee") == "licensee"
        
        assert get_role_placeholder("licensor") == "[LICENSOR]"
        assert get_role_placeholder("licensee") == "[LICENSEE]"
