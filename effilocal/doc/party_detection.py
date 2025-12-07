"""
Party detection utilities for contract documents.

This module provides functions for detecting and extracting party information
from Word documents, including:
- Defined terms (e.g., "Vendor", "Company")
- Company names from preambles
- Comment prefixes (e.g., "For Client:")
- Semantic role mapping

Typical usage:
    from effilocal.doc.party_detection import (
        extract_defined_party_terms,
        extract_company_names,
        extract_party_from_comment_prefixes,
        infer_party_role,
        PartyInfo,
    )
    
    defined_terms = extract_defined_party_terms(doc)
    company_names = extract_company_names(doc)
    prefixes = extract_party_from_comment_prefixes(comments)
"""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field
from typing import Any

from docx import Document


# =============================================================================
# Constants
# =============================================================================


# Mapping of common defined terms to semantic roles
DEFINED_TERM_TO_ROLE: dict[str, str] = {
    # Service/supply relationships
    "vendor": "supplier",
    "supplier": "supplier",
    "provider": "supplier",
    "contractor": "supplier",
    "consultant": "supplier",
    "company": "customer",  # Often the customer in service agreements
    "customer": "customer",
    "client": "customer",
    "buyer": "customer",
    "purchaser": "customer",
    # Licensing relationships
    "licensor": "licensor",
    "licensee": "licensee",
    # Property relationships
    "landlord": "landlord",
    "lessor": "landlord",
    "tenant": "tenant",
    "lessee": "tenant",
    # Employment/agency
    "employer": "employer",
    "employee": "employee",
    "principal": "principal",
    "agent": "agent",
    # Financial
    "lender": "lender",
    "borrower": "borrower",
    "creditor": "creditor",
    "debtor": "debtor",
    # Sale
    "seller": "seller",
}


# =============================================================================
# Extraction Functions
# =============================================================================


def extract_defined_party_terms(doc: Document) -> list[str]:
    """
    Extract likely defined terms for parties from contract text.
    
    Looks for capitalized words that commonly precede shall/will/may,
    which typically indicate party definitions (e.g., "Vendor", "Company", "Licensor").
    
    Args:
        doc: The Word document to analyze
        
    Returns:
        List of likely party defined terms, sorted by frequency
    """
    party_candidates: Counter[str] = Counter()
    
    # Pattern: Capitalized word followed by shall/will/may
    # e.g., "Vendor shall", "Company will", "Licensor may"
    # Use (?:The\s+)? to optionally match "The " but not capture it
    pattern = re.compile(r'\b(?:[Tt]he\s+)?([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s+(?:shall|will|may)\b')
    
    # Words to filter out
    filter_words = {'the', 'each', 'either', 'neither', 'any', 'no', 'such', 'this', 'that'}
    
    for para in doc.paragraphs:
        text = para.text
        matches = pattern.findall(text)
        for match in matches:
            if match.lower() not in filter_words:
                party_candidates[match] += 1
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = pattern.findall(cell.text)
                for match in matches:
                    if match.lower() not in filter_words:
                        party_candidates[match] += 1
    
    # Return terms sorted by frequency (most common first)
    return [term for term, _ in party_candidates.most_common()]


def extract_company_names(doc: Document, max_paragraphs: int = 20) -> list[str]:
    """
    Extract company names from the early paragraphs of a contract.
    
    Looks for patterns with common company suffixes like Ltd, Limited, LLP, Inc, LLC, CIC.
    
    Args:
        doc: The Word document to analyze
        max_paragraphs: How many paragraphs to scan (default: 20, covers preamble)
        
    Returns:
        List of detected company names
    """
    company_names: list[str] = []
    
    # Pattern for company names with legal suffixes
    suffixes = r'(?:[Ll]td\.?|[Ll]imited|LLP|[Ii]nc\.?|LLC|CIC|[Cc]orp\.?|[Cc]orporation|PLC|plc|S\.?A\.?|GmbH|B\.?V\.?|N\.?V\.?)'
    
    # Pattern: One or more capitalized words followed by a legal suffix
    pattern = re.compile(
        rf'([A-Z][A-Za-z0-9]*(?:\s+[A-Z][A-Za-z0-9]*)*)\s*,?\s*{suffixes}\b'
    )
    
    paragraphs_scanned = 0
    for para in doc.paragraphs:
        if paragraphs_scanned >= max_paragraphs:
            break
        
        text = para.text
        if not text.strip():
            continue
            
        paragraphs_scanned += 1
        
        matches = pattern.findall(text)
        for match in matches:
            # Clean up the match
            name = match.strip().rstrip(',')
            # Remove common leading words that aren't part of company names
            for prefix in ("Between", "between", "And", "and"):
                if name.startswith(prefix + " "):
                    name = name[len(prefix) + 1:]
            if name and len(name) > 2:
                company_names.append(name)
    
    # Remove duplicates while preserving order
    seen = set()
    unique_names = []
    for name in company_names:
        if name.lower() not in seen:
            seen.add(name.lower())
            unique_names.append(name)
    
    return unique_names


def extract_full_company_names(doc: Document, max_paragraphs: int = 20) -> list[str]:
    """
    Extract full company names including legal suffix from the preamble.
    
    Looks for patterns like "NBCUniversal Media, LLC" or "Didimo, Inc."
    Returns the complete legal name for replacement purposes.
    
    Args:
        doc: The Word document to analyze
        max_paragraphs: How many paragraphs to scan
        
    Returns:
        List of full company names with their legal suffixes
    """
    full_names: list[str] = []
    
    # Legal entity suffixes
    suffix_bases = [
        'Ltd', 'Limited', 'LLP', 'Inc', 'LLC', 'CIC', 'Corp', 'Corporation',
        'PLC', 'plc', 'SA', 'S.A.', 'GmbH', 'BV', 'B.V.', 'NV', 'N.V.'
    ]
    suffixes = '(' + '|'.join(re.escape(s) for s in suffix_bases) + r')\.?'
    
    # Pattern: Company name followed by optional comma and legal suffix
    pattern = re.compile(
        rf'([A-Z][A-Za-z0-9]*(?:\s+[A-Z][A-Za-z0-9]*)*)(\s*,\s*|\s+){suffixes}(?=\s|[,"\'\(\)]|$)'
    )
    
    paragraphs_scanned = 0
    for para in doc.paragraphs:
        if paragraphs_scanned >= max_paragraphs:
            break
        
        text = para.text
        if not text.strip():
            continue
            
        paragraphs_scanned += 1
        
        for match in pattern.finditer(text):
            name_part = match.group(1).strip()
            separator = match.group(2)
            suffix_part = match.group(3)
            
            # Check if there's a period after the suffix in the full match
            full_match = match.group(0)
            has_trailing_period = full_match.rstrip().endswith('.')
            
            # Remove common leading words
            for prefix in ("Between", "between", "And", "and"):
                if name_part.startswith(prefix + " "):
                    name_part = name_part[len(prefix) + 1:]
            
            if name_part and len(name_part) > 2:
                # Build full name preserving original format
                if ',' in separator:
                    full_name = f"{name_part}, {suffix_part}"
                else:
                    full_name = f"{name_part} {suffix_part}"
                
                # Add trailing period if original had it
                if has_trailing_period and not full_name.endswith('.'):
                    full_name += '.'
                    
                full_names.append(full_name)
    
    # Remove duplicates while preserving order
    seen = set()
    unique_names = []
    for name in full_names:
        if name.lower() not in seen:
            seen.add(name.lower())
            unique_names.append(name)
    
    return unique_names


def extract_company_to_defined_term_mapping(doc: Document, max_paragraphs: int = 20) -> dict[str, str]:
    """
    Extract mapping of company names to their defined terms from the preamble.
    
    Looks for patterns like:
    - 'NBCUniversal Media, LLC ... ("Company" or "NBCUniversal")'
    - 'Didimo, Inc. ... ("Vendor")'
    
    Args:
        doc: The Word document to analyze
        max_paragraphs: How many paragraphs to scan (default: 20, covers preamble)
        
    Returns:
        Dictionary mapping company names to their defined terms
        e.g., {"NBCUniversal": "Company", "Didimo": "Vendor"}
    """
    mapping: dict[str, str] = {}
    
    # Legal entity suffixes
    suffixes = r'(?:[Ll]td\.?|[Ll]imited|LLP|[Ii]nc\.?|LLC|CIC|[Cc]orp\.?|[Cc]orporation|PLC|plc|S\.?A\.?|GmbH|B\.?V\.?|N\.?V\.?)'
    
    # Quote characters: straight quotes and curly quotes (Unicode U+201C/U+201D)
    quotes = r'["\'\u201C\u201D]'
    
    # Pattern to match: "Company Name, LLC ... ("DefinedTerm")"
    pattern = re.compile(
        rf'([A-Z][A-Za-z0-9]*(?:\s+[A-Z][A-Za-z0-9]*)*)\s*,?\s*{suffixes}'  # Company name with suffix
        rf'[^()]*'  # Skip jurisdiction/description
        rf'\(\s*{quotes}([A-Z][A-Za-z]+){quotes}'  # First quoted defined term in parentheses
        rf'(?:\s+or\s+{quotes}[A-Za-z]+{quotes})?'  # Optional "or Alias"
        rf'\s*\)'  # Close paren
    )
    
    paragraphs_scanned = 0
    for para in doc.paragraphs:
        if paragraphs_scanned >= max_paragraphs:
            break
        
        text = para.text
        if not text.strip():
            continue
            
        paragraphs_scanned += 1
        
        matches = pattern.findall(text)
        for company_name, defined_term in matches:
            # Clean up company name
            company_name = company_name.strip()
            for prefix in ("Between", "between", "And", "and"):
                if company_name.startswith(prefix + " "):
                    company_name = company_name[len(prefix) + 1:]
            
            if company_name and defined_term:
                mapping[company_name] = defined_term
    
    return mapping


def extract_party_alternate_names(doc: Document, max_paragraphs: int = 20) -> dict[str, list[str]]:
    """
    Extract all alternate names for each party from the preamble.
    
    Looks for patterns like:
    - '("Company" or "NBCUniversal")' → {"Company": ["Company", "NBCUniversal"]}
    - '("Vendor")' → {"Vendor": ["Vendor"]}
    
    Args:
        doc: The Word document to analyze
        max_paragraphs: How many paragraphs to scan
        
    Returns:
        Dictionary mapping primary defined term to list of all names for that party
    """
    alternates: dict[str, list[str]] = {}
    
    # Quote characters: straight quotes and curly quotes
    quotes = r'["\'\u201C\u201D]'
    
    # Pattern to match: ("DefinedTerm") or ("DefinedTerm" or "Alternate")
    pattern = re.compile(
        rf'\(\s*{quotes}([A-Z][A-Za-z]+){quotes}'  # First quoted term
        rf'(?:\s+or\s+{quotes}([A-Za-z]+){quotes})?'  # Optional "or Alternate"
        rf'\s*\)'  # Close paren
    )
    
    paragraphs_scanned = 0
    for para in doc.paragraphs:
        if paragraphs_scanned >= max_paragraphs:
            break
        
        text = para.text
        if not text.strip():
            continue
            
        paragraphs_scanned += 1
        
        for match in pattern.finditer(text):
            primary = match.group(1)
            alternate = match.group(2)
            
            if primary not in alternates:
                alternates[primary] = [primary]
            
            if alternate and alternate not in alternates[primary]:
                alternates[primary].append(alternate)
    
    return alternates


def extract_party_from_comment_prefixes(comments: list[dict[str, Any]]) -> list[str]:
    """
    Extract unique party identifiers from "For X:" prefixes in comments.
    
    Looks for patterns like "For Didimo:", "For NBC:", etc.
    
    Args:
        comments: List of comment dictionaries with 'text' field
        
    Returns:
        List of unique party identifiers found in comment prefixes
    """
    prefixes: set[str] = set()
    
    # Pattern: "For X:" at the start of a comment
    pattern = re.compile(r'^For\s+([^:]+):', re.IGNORECASE)
    
    for comment in comments:
        text = comment.get("text", "")
        match = pattern.match(text.strip())
        if match:
            prefixes.add(match.group(1).strip())
    
    return list(prefixes)


# =============================================================================
# Similarity and Matching
# =============================================================================


def compute_similarity(s1: str, s2: str) -> float:
    """
    Compute similarity between two strings using token overlap and substring matching.
    
    Args:
        s1: First string
        s2: Second string
        
    Returns:
        Similarity score between 0 and 1
    """
    s1_lower = s1.lower()
    s2_lower = s2.lower()
    
    # Exact match
    if s1_lower == s2_lower:
        return 1.0
    
    # One contains the other
    if s1_lower in s2_lower or s2_lower in s1_lower:
        return 0.9
    
    # Token overlap (Jaccard similarity)
    tokens1 = set(s1_lower.split())
    tokens2 = set(s2_lower.split())
    
    if not tokens1 or not tokens2:
        return 0.0
    
    intersection = tokens1 & tokens2
    union = tokens1 | tokens2
    jaccard = len(intersection) / len(union) if union else 0
    
    # First word match bonus (common for company names)
    first_word_match = 0.3 if s1_lower.split()[0] == s2_lower.split()[0] else 0
    
    return min(jaccard + first_word_match, 1.0)


def match_prefixes_to_parties(
    prefixes: list[str],
    defined_terms: list[str],
    company_names: list[str]
) -> dict[str, str]:
    """
    Match comment prefixes to detected party names/terms.
    
    Args:
        prefixes: Comment prefixes (e.g., ["Didimo", "NBC"])
        defined_terms: Defined terms from contract (e.g., ["Vendor", "Company"])
        company_names: Company names from preamble (e.g., ["Didimo, Inc", "NBC Universal"])
        
    Returns:
        Mapping of prefix -> best matching party identifier
    """
    all_candidates = defined_terms + company_names
    
    matches: dict[str, str] = {}
    
    for prefix in prefixes:
        best_match = prefix  # Default to the prefix itself
        best_score = 0.0
        
        for candidate in all_candidates:
            score = compute_similarity(prefix, candidate)
            if score > best_score:
                best_score = score
                best_match = candidate
        
        # Only use the match if score is reasonable
        if best_score >= 0.3:
            matches[prefix] = best_match
        else:
            matches[prefix] = prefix  # Keep original if no good match
    
    return matches


# =============================================================================
# Role Inference
# =============================================================================


def infer_party_role(defined_term: str) -> str:
    """
    Infer the semantic role of a party from its defined term.
    
    This helps the LLM understand the contractual relationship.
    
    Args:
        defined_term: The defined term in the contract (e.g., "Vendor", "Company")
        
    Returns:
        Semantic role (e.g., "supplier", "customer") or "party" if unknown
    """
    term_lower = defined_term.lower().strip()
    return DEFINED_TERM_TO_ROLE.get(term_lower, "party")


def get_role_placeholder(role: str) -> str:
    """
    Convert a semantic role to its placeholder format.
    
    This is used for anonymization - replacing company names with semantic
    role placeholders that preserve meaning for LLM training.
    
    Args:
        role: The semantic role (e.g., "supplier", "customer", "party")
        
    Returns:
        Placeholder string (e.g., "[SUPPLIER]", "[CUSTOMER]", "[PARTY]")
    """
    return f"[{role.upper()}]"


# =============================================================================
# Data Classes
# =============================================================================


@dataclass
class PartyInfo:
    """Information about the parties in a contract."""
    
    client_prefix: str  # The prefix used in comments for client (e.g., "Didimo")
    counterparty_prefix: str  # The prefix used in comments for counterparty (e.g., "NBC")
    client_defined_term: str  # The primary defined term in contract (e.g., "Vendor")
    counterparty_defined_term: str  # The primary defined term (e.g., "Company")
    original_provided_by: str  # "client" or "counterparty"
    client_role: str = "party"  # Semantic role (e.g., "supplier", "licensor")
    counterparty_role: str = "party"  # Semantic role (e.g., "customer", "licensee")
    # Alternate names for each party (e.g., ["Company", "NBCUniversal"])
    client_alternate_names: list[str] = field(default_factory=list)
    counterparty_alternate_names: list[str] = field(default_factory=list)
    
    @property
    def client_placeholder(self) -> str:
        """Get the placeholder string for the client based on their role."""
        return get_role_placeholder(self.client_role)
    
    @property
    def counterparty_placeholder(self) -> str:
        """Get the placeholder string for the counterparty based on their role."""
        return get_role_placeholder(self.counterparty_role)
    
    @property
    def all_client_names(self) -> list[str]:
        """Get all names that should be replaced for the client."""
        names = [self.client_prefix, self.client_defined_term]
        names.extend(self.client_alternate_names)
        # Deduplicate while preserving order
        seen = set()
        result = []
        for name in names:
            if name and name not in seen:
                seen.add(name)
                result.append(name)
        return result
    
    @property
    def all_counterparty_names(self) -> list[str]:
        """Get all names that should be replaced for the counterparty."""
        names = [self.counterparty_prefix, self.counterparty_defined_term]
        names.extend(self.counterparty_alternate_names)
        # Deduplicate while preserving order
        seen = set()
        result = []
        for name in names:
            if name and name not in seen:
                seen.add(name)
                result.append(name)
        return result
