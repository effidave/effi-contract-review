#!/usr/bin/env python3
"""
Comprehensive tests for party detection and YAML header generation.

These tests verify:
- Party detection from contract text (defined terms, company names)
- Comment prefix extraction and matching
- PartyInfo dataclass with role information
- YAML header generation with parties mapping
- Anonymization that preserves defined terms
"""

import sys
from pathlib import Path
from io import BytesIO
from typing import Any
from unittest.mock import patch, MagicMock

import pytest
from docx import Document

# Add project root to path
_PROJECT_ROOT = Path(__file__).parent.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))


# =============================================================================
# Tests for PartyInfo dataclass
# =============================================================================

class TestPartyInfo:
    """Tests for the enhanced PartyInfo dataclass."""
    
    def test_party_info_has_client_role(self) -> None:
        """PartyInfo should include client_role (e.g., 'supplier', 'customer')."""
        from scripts.generate_review_example import PartyInfo
        
        info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_role="supplier",
            counterparty_role="customer"
        )
        
        assert info.client_role == "supplier"
        assert info.counterparty_role == "customer"
    
    def test_party_info_has_all_required_fields(self) -> None:
        """PartyInfo should have all fields for complete party mapping."""
        from scripts.generate_review_example import PartyInfo
        
        info = PartyInfo(
            client_prefix="Acme",
            counterparty_prefix="BigCorp",
            client_defined_term="Licensor",
            counterparty_defined_term="Licensee",
            original_provided_by="client",
            client_role="licensor",
            counterparty_role="licensee"
        )
        
        # All fields should be accessible
        assert info.client_prefix == "Acme"
        assert info.counterparty_prefix == "BigCorp"
        assert info.client_defined_term == "Licensor"
        assert info.counterparty_defined_term == "Licensee"
        assert info.original_provided_by == "client"
        assert info.client_role == "licensor"
        assert info.counterparty_role == "licensee"


# =============================================================================
# Tests for extract_defined_party_terms function
# =============================================================================

class TestExtractDefinedPartyTerms:
    """Tests for extracting defined terms from contract text."""
    
    def test_extracts_vendor_and_company(self) -> None:
        """Should extract 'Vendor' and 'Company' from typical contract text."""
        from scripts.generate_review_example import extract_defined_party_terms
        
        doc = Document()
        doc.add_paragraph("Vendor shall provide the services to Company.")
        doc.add_paragraph("Company will pay Vendor within 30 days.")
        doc.add_paragraph("Vendor may terminate upon notice.")
        
        terms = extract_defined_party_terms(doc)
        
        assert "Vendor" in terms
        assert "Company" in terms
    
    def test_returns_terms_sorted_by_frequency(self) -> None:
        """Most frequently occurring terms should be first."""
        from scripts.generate_review_example import extract_defined_party_terms
        
        doc = Document()
        # Vendor appears 3 times with shall/will/may
        doc.add_paragraph("Vendor shall do X. Vendor will do Y. Vendor may do Z.")
        # Company appears 1 time
        doc.add_paragraph("Company shall pay.")
        
        terms = extract_defined_party_terms(doc)
        
        # Vendor should come before Company due to higher frequency
        vendor_idx = terms.index("Vendor") if "Vendor" in terms else -1
        company_idx = terms.index("Company") if "Company" in terms else -1
        
        assert vendor_idx != -1
        assert company_idx != -1
        assert vendor_idx < company_idx
    
    def test_filters_common_false_positives(self) -> None:
        """Should filter out 'The', 'Each', 'Either' etc."""
        from scripts.generate_review_example import extract_defined_party_terms
        
        doc = Document()
        doc.add_paragraph("The parties shall agree. Each party will comply.")
        doc.add_paragraph("Either party may terminate. Such party shall notify.")
        
        terms = extract_defined_party_terms(doc)
        
        # Common words should be filtered
        lowercase_terms = [t.lower() for t in terms]
        assert "the" not in lowercase_terms
        assert "each" not in lowercase_terms
        assert "either" not in lowercase_terms
        assert "such" not in lowercase_terms


# =============================================================================
# Tests for extract_company_names function
# =============================================================================

class TestExtractCompanyNames:
    """Tests for extracting company names from contract preamble."""
    
    def test_extracts_names_with_ltd_suffix(self) -> None:
        """Should extract company names ending in Ltd or Limited."""
        from scripts.generate_review_example import extract_company_names
        
        doc = Document()
        doc.add_paragraph("This agreement is between Acme Ltd and BigCorp Limited.")
        
        names = extract_company_names(doc)
        
        assert any("Acme" in name for name in names)
        assert any("BigCorp" in name for name in names)
    
    def test_extracts_names_with_inc_llc_suffix(self) -> None:
        """Should extract company names ending in Inc, LLC, Corp."""
        from scripts.generate_review_example import extract_company_names
        
        doc = Document()
        doc.add_paragraph("Between TechCo Inc. and StartupCo, LLC.")
        doc.add_paragraph("And MegaCorp Corporation agrees.")
        
        names = extract_company_names(doc)
        
        assert any("TechCo" in name for name in names)
        assert any("StartupCo" in name for name in names)
    
    def test_limits_to_first_n_paragraphs(self) -> None:
        """Should only scan first N paragraphs (preamble area)."""
        from scripts.generate_review_example import extract_company_names
        
        doc = Document()
        # Add company name in first paragraph
        doc.add_paragraph("This agreement between EarlyCo Ltd.")
        # Add 25 blank paragraphs
        for _ in range(25):
            doc.add_paragraph("Lorem ipsum dolor sit amet.")
        # Add company name at end
        doc.add_paragraph("LateCo Inc. is also mentioned here.")
        
        names = extract_company_names(doc, max_paragraphs=20)
        
        assert any("EarlyCo" in name for name in names)
        # LateCo should NOT be found (beyond max_paragraphs)
        assert not any("LateCo" in name for name in names)


# =============================================================================
# Tests for extract_full_company_names function
# =============================================================================

class TestExtractFullCompanyNames:
    """Tests for extracting full company names including legal suffixes."""
    
    def test_extracts_llc_with_comma(self) -> None:
        """Should extract 'NBCUniversal Media, LLC' with the comma and suffix."""
        from scripts.generate_review_example import extract_full_company_names
        
        doc = Document()
        doc.add_paragraph('NBCUniversal Media, LLC ("Company")')
        
        names = extract_full_company_names(doc)
        
        assert "NBCUniversal Media, LLC" in names
    
    def test_extracts_inc_with_comma(self) -> None:
        """Should extract 'Didimo, Inc.' with the comma and suffix."""
        from scripts.generate_review_example import extract_full_company_names
        
        doc = Document()
        doc.add_paragraph('Didimo, Inc. ("Vendor")')
        
        names = extract_full_company_names(doc)
        
        assert "Didimo, Inc." in names
    
    def test_extracts_ltd_without_comma(self) -> None:
        """Should extract 'Acme Ltd' without comma."""
        from scripts.generate_review_example import extract_full_company_names
        
        doc = Document()
        doc.add_paragraph('Acme Ltd and BigCorp Limited agree.')
        
        names = extract_full_company_names(doc)
        
        assert "Acme Ltd" in names
        assert "BigCorp Limited" in names
    
    def test_extracts_multiple_companies(self) -> None:
        """Should extract multiple company names from preamble."""
        from scripts.generate_review_example import extract_full_company_names
        
        doc = Document()
        doc.add_paragraph('This Agreement is entered into between:')
        doc.add_paragraph('NBCUniversal Media, LLC ("Company") and')
        doc.add_paragraph('Didimo, Inc. ("Vendor")')
        
        names = extract_full_company_names(doc)
        
        assert "NBCUniversal Media, LLC" in names
        assert "Didimo, Inc." in names
    
    def test_extracts_llc_followed_by_comma(self) -> None:
        """Should extract 'NBCUniversal Media, LLC' even when followed by comma (preamble style)."""
        from scripts.generate_review_example import extract_full_company_names
        
        doc = Document()
        # Real preamble format: "NBCUniversal Media, LLC, a Delaware limited liability company"
        doc.add_paragraph(
            'This Agreement is entered into by NBCUniversal Media, LLC, a Delaware limited liability company ("Company")'
        )
        
        names = extract_full_company_names(doc)
        
        assert "NBCUniversal Media, LLC" in names


# =============================================================================
# Tests for extract_comment_prefixes function
# =============================================================================

class TestExtractCommentPrefixes:
    """Tests for extracting 'For X:' prefixes from comments."""
    
    def test_extracts_prefixes_from_comments(self) -> None:
        """Should extract party names from 'For X:' pattern."""
        from scripts.generate_review_example import extract_comment_prefixes
        
        comments = [
            {"text": "For Didimo: Please review this clause."},
            {"text": "For NBC: We suggest adding..."},
            {"text": "For Didimo: Another note."},
        ]
        
        prefixes = extract_comment_prefixes(comments)
        
        assert "Didimo" in prefixes
        assert "NBC" in prefixes
    
    def test_returns_unique_prefixes(self) -> None:
        """Should not include duplicates."""
        from scripts.generate_review_example import extract_comment_prefixes
        
        comments = [
            {"text": "For Client: Note 1"},
            {"text": "For Client: Note 2"},
            {"text": "For Client: Note 3"},
        ]
        
        prefixes = extract_comment_prefixes(comments)
        
        assert prefixes.count("Client") == 1
    
    def test_ignores_comments_without_prefix(self) -> None:
        """Should ignore comments that don't start with 'For X:'."""
        from scripts.generate_review_example import extract_comment_prefixes
        
        comments = [
            {"text": "For Client: Valid prefix"},
            {"text": "This comment has no prefix"},
            {"text": "Note: Not a For X prefix"},
        ]
        
        prefixes = extract_comment_prefixes(comments)
        
        assert "Client" in prefixes
        assert len(prefixes) == 1


# =============================================================================
# Tests for infer_party_role function
# =============================================================================

class TestInferPartyRole:
    """Tests for inferring party role from defined term."""
    
    def test_vendor_infers_supplier_role(self) -> None:
        """'Vendor' should infer 'supplier' role."""
        from scripts.generate_review_example import infer_party_role
        
        role = infer_party_role("Vendor")
        assert role == "supplier"
    
    def test_company_infers_customer_role(self) -> None:
        """'Company' should infer 'customer' role."""
        from scripts.generate_review_example import infer_party_role
        
        role = infer_party_role("Company")
        assert role == "customer"
    
    def test_licensor_infers_licensor_role(self) -> None:
        """'Licensor' should infer 'licensor' role."""
        from scripts.generate_review_example import infer_party_role
        
        role = infer_party_role("Licensor")
        assert role == "licensor"
    
    def test_licensee_infers_licensee_role(self) -> None:
        """'Licensee' should infer 'licensee' role."""
        from scripts.generate_review_example import infer_party_role
        
        role = infer_party_role("Licensee")
        assert role == "licensee"
    
    def test_supplier_infers_supplier_role(self) -> None:
        """'Supplier' should infer 'supplier' role."""
        from scripts.generate_review_example import infer_party_role
        
        role = infer_party_role("Supplier")
        assert role == "supplier"
    
    def test_customer_infers_customer_role(self) -> None:
        """'Customer' should infer 'customer' role."""
        from scripts.generate_review_example import infer_party_role
        
        role = infer_party_role("Customer")
        assert role == "customer"
    
    def test_landlord_infers_landlord_role(self) -> None:
        """'Landlord' should infer 'landlord' role."""
        from scripts.generate_review_example import infer_party_role
        
        role = infer_party_role("Landlord")
        assert role == "landlord"
    
    def test_tenant_infers_tenant_role(self) -> None:
        """'Tenant' should infer 'tenant' role."""
        from scripts.generate_review_example import infer_party_role
        
        role = infer_party_role("Tenant")
        assert role == "tenant"
    
    def test_unknown_term_returns_party(self) -> None:
        """Unknown terms should return 'party' as default."""
        from scripts.generate_review_example import infer_party_role
        
        role = infer_party_role("UnknownTerm")
        assert role == "party"
    
    def test_case_insensitive_matching(self) -> None:
        """Role inference should be case-insensitive."""
        from scripts.generate_review_example import infer_party_role
        
        assert infer_party_role("VENDOR") == "supplier"
        assert infer_party_role("vendor") == "supplier"
        assert infer_party_role("Vendor") == "supplier"


# =============================================================================
# Tests for generate_yaml_header function
# =============================================================================

class TestGenerateYamlHeader:
    """Tests for YAML header generation with parties mapping."""
    
    def test_includes_parties_mapping(self) -> None:
        """YAML header should include parties mapping with roles."""
        from scripts.generate_review_example import PartyInfo, generate_yaml_header
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_role="supplier",
            counterparty_role="customer"
        )
        
        yaml = generate_yaml_header(
            document_type="original_agreement",
            party_info=party_info
        )
        
        assert "parties:" in yaml
        assert "client:" in yaml
        assert "counterparty:" in yaml
    
    def test_excludes_defined_terms_from_yaml(self) -> None:
        """Parties mapping should NOT include defined_term field (it's redundant)."""
        from scripts.generate_review_example import PartyInfo, generate_yaml_header
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_role="supplier",
            counterparty_role="customer"
        )
        
        yaml = generate_yaml_header(
            document_type="test",
            party_info=party_info
        )
        
        # defined_term is redundant since we replace terms with role placeholders
        assert "defined_term:" not in yaml
    
    def test_includes_role_in_parties(self) -> None:
        """Parties mapping should include role field."""
        from scripts.generate_review_example import PartyInfo, generate_yaml_header
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_role="supplier",
            counterparty_role="customer"
        )
        
        yaml = generate_yaml_header(
            document_type="test",
            party_info=party_info
        )
        
        assert "role:" in yaml
        assert "supplier" in yaml
        assert "customer" in yaml
    
    def test_includes_original_provided_by(self) -> None:
        """YAML should include original_provided_by field."""
        from scripts.generate_review_example import PartyInfo, generate_yaml_header
        
        party_info = PartyInfo(
            client_prefix="Acme",
            counterparty_prefix="BigCorp",
            client_defined_term="Licensor",
            counterparty_defined_term="Licensee",
            original_provided_by="client",
            client_role="licensor",
            counterparty_role="licensee"
        )
        
        yaml = generate_yaml_header(
            document_type="test",
            party_info=party_info
        )
        
        assert "original_provided_by:" in yaml
        assert "client" in yaml


# =============================================================================
# Tests for anonymize_text function with defined terms
# =============================================================================

class TestAnonymizeTextPreservesDefinedTerms:
    """Tests that anonymization only replaces identifiers, not defined terms."""
    
    def test_replaces_party_identifier_not_defined_term(self) -> None:
        """Should replace 'Didimo' but NOT 'Vendor'."""
        from scripts.generate_review_example import anonymize_text
        
        text = "For Didimo: The Vendor shall provide services."
        
        result = anonymize_text(
            text,
            client_names=["Didimo"],
            counterparty_names=["NBC"]
        )
        
        assert "[CLIENT]" in result
        assert "Didimo" not in result
        assert "Vendor" in result  # Preserved!
    
    def test_replaces_counterparty_identifier_not_defined_term(self) -> None:
        """Should replace 'NBC' but NOT 'Company'."""
        from scripts.generate_review_example import anonymize_text
        
        text = "For NBC: The Company will pay within 30 days."
        
        result = anonymize_text(
            text,
            client_names=["Didimo"],
            counterparty_names=["NBC"]
        )
        
        assert "[COUNTERPARTY]" in result
        assert "NBC" not in result
        assert "Company" in result  # Preserved!
    
    def test_replaces_multiple_identifier_variations(self) -> None:
        """Should replace all variations of party identifiers."""
        from scripts.generate_review_example import anonymize_text
        
        text = "NBCUniversal and NBC agree. The Company (NBC) shall..."
        
        result = anonymize_text(
            text,
            client_names=["Didimo"],
            counterparty_names=["NBC", "NBCUniversal"]
        )
        
        assert "NBC" not in result
        assert "NBCUniversal" not in result
        assert "Company" in result  # Defined term preserved


# =============================================================================
# Tests for complete party detection workflow
# =============================================================================

class TestPartyDetectionWorkflow:
    """Integration tests for the complete party detection workflow."""
    
    def test_detect_parties_returns_party_info_with_roles(self) -> None:
        """detect_and_confirm_parties should return PartyInfo with roles."""
        from scripts.generate_review_example import (
            detect_and_confirm_parties, 
            infer_party_role
        )
        
        # Create a test document with clear defined terms
        doc = Document()
        doc.add_paragraph("This agreement is made between the parties.")
        doc.add_paragraph("Vendor shall provide services to Company.")
        doc.add_paragraph("Company will pay Vendor according to the schedule.")
        
        # Use comment prefixes that are found in the document
        comments = [
            {"text": "For Vendor: Accept this clause"},
            {"text": "For Company: Reject this clause"},
        ]
        
        # Mock user input: select first party as client, counterparty provided original
        with patch('builtins.input', side_effect=["1", "2"]):
            party_info = detect_and_confirm_parties(doc, comments)
        
        # The key behavior: roles should be inferred from defined terms
        assert party_info.client_role is not None
        assert party_info.counterparty_role is not None
        
        # Both parties should be one of Vendor/Company (order may vary)
        assert party_info.client_prefix in ["Vendor", "Company"]
        assert party_info.counterparty_prefix in ["Vendor", "Company"]
        assert party_info.client_prefix != party_info.counterparty_prefix
        
        # Roles should be inferred from the defined terms
        assert party_info.client_role == infer_party_role(party_info.client_defined_term)
        assert party_info.counterparty_role == infer_party_role(party_info.counterparty_defined_term)


# =============================================================================
# Tests for YAML output in generated markdown
# =============================================================================

class TestMarkdownYamlOutput:
    """Tests that generated markdown files have correct YAML headers."""
    
    def test_original_agreement_yaml_has_parties(self) -> None:
        """02_original_agreement.md should have parties mapping in YAML."""
        from scripts.generate_review_example import (
            generate_original_agreement_md,
            PartyInfo
        )
        
        # Create minimal docx
        doc = Document()
        doc.add_paragraph("This is a test agreement.")
        doc.add_paragraph("Vendor shall provide. Company will pay.")
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_role="supplier",
            counterparty_role="customer"
        )
        
        markdown = generate_original_agreement_md(
            docx_bytes,
            source_name="test",
            client_names=["Didimo"],
            counterparty_names=["NBC"],
            party_info=party_info
        )
        
        # Check YAML header contains parties
        assert "parties:" in markdown
        assert "defined_term:" not in markdown  # defined_term is redundant
        assert "role:" in markdown
    
    def test_comments_yaml_has_parties(self) -> None:
        """03_review_comments.md should have parties mapping in YAML."""
        from scripts.generate_review_example import (
            generate_comments_md,
            PartyInfo
        )
        
        categorized_comments = {
            "client": [{"text": "For Client: note", "author": "L", "date": "2025"}],
            "counterparty": [],
            "other": []
        }
        
        party_info = PartyInfo(
            client_prefix="Client",
            counterparty_prefix="Counter",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_role="supplier",
            counterparty_role="customer"
        )
        
        markdown = generate_comments_md(
            categorized_comments,
            client_names=["Client"],
            counterparty_names=["Counter"],
            party_info=party_info
        )
        
        assert "parties:" in markdown
        assert "supplier" in markdown
        assert "customer" in markdown
    
    def test_track_changes_yaml_has_parties(self) -> None:
        """04_tracked_changes.md should have parties mapping in YAML."""
        from scripts.generate_review_example import (
            generate_track_changes_md,
            PartyInfo
        )
        
        party_info = PartyInfo(
            client_prefix="Didimo",
            counterparty_prefix="NBC",
            client_defined_term="Vendor",
            counterparty_defined_term="Company",
            original_provided_by="counterparty",
            client_role="supplier",
            counterparty_role="customer"
        )
        
        markdown = generate_track_changes_md(
            diffs=[],
            client_names=["Didimo"],
            counterparty_names=["NBC"],
            party_info=party_info
        )
        
        assert "parties:" in markdown
        assert "SUPPLIER" in markdown  # Role placeholder, not defined term
        assert "CUSTOMER" in markdown  # Role placeholder, not defined term
