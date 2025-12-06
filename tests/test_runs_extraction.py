#!/usr/bin/env python3
"""
Tests for run extraction with revision formatting.

Sprint 3, Phase 2: Track Changes - Runs with Insert/Delete Formats

These tests verify that paragraph runs are extracted with revision information
using the text-based model:
- Normal/insert runs have 'text' field with their content
- Delete runs have 'deleted_text' field with deleted content
- Author and date information is preserved in run metadata
- Combined formatting (bold+insert, italic+delete) works correctly

The add_runs_to_block() function uses AmendedParagraph for extraction.
"""

import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))


class TestAddRunsToBlock:
    """Tests for adding runs to block dictionary."""
    
    @pytest.fixture
    def sample_doc(self):
        return Document()
    
    @pytest.fixture
    def doc_with_insert(self, sample_doc):
        """Create document with an insertion (w:ins)."""
        doc = sample_doc
        para = doc.add_paragraph()
        
        # Add normal text
        run1 = para.add_run("Normal text ")
        
        # Create w:ins element with inserted text
        p_elem = para._p
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), '1')
        ins.set(qn('w:author'), 'John Smith')
        ins.set(qn('w:date'), '2024-01-15T10:30:00Z')
        
        # Create run inside ins
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = 'inserted text'
        r.append(t)
        ins.append(r)
        p_elem.append(ins)
        
        # Add more normal text
        run3 = para.add_run(" after insert")
        
        return doc
    
    @pytest.fixture
    def doc_with_delete(self, sample_doc):
        """Create document with a deletion (w:del)."""
        doc = sample_doc
        para = doc.add_paragraph()
        
        # Add normal text
        run1 = para.add_run("Before ")
        
        # Create w:del element with deleted text
        p_elem = para._p
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), '2')
        del_elem.set(qn('w:author'), 'Jane Doe')
        del_elem.set(qn('w:date'), '2024-01-16T14:45:00Z')
        
        # Create run inside del with w:delText
        r = OxmlElement('w:r')
        del_text = OxmlElement('w:delText')
        del_text.text = 'deleted'
        r.append(del_text)
        del_elem.append(r)
        p_elem.append(del_elem)
        
        # Add more normal text
        run3 = para.add_run(" after")
        
        return doc
    
    def test_import_add_runs_to_block(self):
        """Test that add_runs_to_block can be imported."""
        try:
            from effilocal.doc.runs import add_runs_to_block
        except ImportError as e:
            pytest.fail(f"Cannot import add_runs_to_block: {e}")
    
    def test_add_runs_to_block_structure(self, sample_doc):
        """Test that runs are added with correct structure (text-based model)."""
        from effilocal.doc.runs import add_runs_to_block
        
        para = sample_doc.add_paragraph("Test text")
        block = {
            'id': 'test-001',
            'type': 'paragraph',
            'text': 'Test text'
        }
        
        add_runs_to_block(block, para)
        
        assert 'runs' in block
        assert isinstance(block['runs'], list)
        assert len(block['runs']) >= 1
        # Text-based model: runs have 'text' field, not 'start'/'end'
        assert 'text' in block['runs'][0]
        assert 'formats' in block['runs'][0]
    
    def test_add_runs_preserves_existing_block_fields(self, sample_doc):
        """Test that adding runs doesn't remove existing block fields."""
        from effilocal.doc.runs import add_runs_to_block
        
        para = sample_doc.add_paragraph("Test text")
        block = {
            'id': 'test-001',
            'type': 'paragraph',
            'text': 'Test text',
            'style': 'Normal'
        }
        
        add_runs_to_block(block, para)
        
        assert block['id'] == 'test-001'
        assert block['type'] == 'paragraph'
        assert block['text'] == 'Test text'
        assert block['style'] == 'Normal'
    
    def test_add_runs_with_insert(self, doc_with_insert):
        """Test that insert runs have 'text' field and 'insert' format."""
        from effilocal.doc.runs import add_runs_to_block
        
        para = doc_with_insert.paragraphs[0]
        block = {
            'id': 'test-002',
            'type': 'paragraph',
            'text': para.text
        }
        
        add_runs_to_block(block, para)
        
        # Find runs with 'insert' format
        insert_runs = [r for r in block['runs'] if 'insert' in r.get('formats', [])]
        assert len(insert_runs) >= 1, "Should have at least one run with 'insert' format"
        
        insert_run = insert_runs[0]
        assert 'text' in insert_run, "Insert run should have 'text' field"
        assert insert_run['text'] == 'inserted text'
        assert insert_run['author'] == 'John Smith'
    
    def test_add_runs_with_delete(self, doc_with_delete):
        """Test that delete runs have 'deleted_text' field."""
        from effilocal.doc.runs import add_runs_to_block
        
        para = doc_with_delete.paragraphs[0]
        block = {
            'id': 'test-003',
            'type': 'paragraph',
            'text': para.text
        }
        
        add_runs_to_block(block, para)
        
        # Find runs with 'delete' format
        delete_runs = [r for r in block['runs'] if 'delete' in r.get('formats', [])]
        assert len(delete_runs) >= 1, "Should have at least one run with 'delete' format"
        
        delete_run = delete_runs[0]
        assert 'deleted_text' in delete_run, "Delete run should have 'deleted_text' field"
        assert delete_run['deleted_text'] == 'deleted'
        assert delete_run['author'] == 'Jane Doe'
    
    def test_add_runs_empty_paragraph(self, sample_doc):
        """Test adding runs from empty paragraph creates default run if block has text."""
        from effilocal.doc.runs import add_runs_to_block
        
        para = sample_doc.add_paragraph()  # Empty paragraph
        block = {
            'id': 'test-004',
            'type': 'paragraph',
            'text': ''
        }
        
        add_runs_to_block(block, para)
        
        # Empty paragraph with empty text should have empty runs
        assert 'runs' in block


class TestAddRunsToBlockFromFile:
    """Tests using actual tracked changes document."""
    
    @pytest.fixture
    def tracked_doc_path(self):
        """Path to test document with track changes."""
        return Path(r"C:\Users\DavidSant\effi-contract-review\EL_Projects\Test Project\drafts\current_drafts\Norton R&D Services Agreement (DRAFT) - HJ9 (TRACKED).docx")
    
    def test_add_runs_from_real_document(self, tracked_doc_path):
        """Test adding runs from real document with track changes."""
        if not tracked_doc_path.exists():
            pytest.skip(f"Test document not found: {tracked_doc_path}")
        
        from effilocal.doc.runs import add_runs_to_block
        
        doc = Document(str(tracked_doc_path))
        
        # Find paragraphs with revisions
        all_runs = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            block = {'id': 'test', 'type': 'paragraph', 'text': para.text}
            add_runs_to_block(block, para)
            all_runs.extend(block.get('runs', []))
        
        # Should find some runs with insert or delete format
        revision_runs = [r for r in all_runs if 'insert' in r.get('formats', []) or 'delete' in r.get('formats', [])]
        
        assert len(revision_runs) > 0, "Should find revision runs in tracked document"


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
