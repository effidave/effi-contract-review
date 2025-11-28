from __future__ import annotations

import json
from pathlib import Path
from uuid import uuid4

from docx import Document

from effilocal.flows.analyze_doc import analyze


def _build_doc(tmp_path: Path) -> Path:
    doc = Document()

    doc.add_paragraph("Drafting note: Delete this clause before signing.")
    doc.add_paragraph("Note: Delete this clause once parties agree.")
    doc.add_paragraph("For negotiation only – remove before execution.")
    doc.add_paragraph("Note the supplier must deliver on time.")
    doc.add_paragraph("Regular clause text.")

    doc_path = tmp_path / "drafting_notes.docx"
    doc.save(doc_path)
    return doc_path


def test_drafting_notes_tagged(tmp_path: Path) -> None:
    doc_path = _build_doc(tmp_path)
    out_dir = tmp_path / "analysis"
    analyze(doc_path, doc_id=str(uuid4()), out_dir=out_dir)

    blocks_path = out_dir / "blocks.jsonl"
    blocks = [
        json.loads(line)
        for line in blocks_path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    roles_by_text = {block["text"]: block.get("role") for block in blocks}

    assert roles_by_text["Drafting note: Delete this clause before signing."] == "drafting_note"
    assert roles_by_text["Note: Delete this clause once parties agree."] == "drafting_note"
    assert roles_by_text["For negotiation only – remove before execution."] == "drafting_note"
    assert roles_by_text["Note the supplier must deliver on time."] is None
    assert roles_by_text["Regular clause text."] is None
