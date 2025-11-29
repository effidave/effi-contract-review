"""Test clause-based paragraph insertion tools."""
import pytest
import tempfile
import shutil
from pathlib import Path
from docx import Document
from word_document_server.tools.content_tools import (
    add_paragraph_after_clause,
    add_paragraphs_after_clause,
)


@pytest.fixture
def fixtures_dir():
    """Return path to test fixtures directory."""
    return Path(__file__).parent / "fixtures"


@pytest.fixture
def temp_numbered_doc(fixtures_dir):
    """Create a temporary copy of the numbered document for testing."""
    source = fixtures_dir / "numbering_decimal.docx"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    shutil.copy2(source, tmp_path)
    yield tmp_path
    # Cleanup
    if tmp_path.exists():
        tmp_path.unlink()


@pytest.mark.asyncio
async def test_add_paragraph_after_clause_simple(temp_numbered_doc):
    """Test adding a single paragraph after a top-level clause."""
    result = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "1",
        "New paragraph after clause 1",
        inherit_numbering=True
    )
    
    assert "Paragraph added after clause '1'" in result
    assert "with inherited numbering" in result
    
    # Verify the paragraph was added
    doc = Document(str(temp_numbered_doc))
    # Find paragraph with our new text
    found = False
    for para in doc.paragraphs:
        if "New paragraph after clause 1" in para.text:
            found = True
            break
    assert found, "New paragraph text not found in document"


@pytest.mark.asyncio
async def test_add_paragraph_after_nested_clause(temp_numbered_doc):
    """Test adding a paragraph after a nested clause like 3.1."""
    result = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "3.1",
        "New content under 3.1",
        inherit_numbering=True
    )
    
    assert "Paragraph added after clause '3.1'" in result
    
    # Verify
    doc = Document(str(temp_numbered_doc))
    found = False
    for para in doc.paragraphs:
        if "New content under 3.1" in para.text:
            found = True
            break
    assert found


@pytest.mark.asyncio
async def test_add_paragraph_after_deep_nested_clause(temp_numbered_doc):
    """Test adding a paragraph after a deeply nested clause like 3.1.1."""
    result = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "3.1.1",
        "New deep nested content",
        inherit_numbering=True
    )
    
    assert "Paragraph added after clause '3.1.1'" in result
    
    # Verify
    doc = Document(str(temp_numbered_doc))
    found = False
    for para in doc.paragraphs:
        if "New deep nested content" in para.text:
            found = True
            break
    assert found


@pytest.mark.asyncio
async def test_add_paragraph_without_inheritance(temp_numbered_doc):
    """Test adding a paragraph without inheriting numbering."""
    result = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "1",
        "Non-numbered content",
        inherit_numbering=False
    )
    
    assert "Paragraph added after clause '1'" in result
    assert "with inherited numbering" not in result


@pytest.mark.asyncio
async def test_add_paragraph_with_custom_style(temp_numbered_doc):
    """Test adding a paragraph with a custom style."""
    result = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "1",
        "Styled content",
        style="Normal",
        inherit_numbering=True
    )
    
    assert "Paragraph added" in result


@pytest.mark.asyncio
async def test_add_paragraph_clause_not_found(temp_numbered_doc):
    """Test error when clause number doesn't exist."""
    result = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "99.99.99",
        "This should fail",
        inherit_numbering=True
    )
    
    assert "not found" in result.lower()


@pytest.mark.asyncio
async def test_add_paragraphs_multiple(temp_numbered_doc):
    """Test adding multiple paragraphs after a clause."""
    paragraphs = [
        "First new item",
        "Second new item",
        "Third new item"
    ]
    
    result = await add_paragraphs_after_clause(
        str(temp_numbered_doc),
        "3.1",
        paragraphs,
        inherit_numbering=True
    )
    
    assert "Added 3 paragraph(s)" in result or "Paragraph added" in result
    
    # Verify all paragraphs were added
    doc = Document(str(temp_numbered_doc))
    found_count = 0
    for para in doc.paragraphs:
        if any(text in para.text for text in paragraphs):
            found_count += 1
    
    assert found_count == len(paragraphs), f"Expected {len(paragraphs)} paragraphs, found {found_count}"


@pytest.mark.asyncio
async def test_add_paragraphs_single(temp_numbered_doc):
    """Test that single paragraph works with the plural function."""
    result = await add_paragraphs_after_clause(
        str(temp_numbered_doc),
        "1",
        ["Single paragraph"],
        inherit_numbering=True
    )
    
    assert "Paragraph added" in result or "Added 1 paragraph" in result


@pytest.mark.asyncio
async def test_add_paragraphs_empty_list(temp_numbered_doc):
    """Test error handling for empty paragraph list."""
    result = await add_paragraphs_after_clause(
        str(temp_numbered_doc),
        "1",
        [],
        inherit_numbering=True
    )
    
    assert "must be a non-empty list" in result


@pytest.mark.asyncio
async def test_add_paragraph_nonexistent_file():
    """Test error when file doesn't exist."""
    result = await add_paragraph_after_clause(
        "nonexistent.docx",
        "1",
        "Test",
        inherit_numbering=True
    )
    
    assert "does not exist" in result


@pytest.mark.asyncio
async def test_clause_number_normalization(temp_numbered_doc):
    """Test that clause numbers are normalized (trailing dots removed)."""
    # Both "1." and "1" should find the same clause
    result1 = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "1.",
        "Test with dot",
        inherit_numbering=False
    )
    
    assert "Paragraph added" in result1
    assert "not found" not in result1.lower()


@pytest.mark.asyncio
async def test_numbering_inheritance_properties(temp_numbered_doc):
    """Test that numId and ilvl are properly inherited."""
    result = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "3.1",
        "Test inheritance",
        inherit_numbering=True
    )
    
    assert "with inherited numbering" in result
    assert "numId=" in result or "source=" in result
    assert "level=" in result or "style=" in result
