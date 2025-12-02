
"""Thin facade for parsing `.docx` files.

This module simply loads the document, instantiates `AnalysisPipeline`, and delegates
paragraph/table iteration to the pipeline. Custom callers may inject a preconfigured
pipeline via the optional `pipeline` parameters when they need finer control.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from effilocal.doc.constants import DEFAULT_FALLBACK_HEADING_LABEL
from effilocal.doc.numbering import load_numbering
from effilocal.doc.numbering_inspector import NumberingInspector
from effilocal.doc.pipeline import AnalysisPipeline, Block
from effilocal.doc.trackers import AttachmentTracker


def load_docx(path: str | Path) -> DocxDocument:
    """Load a `.docx` file into a python-docx Document."""

    path_obj = Path(path)
    return Document(str(path_obj))


def iter_paragraph_blocks(
    docx_path: str | Path,
    *,
    fallback_heading_label: str | None = DEFAULT_FALLBACK_HEADING_LABEL,
    pipeline: AnalysisPipeline | None = None,
) -> Iterator[Block]:
    """Yield paragraph and heading blocks from a `.docx` file."""

    docx_path = Path(docx_path)
    document = load_docx(docx_path)
    pipeline = pipeline or _create_pipeline(docx_path, fallback_heading_label)

    for paragraph in document.paragraphs:
        block = pipeline.process_paragraph(paragraph)
        if block is not None:
            yield block


def iter_table_blocks(
    docx_path: str | Path,
    *,
    fallback_heading_label: str | None = DEFAULT_FALLBACK_HEADING_LABEL,
    attachment_tracker: AttachmentTracker | None = None,
    pipeline: AnalysisPipeline | None = None,
) -> Iterator[Block]:
    """Yield table cell blocks from a `.docx` file."""

    if pipeline is not None and attachment_tracker is not None:
        raise ValueError("Specify either `pipeline` or `attachment_tracker`, not both.")

    docx_path = Path(docx_path)
    document = load_docx(docx_path)
    pipeline = pipeline or _create_pipeline(
        docx_path,
        fallback_heading_label,
        attachment_tracker=attachment_tracker,
    )

    for table in document.tables:
        yield from pipeline.process_table(table)


def iter_blocks(
    docx_path: str | Path,
    *,
    fallback_heading_label: str | None = DEFAULT_FALLBACK_HEADING_LABEL,
    pipeline: AnalysisPipeline | None = None,
) -> Iterator[Block]:
    """Yield all blocks (paragraphs, headings, table cells) in document order."""

    docx_path = Path(docx_path)
    document = load_docx(docx_path)
    pipeline = pipeline or _create_pipeline(docx_path, fallback_heading_label)

    for element in _iter_document_elements(document):
        if isinstance(element, Paragraph):
            block = pipeline.process_paragraph(element)
            if block is not None:
                yield block
        elif isinstance(element, Table):
            yield from pipeline.process_table(element)


def _create_pipeline(
    docx_path: Path,
    fallback_heading_label: str | None,
    *,
    attachment_tracker: AttachmentTracker | None = None,
) -> AnalysisPipeline:
    numbering_defs = load_numbering(docx_path)
    inspector = NumberingInspector.from_docx(docx_path)
    return AnalysisPipeline(
        numbering_defs=numbering_defs,
        numbering_inspector=inspector,
        fallback_heading_label=fallback_heading_label,
        attachment_tracker=attachment_tracker,
    )




def _iter_document_elements(document: DocxDocument) -> Iterator[Paragraph | Table]:
    """Iterate over document elements, including those inside SDT content controls.
    
    SDT (Structured Document Tag) elements are content controls that may wrap
    paragraphs and tables. This function unwraps them to yield the inner content.
    """
    from docx.oxml.ns import qn
    
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)
        elif child.tag == qn('w:sdt'):
            # SDT content control - look inside for paragraphs/tables
            sdt_content = child.find(qn('w:sdtContent'))
            if sdt_content is not None:
                for inner in sdt_content.iterchildren():
                    if isinstance(inner, CT_P):
                        yield Paragraph(inner, document)
                    elif isinstance(inner, CT_Tbl):
                        yield Table(inner, document)



