from __future__ import annotations

import json
from pathlib import Path
from uuid import uuid4

from docx import Document

from effilocal.flows.analyze_doc import analyze


def _build_doc(tmp_path: Path) -> Path:
    doc = Document()

    clause = doc.add_paragraph("Supplier shall provide the Services.")
    clause.style = "List Number"

    doc.add_paragraph("The services include hosting.")
    doc.add_paragraph("Support will be 24/7.")
    doc.add_paragraph("This paragraph should not be a continuation.")

    clause2 = doc.add_paragraph("Supplier shall maintain security controls.")
    clause2.style = "List Number"

    doc_path = tmp_path / "multipart.docx"
    doc.save(doc_path)
    return doc_path


def _load_blocks(out_dir: Path) -> list[dict[str, object]]:
    path = out_dir / "blocks.jsonl"
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def test_continuations_within_window(tmp_path: Path) -> None:
    doc_path = _build_doc(tmp_path)
    out_dir = tmp_path / "analysis"
    analyze(doc_path, doc_id=str(uuid4()), out_dir=out_dir)

    blocks = _load_blocks(out_dir)
    clause = next(block for block in blocks if block["text"] == "Supplier shall provide the Services.")
    cont1 = next(block for block in blocks if block["text"] == "The services include hosting.")
    cont2 = next(block for block in blocks if block["text"] == "Support will be 24/7.")
    extra = next(block for block in blocks if block["text"] == "This paragraph should not be a continuation.")

    clause_id = clause["id"]
    assert clause["clause_group_id"] == clause_id
    assert cont1["continuation_of"] == clause_id
    assert cont1["clause_group_id"] == clause_id
    assert cont2["continuation_of"] == clause_id
    assert cont2["clause_group_id"] == clause_id
    assert extra.get("continuation_of") is None
    assert extra["clause_group_id"] == extra["id"]
