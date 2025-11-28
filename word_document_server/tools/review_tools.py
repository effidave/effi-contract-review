"""
Combined review tools for Word Document Server.

These tools support common legal review workflows that bundle edits and
commenting actions into single Model Context Protocol operations.
"""
from __future__ import annotations

import os
import json
import re
from collections import defaultdict
from typing import Optional, Dict, Any, Tuple, Iterable, List
from docx import Document
from docx.text.run import Run
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.comments import CommentsPart

from word_document_server.utils.file_utils import ensure_docx_extension, check_file_writeable


# Default metadata to keep internal/external comments distinguishable
DEFAULT_INTERNAL_AUTHOR = "Internal Counsel"
DEFAULT_INTERNAL_INITIALS = "IC"
DEFAULT_EXTERNAL_AUTHOR = "Counterparty Comment"
DEFAULT_EXTERNAL_INITIALS = "EC"

EFFI_CODE_PREFIX = "EFFI-C-"
EFFI_CODE_PATTERN = re.compile(rf"{EFFI_CODE_PREFIX}[A-Za-z0-9\-]+")

TODO_INSTRUCTION_MAP = {
    "[E": "edit this section to deal with the following issue(s)",
    "[NE": "leave this section as it currently is",
    "+": "and also",
    "IC": "add an internal comment for the client to explain the following issue(s) and how the edit helps",
    "EC": "add an external comment for the counterparty to raise the issue(s) and explain the edit [n.b. the external comment should be suitable to share with a counterparty in a legal negotiation]",
    "]": "",
}
TODO_CODE_PATTERN = re.compile(r"\[(N?E)(\+IC)?(\+EC)?\]")

SUPPORTED_INSTRUCTION_CODES: Tuple[str, ...] = (
    "NE",
    "NE+IC",
    "NE+EC",
    "NE+IC+EC",
    "E",
    "E+IC",
    "E+EC",
    "E+IC+EC",
)


def _iter_run_contexts(doc: Document) -> Iterable[Tuple[Run, Dict[str, Any]]]:
    """Yield every run in the document along with contextual metadata."""
    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        for run_index, run in enumerate(paragraph.runs):
            yield run, {
                "in_table": False,
                "paragraph_index": paragraph_index,
                "run_index": run_index,
                "paragraph_text": paragraph.text,
            }

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_in_cell_index, paragraph in enumerate(cell.paragraphs):
                    for run_index, run in enumerate(paragraph.runs):
                        yield run, {
                            "in_table": True,
                            "table_index": table_index,
                            "row_index": row_index,
                            "cell_index": cell_index,
                            "paragraph_in_cell_index": paragraph_in_cell_index,
                            "run_index": run_index,
                            "paragraph_text": paragraph.text,
                        }


def _get_or_add_comments_part(doc_part):
    """Return the comments part; create it if it does not exist."""
    try:
        return doc_part.part_related_by(RT.COMMENTS)
    except KeyError:
        package = doc_part.package
        assert package is not None
        comments_part = CommentsPart.default(package)
        doc_part.relate_to(comments_part, RT.COMMENTS)
        return comments_part


def _wrap_run_with_comment(run: Run, cid: int) -> None:
    """Wrap a run with a comment range and add a comment reference."""
    run_element = run._r

    start = OxmlElement('w:commentRangeStart')
    start.set(qn('w:id'), str(cid))
    run_element.addprevious(start)

    end = OxmlElement('w:commentRangeEnd')
    end.set(qn('w:id'), str(cid))
    run_element.addnext(end)

    reference_run = OxmlElement('w:r')
    comment_ref = OxmlElement('w:commentReference')
    comment_ref.set(qn('w:id'), str(cid))
    reference_run.append(comment_ref)
    end.addnext(reference_run)


def _get_comments_part_if_exists(doc: Document):
    """Return the comments part if present; otherwise None."""
    try:
        return doc.part.part_related_by(RT.COMMENTS)
    except KeyError:
        return None


def _expand_instruction_codes(text: str) -> str:
    """Replace bracketed instruction codes with human-readable guidance."""

    def replacer(match: re.Match) -> str:
        parts: List[str] = []
        base = match.group(1)
        if base == "NE":
            parts.append(TODO_INSTRUCTION_MAP["[NE"])
        else:
            parts.append(TODO_INSTRUCTION_MAP["[E"])

        if match.group(2):
            parts.append(TODO_INSTRUCTION_MAP["+"])
            parts.append(TODO_INSTRUCTION_MAP["IC"])

        if match.group(3):
            parts.append(TODO_INSTRUCTION_MAP["+"])
            parts.append(TODO_INSTRUCTION_MAP["EC"])

        return " ".join(parts)

    return TODO_CODE_PATTERN.sub(replacer, text)


def _get_comment_ids_for_run(run: Run) -> List[str]:
    """Return the comment ids that cover the provided run."""
    parent = run._r.getparent()
    if parent is None:
        return []

    active: List[str] = []
    ids_for_run: List[str] = []
    comment_start_tag = qn('w:commentRangeStart')
    comment_end_tag = qn('w:commentRangeEnd')

    for child in parent.iterchildren():
        if child.tag == comment_start_tag:
            cid = child.get(qn('w:id'))
            if cid is not None:
                active.append(cid)
        elif child.tag == comment_end_tag:
            cid = child.get(qn('w:id'))
            if cid in active:
                active.remove(cid)
        elif child is run._r:
            ids_for_run.extend(active)
            break

    return ids_for_run


def _location_from_context(context: Dict[str, Any]) -> Dict[str, Any]:
    """Convert a run context dictionary into a concise locator."""
    location: Dict[str, Any] = {
        "in_table": context.get("in_table", False),
        "run_index": context.get("run_index"),
    }

    if context.get("in_table"):
        location.update({
            "table_index": context.get("table_index"),
            "row_index": context.get("row_index"),
            "cell_index": context.get("cell_index"),
            "paragraph_in_cell_index": context.get("paragraph_in_cell_index"),
        })
    else:
        location["paragraph_index"] = context.get("paragraph_index")

    if "match_index_in_run" in context:
        location["match_index_in_run"] = context["match_index_in_run"]
    if "match_length" in context:
        location["match_length"] = context["match_length"]

    return location


def _collect_comment_run_data(doc: Document) -> Dict[str, Dict[str, Any]]:
    """
    Collect run texts and locations associated with each comment id.

    Returns a mapping {comment_id: {"segments": [{"text": str, "location": {...}}], "paragraph_preview": str}}
    """
    comment_runs: Dict[str, Dict[str, Any]] = defaultdict(
        lambda: {
            "segments": [],
            "paragraph_preview": None,
        }
    )

    for run, context in _iter_run_contexts(doc):
        comment_ids = _get_comment_ids_for_run(run)
        if not comment_ids:
            continue

        run_text = run.text or ""
        for cid in comment_ids:
            entry = comment_runs[cid]
            entry["segments"].append({
                "text": run_text,
                "location": _location_from_context(context),
            })
            if entry["paragraph_preview"] is None:
                entry["paragraph_preview"] = context.get("paragraph_text")

    return comment_runs


def _extract_identifier_from_text(text: str) -> Optional[str]:
    match = EFFI_CODE_PATTERN.search(text or "")
    return match.group(0) if match else None


def _replace_comment_text(comment, text: str) -> None:
    paragraphs = comment.paragraphs
    if paragraphs:
        paragraphs[0].text = text
    else:
        comment.add_paragraph(text=text)


def _set_comment_text_with_identifier(comment, body_text: str) -> str:
    identifier = f"{EFFI_CODE_PREFIX}{comment.comment_id}"
    body = (body_text or "").strip()
    final_text = f"{identifier} {body}".strip()
    _replace_comment_text(comment, final_text)
    return identifier


def _ensure_comment_identifier_in_existing_comment(comment) -> Tuple[str, bool]:
    text = comment.text or ""
    identifier = _extract_identifier_from_text(text)
    updated = False

    if identifier:
        if not text.startswith(identifier):
            remaining = (text[:text.find(identifier)] + text[text.find(identifier) + len(identifier):]).strip()
            new_text = f"{identifier} {remaining}".strip()
            _replace_comment_text(comment, new_text)
            text = new_text
            updated = True
        return identifier, updated

    identifier = f"{EFFI_CODE_PREFIX}{comment.comment_id}"
    remaining = text.strip()
    new_text = f"{identifier} {remaining}".strip() if remaining else identifier
    _replace_comment_text(comment, new_text)
    return identifier, True


def _build_comment_lookup(comments_part) -> Dict[str, Dict[str, Any]]:
    """Map comment ids to their metadata (text, author, initials, timestamp)."""
    lookup: Dict[str, Dict[str, Any]] = {}
    if comments_part is None:
        return lookup

    for comment in comments_part.comments:
        timestamp = comment.timestamp
        identifier = _extract_identifier_from_text(comment.text or "") or f"{EFFI_CODE_PREFIX}{comment.comment_id}"
        lookup[str(comment.comment_id)] = {
            "comment_id": str(comment.comment_id),
            "text": comment.text,
            "author": comment.author,
            "initials": comment.initials,
            "identifier": identifier,
            "timestamp": timestamp.isoformat() if timestamp else None,
        }
    return lookup


def _find_run_matches(doc: Document, search_text: str) -> List[Tuple[Run, Dict[str, Any]]]:
    """Locate all runs containing the requested text, including multiple matches per run."""
    matches: List[Tuple[Run, Dict[str, Any]]] = []
    if not search_text:
        return matches

    for run, context in _iter_run_contexts(doc):
        run_text = run.text or ""
        start = 0
        while True:
            idx = run_text.find(search_text, start)
            if idx == -1:
                break
            augmented = dict(context)
            augmented["match_index_in_run"] = idx
            augmented["match_length"] = len(search_text)
            augmented["run_text_before"] = run_text
            matches.append((run, augmented))
            start = idx + len(search_text) if len(search_text) > 0 else idx + 1
    return matches


async def get_todo_comments(filename: str) -> str:
    """Collect TODO comments, expanding action codes and surfacing run locators."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"}, indent=2)

    try:
        doc = Document(filename)
    except Exception as exc:
        return json.dumps({"success": False, "error": f"Failed to open document: {exc}"}, indent=2)

    comments_part = _get_comments_part_if_exists(doc)
    if comments_part is None:
        return json.dumps({"success": True, "total": 0, "todos": []}, indent=2)

    comment_runs = _collect_comment_run_data(doc)
    results: List[Dict[str, Any]] = []
    doc_modified = False

    for comment in comments_part.comments:
        identifier, updated = _ensure_comment_identifier_in_existing_comment(comment)
        if updated:
            doc_modified = True

        raw_text = (comment.text or "").strip()
        if not raw_text:
            continue

        text_after_identifier = raw_text[len(identifier):].lstrip() if raw_text.startswith(identifier) else raw_text
        stripped = text_after_identifier.lstrip()
        if not stripped.upper().startswith('TODO'):
            continue

        remainder = stripped[4:]
        remainder = remainder.lstrip(" :-\u2013\u2014")
        expanded_todo = _expand_instruction_codes(remainder).strip()

        cid = str(comment.comment_id)
        run_bundle = comment_runs.get(cid, {"segments": [], "paragraph_preview": ""})
        segments = run_bundle.get("segments", [])
        paragraph_preview = run_bundle.get("paragraph_preview") or ""

        source_runs = []
        for seg in segments:
            location = seg.get("location") or {}
            source_runs.append({
                "text": seg.get("text", ""),
                "location": dict(location),
            })

        run_identifiers = [dict(seg["location"]) for seg in segments if seg.get("location")]
        source_text = ''.join(seg.get("text", "") for seg in segments).strip()

        results.append({
            "comment_id": cid,
            "identifier": identifier,
            "author": comment.author,
            "initials": comment.initials,
            "raw_comment": raw_text,
            "to_do": expanded_todo,
            "source_text": source_text,
            "source_runs": source_runs,
            "primary_identifier": run_identifiers[0] if run_identifiers else None,
            "run_identifiers": run_identifiers,
            "paragraph_preview": paragraph_preview,
        })

    if doc_modified:
        doc.save(filename)

    return json.dumps({
        "success": True,
        "total": len(results),
        "todos": results,
    }, indent=2)


def _perform_review_action(
    filename: str,
    search_text: str,
    action_code: str,
    replacement_text: Optional[str] = None,
    comment_text: Optional[str] = None,
    comment_type: Optional[str] = None,
    comment_author: Optional[str] = None,
    comment_initials: Optional[str] = None,
    comment_entries: Optional[List[Dict[str, Optional[str]]]] = None,
) -> str:
    """Execute the requested review action and return a JSON payload."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"}, indent=2)

    if not search_text:
        return json.dumps({"success": False, "error": "search_text cannot be empty"}, indent=2)

    if replacement_text is not None and replacement_text == "":
        return json.dumps({"success": False, "error": "replacement_text cannot be empty"}, indent=2)

    if comment_entries is not None and (
        comment_text is not None
        or comment_type is not None
        or comment_author is not None
        or comment_initials is not None
    ):
        return json.dumps({
            "success": False,
            "error": "Specify either single comment parameters or comment_entries, not both."
        }, indent=2)

    normalized_comment_entries: List[Dict[str, Optional[str]]] = []

    if comment_entries is not None:
        for entry in comment_entries:
            if entry is None:
                return json.dumps({"success": False, "error": "comment_entries cannot contain null entries"}, indent=2)
            text_value = (entry.get("text") or "").strip()
            if not text_value:
                return json.dumps({"success": False, "error": "Each comment entry must include non-empty text"}, indent=2)
            normalized_entry = {
                "text": text_value,
                "type": entry.get("type"),
                "author": entry.get("author"),
                "initials": entry.get("initials"),
            }
            normalized_comment_entries.append(normalized_entry)
    elif comment_text is not None:
        stripped_comment_text = comment_text.strip()
        if not stripped_comment_text:
            return json.dumps({"success": False, "error": "comment_text cannot be empty"}, indent=2)
        normalized_comment_entries.append({
            "text": stripped_comment_text,
            "type": comment_type,
            "author": comment_author,
            "initials": comment_initials,
        })

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return json.dumps({"success": False, "error": f"Cannot modify document: {error_message}"}, indent=2)

    try:
        doc = Document(filename)
    except Exception as exc:
        return json.dumps({"success": False, "error": f"Failed to open document: {exc}"}, indent=2)

    matches = _find_run_matches(doc, search_text)
    if not matches:
        return json.dumps({
            "success": False,
            "error": f"Text '{search_text}' not found in document."
        }, indent=2)

    if len(matches) > 1:
        comments_part_existing = _get_comments_part_if_exists(doc)
        comment_lookup = _build_comment_lookup(comments_part_existing)

        disambiguation_options: List[Dict[str, Any]] = []
        for option_id, (candidate_run, candidate_context) in enumerate(matches, start=1):
            comment_ids = _get_comment_ids_for_run(candidate_run)
            comment_details = [
                comment_lookup.get(cid, {"comment_id": cid})
                for cid in comment_ids
            ]
            disambiguation_options.append({
                "option_id": option_id,
                "location": _location_from_context(candidate_context),
                "run_text": candidate_context.get("run_text_before"),
                "paragraph_preview": candidate_context.get("paragraph_text"),
                "comment_ids": comment_ids,
                "comments": comment_details,
            })

        return json.dumps({
            "success": False,
            "requires_disambiguation": True,
            "matches": disambiguation_options,
            "message": (
                f"search_text matched {len(matches)} locations. "
                "Select the appropriate option (by option_id) or refine the search_text."
            ),
        }, indent=2)

    run, context = matches[0]

    comment_records: List[Dict[str, Any]] = []

    if replacement_text is not None:
        original_text = run.text or ""
        match_index = context.get("match_index_in_run", original_text.find(search_text))
        if match_index is None or match_index < 0:
            match_index = original_text.find(search_text)
        context["original_run_text"] = original_text
        run.text = (
            original_text[:match_index]
            + replacement_text
            + original_text[match_index + len(search_text):]
        )
        context["run_text_after"] = run.text
    else:
        context["original_run_text"] = run.text
        context["run_text_after"] = run.text

    if normalized_comment_entries:
        comments_part = _get_or_add_comments_part(doc.part)
        for entry in normalized_comment_entries:
            comment = comments_part.comments.add_comment(
                text="",
                author=entry.get("author") or "",
                initials=entry.get("initials") if entry.get("initials") is not None else "",
            )
            comment_id = comment.comment_id
            identifier = _set_comment_text_with_identifier(comment, entry["text"])
            _wrap_run_with_comment(run, comment_id)
            comment_records.append({
                "id": comment_id,
                "identifier": identifier,
                "audience": entry.get("type"),
                "author": entry.get("author"),
                "initials": entry.get("initials"),
                "text": entry.get("text"),
            })

    try:
        doc.save(filename)
    except Exception as exc:
        return json.dumps({"success": False, "error": f"Failed to save document: {exc}"}, indent=2)

    paragraph_preview = context.get("paragraph_text") or ""
    if len(paragraph_preview) > 160:
        paragraph_preview = paragraph_preview[:157] + "..."

    location = _location_from_context(context)

    result: Dict[str, Any] = {
        "success": True,
        "action": action_code,
        "filename": filename,
        "search_text": search_text,
        "did_edit": replacement_text is not None,
        "did_comment": bool(comment_records),
        "location": location,
        "paragraph_text_preview": paragraph_preview,
        "original_run_text": context.get("original_run_text"),
        "updated_run_text": context.get("run_text_after"),
    }

    if replacement_text is not None:
        result["replacement_text"] = replacement_text

    if comment_records:
        if len(comment_records) == 1:
            result["comment"] = comment_records[0]
        else:
            result["comments"] = comment_records

    return json.dumps(result, indent=2)


def _json_error(message: str) -> str:
    """Return a standardized JSON error payload."""
    return json.dumps({"success": False, "error": message}, indent=2)


def _canonicalize_instruction_code(instruction_code: str) -> Optional[str]:
    """Normalize instruction codes and tolerate bracketed forms."""
    if instruction_code is None:
        return None

    normalized = str(instruction_code).strip()
    if not normalized:
        return None

    if normalized.startswith("[") and normalized.endswith("]"):
        normalized = normalized[1:-1].strip()

    if not normalized:
        return None

    parts = [part for part in normalized.upper().split("+") if part]
    if not parts:
        return None

    base = parts[0]
    if base not in {"E", "NE"}:
        return None

    suffixes = {suffix for suffix in parts[1:]}
    if any(suffix not in {"IC", "EC"} for suffix in suffixes):
        return None

    canonical = base
    if "IC" in suffixes:
        canonical += "+IC"
    if "EC" in suffixes:
        canonical += "+EC"
    return canonical


def _build_internal_comment_entry(
    comment_text: Optional[str],
    comment_author: Optional[str],
    comment_initials: Optional[str],
    instruction_code: str,
) -> Tuple[Optional[Dict[str, Optional[str]]], Optional[str]]:
    stripped = (comment_text or "").strip()
    if not stripped:
        return None, _json_error(f"internal_comment_text is required for instruction {instruction_code}")

    return {
        "text": stripped,
        "type": "internal",
        "author": comment_author or DEFAULT_INTERNAL_AUTHOR,
        "initials": comment_initials or DEFAULT_INTERNAL_INITIALS,
    }, None


def _build_external_comment_entry(
    comment_text: Optional[str],
    comment_author: Optional[str],
    comment_initials: Optional[str],
    instruction_code: str,
) -> Tuple[Optional[Dict[str, Optional[str]]], Optional[str]]:
    stripped = (comment_text or "").strip()
    if not stripped:
        return None, _json_error(f"external_comment_text is required for instruction {instruction_code}")

    return {
        "text": stripped,
        "type": "external",
        "author": comment_author or DEFAULT_EXTERNAL_AUTHOR,
        "initials": comment_initials or DEFAULT_EXTERNAL_INITIALS,
    }, None


def _reject_internal_comment_params(
    instruction_code: str,
    comment_text: Optional[str],
    comment_author: Optional[str],
    comment_initials: Optional[str],
) -> Optional[str]:
    if any(value is not None for value in (comment_text, comment_author, comment_initials)):
        return _json_error(f"Instruction {instruction_code} does not accept internal comment parameters.")
    return None


def _reject_external_comment_params(
    instruction_code: str,
    comment_text: Optional[str],
    comment_author: Optional[str],
    comment_initials: Optional[str],
) -> Optional[str]:
    if any(value is not None for value in (comment_text, comment_author, comment_initials)):
        return _json_error(f"Instruction {instruction_code} does not accept external comment parameters.")
    return None


def _execute_ne(
    filename: str,
    search_text: str,
    replacement_text: Optional[str],
    internal_comment_text: Optional[str],
    external_comment_text: Optional[str],
    internal_comment_author: Optional[str],
    internal_comment_initials: Optional[str],
    external_comment_author: Optional[str],
    external_comment_initials: Optional[str],
) -> str:
    if replacement_text is not None:
        return _json_error("Instruction NE does not support replacement_text. Use an E-based code to perform edits.")

    error = _reject_internal_comment_params("NE", internal_comment_text, internal_comment_author, internal_comment_initials)
    if error:
        return error

    error = _reject_external_comment_params("NE", external_comment_text, external_comment_author, external_comment_initials)
    if error:
        return error

    return _perform_review_action(
        filename=filename,
        search_text=search_text,
        action_code="NE",
    )


def _execute_ne_ic(
    filename: str,
    search_text: str,
    replacement_text: Optional[str],
    internal_comment_text: Optional[str],
    external_comment_text: Optional[str],
    internal_comment_author: Optional[str],
    internal_comment_initials: Optional[str],
    external_comment_author: Optional[str],
    external_comment_initials: Optional[str],
) -> str:
    if replacement_text is not None:
        return _json_error("Instruction NE+IC cannot be used for edits. Use E+IC if the text should be modified.")

    error = _reject_external_comment_params("NE+IC", external_comment_text, external_comment_author, external_comment_initials)
    if error:
        return error

    internal_entry, internal_error = _build_internal_comment_entry(
        internal_comment_text,
        internal_comment_author,
        internal_comment_initials,
        "NE+IC",
    )
    if internal_error:
        return internal_error
    assert internal_entry is not None

    return _perform_review_action(
        filename=filename,
        search_text=search_text,
        action_code="NE+IC",
        comment_entries=[internal_entry],
    )


def _execute_ne_ec(
    filename: str,
    search_text: str,
    replacement_text: Optional[str],
    internal_comment_text: Optional[str],
    external_comment_text: Optional[str],
    internal_comment_author: Optional[str],
    internal_comment_initials: Optional[str],
    external_comment_author: Optional[str],
    external_comment_initials: Optional[str],
) -> str:
    if replacement_text is not None:
        return _json_error("Instruction NE+EC cannot be used for edits. Use E+EC if the text should be modified.")

    error = _reject_internal_comment_params("NE+EC", internal_comment_text, internal_comment_author, internal_comment_initials)
    if error:
        return error

    external_entry, external_error = _build_external_comment_entry(
        external_comment_text,
        external_comment_author,
        external_comment_initials,
        "NE+EC",
    )
    if external_error:
        return external_error
    assert external_entry is not None

    return _perform_review_action(
        filename=filename,
        search_text=search_text,
        action_code="NE+EC",
        comment_entries=[external_entry],
    )


def _execute_ne_ic_ec(
    filename: str,
    search_text: str,
    replacement_text: Optional[str],
    internal_comment_text: Optional[str],
    external_comment_text: Optional[str],
    internal_comment_author: Optional[str],
    internal_comment_initials: Optional[str],
    external_comment_author: Optional[str],
    external_comment_initials: Optional[str],
) -> str:
    if replacement_text is not None:
        return _json_error("Instruction NE+IC+EC cannot be used for edits. Use E+IC+EC if the text should be modified.")

    internal_entry, internal_error = _build_internal_comment_entry(
        internal_comment_text,
        internal_comment_author,
        internal_comment_initials,
        "NE+IC+EC",
    )
    if internal_error:
        return internal_error
    assert internal_entry is not None

    external_entry, external_error = _build_external_comment_entry(
        external_comment_text,
        external_comment_author,
        external_comment_initials,
        "NE+IC+EC",
    )
    if external_error:
        return external_error
    assert external_entry is not None

    return _perform_review_action(
        filename=filename,
        search_text=search_text,
        action_code="NE+IC+EC",
        comment_entries=[internal_entry, external_entry],
    )


def _execute_e(
    filename: str,
    search_text: str,
    replacement_text: Optional[str],
    internal_comment_text: Optional[str],
    external_comment_text: Optional[str],
    internal_comment_author: Optional[str],
    internal_comment_initials: Optional[str],
    external_comment_author: Optional[str],
    external_comment_initials: Optional[str],
) -> str:
    if replacement_text is None:
        return _json_error("replacement_text is required for instruction E.")

    error = _reject_internal_comment_params("E", internal_comment_text, internal_comment_author, internal_comment_initials)
    if error:
        return error

    error = _reject_external_comment_params("E", external_comment_text, external_comment_author, external_comment_initials)
    if error:
        return error

    return _perform_review_action(
        filename=filename,
        search_text=search_text,
        action_code="E",
        replacement_text=replacement_text,
    )


def _execute_e_ic(
    filename: str,
    search_text: str,
    replacement_text: Optional[str],
    internal_comment_text: Optional[str],
    external_comment_text: Optional[str],
    internal_comment_author: Optional[str],
    internal_comment_initials: Optional[str],
    external_comment_author: Optional[str],
    external_comment_initials: Optional[str],
) -> str:
    if replacement_text is None:
        return _json_error("replacement_text is required for instruction E+IC.")

    error = _reject_external_comment_params("E+IC", external_comment_text, external_comment_author, external_comment_initials)
    if error:
        return error

    internal_entry, internal_error = _build_internal_comment_entry(
        internal_comment_text,
        internal_comment_author,
        internal_comment_initials,
        "E+IC",
    )
    if internal_error:
        return internal_error
    assert internal_entry is not None

    return _perform_review_action(
        filename=filename,
        search_text=search_text,
        action_code="E+IC",
        replacement_text=replacement_text,
        comment_entries=[internal_entry],
    )


def _execute_e_ec(
    filename: str,
    search_text: str,
    replacement_text: Optional[str],
    internal_comment_text: Optional[str],
    external_comment_text: Optional[str],
    internal_comment_author: Optional[str],
    internal_comment_initials: Optional[str],
    external_comment_author: Optional[str],
    external_comment_initials: Optional[str],
) -> str:
    if replacement_text is None:
        return _json_error("replacement_text is required for instruction E+EC.")

    error = _reject_internal_comment_params("E+EC", internal_comment_text, internal_comment_author, internal_comment_initials)
    if error:
        return error

    external_entry, external_error = _build_external_comment_entry(
        external_comment_text,
        external_comment_author,
        external_comment_initials,
        "E+EC",
    )
    if external_error:
        return external_error
    assert external_entry is not None

    return _perform_review_action(
        filename=filename,
        search_text=search_text,
        action_code="E+EC",
        replacement_text=replacement_text,
        comment_entries=[external_entry],
    )


def _execute_e_ic_ec(
    filename: str,
    search_text: str,
    replacement_text: Optional[str],
    internal_comment_text: Optional[str],
    external_comment_text: Optional[str],
    internal_comment_author: Optional[str],
    internal_comment_initials: Optional[str],
    external_comment_author: Optional[str],
    external_comment_initials: Optional[str],
) -> str:
    if replacement_text is None:
        return _json_error("replacement_text is required for instruction E+IC+EC.")

    internal_entry, internal_error = _build_internal_comment_entry(
        internal_comment_text,
        internal_comment_author,
        internal_comment_initials,
        "E+IC+EC",
    )
    if internal_error:
        return internal_error
    assert internal_entry is not None

    external_entry, external_error = _build_external_comment_entry(
        external_comment_text,
        external_comment_author,
        external_comment_initials,
        "E+IC+EC",
    )
    if external_error:
        return external_error
    assert external_entry is not None

    return _perform_review_action(
        filename=filename,
        search_text=search_text,
        action_code="E+IC+EC",
        replacement_text=replacement_text,
        comment_entries=[internal_entry, external_entry],
    )


INSTRUCTION_EXECUTORS = {
    "NE": _execute_ne,
    "NE+IC": _execute_ne_ic,
    "NE+EC": _execute_ne_ec,
    "NE+IC+EC": _execute_ne_ic_ec,
    "E": _execute_e,
    "E+IC": _execute_e_ic,
    "E+EC": _execute_e_ec,
    "E+IC+EC": _execute_e_ic_ec,
}


async def execute_review_instruction(
    filename: str,
    search_text: str,
    instruction_code: str,
    replacement_text: Optional[str] = None,
    internal_comment_text: Optional[str] = None,
    external_comment_text: Optional[str] = None,
    internal_comment_author: Optional[str] = None,
    internal_comment_initials: Optional[str] = None,
    external_comment_author: Optional[str] = None,
    external_comment_initials: Optional[str] = None,
) -> str:
    """Execute a combined edit/comment workflow based on a TODO instruction code."""
    canonical_code = _canonicalize_instruction_code(instruction_code)
    if canonical_code is None or canonical_code not in INSTRUCTION_EXECUTORS:
        supported = ", ".join(SUPPORTED_INSTRUCTION_CODES)
        return _json_error(
            f"Unsupported instruction_code '{instruction_code}'. Supported values: {supported}"
        )

    executor = INSTRUCTION_EXECUTORS[canonical_code]
    return executor(
        filename=filename,
        search_text=search_text,
        replacement_text=replacement_text,
        internal_comment_text=internal_comment_text,
        external_comment_text=external_comment_text,
        internal_comment_author=internal_comment_author,
        internal_comment_initials=internal_comment_initials,
        external_comment_author=external_comment_author,
        external_comment_initials=external_comment_initials,
    )
