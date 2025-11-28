from __future__ import annotations

import json
from pathlib import Path
from uuid import uuid4

from docx import Document

from effilocal.flows.analyze_doc import analyze


def _build_doc(tmp_path: Path) -> Path:
    doc = Document()

    heading = doc.add_paragraph("Agreement Details")
    heading.style = "Heading 1"
    doc.add_paragraph("Summary of the arrangement.")

    heading = doc.add_paragraph("Background")
    heading.style = "Heading 1"
    doc.add_paragraph("Background information.")

    heading = doc.add_paragraph("Effective Date")
    heading.style = "Heading 1"
    doc.add_paragraph("Effective on 1 January.")

    heading = doc.add_paragraph("Parties")
    heading.style = "Heading 1"
    doc.add_paragraph("Party information.")

    heading = doc.add_paragraph("Signatures")
    heading.style = "Heading 1"
    doc.add_paragraph("Signature page content.")

    heading = doc.add_paragraph("1 Scope")
    heading.style = "Heading 1"
    doc.add_paragraph("Scope clause content.")

    heading = doc.add_paragraph("Definitions")
    heading.style = "Heading 1"
    doc.add_paragraph("Definitions content.")

    doc_path = tmp_path / "section_roles.docx"
    doc.save(doc_path)
    return doc_path


def test_section_roles_assigned(tmp_path: Path) -> None:
    doc_path = _build_doc(tmp_path)
    out_dir = tmp_path / "analysis"
    doc_id = str(uuid4())

    analyze(doc_path, doc_id=doc_id, out_dir=out_dir)

    sections_path = out_dir / "sections.json"
    payload = json.loads(sections_path.read_text(encoding="utf-8"))
    root_children = payload["root"]["children"]

    roles_by_title = {child["title"]: child.get("role") for child in root_children}

    assert roles_by_title["Agreement Details"] == "order_details"
    assert roles_by_title["Background"] == "front_matter"
    assert roles_by_title["Effective Date"] == "agreement_date"
    assert roles_by_title["Parties"] == "parties"
    assert roles_by_title["Signatures"] == "signatures"
    assert roles_by_title["1 Scope"] == "main_body"
    assert roles_by_title["Definitions"] == "definitions"

    main_body_count = sum(1 for role in roles_by_title.values() if role == "main_body")
    assert main_body_count == 1
