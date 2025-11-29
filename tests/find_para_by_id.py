"""Find paragraph by para_id."""
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

doc = Document(Path(__file__).parent / "fixtures" / "real_world" / "schedules.docx")

target_para_id = "0407B052"
print(f"Looking for paragraph with para_id={target_para_id}:")

for i, p in enumerate(doc.paragraphs):
    pid = p._element.get(qn('w14:paraId'))
    if pid == target_para_id:
        print(f"  Para {i}: para_id={pid}")
        print(f"  Text: {p.text[:80]}")
        
        # Check the next paragraph
        if i + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i + 1]
            next_pid = next_para._element.get(qn('w14:paraId'))
            print(f"\n  Next para {i+1}: para_id={next_pid}")
            print(f"  Text: {next_para.text[:80]}")
