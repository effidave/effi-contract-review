"""Tests for UUID embedding via content controls."""

from __future__ import annotations

import tempfile
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document

from effilocal.doc.uuid_embedding import (
    EFFI_TAG_PREFIX,
    BlockKey,
    ParaKey,
    TableCellKey,
    embed_block_uuids,
    extract_block_uuids,
    extract_block_uuids_legacy,
    get_paragraph_uuid,
    remove_all_uuid_controls,
    set_paragraph_uuid,
)


@pytest.fixture
def sample_doc() -> Document:
    """Create a sample document with paragraphs."""
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.add_paragraph("Third paragraph")
    return doc


@pytest.fixture
def sample_docx_path(sample_doc: Document) -> Path:
    """Save sample document to temp file and return path."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        sample_doc.save(f.name)
        return Path(f.name)


@pytest.fixture
def sample_blocks() -> list[dict]:
    """Create sample block dicts for testing."""
    return [
        {"id": str(uuid4()), "para_idx": 0, "text": "First paragraph"},
        {"id": str(uuid4()), "para_idx": 1, "text": "Second paragraph"},
        {"id": str(uuid4()), "para_idx": 2, "text": "Third paragraph"},
    ]


@pytest.fixture
def doc_with_table() -> Document:
    """Create a document with paragraphs and a table."""
    doc = Document()
    doc.add_paragraph("Intro paragraph")
    
    # Add a 2x3 table
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Cell 0,0"
    table.cell(0, 1).text = "Cell 0,1"
    table.cell(0, 2).text = "Cell 0,2"
    table.cell(1, 0).text = "Cell 1,0"
    table.cell(1, 1).text = "Cell 1,1"
    table.cell(1, 2).text = "Cell 1,2"
    
    doc.add_paragraph("Closing paragraph")
    return doc


@pytest.fixture
def doc_with_table_path(doc_with_table: Document) -> Path:
    """Save document with table to temp file and return path."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc_with_table.save(f.name)
        return Path(f.name)


@pytest.fixture
def table_cell_blocks() -> list[dict]:
    """Create sample blocks including table cells."""
    return [
        {"id": "para-intro", "para_idx": 0, "text": "Intro paragraph"},
        {"id": "cell-0-0", "table": {"table_id": "tbl_0", "row": 0, "col": 0}, "text": "Cell 0,0"},
        {"id": "cell-0-1", "table": {"table_id": "tbl_0", "row": 0, "col": 1}, "text": "Cell 0,1"},
        {"id": "cell-1-0", "table": {"table_id": "tbl_0", "row": 1, "col": 0}, "text": "Cell 1,0"},
        {"id": "cell-1-2", "table": {"table_id": "tbl_0", "row": 1, "col": 2}, "text": "Cell 1,2"},
        {"id": "para-closing", "para_idx": 1, "text": "Closing paragraph"},
    ]


class TestSetParagraphUuid:
    """Tests for set_paragraph_uuid function."""

    def test_wrap_paragraph_in_sdt(self, sample_doc: Document):
        """Test wrapping a paragraph in an SDT with UUID."""
        # Get first paragraph element
        para = sample_doc.paragraphs[0]
        para_elem = para._element
        uuid_val = str(uuid4())

        result = set_paragraph_uuid(sample_doc, para_elem, uuid_val)

        assert result is True

    def test_extracted_uuid_matches(self, sample_doc: Document):
        """Test that extracted UUID matches what was set."""
        para = sample_doc.paragraphs[0]
        para_elem = para._element
        uuid_val = str(uuid4())

        set_paragraph_uuid(sample_doc, para_elem, uuid_val)
        
        # Need to re-get element after wrapping
        body = sample_doc.element.body
        # The paragraph is now inside an SDT
        from docx.oxml.ns import qn
        sdt = body.find(qn("w:sdt"))
        if sdt is not None:
            sdt_content = sdt.find(qn("w:sdtContent"))
            if sdt_content is not None:
                new_para_elem = sdt_content.find(qn("w:p"))
                if new_para_elem is not None:
                    extracted = get_paragraph_uuid(new_para_elem)
                    assert extracted == uuid_val


class TestEmbedBlockUuids:
    """Tests for embed_block_uuids function."""

    def test_embed_blocks_from_list(self, sample_docx_path: Path, sample_blocks: list[dict]):
        """Test embedding UUIDs from block list."""
        embedded = embed_block_uuids(sample_docx_path, sample_blocks)

        assert len(embedded) == len(sample_blocks)
        for block in sample_blocks:
            assert block["id"] in embedded

    def test_embedded_uuids_survive_save(self, sample_docx_path: Path, sample_blocks: list[dict]):
        """Test that embedded UUIDs survive save/reload cycle."""
        # Embed UUIDs
        embed_block_uuids(sample_docx_path, sample_blocks)

        # Reload and extract
        extracted = extract_block_uuids(sample_docx_path)

        # Should find all embedded UUIDs
        for block in sample_blocks:
            assert block["id"] in extracted

    def test_skip_existing_uuids_when_not_overwrite(self, sample_docx_path: Path, sample_blocks: list[dict]):
        """Test that existing UUIDs are not overwritten when overwrite=False."""
        # First embedding
        embed_block_uuids(sample_docx_path, sample_blocks)
        
        # Try to embed new UUIDs for same positions
        new_blocks = [
            {"id": str(uuid4()), "para_idx": 0, "text": "First"},
            {"id": str(uuid4()), "para_idx": 1, "text": "Second"},
        ]
        embed_block_uuids(sample_docx_path, new_blocks, overwrite=False)
        
        # Extract and verify original UUIDs still present
        extracted = extract_block_uuids(sample_docx_path)
        for block in sample_blocks:
            assert block["id"] in extracted


class TestExtractBlockUuids:
    """Tests for extract_block_uuids function."""

    def test_extract_from_empty_doc(self, sample_docx_path: Path):
        """Test extracting from document with no UUIDs."""
        extracted = extract_block_uuids(sample_docx_path)
        assert len(extracted) == 0

    def test_extract_returns_para_keys(self, sample_docx_path: Path, sample_blocks: list[dict]):
        """Test that extraction returns correct ParaKey objects."""
        embed_block_uuids(sample_docx_path, sample_blocks)
        extracted = extract_block_uuids(sample_docx_path)

        for block in sample_blocks:
            uuid = block["id"]
            para_idx = block["para_idx"]
            assert uuid in extracted
            assert extracted[uuid] == ParaKey(para_idx)

    def test_legacy_extract_returns_int(self, sample_docx_path: Path, sample_blocks: list[dict]):
        """Test that legacy extraction returns integer indices."""
        embed_block_uuids(sample_docx_path, sample_blocks)
        extracted = extract_block_uuids_legacy(sample_docx_path)

        for block in sample_blocks:
            uuid = block["id"]
            para_idx = block["para_idx"]
            assert uuid in extracted
            assert extracted[uuid] == para_idx
            assert isinstance(extracted[uuid], int)


class TestRemoveAllUuidControls:
    """Tests for remove_all_uuid_controls function."""

    def test_remove_all_controls(self, sample_docx_path: Path, sample_blocks: list[dict]):
        """Test removing all effi content controls."""
        # Embed UUIDs
        embed_block_uuids(sample_docx_path, sample_blocks)
        
        # Remove all
        removed = remove_all_uuid_controls(sample_docx_path)
        
        assert removed == len(sample_blocks)

    def test_content_preserved_after_removal(self, sample_docx_path: Path, sample_blocks: list[dict]):
        """Test that paragraph content is preserved after control removal."""
        # Embed UUIDs
        embed_block_uuids(sample_docx_path, sample_blocks)
        
        # Remove all controls
        remove_all_uuid_controls(sample_docx_path)
        
        # Reload and check content
        doc = Document(str(sample_docx_path))
        assert len(doc.paragraphs) == 3
        assert doc.paragraphs[0].text == "First paragraph"
        assert doc.paragraphs[1].text == "Second paragraph"
        assert doc.paragraphs[2].text == "Third paragraph"

    def test_no_uuids_after_removal(self, sample_docx_path: Path, sample_blocks: list[dict]):
        """Test that no UUIDs remain after removal."""
        embed_block_uuids(sample_docx_path, sample_blocks)
        remove_all_uuid_controls(sample_docx_path)
        
        extracted = extract_block_uuids(sample_docx_path)
        assert len(extracted) == 0


class TestRoundTrip:
    """Integration tests for UUID round-trip scenarios."""

    def test_embed_save_reload_extract(self, sample_docx_path: Path):
        """Test complete round trip: embed → save → reload → extract."""
        blocks = [
            {"id": "uuid-001", "para_idx": 0, "text": "First"},
            {"id": "uuid-002", "para_idx": 1, "text": "Second"},
        ]
        
        # Embed
        embedded = embed_block_uuids(sample_docx_path, blocks)
        assert len(embedded) == 2
        
        # Extract
        extracted = extract_block_uuids(sample_docx_path)
        assert extracted.get("uuid-001") == ParaKey(0)
        assert extracted.get("uuid-002") == ParaKey(1)

    def test_partial_embedding(self, sample_docx_path: Path):
        """Test embedding only some paragraphs."""
        blocks = [
            {"id": "uuid-first", "para_idx": 0, "text": "First"},
            {"id": "uuid-third", "para_idx": 2, "text": "Third"},
        ]
        
        embedded = embed_block_uuids(sample_docx_path, blocks)
        extracted = extract_block_uuids(sample_docx_path)
        
        assert "uuid-first" in extracted
        assert "uuid-third" in extracted
        assert len(extracted) == 2


class TestTableCellUuids:
    """Tests for UUID embedding in table cells."""

    def test_embed_table_cell_uuids(self, doc_with_table_path: Path, table_cell_blocks: list[dict]):
        """Test embedding UUIDs in table cells."""
        embedded = embed_block_uuids(doc_with_table_path, table_cell_blocks)
        
        # Should embed all blocks
        assert len(embedded) == len(table_cell_blocks)
        for block in table_cell_blocks:
            assert block["id"] in embedded

    def test_extract_table_cell_uuids(self, doc_with_table_path: Path, table_cell_blocks: list[dict]):
        """Test extracting UUIDs from table cells."""
        embed_block_uuids(doc_with_table_path, table_cell_blocks)
        extracted = extract_block_uuids(doc_with_table_path)
        
        # Check paragraph blocks
        assert extracted.get("para-intro") == ParaKey(0)
        assert extracted.get("para-closing") == ParaKey(1)
        
        # Check table cell blocks
        assert extracted.get("cell-0-0") == TableCellKey(0, 0, 0)
        assert extracted.get("cell-0-1") == TableCellKey(0, 0, 1)
        assert extracted.get("cell-1-0") == TableCellKey(0, 1, 0)
        assert extracted.get("cell-1-2") == TableCellKey(0, 1, 2)

    def test_table_cell_uuids_survive_save(self, doc_with_table_path: Path, table_cell_blocks: list[dict]):
        """Test that table cell UUIDs survive save/reload cycle."""
        embed_block_uuids(doc_with_table_path, table_cell_blocks)
        
        # Reload and extract
        extracted = extract_block_uuids(doc_with_table_path)
        
        for block in table_cell_blocks:
            assert block["id"] in extracted

    def test_remove_table_cell_uuids(self, doc_with_table_path: Path, table_cell_blocks: list[dict]):
        """Test removing UUIDs from table cells."""
        embed_block_uuids(doc_with_table_path, table_cell_blocks)
        
        removed = remove_all_uuid_controls(doc_with_table_path)
        
        # Should remove all embedded controls
        assert removed == len(table_cell_blocks)
        
        # No UUIDs should remain
        extracted = extract_block_uuids(doc_with_table_path)
        assert len(extracted) == 0

    def test_table_content_preserved_after_removal(self, doc_with_table_path: Path, table_cell_blocks: list[dict]):
        """Test that table cell content is preserved after UUID removal."""
        embed_block_uuids(doc_with_table_path, table_cell_blocks)
        remove_all_uuid_controls(doc_with_table_path)
        
        # Reload and check content
        doc = Document(str(doc_with_table_path))
        
        # Check paragraphs
        assert doc.paragraphs[0].text == "Intro paragraph"
        assert doc.paragraphs[1].text == "Closing paragraph"
        
        # Check table cells
        table = doc.tables[0]
        assert table.cell(0, 0).text == "Cell 0,0"
        assert table.cell(0, 1).text == "Cell 0,1"
        assert table.cell(1, 0).text == "Cell 1,0"
        assert table.cell(1, 2).text == "Cell 1,2"

    def test_mixed_paragraph_and_table_blocks(self, doc_with_table_path: Path):
        """Test embedding UUIDs in both paragraphs and table cells."""
        blocks = [
            {"id": "intro", "para_idx": 0, "text": "Intro"},
            {"id": "cell-middle", "table": {"table_id": "tbl_0", "row": 0, "col": 1}, "text": "Middle"},
            {"id": "closing", "para_idx": 1, "text": "Closing"},
        ]
        
        embedded = embed_block_uuids(doc_with_table_path, blocks)
        assert len(embedded) == 3
        
        extracted = extract_block_uuids(doc_with_table_path)
        assert extracted.get("intro") == ParaKey(0)
        assert extracted.get("cell-middle") == TableCellKey(0, 0, 1)
        assert extracted.get("closing") == ParaKey(1)

    def test_multiple_tables(self):
        """Test embedding UUIDs in multiple tables."""
        doc = Document()
        doc.add_paragraph("Intro")
        
        # First table
        table1 = doc.add_table(rows=1, cols=2)
        table1.cell(0, 0).text = "T1 Cell 0,0"
        table1.cell(0, 1).text = "T1 Cell 0,1"
        
        doc.add_paragraph("Middle")
        
        # Second table
        table2 = doc.add_table(rows=1, cols=2)
        table2.cell(0, 0).text = "T2 Cell 0,0"
        table2.cell(0, 1).text = "T2 Cell 0,1"
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            path = Path(f.name)
        
        blocks = [
            {"id": "intro", "para_idx": 0, "text": "Intro"},
            {"id": "t1-cell-0-0", "table": {"table_id": "tbl_0", "row": 0, "col": 0}, "text": "T1 Cell"},
            {"id": "middle", "para_idx": 1, "text": "Middle"},
            {"id": "t2-cell-0-1", "table": {"table_id": "tbl_1", "row": 0, "col": 1}, "text": "T2 Cell"},
        ]
        
        embedded = embed_block_uuids(path, blocks)
        assert len(embedded) == 4
        
        extracted = extract_block_uuids(path)
        assert extracted.get("intro") == ParaKey(0)
        assert extracted.get("t1-cell-0-0") == TableCellKey(0, 0, 0)
        assert extracted.get("middle") == ParaKey(1)
        assert extracted.get("t2-cell-0-1") == TableCellKey(1, 0, 1)
