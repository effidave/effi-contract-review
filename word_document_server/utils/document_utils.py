"""
Document utility functions for Word Document Server.
"""
import json
from typing import Dict, List, Any
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn


def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """Get properties of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    try:
        doc = Document(doc_path)
        core_props = doc.core_properties

        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


def edit_run_text(doc_path: str, paragraph_index: int, run_index: int, new_text: str, start_offset: int = None, end_offset: int = None) -> str:
    """
    Edit text within a specific run of a paragraph.
    
    Args:
        doc_path: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        run_index: Index of the run within the paragraph (0-based)
        new_text: Text to insert/replace
        start_offset: Optional start position within the run (0-based)
        end_offset: Optional end position within the run (0-based, exclusive)
                   If not provided, replaces the entire run text
        
    Returns:
        Status message
    """
    import os
    from docx import Document
    from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
    
    doc_path = ensure_docx_extension(doc_path)
    
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    is_writeable, error_message = check_file_writeable(doc_path)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    try:
        doc = Document(doc_path)
        
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
        
        para = doc.paragraphs[paragraph_index]
        
        if run_index < 0 or run_index >= len(para.runs):
            return f"Invalid run index: {run_index}. Paragraph {paragraph_index} has {len(para.runs)} runs."
        
        run = para.runs[run_index]
        old_text = run.text
        
        # Default to replacing entire run
        if start_offset is None:
            start_offset = 0
        if end_offset is None:
            end_offset = len(old_text)
        
        # Validate offsets
        if start_offset < 0 or start_offset > len(old_text):
            return f"Invalid start_offset: {start_offset}. Run text length is {len(old_text)}."
        if end_offset < 0 or end_offset > len(old_text):
            return f"Invalid end_offset: {end_offset}. Run text length is {len(old_text)}."
        if start_offset > end_offset:
            return f"Invalid offsets: start_offset ({start_offset}) > end_offset ({end_offset})."
        
        # Perform the replacement
        run.text = old_text[:start_offset] + new_text + old_text[end_offset:]
        
        doc.save(doc_path)
        
        return f"Run {run_index} in paragraph {paragraph_index} updated: '{old_text[start_offset:end_offset]}' → '{new_text}'"
    
    except Exception as e:
        return f"Failed to edit run: {str(e)}"


def extract_document_text(doc_path: str) -> str:
    """Extract all text from a Word document."""
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text.append(paragraph.text)
        
        return "\n".join(text)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append({
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal"
            })
        
        # Get tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # Get sample of table data
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}


def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.
    
    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text
        
    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)
            
    return matching_paragraphs


def find_and_replace_text(doc, old_text, new_text, whole_word_only=False):
    """
    Find and replace text throughout the document, skipping Table of Contents (TOC) paragraphs.
    
    **BEHAVIOR**: This function replaces text ONLY when it exists entirely within a single run.
    If search text spans multiple runs exactly (with whitespace matching), it is detected and 
    reported as a split match but NOT automatically replaced. The caller must use edit_run_text 
    to handle those cases precisely.
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        whole_word_only: If True, only replace whole words (surrounded by spaces, punctuation, or boundaries)
        
    Returns:
        Tuple of (replaced_count, snippets, split_matches) where:
        - replaced_count: Number of in-run replacements made
        - snippets: List of dicts with 'before', 'after', 'location' for replaced occurrences
        - split_matches: List of dicts describing matches that span multiple runs
    """
    import re
    
    count = 0
    snippets = []
    split_matches = []
    
    def should_replace(text, search_text, pos):
        """Check if replacement at this position should occur (for whole_word_only mode)."""
        if not whole_word_only:
            return True
        
        # Check character before
        if pos > 0:
            before_char = text[pos - 1]
            if before_char.isalnum() or before_char == '_':
                return False
        
        # Check character after
        end_pos = pos + len(search_text)
        if end_pos < len(text):
            after_char = text[end_pos]
            if after_char.isalnum() or after_char == '_':
                return False
        
        return True
    
    def build_run_map(runs):
        """Build a map of run boundaries: (run_index, start_char, end_char, text)."""
        run_map = []
        char_count = 0
        for run_idx, run in enumerate(runs):
            run_len = len(run.text)
            run_map.append((run_idx, char_count, char_count + run_len, run.text))
            char_count += run_len
        return run_map
    
    def find_runs_for_match(run_map, match_start, match_end):
        """Find which runs a match spans. Returns list of (run_idx, offset_in_run_start, offset_in_run_end)."""
        spans = []
        for run_idx, run_start, run_end, _ in run_map:
            # Check overlap
            overlap_start = max(match_start, run_start)
            overlap_end = min(match_end, run_end)
            if overlap_start < overlap_end:
                # This run contains part of the match
                offset_in_start = overlap_start - run_start
                offset_in_end = overlap_end - run_start
                spans.append((run_idx, offset_in_start, offset_in_end))
        return spans
    
    def replace_in_paragraph(para, old_text, new_text, whole_word_only, location, para_idx):
        """
        Find and replace text in a paragraph. Only replace in-run matches; report split matches.
        
        Returns: (count, snippets, split_matches) for this paragraph
        """
        local_count = 0
        local_snippets = []
        local_split_matches = []
        
        full_text = para.text
        if old_text not in full_text:
            return local_count, local_snippets, local_split_matches
        
        runs = para.runs
        if not runs:
            return local_count, local_snippets, local_split_matches
        
        # Find all occurrences
        pos = 0
        while True:
            pos = full_text.find(old_text, pos)
            if pos == -1:
                break
            
            if not should_replace(full_text, old_text, pos):
                pos += 1
                continue
            
            # Rebuild run_map after each replacement to account for offset changes
            run_map = build_run_map(runs)
            
            # Determine which runs this match spans
            match_start = pos
            match_end = pos + len(old_text)
            run_spans = find_runs_for_match(run_map, match_start, match_end)
            
            if len(run_spans) == 1:
                # Match is entirely within one run - replace it
                run_idx, offset_start, offset_end = run_spans[0]
                run = runs[run_idx]
                run.text = run.text[:offset_start] + new_text + run.text[offset_end:]
                
                local_count += 1
                local_snippets.append({
                    "before": full_text[max(0, pos - 20):min(len(full_text), pos + len(old_text) + 20)],
                    "after": run.text[max(0, offset_start - 20):min(len(run.text), offset_start + len(new_text) + 20)],
                    "location": location
                })
                
                # Update full_text for next iteration
                full_text = full_text[:pos] + new_text + full_text[match_end:]
                pos += len(new_text)
            else:
                # Match spans multiple runs - verify concatenated text equals old_text exactly
                concat_text = ""
                run_details = []
                for run_idx, offset_start, offset_end in run_spans:
                    _, _, _, run_text = run_map[run_idx]
                    fragment = run_text[offset_start:offset_end]
                    concat_text += fragment
                    run_details.append({
                        "run_index": run_idx,
                        "offset_start": offset_start,
                        "offset_end": offset_end,
                        "text": fragment
                    })
                
                # Only report as split match if concatenation exactly matches
                if concat_text == old_text:
                    local_split_matches.append({
                        "paragraph_index": para_idx,
                        "runs": run_details,
                        "full_match": old_text
                    })
                
                pos += 1
        
        return local_count, local_snippets, local_split_matches
    
    # Search in paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        # Skip TOC paragraphs
        if para.style and para.style.name.startswith("TOC"):
            continue
        
        para_count, para_snippets, para_splits = replace_in_paragraph(para, old_text, new_text, whole_word_only, "paragraph", para_idx)
        count += para_count
        snippets.extend(para_snippets)
        split_matches.extend(para_splits)
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for cell_para_idx, para in enumerate(cell.paragraphs):
                    # Skip TOC paragraphs in tables
                    if para.style and para.style.name.startswith("TOC"):
                        continue
                    
                    para_count, para_snippets, para_splits = replace_in_paragraph(para, old_text, new_text, whole_word_only, "table", cell_para_idx)
                    count += para_count
                    snippets.extend(para_snippets)
                    split_matches.extend(para_splits)
    
    return count, snippets, split_matches


def get_document_xml(doc_path: str) -> str:
    """Extract and return the raw XML structure of the Word document (word/document.xml)."""
    import os
    import zipfile
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        with zipfile.ZipFile(doc_path) as docx_zip:
            with docx_zip.open('word/document.xml') as xml_file:
                return xml_file.read().decode('utf-8')
    except Exception as e:
        return f"Failed to extract XML: {str(e)}"


def insert_header_near_text(doc_path: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search."""
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        new_para = doc.add_paragraph(header_title, style=header_style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} paragraph (index {anchor_index})."
        else:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert header: {str(e)}"


def insert_line_or_paragraph_near_text(doc_path: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """
    Insert a new line or paragraph (with specified or matched style) before or after the target paragraph.
    You can specify the target by text (first match) or by paragraph index.
    Skips paragraphs whose style name starts with 'TOC' if using text search.
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Determine style: use provided or match target
        style = line_style if line_style else para.style
        new_para = doc.add_paragraph(line_text, style=style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Line/paragraph inserted {position} paragraph (index {anchor_index}) with style '{style}'."
        else:
            return f"Line/paragraph inserted {position} the target paragraph with style '{style}'."
    except Exception as e:
        return f"Failed to insert line/paragraph: {str(e)}"


def insert_numbered_list_near_text(
    doc_path: str,
    target_text: str = None,
    list_items: list = None,
    position: str = 'after',
    target_paragraph_index: int = None,
    list_style: str = 'bullet',
) -> str:
    """
    Insert a numbered list before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search.
    Args:
        doc_path: Path to the Word document
        target_text: Text to search for in paragraphs (optional if using index)
        list_items: List of strings, each as a list item
        position: 'before' or 'after' (default: 'after')
        target_paragraph_index: Optional paragraph index to use as anchor
        list_style: 'bullet' or 'number' to choose between bullet or numbered styles
    Returns:
        Status message
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Robust style selection for numbered/bulleted list
        style_name = None
        style_preference = []
        normalized_style = (list_style or 'bullet').lower()
        list_label = 'Bulleted'
        if normalized_style == 'number':
            style_preference = ['List Number', 'Numbered List', 'List Paragraph', 'Normal']
            list_label = 'Numbered'
        else:
            style_preference = ['List Bullet', 'Bullet List', 'List Paragraph', 'Normal']

        for candidate in style_preference:
            try:
                _ = doc.styles[candidate]
                style_name = candidate
                break
            except KeyError:
                continue
        if not style_name:
            style_name = None  # fallback to default
        new_paras = []
        for item in (list_items or []):
            p = doc.add_paragraph(item, style=style_name)
            new_paras.append(p)
        # Move the new paragraphs to the correct position
        for p in reversed(new_paras):
            if position == 'before':
                para._element.addprevious(p._element)
            else:
                para._element.addnext(p._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"{list_label} list inserted {position} paragraph (index {anchor_index})."
        else:
            return f"{list_label} list inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert numbered list: {str(e)}"


def is_toc_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de tabla de contenido (TOC)."""
    return para.style and para.style.name.upper().startswith("TOC")


def is_heading_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de encabezado (Heading 1, Heading 2, etc)."""
    return para.style and para.style.name.lower().startswith("heading")


# --- Helper: Get style name from a <w:p> element ---
def get_paragraph_style(el):
    from docx.oxml.ns import qn
    pPr = el.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None and 'w:val' in pStyle.attrib:
            return pStyle.attrib['w:val']
    return None

# --- Main: Delete everything under a header until next heading/TOC ---
def delete_block_under_header(doc, header_text):
    """
    Remove all elements (paragraphs, tables, etc.) after the header (by text) and before the next heading/TOC (by style).
    Returns: (header_element, elements_removed)
    """
    # Find the header paragraph by text (like delete_paragraph finds by index)
    header_para = None
    header_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == header_text.strip().lower():
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return None, 0
    
    # Find the next heading/TOC paragraph to determine the end of the block
    end_idx = None
    for i in range(header_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style and para.style.name.lower().startswith(('heading', 'título', 'toc')):
            end_idx = i
            break
    
    # If no next heading found, delete until end of document
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    
    # Remove paragraphs by index (like delete_paragraph does)
    removed_count = 0
    for i in range(header_idx + 1, end_idx):
        if i < len(doc.paragraphs):  # Safety check
            para = doc.paragraphs[header_idx + 1]  # Always remove the first paragraph after header
            p = para._p
            p.getparent().remove(p)
            removed_count += 1
    
    return header_para._p, removed_count

# --- Usage in replace_paragraph_block_below_header ---
def replace_paragraph_block_below_header(
    doc_path: str,
    header_text: str,
    new_paragraphs: list,
    detect_block_end_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Reemplaza todo el contenido debajo de una cabecera (por texto), hasta el siguiente encabezado/TOC (por estilo).
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    
    doc = Document(doc_path)
    
    # Find the header paragraph first
    header_para = None
    header_idx = None
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text.strip().lower()
        is_toc = is_toc_paragraph(para)
        if para_text == header_text.strip().lower() and not is_toc:
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return f"Header '{header_text}' not found in document."
    
    # Delete everything under the header using the same document instance
    header_el, removed_count = delete_block_under_header(doc, header_text)
    
    # Now insert new paragraphs after the header (which should still be in the document)
    style_to_use = new_paragraph_style or "Normal"
    
    # Find the header again after deletion (it should still be there)
    current_para = header_para
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        current_para._element.addnext(new_para._element)
        current_para = new_para
    
    doc.save(doc_path)
    return f"Replaced content under '{header_text}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {removed_count} elements."


def replace_block_between_manual_anchors(
    doc_path: str,
    start_anchor_text: str,
    new_paragraphs: list,
    end_anchor_text: str = None,
    match_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Replace all content (paragraphs, tables, etc.) between start_anchor_text and end_anchor_text (or next logical header if not provided).
    If end_anchor_text is None, deletes until next visually distinct paragraph (bold, all caps, or different font size), or end of document.
    Inserts new_paragraphs after the start anchor.
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    doc = Document(doc_path)
    body = doc.element.body
    elements = list(body)
    start_idx = None
    end_idx = None
    # Find start anchor
    paragraph_tag = qn('w:p')

    for i, el in enumerate(elements):
        if el.tag == paragraph_tag:
            p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
            if match_fn:
                if match_fn(p_text, el):
                    start_idx = i
                    break
            elif p_text == start_anchor_text.strip():
                start_idx = i
                break
    if start_idx is None:
        return f"Start anchor '{start_anchor_text}' not found."
    # Find end anchor
    if end_anchor_text:
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == paragraph_tag:
                p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
                if match_fn:
                    if match_fn(p_text, el, is_end=True):
                        end_idx = i
                        break
                elif p_text == end_anchor_text.strip():
                    end_idx = i
                    break
    else:
        # Heuristic: next visually distinct paragraph (bold, all caps, or different font size), or end of document
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == paragraph_tag:
                # Check for bold, all caps, or font size
                runs = [node for node in el.iter() if node.tag.endswith('}r')]
                for run in runs:
                    rpr = run.find(qn('w:rPr'))
                    if rpr is not None:
                        if rpr.find(qn('w:b')) is not None or rpr.find(qn('w:caps')) is not None or rpr.find(qn('w:sz')) is not None:
                            end_idx = i
                            break
                if end_idx is not None:
                    break
    # Mark elements for removal
    to_remove = []
    for i in range(start_idx + 1, end_idx if end_idx is not None else len(elements)):
        to_remove.append(elements[i])
    for el in to_remove:
        body.remove(el)
    doc.save(doc_path)
    # Reload and find start anchor for insertion
    doc = Document(doc_path)
    paras = doc.paragraphs
    anchor_idx = None
    for i, para in enumerate(paras):
        if para.text.strip() == start_anchor_text.strip():
            anchor_idx = i
            break
    if anchor_idx is None:
        return f"Start anchor '{start_anchor_text}' not found after deletion (unexpected)."
    anchor_para = paras[anchor_idx]
    style_to_use = new_paragraph_style or "Normal"
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        anchor_para._element.addnext(new_para._element)
        anchor_para = new_para
    doc.save(doc_path)
    return f"Replaced content between '{start_anchor_text}' and '{end_anchor_text or 'next logical header'}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {len(to_remove)} elements."
