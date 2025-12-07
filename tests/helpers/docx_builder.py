"""
Test helper for creating docx documents with proper w14:paraId attributes.

python-docx doesn't set w14:paraId on paragraphs by default, but our analysis
code depends on these IDs. This helper wraps Document and adds para_ids
automatically when adding paragraphs.

Usage:
    from tests.helpers.docx_builder import DocBuilder
    
    doc = DocBuilder()
    doc.add_paragraph("First paragraph")  # Auto-assigns w14:paraId
    doc.add_paragraph("Second paragraph")
    
    buffer = io.BytesIO()
    doc.save(buffer)
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml.ns import qn

from effilocal.doc.uuid_embedding import generate_para_id, set_paragraph_para_id

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph


class DocBuilder:
    """
    A wrapper around python-docx Document that auto-assigns w14:paraId.
    
    Every paragraph added via add_paragraph() will have a unique para_id
    and text_id set, matching what Word would create.
    """
    
    def __init__(self, docx_path: str | None = None):
        """
        Create a new test document.
        
        Args:
            docx_path: Optional path to existing docx to load.
        """
        self._doc = Document(docx_path) if docx_path else Document()
        self._existing_ids: set[str] = set()
        
        # Collect any existing para_ids from loaded document
        self._collect_existing_ids()
    
    def _collect_existing_ids(self) -> None:
        """Collect all existing w14:paraId values from the document."""
        for para in self._doc.paragraphs:
            para_id = para._element.get(qn("w14:paraId"))
            if para_id:
                self._existing_ids.add(para_id.upper())
    
    def add_paragraph(self, text: str = "", style: str | None = None) -> Paragraph:
        """
        Add a paragraph with auto-generated w14:paraId.
        
        Args:
            text: The paragraph text.
            style: Optional style name to apply.
            
        Returns:
            The created Paragraph object.
        """
        para = self._doc.add_paragraph(text, style)
        
        # Generate and set para_id
        para_id = generate_para_id(self._existing_ids)
        set_paragraph_para_id(para._element, para_id)
        self._existing_ids.add(para_id.upper())
        
        return para
    
    def add_heading(self, text: str, level: int = 1) -> Paragraph:
        """
        Add a heading with auto-generated w14:paraId.
        
        Args:
            text: The heading text.
            level: Heading level (0-9, where 0 is Title).
            
        Returns:
            The created Paragraph object.
        """
        para = self._doc.add_heading(text, level)
        
        # Generate and set para_id
        para_id = generate_para_id(self._existing_ids)
        set_paragraph_para_id(para._element, para_id)
        self._existing_ids.add(para_id.upper())
        
        return para
    
    def save(self, path_or_stream) -> None:
        """Save the document to a file or stream."""
        self._doc.save(path_or_stream)
    
    def to_bytes(self) -> bytes:
        """Save the document to bytes."""
        buffer = io.BytesIO()
        self._doc.save(buffer)
        return buffer.getvalue()
    
    @property
    def paragraphs(self):
        """Access underlying document paragraphs."""
        return self._doc.paragraphs
    
    @property
    def element(self):
        """Access underlying document element."""
        return self._doc.element


def create_test_doc_bytes(*paragraphs: str) -> bytes:
    """
    Convenience function to create a docx with paragraphs and return bytes.
    
    Args:
        paragraphs: Text for each paragraph to add.
        
    Returns:
        The docx file as bytes.
        
    Example:
        docx_bytes = create_test_doc_bytes(
            "INDEMNIFICATION. The party shall...",
            "LIMITATION OF LIABILITY. Subject to..."
        )
    """
    doc = DocBuilder()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc.to_bytes()
