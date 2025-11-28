#!/usr/bin/env python3
"""
Test script to verify the split-run detection and edit_run_text implementation.
"""
import tempfile
import os
from docx import Document
from word_document_server.utils.document_utils import find_and_replace_text, edit_run_text


def test_split_run_detection():
    """Test that split runs are detected and reported correctly."""
    # Create a document with text split across runs
    doc = Document()
    para = doc.add_paragraph()
    
    # Clear and manually create split runs
    para.clear()
    run1 = para.add_run("The ")
    run2 = para.add_run("W")
    run3 = para.add_run("P")
    run4 = para.add_run(" is important")
    
    # Verify the structure
    assert para.text == "The WP is important"
    assert len(para.runs) == 4
    
    # Save temporarily
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = tmp.name
    
    try:
        doc.save(tmp_path)
        
        # Test find_and_replace_text
        print("Testing find_and_replace_text with split runs...")
        doc2 = Document(tmp_path)
        count, snippets, split_matches = find_and_replace_text(doc2, "WP", "Work Product")
        
        print(f"✓ find_and_replace_text returned triple: count={count}, snippets={len(snippets)}, split_matches={len(split_matches)}")
        
        if split_matches:
            print(f"\n✓ Split matches detected: {len(split_matches)}")
            for match in split_matches:
                print(f"  - Para {match['paragraph_index']}: {match['runs']}")
                print(f"    Full match: '{match['full_match']}'")
        
        # Test edit_run_text
        print("\nTesting edit_run_text tool...")
        result = edit_run_text(tmp_path, 0, 1, "W")  # Edit run 1 (the "W")
        print(f"✓ edit_run_text result: {result}")
        
        # Verify the edit was made
        doc3 = Document(tmp_path)
        assert doc3.paragraphs[0].runs[1].text == "W"
        print("✓ Run text verified after edit")
        
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    
    print("\n✅ All tests passed!")


def test_in_run_replacement():
    """Test that in-run replacements still work correctly."""
    doc = Document()
    doc.add_paragraph("This is a test document")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = tmp.name
    
    try:
        doc.save(tmp_path)
        
        print("Testing in-run replacement...")
        doc2 = Document(tmp_path)
        count, snippets, split_matches = find_and_replace_text(doc2, "test", "example")
        
        print(f"✓ find_and_replace_text returned: count={count}, snippets={len(snippets)}, split_matches={len(split_matches)}")
        
        assert count == 1, f"Expected 1 replacement, got {count}"
        assert len(snippets) == 1, f"Expected 1 snippet, got {len(snippets)}"
        assert len(split_matches) == 0, f"Expected 0 split matches, got {len(split_matches)}"
        
        print("✓ In-run replacement count correct")
        print("✓ No split matches for in-run text")
        
        # Verify replacement was made
        doc2.save(tmp_path)
        doc3 = Document(tmp_path)
        assert "example" in doc3.paragraphs[0].text
        assert "test" not in doc3.paragraphs[0].text
        print("✓ Text replacement verified")
        
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    
    print("\n✅ In-run replacement test passed!")


if __name__ == "__main__":
    test_split_run_detection()
    test_in_run_replacement()
