#!/usr/bin/env python3
"""
Comprehensive tests for generate_review_example.py script.

These tests verify the complete workflow of extracting legal review examples
from email + document pairs, including:
- Email parsing and attachment extraction
- Comment categorization by recipient
- Track changes extraction
- Markdown generation
"""

import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

# Test data paths
TEST_DATA_DIR = Path(__file__).parent.parent / "EL_Precedents" / "analysis" / "quick_contract_review"
INSTRUCTIONS_MSG = TEST_DATA_DIR / "email_instructions.msg"
ADVICE_MSG = TEST_DATA_DIR / "email_advice.msg"


# =============================================================================
# Fixtures
# =============================================================================

@pytest.fixture
def sample_msg_paths() -> tuple[Path, Path]:
    """Return paths to sample MSG files for integration tests."""
    return INSTRUCTIONS_MSG, ADVICE_MSG


@pytest.fixture
def temp_output_dir() -> Path:
    """Create a temporary directory for test outputs."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


# =============================================================================
# Tests for EmailData dataclass
# =============================================================================

class TestEmailData:
    """Tests for EmailData dataclass structure."""
    
    def test_email_data_has_required_fields(self) -> None:
        """EmailData should have subject, sender, recipients, date, body, attachments."""
        from scripts.generate_review_example import EmailData
        
        email = EmailData(
            subject="Test Subject",
            sender="sender@test.com",
            recipients="recipient@test.com",
            date="2025-12-01",
            body="Test body",
            attachments=[]
        )
        
        assert email.subject == "Test Subject"
        assert email.sender == "sender@test.com"
        assert email.recipients == "recipient@test.com"
        assert email.date == "2025-12-01"
        assert email.body == "Test body"
        assert email.attachments == []


# =============================================================================
# Tests for Attachment dataclass
# =============================================================================

class TestAttachment:
    """Tests for Attachment dataclass structure."""
    
    def test_attachment_has_filename_and_data(self) -> None:
        """Attachment should have filename and binary data."""
        from scripts.generate_review_example import Attachment
        
        attachment = Attachment(
            filename="test.docx",
            data=b"test data"
        )
        
        assert attachment.filename == "test.docx"
        assert attachment.data == b"test data"
    
    def test_attachment_is_docx_property(self) -> None:
        """Attachment should have is_docx property."""
        from scripts.generate_review_example import Attachment
        
        docx_att = Attachment(filename="test.docx", data=b"")
        pdf_att = Attachment(filename="test.pdf", data=b"")
        
        assert docx_att.is_docx is True
        assert pdf_att.is_docx is False


# =============================================================================
# Tests for parse_msg_file function
# =============================================================================

class TestParseMsgFile:
    """Tests for parsing Outlook MSG files."""
    
    def test_parse_msg_file_returns_email_data(self, sample_msg_paths: tuple[Path, Path]) -> None:
        """parse_msg_file should return EmailData with populated fields."""
        from scripts.generate_review_example import parse_msg_file
        
        instructions_path, _ = sample_msg_paths
        email_data = parse_msg_file(instructions_path)
        
        assert email_data.subject is not None
        assert len(email_data.subject) > 0
        assert email_data.sender is not None
        assert email_data.body is not None
    
    def test_parse_msg_file_extracts_docx_attachments(self, sample_msg_paths: tuple[Path, Path]) -> None:
        """parse_msg_file should extract .docx attachments with data."""
        from scripts.generate_review_example import parse_msg_file
        
        instructions_path, _ = sample_msg_paths
        email_data = parse_msg_file(instructions_path)
        
        docx_attachments = [a for a in email_data.attachments if a.is_docx]
        
        assert len(docx_attachments) >= 1
        assert docx_attachments[0].data is not None
        assert len(docx_attachments[0].data) > 0
    
    def test_parse_msg_file_raises_for_nonexistent_file(self) -> None:
        """parse_msg_file should raise FileNotFoundError for missing files."""
        from scripts.generate_review_example import parse_msg_file
        
        with pytest.raises(FileNotFoundError):
            parse_msg_file(Path("/nonexistent/file.msg"))


# =============================================================================
# Tests for select_docx_attachment function
# =============================================================================

class TestSelectDocxAttachment:
    """Tests for selecting .docx attachments from email."""
    
    def test_auto_selects_single_docx(self) -> None:
        """Should auto-select when only one .docx attachment exists."""
        from scripts.generate_review_example import Attachment, select_docx_attachment
        
        attachments = [
            Attachment(filename="image.png", data=b"png"),
            Attachment(filename="contract.docx", data=b"docx"),
        ]
        
        selected = select_docx_attachment(attachments, "Select document:")
        
        assert selected.filename == "contract.docx"
    
    def test_uses_provided_index(self) -> None:
        """Should use preselected_index when provided."""
        from scripts.generate_review_example import Attachment, select_docx_attachment
        
        attachments = [
            Attachment(filename="first.docx", data=b"1"),
            Attachment(filename="second.docx", data=b"2"),
        ]
        
        selected = select_docx_attachment(attachments, "Select:", preselected_index=2)
        
        assert selected.filename == "second.docx"
    
    def test_raises_when_no_docx_attachments(self) -> None:
        """Should raise ValueError when no .docx attachments exist."""
        from scripts.generate_review_example import Attachment, select_docx_attachment
        
        attachments = [
            Attachment(filename="image.png", data=b"png"),
        ]
        
        with pytest.raises(ValueError, match="No .docx attachments found"):
            select_docx_attachment(attachments, "Select:")


# =============================================================================
# Tests for categorize_comments_by_prefix function
# =============================================================================

class TestCategorizeCommentsByPrefix:
    """Tests for categorizing comments by prefix."""
    
    def test_groups_comments_by_client_prefix(self, sample_msg_paths: tuple[Path, Path]) -> None:
        """Comments starting with 'For {client}:' should be in client group."""
        from effilocal.mcp_server.core.comments import extract_all_comments
        from scripts.generate_review_example import categorize_comments_by_prefix, parse_msg_file, select_docx_attachment
        
        _, advice_path = sample_msg_paths
        email_data = parse_msg_file(advice_path)
        docx_att = select_docx_attachment(email_data.attachments, "Select:")
        doc = Document(BytesIO(docx_att.data))
        
        all_comments = extract_all_comments(doc)
        # Function expects just the name, not "For name" - it builds "For X:" pattern internally
        categorized = categorize_comments_by_prefix(all_comments, "Didimo", "NBC")
        
        assert "client" in categorized
        assert "counterparty" in categorized
        assert len(categorized["client"]) > 0
        
        # All client comments should start with "For Didimo"
        for comment in categorized["client"]:
            assert comment["text"].startswith("For Didimo")
    
    def test_groups_comments_by_counterparty_prefix(self, sample_msg_paths: tuple[Path, Path]) -> None:
        """Comments starting with 'For {counterparty}:' should be in counterparty group."""
        from effilocal.mcp_server.core.comments import extract_all_comments
        from scripts.generate_review_example import categorize_comments_by_prefix, parse_msg_file, select_docx_attachment
        
        _, advice_path = sample_msg_paths
        email_data = parse_msg_file(advice_path)
        docx_att = select_docx_attachment(email_data.attachments, "Select:")
        doc = Document(BytesIO(docx_att.data))
        
        all_comments = extract_all_comments(doc)
        categorized = categorize_comments_by_prefix(all_comments, "Didimo", "NBC")
        
        assert len(categorized["counterparty"]) > 0
        
        # All counterparty comments should start with "For NBC"
        for comment in categorized["counterparty"]:
            assert comment["text"].startswith("For NBC")


# =============================================================================
# Tests for process_track_changes function
# =============================================================================

class TestProcessTrackChanges:
    """Tests for extracting tracked insertions and deletions."""
    
    def test_returns_list_of_changes(self, sample_msg_paths: tuple[Path, Path]) -> None:
        """process_track_changes should return a list of change dictionaries."""
        from scripts.generate_review_example import parse_msg_file, process_track_changes, select_docx_attachment
        
        _, advice_path = sample_msg_paths
        email_data = parse_msg_file(advice_path)
        docx_att = select_docx_attachment(email_data.attachments, "Select:")
        doc = Document(BytesIO(docx_att.data))
        
        changes = process_track_changes(doc)
        
        assert isinstance(changes, list)
    
    def test_change_has_required_fields(self, sample_msg_paths: tuple[Path, Path]) -> None:
        """Each change should be a ParagraphDiff with before_text and after_text."""
        from scripts.generate_review_example import ParagraphDiff, parse_msg_file, process_track_changes, select_docx_attachment
        
        _, advice_path = sample_msg_paths
        email_data = parse_msg_file(advice_path)
        docx_att = select_docx_attachment(email_data.attachments, "Select:")
        doc = Document(BytesIO(docx_att.data))
        
        changes = process_track_changes(doc)
        
        if len(changes) > 0:
            change = changes[0]
            # Now uses ParagraphDiff dataclass
            assert isinstance(change, ParagraphDiff)
            assert hasattr(change, "before_text")
            assert hasattr(change, "after_text")
            assert hasattr(change, "para_id")


# =============================================================================
# Tests for markdown generation functions
# =============================================================================

class TestGenerateInstructionsMd:
    """Tests for instructions markdown generation."""
    
    def test_includes_email_metadata(self) -> None:
        """Generated markdown should include From, To (redacted), Date, Subject."""
        from scripts.generate_review_example import Attachment, EmailData, generate_instructions_md
        
        email_data = EmailData(
            subject="Review Contract",
            sender="client@example.com",
            recipients="lawyer@example.com",
            date="2025-12-01",
            body="Please review attached.",
            attachments=[Attachment(filename="contract.docx", data=b"")]
        )
        
        markdown = generate_instructions_md(email_data)
        
        # Emails are now anonymized to [REDACTED]
        assert "[REDACTED]" in markdown
        assert "2025-12-01" in markdown
        assert "Review Contract" in markdown
    
    def test_includes_email_body(self) -> None:
        """Generated markdown should include email body text."""
        from scripts.generate_review_example import Attachment, EmailData, generate_instructions_md
        
        email_data = EmailData(
            subject="Subject",
            sender="sender@test.com",
            recipients="recipient@test.com",
            date="2025-12-01",
            body="This is the email body with instructions.",
            attachments=[]
        )
        
        markdown = generate_instructions_md(email_data)
        
        assert "This is the email body with instructions" in markdown


class TestGenerateCommentsMd:
    """Tests for comments markdown generation."""
    
    def test_separates_client_and_counterparty_sections(self) -> None:
        """Should have separate sections for client and counterparty comments."""
        from scripts.generate_review_example import generate_comments_md
        
        # Comments are now anonymized, so we use placeholder names
        categorized_comments = {
            "client": [{"text": "For [CLIENT]: note", "author": "Lawyer", "date": "2025-12-01", 
                       "reference_text": "ref"}],
            "counterparty": [{"text": "For [COUNTERPARTY]: note", "author": "Lawyer", "date": "2025-12-01",
                             "reference_text": "ref"}],
            "other": []
        }
        
        # generate_comments_md expects list of names, not single name
        markdown = generate_comments_md(categorized_comments, ["Client"], ["Counter"])
        
        # Check for section headers - they use placeholders now
        assert "[CLIENT]" in markdown
        assert "[COUNTERPARTY]" in markdown
    
    def test_includes_reference_text(self) -> None:
        """Each comment should show reference text."""
        from scripts.generate_review_example import generate_comments_md
        
        categorized_comments = {
            "client": [{
                "text": "For [CLIENT]: important note",
                "author": "Lawyer",
                "date": "2025-12-01",
                "reference_text": "the specific phrase"
            }],
            "counterparty": [],
            "other": []
        }
        
        markdown = generate_comments_md(categorized_comments, ["Client"], ["Counter"])
        
        # Reference text should be shown (paragraph_text was removed)
        assert "the specific phrase" in markdown


class TestGenerateTrackChangesMd:
    """Tests for track changes markdown generation."""
    
    def test_formats_paragraph_diffs(self) -> None:
        """Should format ParagraphDiff objects with before/after text."""
        from scripts.generate_review_example import ParagraphDiff, generate_track_changes_md
        
        # Now uses ParagraphDiff dataclass instead of dict
        diffs = [
            ParagraphDiff(
                before_text="original text",
                after_text="modified text",
                authors={"Lawyer"},
                dates={"2025-12-01"},
                insertions=["modified"],
                deletions=["original"],
                clause_number="",
                para_id="12345678"
            )
        ]
        
        markdown = generate_track_changes_md(diffs)
        
        # Should contain the diff content
        assert "original" in markdown or "modified" in markdown


class TestGenerateAdviceNoteMd:
    """Tests for advice note markdown generation."""
    
    def test_includes_email_metadata(self) -> None:
        """Generated markdown should include email metadata."""
        from scripts.generate_review_example import Attachment, EmailData, generate_advice_note_md
        
        email_data = EmailData(
            subject="RE: Review Contract",
            sender="lawyer@example.com",
            recipients="client@example.com",
            date="2025-12-02",
            body="Please find my comments attached.",
            attachments=[]
        )
        
        markdown = generate_advice_note_md(email_data)
        
        # Advice note includes sender/recipients (unlike instructions which uses [REDACTED])
        assert "lawyer@example.com" in markdown
        assert "client@example.com" in markdown
        assert "RE: Review Contract" in markdown


# =============================================================================
# Tests for main workflow
# =============================================================================

class TestMainWorkflow:
    """Integration tests for the complete workflow."""
    
    def test_generates_all_output_files(
        self, sample_msg_paths: tuple[Path, Path], temp_output_dir: Path
    ) -> None:
        """Main function should generate all output files."""
        from scripts.generate_review_example import generate_review_example
        
        instructions_path, advice_path = sample_msg_paths
        
        # Mock user input to select Didimo as client (option 2)
        with patch('builtins.input', return_value="2"):
            generate_review_example(
                incoming_path=instructions_path,
                outgoing_path=advice_path,
                output_dir=temp_output_dir
            )
        
        # Current output files (merged comments + edits into 03_review_edits.md)
        expected_files = [
            "00_mappings.md",
            "01_instructions.md",
            "02_original_agreement.md",
            "03_review_edits.md",  # Merged comments + track changes
            "04_advice_note.md"
        ]
        
        for filename in expected_files:
            output_file = temp_output_dir / filename
            assert output_file.exists(), f"Expected {filename} to be created"
            assert output_file.stat().st_size > 0, f"Expected {filename} to have content"
    
    def test_original_agreement_contains_document_text(
        self, sample_msg_paths: tuple[Path, Path], temp_output_dir: Path
    ) -> None:
        """02_original_agreement.md should contain document text."""
        from scripts.generate_review_example import generate_review_example
        
        instructions_path, advice_path = sample_msg_paths
        
        # Mock user input to select Didimo as client (option 2)
        with patch('builtins.input', return_value="2"):
            generate_review_example(
                incoming_path=instructions_path,
                outgoing_path=advice_path,
                output_dir=temp_output_dir
            )
        
        original_md = (temp_output_dir / "02_original_agreement.md").read_text(encoding="utf-8")
        
        # Should contain some agreement-related content
        assert len(original_md) > 100
    
    def test_review_edits_has_anonymized_parties(
        self, sample_msg_paths: tuple[Path, Path], temp_output_dir: Path
    ) -> None:
        """03_review_edits.md should have anonymized party placeholders."""
        from scripts.generate_review_example import generate_review_example
        
        instructions_path, advice_path = sample_msg_paths
        
        # Mock user input to select Didimo as client (option 2)
        with patch('builtins.input', return_value="2"):
            generate_review_example(
                incoming_path=instructions_path,
                outgoing_path=advice_path,
                output_dir=temp_output_dir
            )
        
        edits_md = (temp_output_dir / "03_review_edits.md").read_text(encoding="utf-8")
        
        # Should use anonymized placeholders (role-based like [SUPPLIER]/[CUSTOMER])
        # The exact placeholders depend on the detected party roles
        has_placeholders = (
            "[CLIENT]" in edits_md or 
            "[COUNTERPARTY]" in edits_md or
            "[SUPPLIER]" in edits_md or 
            "[CUSTOMER]" in edits_md
        )
        assert has_placeholders, "Expected anonymized party placeholders in output"


# =============================================================================
# Tests for CLI argument parsing
# =============================================================================

# =============================================================================
# Tests for _extract_clause_title_from_text function
# =============================================================================

class TestExtractClauseTitleFromText:
    """Tests for extracting clause titles from paragraph text."""
    
    def test_all_caps_multi_word_title(self) -> None:
        """Should extract multi-word ALL CAPS titles like 'LIMITATION OF LIABILITY.'"""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        assert _extract_clause_title_from_text("LIMITATION OF LIABILITY.") == "LIMITATION OF LIABILITY"
        assert _extract_clause_title_from_text("DATA SECURITY.") == "DATA SECURITY"
        assert _extract_clause_title_from_text("REPRESENTATIONS, WARRANTIES AND COVENANTS.") == "REPRESENTATIONS, WARRANTIES AND COVENANTS"
    
    def test_all_caps_single_word_title(self) -> None:
        """Should extract single-word ALL CAPS titles like 'INDEMNIFICATION.'"""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        assert _extract_clause_title_from_text("INDEMNIFICATION.") == "INDEMNIFICATION"
        assert _extract_clause_title_from_text("TERMINATION.") == "TERMINATION"
        assert _extract_clause_title_from_text("NOTICES.") == "NOTICES"
    
    def test_title_case_single_word(self) -> None:
        """Should extract single-word Title Case titles like 'Indemnification.'"""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        assert _extract_clause_title_from_text("Indemnification.") == "Indemnification"
        assert _extract_clause_title_from_text("Termination.") == "Termination"
        assert _extract_clause_title_from_text("Notices.") == "Notices"
    
    def test_title_case_multi_word(self) -> None:
        """Should extract multi-word Title Case titles like 'Limitation of Liability.'"""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        assert _extract_clause_title_from_text("Limitation of Liability.") == "Limitation of Liability"
        assert _extract_clause_title_from_text("Fees and Payment.") == "Fees and Payment"
        assert _extract_clause_title_from_text("Term of Agreement.") == "Term of Agreement"
    
    def test_title_with_ampersand(self) -> None:
        """Should handle titles with & symbol."""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        assert _extract_clause_title_from_text("FEES & EXPENSES.") == "FEES & EXPENSES"
        assert _extract_clause_title_from_text("Terms & Conditions.") == "Terms & Conditions"
    
    def test_title_followed_by_content(self) -> None:
        """Should extract title when followed by clause content."""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        result = _extract_clause_title_from_text("INDEMNIFICATION. The Vendor shall indemnify...")
        assert result == "INDEMNIFICATION"
        
        result = _extract_clause_title_from_text("Termination. Either party may terminate...")
        assert result == "Termination"
    
    def test_non_title_text_returns_empty(self) -> None:
        """Should return empty string for regular paragraph text."""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        assert _extract_clause_title_from_text("Some normal text without a title.") == ""
        assert _extract_clause_title_from_text("The parties agree to the following terms.") == ""
        assert _extract_clause_title_from_text("This Agreement is entered into as of...") == ""
    
    def test_empty_or_none_input(self) -> None:
        """Should handle empty or None input gracefully."""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        assert _extract_clause_title_from_text("") == ""
        assert _extract_clause_title_from_text(None) == ""
    
    def test_short_all_caps_rejected(self) -> None:
        """Should reject very short ALL CAPS text (less than 3 chars)."""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        # "AB." should not be matched as a title (too short)
        assert _extract_clause_title_from_text("AB.") == ""
        # But "ABC." should work
        assert _extract_clause_title_from_text("ABC.") == "ABC"
    
    def test_title_with_following_clause_content(self) -> None:
        """Should extract title from text like 'INDEMNIFICATION.  Vendor shall...'"""
        from scripts.generate_review_example import _extract_clause_title_from_text
        
        # This is the real-world case: reference_text starts with title followed by clause body
        text = "INDEMNIFICATION.  Vendor shall indemnify, defend, and hold Company"
        assert _extract_clause_title_from_text(text) == "INDEMNIFICATION"
        
        text = "DATA SECURITY. Supplier shall comply with Exhibit 3 to this Trial."
        assert _extract_clause_title_from_text(text) == "DATA SECURITY"


class TestCLIArgumentParsing:
    """Tests for command-line argument parsing."""
    
    def test_required_arguments(self) -> None:
        """CLI should require --incoming and --outgoing."""
        from scripts.generate_review_example import create_argument_parser
        
        parser = create_argument_parser()
        
        # Should raise on missing required args
        with pytest.raises(SystemExit):
            parser.parse_args([])
    
    def test_optional_index_arguments(self) -> None:
        """CLI should accept optional --original-index and --edited-index."""
        from scripts.generate_review_example import create_argument_parser
        
        parser = create_argument_parser()
        
        # Parties are now auto-detected, so --client/--counterparty no longer exist
        args = parser.parse_args([
            "--incoming", "in.msg",
            "--outgoing", "out.msg",
            "--original-index", "1",
            "--edited-index", "2"
        ])
        
        assert args.original_index == 1
        assert args.edited_index == 2
    
    def test_optional_single_file_flag(self) -> None:
        """CLI should accept --single-file flag."""
        from scripts.generate_review_example import create_argument_parser
        
        parser = create_argument_parser()
        
        args = parser.parse_args([
            "--incoming", "in.msg",
            "--outgoing", "out.msg",
            "--single-file"
        ])
        
        assert args.single_file is True
