#!/usr/bin/env python
"""Script to manage tracked changes (revisions) in Word documents.

Usage:
    manage_revisions.py get_revisions <docx_path>
    manage_revisions.py accept_revision <docx_path> <revision_id>
    manage_revisions.py reject_revision <docx_path> <revision_id>
    manage_revisions.py accept_all <docx_path>
    manage_revisions.py reject_all <docx_path>

Outputs JSON to stdout.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from docx import Document

from effilocal.mcp_server.core.revisions import (
    extract_all_revisions,
    accept_revision,
    reject_revision,
    accept_all_revisions,
    reject_all_revisions,
)


def get_revisions(docx_path: Path) -> dict:
    """Extract all revisions from a Word document."""
    if not docx_path.exists():
        return {"success": False, "error": f"File not found: {docx_path}"}
    
    try:
        doc = Document(str(docx_path))
        revisions = extract_all_revisions(doc)
        
        return {
            "success": True,
            "revisions": revisions,
            "total_revisions": len(revisions),
        }
    except Exception as e:
        return {"success": False, "error": f"Failed to extract revisions: {str(e)}"}


def accept_revision_cmd(docx_path: Path, revision_id: str) -> dict:
    """Accept a specific revision."""
    if not docx_path.exists():
        return {"success": False, "error": f"File not found: {docx_path}"}
    
    try:
        doc = Document(str(docx_path))
        success = accept_revision(doc, revision_id)
        
        if success:
            doc.save(str(docx_path))
            return {
                "success": True,
                "message": f"Revision {revision_id} accepted",
                "revision_id": revision_id,
            }
        else:
            return {
                "success": False,
                "error": f"Could not accept revision {revision_id}. Revision may not exist.",
                "revision_id": revision_id,
            }
    except PermissionError:
        return {
            "success": False,
            "error": f"Permission denied - file '{docx_path.name}' may be open in Word",
            "revision_id": revision_id,
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to accept revision: {str(e)}",
            "revision_id": revision_id,
        }


def reject_revision_cmd(docx_path: Path, revision_id: str) -> dict:
    """Reject a specific revision."""
    if not docx_path.exists():
        return {"success": False, "error": f"File not found: {docx_path}"}
    
    try:
        doc = Document(str(docx_path))
        success = reject_revision(doc, revision_id)
        
        if success:
            doc.save(str(docx_path))
            return {
                "success": True,
                "message": f"Revision {revision_id} rejected",
                "revision_id": revision_id,
            }
        else:
            return {
                "success": False,
                "error": f"Could not reject revision {revision_id}. Revision may not exist.",
                "revision_id": revision_id,
            }
    except PermissionError:
        return {
            "success": False,
            "error": f"Permission denied - file '{docx_path.name}' may be open in Word",
            "revision_id": revision_id,
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to reject revision: {str(e)}",
            "revision_id": revision_id,
        }


def accept_all_cmd(docx_path: Path) -> dict:
    """Accept all revisions in the document."""
    if not docx_path.exists():
        return {"success": False, "error": f"File not found: {docx_path}"}
    
    try:
        doc = Document(str(docx_path))
        result = accept_all_revisions(doc)
        
        if result['success']:
            doc.save(str(docx_path))
            return {
                "success": True,
                "message": f"Accepted {result['accepted_count']} revisions",
                "accepted_count": result['accepted_count'],
            }
        else:
            return {
                "success": False,
                "error": "Failed to accept all revisions",
                "errors": result.get('errors'),
            }
    except PermissionError:
        return {
            "success": False,
            "error": f"Permission denied - file '{docx_path.name}' may be open in Word",
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to accept all revisions: {str(e)}",
        }


def reject_all_cmd(docx_path: Path) -> dict:
    """Reject all revisions in the document."""
    if not docx_path.exists():
        return {"success": False, "error": f"File not found: {docx_path}"}
    
    try:
        doc = Document(str(docx_path))
        result = reject_all_revisions(doc)
        
        if result['success']:
            doc.save(str(docx_path))
            return {
                "success": True,
                "message": f"Rejected {result['rejected_count']} revisions",
                "rejected_count": result['rejected_count'],
            }
        else:
            return {
                "success": False,
                "error": "Failed to reject all revisions",
                "errors": result.get('errors'),
            }
    except PermissionError:
        return {
            "success": False,
            "error": f"Permission denied - file '{docx_path.name}' may be open in Word",
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to reject all revisions: {str(e)}",
        }


def main():
    parser = argparse.ArgumentParser(description="Manage tracked changes in Word documents")
    parser.add_argument("command", 
                        choices=["get_revisions", "accept_revision", "reject_revision", "accept_all", "reject_all"],
                        help="Command to execute")
    parser.add_argument("docx_path", type=Path, help="Path to .docx file")
    parser.add_argument("revision_id", nargs="?", help="Revision ID (for accept/reject individual)")

    args = parser.parse_args()

    if args.command == "get_revisions":
        result = get_revisions(args.docx_path)
    elif args.command == "accept_revision":
        if not args.revision_id:
            result = {"success": False, "error": "revision_id is required for accept_revision"}
        else:
            result = accept_revision_cmd(args.docx_path, args.revision_id)
    elif args.command == "reject_revision":
        if not args.revision_id:
            result = {"success": False, "error": "revision_id is required for reject_revision"}
        else:
            result = reject_revision_cmd(args.docx_path, args.revision_id)
    elif args.command == "accept_all":
        result = accept_all_cmd(args.docx_path)
    elif args.command == "reject_all":
        result = reject_all_cmd(args.docx_path)
    else:
        result = {"success": False, "error": f"Unknown command: {args.command}"}

    print(json.dumps(result))


if __name__ == "__main__":
    main()
