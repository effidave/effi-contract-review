from __future__ import annotations

import json
from pathlib import Path
from uuid import uuid4

from docx import Document

from effilocal.flows.analyze_doc import analyze


def _build_doc(tmp_path: Path) -> Path:
    doc = Document()

    clause = doc.add_paragraph("Supplier shall:")
    clause.style = "List Number"

    child = doc.add_paragraph("provide the services")
    child.style = "List Number 2"
    child2 = doc.add_paragraph("notify the customer of outages")
    child2.style = "List Number 2"

    bridge = doc.add_paragraph("The Supplier must keep the customer informed of major incidents.")
    bridge.paragraph_format.left_indent = clause.paragraph_format.left_indent

    doc_path = tmp_path / "clause_list_bridge.docx"
    doc.save(doc_path)
    return doc_path


def _load_blocks(out_dir: Path) -> list[dict[str, object]]:
    path = out_dir / "blocks.jsonl"
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def test_paragraph_after_list_is_continuation(tmp_path: Path) -> None:
    doc_path = _build_doc(tmp_path)
    out_dir = tmp_path / "analysis"
    analyze(doc_path, doc_id=str(uuid4()), out_dir=out_dir)

    blocks = _load_blocks(out_dir)
    clause = next(block for block in blocks if block["text"] == "Supplier shall:")
    bridge = next(block for block in blocks if block["text"] == "The Supplier must keep the customer informed of major incidents.")

    assert bridge["continuation_of"] == clause["id"]
    assert bridge["clause_group_id"] == clause["id"]
