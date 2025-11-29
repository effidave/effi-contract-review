"""Test hierarchical insertion - inserting after parent clause should skip children."""
import pytest
import tempfile
import shutil
from pathlib import Path
from docx import Document
from word_document_server.tools.content_tools import add_paragraph_after_clause
from effilocal.doc.numbering_inspector import NumberingInspector


@pytest.fixture
def fixtures_dir():
    """Return path to test fixtures directory."""
    return Path(__file__).parent / "fixtures"


@pytest.fixture
def temp_numbered_doc(fixtures_dir):
    """Create a temporary copy of the numbered document for testing."""
    source = fixtures_dir / "numbering_decimal.docx"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    shutil.copy2(source, tmp_path)
    yield tmp_path
    # Cleanup
    if tmp_path.exists():
        tmp_path.unlink()


@pytest.mark.asyncio
async def test_insert_after_parent_skips_children(temp_numbered_doc):
    """
    Test that inserting after clause 3 skips over subclauses 3.1, 3.1.1, 3.2.
    
    Document structure (from numbering_decimal.docx):
    1. Section 1
    2. Test paragraph
    3. Test paragraph
    3.1 Section 1.1
    3.1.1 Section 1.1.1
    3.2 Section 1.2
    4. Section 2
    
    When inserting after "3", should insert after "3.2" (before "4").
    """
    # First, verify the structure
    inspector = NumberingInspector.from_docx(temp_numbered_doc)
    rows, _ = inspector.analyze(debug=False)
    
    # Find clauses 3, 3.1, 3.1.1, 3.2, 4
    clause_indices = {}
    for row in rows:
        rendered = row.get("rendered_number", "").strip().rstrip(".")
        if rendered in ["3", "3.1", "3.1.1", "3.2", "4"]:
            clause_indices[rendered] = row.get("idx")
    
    print(f"Clause indices before insertion: {clause_indices}")
    
    # Insert after clause 3
    result = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "3",
        "New clause at level 3 (should become clause 4)",
        inherit_numbering=True
    )
    
    print(f"Result: {result}")
    assert "Paragraph added after clause '3'" in result
    
    # Re-analyze to see the new structure
    inspector = NumberingInspector.from_docx(temp_numbered_doc)
    rows, _ = inspector.analyze(debug=False)
    
    print("\nDocument structure after insertion:")
    for row in rows:
        rendered = row.get("rendered_number", "").strip()
        if rendered:
            idx = row.get("idx")
            text = row.get("text", "")[:50]
            ilvl = row.get("ilvl")
            print(f"  {rendered:10} (idx={idx}, ilvl={ilvl}): {text}")
    
    # Find our new paragraph
    new_para_found = False
    new_para_index = None
    for row in rows:
        text = row.get("text", "")
        if "New clause at level 3" in text:
            new_para_found = True
            new_para_index = row.get("idx")
            new_para_rendered = row.get("rendered_number", "").strip()
            print(f"\nNew paragraph found at index {new_para_index} with number '{new_para_rendered}'")
            break
    
    assert new_para_found, "New paragraph not found in document"
    
    # The new paragraph should appear after clause 3.2 (index should be > 3.2's index)
    assert new_para_index > clause_indices["3.2"], \
        f"New paragraph at index {new_para_index} should be after 3.2 at index {clause_indices['3.2']}"
    
    # It should appear before or at the original position of clause 4
    assert new_para_index <= clause_indices["4"] + 1, \
        f"New paragraph at index {new_para_index} should be before/at original clause 4 position {clause_indices['4']}"


@pytest.mark.asyncio
async def test_insert_after_leaf_clause(temp_numbered_doc):
    """
    Test that inserting after a leaf clause (no children) works normally.
    
    Inserting after "3.1.1" (which has no children) should insert immediately after it.
    """
    # Insert after clause 3.1.1 (leaf node)
    result = await add_paragraph_after_clause(
        str(temp_numbered_doc),
        "3.1.1",
        "New content after 3.1.1",
        inherit_numbering=True
    )
    
    assert "Paragraph added after clause '3.1.1'" in result
    
    # Verify insertion
    doc = Document(str(temp_numbered_doc))
    found = False
    for para in doc.paragraphs:
        if "New content after 3.1.1" in para.text:
            found = True
            break
    assert found, "New paragraph not found"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
